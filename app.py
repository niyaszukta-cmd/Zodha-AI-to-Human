import streamlit as st
import streamlit.components.v1 as components
import requests
import re
import time
import math
import json
import os
import io
from datetime import datetime
try:
    import pdfplumber
    PDF_OK = True
except ImportError:
    PDF_OK = False
try:
    from openpyxl import Workbook
    from openpyxl.styles import (Font, PatternFill, Alignment,
                                  Border, Side, GradientFill)
    from openpyxl.utils import get_column_letter
    XLSX_OK = True
except ImportError:
    XLSX_OK = False

# ══════════════════════════════════════════════════════════════════════════════
# CONFIG
# ══════════════════════════════════════════════════════════════════════════════

# ── Provider model lists ─────────────────────────────────────────────────────
OPENROUTER_MODELS = {
    "anthropic/claude-haiku-4-5":       "Claude Haiku 4.5 · Best undetection ⭐",
    "anthropic/claude-3-5-haiku":       "Claude 3.5 Haiku · Fast & strong",
    "anthropic/claude-3-haiku":         "Claude 3 Haiku · Ultra cheap",
    "google/gemini-flash-1.5":          "Gemini Flash 1.5 · Free tier available",
    "meta-llama/llama-3.3-70b-instruct:free": "Llama 3.3 70B · Free ✓",
    "mistralai/mistral-7b-instruct:free":     "Mistral 7B · Free ✓",
}

GROQ_MODELS = {
    "llama-3.3-70b-versatile": "Llama 3.3 70B · Best quality",
    "llama-3.1-8b-instant":    "Llama 3.1 8B · Fastest",
    "mixtral-8x7b-32768":      "Mixtral 8x7B · Long context",
    "gemma2-9b-it":            "Gemma 2 9B · Balanced",
}

# Active provider — set by admin ("groq" | "openrouter")
_DEFAULT_PROVIDER = "openrouter"

# Admin password — set via environment variable or change here
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "zodha2026")
# Both provider keys read from env — admin can also override via Admin panel
_GROQ_ENV_KEY = os.environ.get("GROQ_API_KEY", "")
_OR_ENV_KEY   = os.environ.get("OPENROUTER_API_KEY", "")

# ══════════════════════════════════════════════════════════════════════════════
# ZERO-DEPENDENCY READABILITY ENGINE
# ══════════════════════════════════════════════════════════════════════════════

def _count_syllables(word):
    word = word.lower().strip('.,!?;:\'\'"()')
    if not word: return 0
    if len(word) <= 3: return 1
    word = re.sub(r'(?:[^laeiouy]es|ed|[^laeiouy]e)$', '', word)
    word = re.sub(r'^y', '', word)
    return max(1, len(re.findall(r'[aeiouy]{1,2}', word)))

def _flesch_reading_ease(text):
    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
    words = re.findall(r'\b[a-zA-Z]+\b', text)
    if not sentences or not words: return 50.0
    syllables = sum(_count_syllables(w) for w in words)
    score = 206.835 - 1.015*(len(words)/len(sentences)) - 84.6*(syllables/len(words))
    return round(max(0.0, min(100.0, score)), 1)

def _flesch_kincaid_grade(text):
    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
    words = re.findall(r'\b[a-zA-Z]+\b', text)
    if not sentences or not words: return 10.0
    syllables = sum(_count_syllables(w) for w in words)
    grade = 0.39*(len(words)/len(sentences)) + 11.8*(syllables/len(words)) - 15.59
    return round(max(0.0, grade), 1)

def _sent_tokenize(text):
    parts = re.split(r'(?<=[.!?])\s+', text.strip())
    return [p for p in parts if p.strip()] or [text]

def _word_tokenize(text):
    return re.findall(r'\b[a-zA-Z]+\b', text.lower())

def compute_scores(text, style="Conversational"):
    """Compute humanness scores with style-aware weighting."""
    if not text or not text.strip(): return {}
    sentences    = _sent_tokenize(text)
    words_alpha  = _word_tokenize(text)
    word_count   = len(words_alpha)
    sent_count   = max(len(sentences), 1)
    flesch       = _flesch_reading_ease(text)
    unique_words = set(words_alpha)
    ttr          = (len(unique_words)/word_count*100) if word_count > 0 else 0

    sent_lengths = [len(_word_tokenize(s)) for s in sentences]
    sl_variation = 0.0
    if len(sent_lengths) > 1:
        mean_sl = sum(sent_lengths)/len(sent_lengths)
        sl_variation = min(100, math.sqrt(sum((l-mean_sl)**2 for l in sent_lengths)/len(sent_lengths))*5)
    avg_sent_len = word_count/sent_count

    burstiness = 0.0
    if len(sent_lengths) > 2:
        diffs = [abs(sent_lengths[i]-sent_lengths[i-1]) for i in range(1, len(sent_lengths))]
        burstiness = min(100,(sum(diffs)/len(diffs))*4)

    contractions = len(re.findall(
        r"\b(i'm|you're|he's|she's|it's|we're|they're|i've|you've|we've|they've|"
        r"i'd|you'd|he'd|she'd|we'd|they'd|i'll|you'll|he'll|she'll|we'll|they'll|"
        r"isn't|aren't|wasn't|weren't|don't|doesn't|didn't|won't|wouldn't|can't|"
        r"couldn't|shouldn't|haven't|hasn't|hadn't|that's|there's|here's|let's)\b",
        text.lower()))
    contraction_score = min(100,(contractions/max(sent_count,1))*40)

    first_person  = len(re.findall(r'\b(i|me|my|myself|we|our|us)\b', text.lower()))
    fp_score      = min(100,(first_person/max(word_count,1))*500)

    passive_count = len(re.findall(r'\b(is|are|was|were|be|been|being)\s+\w+ed\b', text.lower()))
    passive_score = max(0, 100-(passive_count/max(sent_count,1))*60)

    # Extended transition words — includes academic discourse markers
    transition_words = [
        'however','therefore','moreover','furthermore','although',
        'despite','meanwhile','consequently','additionally','nevertheless',
        'on the other hand','in contrast','for instance','in other words',
        'as a result','similarly','in fact','of course','after all',
        'notably','crucially','significantly','interestingly','importantly',
        'that said','more importantly','in practice','taken together',
        'what is striking','perhaps most','it appears','the evidence suggests',
        'one may argue','this suggests','this implies','this indicates'
    ]
    transitions      = sum(1 for t in transition_words if t in text.lower())
    transition_score = min(100, transitions * 6)

    # Sentence opener variety (human writers vary their openers)
    if sent_lengths:
        openers = []
        for s in sentences[:min(10, len(sentences))]:
            words = _word_tokenize(s)
            if words: openers.append(words[0])
        unique_openers = len(set(openers))
        opener_variety = min(100, (unique_openers / max(len(openers), 1)) * 100)
    else:
        opener_variety = 50

    grade = _flesch_kincaid_grade(text)

    # ── Style-aware weighting ──────────────────────────────────────────────
    # Academic text has no contractions/first-person by design — don't penalise
    if style in ("Academic", "Professional"):
        humanness = round(min(100, max(0,
            sl_variation     * 0.28 +   # most important for academic
            burstiness       * 0.22 +
            passive_score    * 0.15 +
            transition_score * 0.15 +
            ttr              * 0.10 +
            opener_variety   * 0.10
        )), 1)
    else:
        humanness = round(min(100, max(0,
            sl_variation      * 0.20 +
            burstiness        * 0.18 +
            ttr               * 0.12 +
            contraction_score * 0.12 +
            transition_score  * 0.12 +
            passive_score     * 0.10 +
            opener_variety    * 0.10 +
            fp_score          * 0.06
        )), 1)

    return {"humanness":humanness,"flesch":round(flesch,1),"ttr":round(ttr,1),
            "sl_variation":round(sl_variation,1),"burstiness":round(burstiness,1),
            "contraction_score":round(contraction_score,1),"passive_score":round(passive_score,1),
            "transition_score":round(transition_score,1),"word_count":word_count,
            "sent_count":sent_count,"avg_sent_len":round(avg_sent_len,1),
            "grade_level":round(grade,1),"unique_words":len(unique_words),
            "opener_variety":round(opener_variety,1)}

def score_color(score):
    if score>=75:   return("#1a3a2e","#6fcf97","Excellent")
    elif score>=55: return("#2d3a1e","#b5d97a","Good")
    elif score>=35: return("#3a2d1a","#e8c97a","Fair")
    else:           return("#3a1a1a","#e87a7a","Needs Work")

def render_score_ring(score, label):
    bg,fg,grade_lbl = score_color(score)
    return f"""<div class="score-ring-wrap">
      <div class="score-ring" style="background:{bg};color:{fg};box-shadow:0 0 0 3px {fg}30,0 4px 16px rgba(0,0,0,0.2);">
        {score:.0f}</div>
      <div class="score-label">{label}</div>
      <div class="score-label" style="color:{fg};font-size:0.68rem;">{grade_lbl}</div>
    </div>"""

def render_metrics(sc):
    chips=[("Words",sc.get("word_count",0)),("Sentences",sc.get("sent_count",0)),
           ("Avg Len",sc.get("avg_sent_len",0)),("Grade",sc.get("grade_level",0)),
           ("Unique",sc.get("unique_words",0)),("Flesch",sc.get("flesch",0)),
           ("Lex Div%",sc.get("ttr",0)),("Rhythm",sc.get("sl_variation",0))]
    html='<div class="metric-row">'
    for name,val in chips:
        html+=f'<div class="metric-chip"><b>{val}</b> {name}</div>'
    return html+"</div>"

def make_copy_btn(copy_id, text, label="📋 Copy", color="#ffffff", bg="#1e5c22", border="#3a8c3f"):
    import html as _html
    safe = _html.escape(text, quote=True)
    html_src = f"""<!DOCTYPE html><html><body style="margin:0;padding:0;background:transparent;">
<textarea id="cp" readonly style="position:absolute;left:-9999px;top:0;width:1px;height:1px;">{safe}</textarea>
<button id="btn" onclick="var el=document.getElementById('cp');el.select();el.setSelectionRange(0,99999);
  var ok=false;try{{ok=document.execCommand('copy');}}catch(e){{}}
  if(!ok&&navigator.clipboard){{navigator.clipboard.writeText(el.value).then(function(){{
    document.getElementById('btn').innerHTML='✅ Copied!';
    setTimeout(function(){{document.getElementById('btn').innerHTML='{label}';}},2000);}});}}
  else{{document.getElementById('btn').innerHTML='✅ Copied!';
    setTimeout(function(){{document.getElementById('btn').innerHTML='{label}';}},2000);}}"
  style="background:{bg};color:{color};border:1.5px solid {border};border-radius:8px;
         padding:0.4rem 1.1rem;cursor:pointer;font-size:0.82rem;
         font-family:'DM Sans',sans-serif;font-weight:500;transition:all 0.2s;white-space:nowrap;">
  {label}</button></body></html>"""
    components.html(html_src, height=46, scrolling=False)

# ══════════════════════════════════════════════════════════════════════════════
# AI PROMPTS
# ══════════════════════════════════════════════════════════════════════════════

# ── Anti-AI pattern list injected into every prompt ─────────────────────
_AI_PATTERNS_TO_AVOID = """
HOW AI DETECTORS CATCH YOU — avoid ALL of these to pass:

PERPLEXITY KILLERS (AI always picks the most predictable word — don't):
✗ "play a crucial role" → say "shape", "drive", "define", "remake", "sit at the heart of"
✗ "in today's rapidly changing world" → cut entirely or anchor to a specific year/event
✗ "it is important to note" / "it is worth mentioning" → delete these entirely
✗ "has been shown to" / "studies have shown" → "research confirms", "the data tell a different story", "the numbers point elsewhere"
✗ "due to the fact that" → "because", "since", "given that"
✗ "in order to" → "to"
✗ "a wide range of" / "a variety of" → pick specific ones or say "several competing"
✗ "state-of-the-art" / "cutting-edge" / "innovative solution" → be specific about what is new
✗ "leverage" (as a verb) → "use", "apply", "draw on", "exploit"
✗ "utilize" → "use"
✗ "facilitate" → "help", "enable", "make possible", "open the door to"
✗ "Furthermore," / "Moreover," / "Additionally," / "In conclusion," → banned entirely
✗ "This paper/study/article examines/explores/investigates" at sentence start
✗ "It is clear that" / "Clearly," → let the evidence speak without announcing it
✗ "Overall," as a paragraph opener
✗ "In summary," or "To summarize," → find a real closing thought instead

BURSTINESS KILLERS (AI has uniform sentence lengths — be wildly varied):
✗ Three or more consecutive sentences between 15-25 words → must break this
✗ Paragraphs where every sentence has similar clause depth
✗ Never having a sentence shorter than 8 words
✗ Never having a sentence longer than 30 words

STRUCTURAL AI TELLS:
✗ Opening a paragraph by stating its topic, then supporting it, then concluding it — too clean
✗ Every paragraph the same length (AI loves 3-5 sentence paragraphs uniformly)
✗ Parallel construction in every sentence of a list
✗ Same subject noun repeated in consecutive sentences instead of pronoun/synonym
✗ Adjective before noun before verb patterns repeated
"""

# ── Style-specific few-shot transformation examples ───────────────────────
_FEW_SHOT = {
    "Academic": """
TRANSFORMATION EXAMPLE — Academic (target: undetectable as AI):
BEFORE: "The study examines the impact of climate change on biodiversity. The research uses quantitative methods. The findings show significant correlations. The results indicate that action is needed."

AFTER: "How climate change dismantles biodiversity — not uniformly, but unevenly, and with a kind of ecological specificity that aggregated models routinely miss — is precisely what this study sets out to trace. Using regression-based modelling of species distribution data across 47 temperate zones, the analysis uncovers a statistically significant compression of range boundaries. Temperature anomalies exceeding 1.5°C correlate strongly with habitat contraction. The implications are not abstract. Without substantive policy intervention, species loss will compound in ways that current conservation frameworks are structurally unprepared to address."

WHY THIS PASSES DETECTORS: Sentence lengths: 42, 18, 8, 4, 19 words — extreme burstiness. Unexpected opener (question/fragment). "Unevenly, and with a kind of" is low-probability phrasing. Short declarative "The implications are not abstract" breaks paragraph rhythm. No forbidden transitions. "Structurally unprepared" is low-frequency but precise.
""",
    "Conversational": """
TRANSFORMATION EXAMPLE — Conversational (target: undetectable as AI):
BEFORE: "Artificial intelligence is transforming many industries. It is being used in healthcare and finance. The technology enables better decision making. This has many benefits for organizations."

AFTER: "Nobody predicted it would happen this fast. AI isn't just nudging industries forward — it's rewriting the rulebook while people are still reading the previous edition. Healthcare is probably the starkest example: diagnostic tools that used to require specialist teams can now flag anomalies before a doctor has reviewed the chart. Finance runs a close second. And the thing that gets lost in headlines about efficiency gains is how much this shifts the actual texture of decision-making. Not just speed. The whole character of the judgment call."

WHY THIS PASSES DETECTORS: Opens with 5-word sentence nobody would predict. Em-dash mid-sentence disrupts rhythm. "Rewriting the rulebook while people are still reading the previous edition" is high-information, low-probability. Final two sentences are fragments (3 words, 6 words) — extreme burstiness. No banned transitions anywhere.
""",
    "Professional": """
TRANSFORMATION EXAMPLE — Professional (target: undetectable as AI):
BEFORE: "The company implemented new strategies. The strategies improved performance. The results were positive. The organization benefited from these changes."

AFTER: "The strategic repositioning completed in Q3 produced visible gains — three divisions posted margin improvement within two quarters of implementation. That said, headline numbers flatter the picture somewhat. The performance uplift traces largely to cost discipline rather than revenue growth, which is a meaningful distinction for anyone thinking about durability. Competitive advantage built on efficiency has a ceiling. The organisation knows this. The harder work — building demand-side momentum — hasn't really started yet."

WHY THIS PASSES DETECTORS: "flatter the picture somewhat" is uncommon, natural phrasing. "For anyone thinking about" is conversational register in professional text — deliberate mix. "Competitive advantage built on efficiency has a ceiling" is an 8-word punch after longer sentences. Two consecutive short declarations followed by longer qualifying sentence — human rhythm pattern.
""",
    "Journalistic": """
TRANSFORMATION EXAMPLE — Journalistic (target: undetectable as AI):
BEFORE: "The economy is experiencing significant changes. Many factors are contributing to these changes. Experts believe this will continue. There are both positive and negative effects."

AFTER: "Something is breaking in the economic consensus — and it isn't obvious yet what replaces it. Interest rates, supply-chain fractures, and a productivity surge nobody quite saw coming have converged in 2024 in ways that leave most forecasting models looking a bit embarrassed. The result is a labour market economists can't agree on. Tight in some sectors, slack in others, strange everywhere. For workers in technology or advanced manufacturing, the outlook remains enviable. For the rest — and it's worth remembering the rest is most people — conditions are harder to read, and harder to live through."

WHY THIS PASSES DETECTORS: "a bit embarrassed" applied to forecasting models is unexpected metaphoric attribution. "Strange everywhere." is a 2-word paragraph. Parenthetical "and it's worth remembering the rest is most people" breaks formal structure. Mixed clause depths in final sentence. No banned transitions.
""",
    "Creative": """
TRANSFORMATION EXAMPLE — Creative (target: undetectable as AI):
BEFORE: "The city was busy with many people. There was a lot of noise. The streets were crowded. It was an interesting place to be."

AFTER: "The city had a metabolism. You felt it before you saw it — the pressure of ten thousand simultaneous decisions pressing against your skin like weather. Horns. A child's shoe on the pavement, somehow separated from its owner. Someone on a phone saying the same word three times, loudly, then laughing. You didn't visit a place like this so much as get metabolised by it, broken down and reconstituted as something slightly different, carrying fragments of it home without meaning to."

WHY THIS PASSES DETECTORS: "had a metabolism" is surprising but precise — low probability. "Pressing against your skin like weather" is novel simile. Noun-only sentence fragments. The final sentence at 43 words after short fragments creates maximum burstiness. "Broken down and reconstituted" medical register in creative context — unexpected vocabulary register shift.
""",
}

STYLE_PROMPTS = {
    "Academic": (
        "You are a senior academic who has published in Nature, The Lancet, and top economics journals. "
        "You have been handed AI-generated text to transform into prose that a respected human scholar "
        "would actually write. Your ONLY goal: defeat AI detection while preserving every fact and term. "
        "\n\nHOW TO DEFEAT ACADEMIC AI DETECTORS:"
        "\n• PERPLEXITY INJECTION: Choose low-probability but correct academic words. "
        "'Unsettles' not 'challenges'. 'Traces' not 'examines'. 'Sits awkwardly with' not 'contradicts'. "
        "'Dismantles' not 'affects'. These are what detectors miss — correct but not LLM-default."
        "\n• EXTREME BURSTINESS: Every paragraph must have one sentence of 5-8 words AND one of 38-50 words. "
        "NEVER two consecutive sentences of similar length."
        "\n• SCHOLAR'S INTRUSIONS: 'What is striking here is', 'The picture that emerges is', "
        "'This, of course, is not the whole story', 'A more careful reading suggests', 'Curiously,', "
        "'What this finding cannot tell us is'. Human scholars use these. LLMs systematically avoid them."
        "\n• LATE-ARRIVING MAIN CLAUSE: Build sentences where the point arrives after qualifiers. "
        "'Although X, and notwithstanding Y, despite Z, the analysis nonetheless finds W.'"
        "\n• HEDGED PRECISION: 'The data tentatively indicate', 'one may cautiously argue', "
        "'this finding, provisional as it is, points toward'"
        "\n• SINGLE-SENTENCE PARAGRAPH PUNCH: At least one paragraph must end with a "
        "1-sentence paragraph after a long analytical one. The contrast is distinctly human."
        "\n• STRICTLY BANNED: contractions, 'Furthermore', 'Moreover', 'Additionally', "
        "'In conclusion', 'It is important to note', 'It is worth noting', "
        "'plays a crucial role', 'state-of-the-art', 'The study examines'"
        "\n\nPreserve 100% of meaning, all data, all citations, all technical terms."
    ),
    "Conversational": (
        "You are rewriting AI text to sound like a knowledgeable person talking to a friend. "
        "Target: score under 5% AI probability on GPTZero and Originality.ai. "
        "\n\nEXACT TECHNIQUES THAT DEFEAT CONVERSATIONAL DETECTORS:"
        "\n• CONTRACTIONS: Mandatory. Every 2nd or 3rd sentence minimum. "
        "it's, don't, you're, can't, that's, isn't, haven't, wouldn't — use them freely"
        "\n• UNPREDICTABLE OPENERS: Rotate through — "
        "'Nobody expected...', 'Here's what actually happened:', 'The thing is,', "
        "'Which raises an obvious question.', 'And that's not even the strange part.', "
        "'Think about what that actually means.', 'Most people get this wrong.'"
        "\n• FRAGMENT SENTENCES: Drop 2-4 word sentences deliberately. "
        "'Not quite.' 'Worth pausing on.' 'Fair enough.' 'That's the real question.'"
        "\n• EM-DASH DISRUPTION: Use em-dashes to break sentence flow mid-thought — "
        "the way a real person would if they corrected themselves or added emphasis."
        "\n• REGISTER MARKERS: 'And honestly,', 'To be fair,', 'Here's the kicker —', "
        "'What's easy to miss here is', 'But this is where it gets interesting.'"
        "\n• WILDLY VARIED LENGTHS: 3 words. Then a sentence that spirals outward for thirty or "
        "forty words because that's how people think when excited about something and want to make "
        "sure you follow every step of the logic. Then short again."
        "\n• NEVER USE: 'Furthermore', 'Moreover', 'Additionally', 'In conclusion', "
        "'plays a crucial role', 'it is important to note', 'it is worth mentioning'"
        "\nPreserve all original meaning and facts."
    ),
    "Professional": (
        "You write for Harvard Business Review and senior strategy memos. Transform this into professional "
        "prose that a real senior executive would write — confident, nuanced, never robotic. "
        "Target: undetectable by Copyleaks and Originality.ai. "
        "\n\nTECHNIQUES THAT FOOL PROFESSIONAL DETECTORS:"
        "\n• ANALYTICAL HEDGING: "
        "'The numbers are encouraging, though they don't tell the full story.', "
        "'That said, the picture is more complicated than it first appears.', "
        "'A distinction worth drawing here:', 'The harder question is', "
        "'Which raises an uncomfortable implication:'"
        "\n• SHORT PUNCHES: After 2-3 long analytical sentences, drop a 6-8 word declarative. "
        "'The data bear this out.' 'That margin matters.' 'The risk is real.' "
        "This is what human executives actually write."
        "\n• ACTIVE + SPECIFIC: 'Revenue grew 12% in the APAC division' not 'Growth was observed'. "
        "'The team delivered' not 'It was delivered by the team.'"
        "\n• HUMAN TRANSITIONS ONLY: "
        "'That said,', 'The implication is', 'This matters for one reason:', "
        "'Taken together, these points suggest', 'In practice,', 'At its core,', "
        "'Which is why' — NEVER 'Furthermore', 'Moreover', 'Additionally'"
        "\n• ONE COLLOQUIAL PHRASE PER PARAGRAPH: 'the short answer is', 'frankly', "
        "'in plain terms', 'to put it directly'. Breaks LLM uniformity of register."
        "\n• OCCASIONAL CONTRACTION: 'it's', 'doesn't', 'that's' — once per paragraph."
        "\nPreserve every fact, metric, and key idea."
    ),
    "Journalistic": (
        "You are a senior correspondent — The Economist, The Atlantic, Foreign Affairs — who has been "
        "handed AI copy and needs to make it undetectable before filing. Target: under 3% AI probability. "
        "\n\nMOVES THAT DEFEAT JOURNALISTIC DETECTORS:"
        "\n• FIRST SENTENCE: 8-12 words maximum. Specific, surprising, impossible to ignore. "
        "NOT 'This article examines' or 'In recent years'. Something that earns the second sentence."
        "\n• THE ECONOMIST RHYTHM: Long sentence building complexity and evidence and qualification. "
        "Short sentence landing the point. Long again. Short. Never three sentences of similar length."
        "\n• SPECIFIC OVER VAGUE ALWAYS: Not 'many experts' but 'the IMF's 2024 report'. "
        "Not 'significant changes' but 'a 14-point swing'. "
        "Not 'some industries' but 'semiconductor fabrication and pharmaceutical logistics'. "
        "Specificity is what LLMs avoid — a human journalist uses it."
        "\n• JOURNALIST'S TOOLKIT: "
        "'The result?' (one-sentence paragraph). 'This is not accidental.' "
        "'Consider what this actually means for'... 'The numbers tell a different story.' "
        "'That conclusion sits uneasily with the data.' "
        "'Which raises the question nobody has yet answered satisfactorily:'"
        "\n• END ON TENSION: Last sentence of every paragraph leaves something unresolved "
        "or questioned. AI paragraphs wrap up cleanly. Human journalism doesn't."
        "\n• COLONS AND DASHES: 'There is one number that matters here: 12.' "
        "'The answer — if there is one — lies elsewhere.'"
        "\nKeep every fact, statistic, name, and date from the original."
    ),
    "Creative": (
        "You write like Joan Didion on a deadline, with Orwell's clarity and a rhythm that sounds like "
        "no one else's. Transform this AI text into something a reader would remember and a detector "
        "would never flag. Target: 0% AI probability. "
        "\n\nCREATIVE GHOST-WRITING TECHNIQUES:"
        "\n• RHYTHM AS WEAPON: Short. Then one that runs and coils and loops, accumulating detail "
        "and qualification the way memory actually works, not the way a language model processes tokens. "
        "Short again. The contrast is everything."
        "\n• NOVEL METAPHOR ONLY: "
        "'The report had a metabolism.' 'The data breathed in short, anxious intervals.' "
        "'The argument was architecturally sound and emotionally empty.' "
        "NEVER: 'navigate the landscape', 'shed light on', 'moving the needle'"
        "\n• FRAGMENT SENTENCES: Deploy deliberately. One word, even. "
        "'Silence.' 'Barely.' 'Not quite.' 'That's not it.' "
        "Fragments almost never appear in AI output — use them."
        "\n• SYNTACTIC SURPRISE: Start sentences with 'What' or 'Which' or 'How' as noun clause: "
        "'What nobody said out loud was the more interesting story.' "
        "'How it ended was less surprising than how long it took.'"
        "\n• SECOND PERSON PULL: 'You felt it before you saw it.' "
        "'You could spend a long time looking at those numbers and not know what they meant.'"
        "\n• ONE-SENTENCE PARAGRAPH MID-PIECE: Changes everything. Human writers do this. LLMs don't."
        "\n• VOCABULARY REGISTER MIXING: Literary + technical, colloquial + formal. "
        "Detectors look for register consistency — humans don't have it."
        "\nPreserve every idea and meaning from the original."
    ),
}

PARAPHRASE_MODES = {
    "Standard": "Rewrite using completely different words and structures while preserving exact meaning. Output ONLY the paraphrased text.",
    "Simplify": "Rewrite in simpler, clearer language. Shorter sentences, common words, active voice. Output ONLY the simplified text.",
    "Formal":   "Rewrite in highly formal academic register. Precise vocabulary, structured sentences. Output ONLY the formal text.",
    "Concise":  "Remove all redundancy — make it as tight as possible while keeping every key idea. Output ONLY the concise text.",
    "Creative": "Paraphrase with vivid expressive language, fresh metaphors, varied rhythm. Output ONLY the paraphrased text.",
}

GRAMMAR_SYSTEM = (
    "You are a professional copy-editor. Proofread and return a JSON response with exactly these keys:\n"
    "- \"corrected\": the fully corrected text\n"
    "- \"issues\": list of objects with keys \"original\", \"corrected\", \"type\", \"explanation\"\n"
    "Types: grammar, spelling, punctuation, style, wordiness, clarity\n"
    "Return ONLY valid JSON."
)

INTENSITY_INSTRUCTIONS = {
    "Light": (
        "LIGHT EDIT: Preserve 85% of structure. Target changes:\n"
        "• Fix 3-4 of the most robotic sentences\n"
        "• Add 2 varied-length sentences (one very short <8 words, one complex >28 words)\n"
        "• Replace 3-4 hollow AI transitions with specific discourse markers\n"
        "• Fix any passive voice in the opening sentences\n"
        "• Do NOT add contractions to academic/professional text"
    ),
    "Moderate": (
        "MODERATE REWRITE: Restructure at least 60% of sentences. Required changes:\n"
        "• Sentence length must vary dramatically — shortest under 10 words, longest over 30 words\n"
        "• Replace every generic transition (Additionally, Furthermore, Moreover, In conclusion) with a specific, contextual one\n"
        "• Convert all passive-voice sentences to active where possible\n"
        "• Add at least 4 discourse markers appropriate to the style\n"
        "• Vary paragraph structure — at least one very short paragraph (1-2 sentences)\n"
        "• Academic/Professional: NO contractions. Conversational/Journalistic: contractions mandatory."
    ),
    "Deep": (
        "DEEP GHOST-MODE TRANSFORMATION: Rebuild every sentence to defeat AI detection. Non-negotiable:\n"
        "• EVERY sentence rebuilt — not a single phrase kept from the original\n"
        "• SENTENCE LENGTH: Must include sentences of 3-5 words AND 38-50 words. "
        "NEVER 3 consecutive sentences between 15-25 words.\n"
        "• PERPLEXITY: Deliberately choose low-probability but natural words. "
        "'Reshape' not 'impact'. 'Trace' not 'examine'. 'Sit at' not 'is located in'. "
        "'Dismantle' not 'affect'. 'Fracture' not 'change'. 'Elusive' not 'difficult to find'.\n"
        "• STRUCTURAL UNPREDICTABILITY: Mix these deliberately:\n"
        "  — Open a sentence mid-thought with a dash — like this — to break rhythm\n"
        "  — Use a 1-2 sentence paragraph for punch between longer ones\n"
        "  — Embed a parenthetical that qualifies (and humanises) mid-sentence\n"
        "  — Drop in a fragment. Just one word, even.\n"
        "  — End a paragraph on a tension, contradiction, or open question\n"
        "• FORBIDDEN TRANSITIONS: 'Furthermore', 'Moreover', 'Additionally', 'In conclusion', "
        "'It is important to note', 'It is worth mentioning', 'Overall', 'In summary'\n"
        "• CLAUSE DEPTH VARIATION: Alternate between simple declarative (Subject + Verb) and "
        "complex subordinate (Although X, which Y, despite Z, the analysis W)\n"
        "• VOCABULARY REGISTER SHIFTS: In academic text, drop in one precise but unexpected colloquial "
        "phrase. In conversational text, use one precise technical term. This mimics human code-switching.\n"
        "• AVOID NOUN REPETITION: After first use, replace with pronoun or synonym\n"
        "• Preserve 100% of original meaning, all data, citations, and technical terms"
    ),
}

_STYLE_REGISTER_RULES = {
    "Academic": (
        "ACADEMIC REGISTER ENFORCEMENT:\n"
        "✗ FORBIDDEN: contractions of any kind (write: it is, do not, we have, cannot)\n"
        "✗ FORBIDDEN: informal phrases, casual asides, colloquialisms\n"
        "✓ REQUIRED: formal scholarly vocabulary\n"
        "✓ REQUIRED: hedging language (the evidence suggests, it appears, arguably, tentatively)\n"
        "✓ REQUIRED: discipline-specific discourse markers (Notably, Crucially, Of significance)\n"
        "✓ REQUIRED: complex subordinate clause structures"
    ),
    "Conversational": (
        "CONVERSATIONAL REGISTER ENFORCEMENT:\n"
        "✓ REQUIRED: contractions in every 2nd or 3rd sentence minimum\n"
        "✓ REQUIRED: at least 2 rhetorical questions\n"
        "✓ REQUIRED: at least 1 em-dash (—) for rhythm\n"
        "✓ REQUIRED: at least 1 fragment sentence for emphasis\n"
        "✓ REQUIRED: first-person reference (I, we, you) at least once\n"
        "✓ REQUIRED: natural connectors (And honestly, Here's the thing, What's more)"
    ),
    "Professional": (
        "PROFESSIONAL REGISTER ENFORCEMENT:\n"
        "✓ REQUIRED: at least 1 analytical caveat or nuance ('though the picture is more complex')\n"
        "✓ REQUIRED: active voice for 80%+ of sentences\n"
        "✓ REQUIRED: at least 1 short solo paragraph (1-2 sentences) for impact\n"
        "✓ REQUIRED: power transitions (That said, More importantly, In practice, Taken together)\n"
        "✗ AVOID: consecutive sentences starting with 'The'"
    ),
    "Journalistic": (
        "JOURNALISTIC REGISTER ENFORCEMENT:\n"
        "✓ REQUIRED: first sentence must be punchy hook, 8-12 words\n"
        "✓ REQUIRED: at least 1 Economist-style rhetorical transition (The result? Consider this:)\n"
        "✓ REQUIRED: end with a tension, contrast, or implication\n"
        "✓ REQUIRED: active voice throughout\n"
        "✗ AVOID: vague generalisations — replace with specifics"
    ),
    "Creative": (
        "CREATIVE REGISTER ENFORCEMENT:\n"
        "✓ REQUIRED: at least 2 fragment sentences for rhythmic emphasis\n"
        "✓ REQUIRED: at least 1 metaphor or concrete image\n"
        "✓ REQUIRED: at least 1 single-sentence paragraph\n"
        "✓ REQUIRED: varied sentence rhythm (short-long-short pattern)\n"
        "✗ AVOID: predictable adjective choices — find surprising but accurate alternatives"
    ),
}

def _build_prompt(style, intensity, chunk):
    system     = STYLE_PROMPTS[style]
    note       = INTENSITY_INSTRUCTIONS[intensity]
    reg        = _STYLE_REGISTER_RULES[style]
    few_shot   = _FEW_SHOT.get(style, "")
    anti_ai    = _AI_PATTERNS_TO_AVOID

    # ── Detect input format ─────────────────────────────────────────
    fmt = detect_format(chunk)
    fmt_type = fmt['format_type']

    # Build format-preservation instruction
    if fmt_type == 'bullets':
        bullet_char = list(fmt['bullet_chars'])[0] if fmt['bullet_chars'] else '•'
        fmt_rule = (
            f"FORMAT PRESERVATION (MANDATORY):\n"
            f"The input text is a BULLETED LIST. Your output MUST:\n"
            f"  • Keep exactly the same number of bullet points as the input.\n"
            f"  • Start each bullet with '{bullet_char}' (same bullet character as input).\n"
            f"  • Rewrite each bullet point in human, natural language — vary sentence length, "
            f"avoid AI patterns — but keep the bullet structure intact.\n"
            f"  • Preserve any introductory sentence or heading before the list.\n"
            f"  • Do NOT merge bullet points or convert them to prose paragraphs.\n"
            f"  • RULE 4 on sentence variation applies WITHIN each bullet point where length allows.\n"
        )
    elif fmt_type == 'numbered':
        fmt_rule = (
            f"FORMAT PRESERVATION (MANDATORY):\n"
            f"The input text is a NUMBERED LIST. Your output MUST:\n"
            f"  • Keep exactly the same number of items, in the same order.\n"
            f"  • Preserve the numbering format (1. / 1) / (1) — match what the input uses).\n"
            f"  • Rewrite each numbered item in natural, varied language.\n"
            f"  • Preserve any introductory sentence or heading before the list.\n"
            f"  • Do NOT merge items or convert them to prose.\n"
        )
    elif fmt_type == 'mixed':
        fmt_rule = (
            f"FORMAT PRESERVATION (MANDATORY):\n"
            f"The input has MIXED formatting (headings + lists). Your output MUST:\n"
            f"  • Preserve all headings in place.\n"
            f"  • Keep all bullet/numbered lists as lists — do not convert to prose.\n"
            f"  • Rewrite content within each section naturally.\n"
            f"  • Maintain the same overall document structure.\n"
        )
    elif fmt_type == 'headings':
        fmt_rule = (
            f"FORMAT PRESERVATION (MANDATORY):\n"
            f"The input contains SECTION HEADINGS. Your output MUST:\n"
            f"  • Keep all headings exactly as they appear (do not rewrite headings).\n"
            f"  • Rewrite only the body text under each heading.\n"
            f"  • Do not add or remove sections.\n"
        )
    else:
        # Pure prose — standard no-list rule
        fmt_rule = (
            "FORMAT RULE: The input is prose. Do NOT introduce bullet points, "
            "numbered lists, or headers that were not in the original.\n"
        )

    user = (
        "TASK: Rewrite the text below to maximise humanness score.\n\n"
        f"WORD COUNT CONSTRAINT (STRICT): The input has {len(chunk.split())} words. Your output MUST be between {max(10, int(len(chunk.split())*0.88))} and {int(len(chunk.split())*1.12)} words — within ±12% of the input length. Do NOT pad, expand, or add new content. Rewrite; do not inflate.\n\n"
        f"INTENSITY LEVEL:\n{note}\n\n"
        f"REGISTER REQUIREMENTS:\n{reg}\n\n"
        f"{anti_ai}\n\n"
        f"{few_shot}\n\n"
        f"{fmt_rule}\n"
        "ABSOLUTE OUTPUT RULES:\n"
        "1. Output ONLY the rewritten text — zero preamble, zero commentary, zero explanation.\n"
        "2. Preserve 100% of original meaning, all numerical data, all citations, all technical terms.\n"
        "3. Preserve the EXACT formatting structure of the input (bullets stay bullets, "
        "numbered lists stay numbered, headings stay headings, prose stays prose).\n"
        "4. SENTENCE LENGTH VARIATION IS NON-NEGOTIABLE: your output MUST contain both very short "
        "sentences (4-8 words) AND long complex sentences (32-45 words). "
        "If every sentence is 15-25 words, you have failed.\n"
        "5. No two consecutive sentences may begin with the same word.\n"
        "6. Do NOT invent new facts, statistics, or examples not present in the original.\n"
        "7. PARAGRAPH BOUNDARIES (NON-NEGOTIABLE):\n"
        "   a) Count the paragraphs in the input. Your output MUST have the same number of paragraphs.\n"
        "   b) The FIRST paragraph of your output must begin with the SAME opening subject/topic "
        "as the first paragraph of the input — do not replace or skip the opening idea.\n"
        "   c) The LAST paragraph of your output must end on the SAME concluding topic/subject "
        "as the last paragraph of the input — do not drop or replace the closing idea.\n"
        "   d) Do NOT merge paragraphs together or split one paragraph into two.\n"
        "   e) Do NOT reorder paragraphs.\n"
        "8. SELF-CHECK before outputting: scan your rewrite for AI patterns listed above "
        "and fix any you find. Verify paragraph count matches input. Verify output format matches input format.\n\n"
        f"TEXT TO REWRITE:\n\"\"\"\n{chunk}\n\"\"\"\n"
    )
    return system, user

def detect_format(text: str) -> dict:
    """
    Detect structural formatting in the input text.
    Returns a dict describing what formatting was found.
    """
    lines = text.split('\n')
    info = {
        'has_bullets':  False,
        'has_numbered': False,
        'has_headings': False,
        'bullet_chars': set(),
        'format_type':  'prose',  # 'prose' | 'bullets' | 'numbered' | 'mixed' | 'headings'
    }
    bullet_count   = 0
    numbered_count = 0
    heading_count  = 0
    for line in lines:
        s = line.strip()
        if not s:
            continue
        # Bullet: lines starting with •, -, *, –, —, ▪, ○, ✓, ✗, ·
        if re.match(r'^[•\-\*–—▪○✓✗·]\s+', s):
            bullet_count += 1
            info['bullet_chars'].add(s[0])
        # Numbered: 1. or 1) or (1) formats
        elif re.match(r'^\(?\d+[.)\]]?\s+', s):
            numbered_count += 1
        # Heading: ALL CAPS line, or markdown ## heading, or ends with :
        elif re.match(r'^#{1,4}\s', s) or (s.isupper() and len(s) > 4 and len(s) < 80):
            heading_count += 1

    info['has_bullets']  = bullet_count  >= 2
    info['has_numbered'] = numbered_count >= 2
    info['has_headings'] = heading_count  >= 1

    if bullet_count + numbered_count >= 3:
        info['format_type'] = 'bullets' if bullet_count >= numbered_count else 'numbered'
    elif bullet_count + numbered_count >= 1 and heading_count >= 1:
        info['format_type'] = 'mixed'
    elif heading_count >= 1:
        info['format_type'] = 'headings'

    return info


def chunk_text(text, max_words=450):
    """
    Split text into chunks for processing.
    Crucially: keeps bullet lists and numbered lists together —
    never splits in the middle of a list block.
    """
    # Normalise line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Detect if this is list-formatted text
    fmt = detect_format(text)

    if fmt['format_type'] in ('bullets', 'numbered', 'mixed'):
        # List mode: group by logical blocks separated by blank lines,
        # but never break mid-list — keep complete list blocks together
        blocks = []
        current_block = []
        for line in text.split('\n'):
            if line.strip() == '':
                if current_block:
                    blocks.append('\n'.join(current_block))
                    current_block = []
            else:
                current_block.append(line)
        if current_block:
            blocks.append('\n'.join(current_block))

        # Now group blocks into chunks under max_words
        chunks, current, current_wc = [], [], 0
        for block in blocks:
            wc = len(block.split())
            if current_wc + wc > max_words and current:
                chunks.append('\n\n'.join(current))
                current, current_wc = [block], wc
            else:
                current.append(block); current_wc += wc
        if current:
            chunks.append('\n\n'.join(current))
        return chunks if chunks else [text]

    else:
        # Prose mode: original paragraph-splitting logic
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        chunks, current, current_wc = [], [], 0
        for para in paragraphs:
            wc = len(para.split())
            if current_wc + wc > max_words and current:
                chunks.append('\n\n'.join(current))
                current, current_wc = [para], wc
            else:
                current.append(para); current_wc += wc
        if current:
            chunks.append('\n\n'.join(current))
        return chunks if chunks else [text]

def _get_provider(model: str) -> str:
    """Infer provider from model string or session state."""
    provider = st.session_state.get("admin_provider", _DEFAULT_PROVIDER)
    return provider


def call_groq(api_key, model, system_prompt, user_prompt, max_tokens=2048, stream=False):
    """
    Dual-provider router — routes to Groq OR OpenRouter depending on
    admin_provider session state.  All call sites remain unchanged.
    """
    provider = _get_provider(model)

    if provider == "groq":
        # ── Groq endpoint ──────────────────────────────────────────
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ],
            "max_tokens": max_tokens,
            "temperature": 0.85,
            "stream": stream,
        }
        resp = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers=headers, json=payload, timeout=120, stream=stream
        )
    else:
        # ── OpenRouter endpoint ────────────────────────────────────
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://zodha.streamlit.app",
            "X-Title": "Zodha Research Writing Pro",
        }
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ],
            "max_tokens": max_tokens,
            "temperature": 0.95,   # higher entropy → better AI-detection bypass
            "stream": stream,
        }
        resp = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers, json=payload, timeout=120, stream=stream
        )

    resp.raise_for_status()
    return resp

def stream_groq(api_key, model, system_prompt, user_prompt, max_tokens=2048):
    resp = call_groq(api_key, model, system_prompt, user_prompt, max_tokens, stream=True)
    for line in resp.iter_lines():
        if line:
            line = line.decode('utf-8')
            if line.startswith('data: '):
                data = line[6:]
                if data == '[DONE]': break
                try:
                    chunk = json.loads(data)
                    delta = chunk['choices'][0]['delta'].get('content','')
                    if delta: yield delta
                except: continue

def humanize_streaming(api_key, model, chunk, style, intensity, placeholder):
    system, user = _build_prompt(style, intensity, chunk)
    # Cap tokens to ~1.2× input to prevent word bloat
    # ~1.4 tokens per word on average; cap at 1.2× word count × 1.5 tokens
    input_words = len(chunk.split())
    max_tok = max(256, min(3500, int(input_words * 2.1)))
    full_text = ""
    import html as _html

    def render_output(text, cursor=False):
        """
        Render output with correct formatting:
        - Bullet/numbered lists: native Streamlit markdown
        - Prose: paragraph-wrapped HTML with justified alignment
        """
        import html as _html2
        fmt = detect_format(text)
        cur = "▌" if cursor else ""

        if fmt['format_type'] in ('bullets', 'numbered', 'mixed', 'headings'):
            placeholder.markdown(text + cur)
        else:
            # Convert double-newline separated paragraphs to <p> tags
            paras = [p.strip() for p in text.split("\n\n") if p.strip()]
            if not paras:
                paras = [text]
            html_paras = "".join(
                f'<p style="margin:0 0 0.8em 0;text-indent:0;">{_html2.escape(p)}</p>'
                for p in paras
            )
            if cursor:
                html_paras = html_paras.rstrip("</p>") + cur + "</p>"
            placeholder.markdown(
                f'<div class="output-box" style="min-height:100px;">{html_paras}</div>',
                unsafe_allow_html=True)

    for token in stream_groq(api_key, model, system, user, max_tokens=max_tok):
        full_text += token
        render_output(full_text, cursor=True)

    render_output(full_text, cursor=False)
    return full_text

def paraphrase_text(api_key, model, text, mode):
    fmt = detect_format(text)
    fmt_type = fmt['format_type']
    fmt_note = ""
    if fmt_type in ('bullets','numbered','mixed'):
        fmt_note = (
            "\n\nFORMAT PRESERVATION: The input is a list. "
            "Keep ALL bullet points or numbered items in the output. "
            "Do NOT convert to prose. Rewrite each item naturally but preserve the list structure exactly."
        )
    elif fmt_type == 'headings':
        fmt_note = (
            "\n\nFORMAT PRESERVATION: The input has headings. "
            "Keep all headings intact. Only rewrite the body text under each heading."
        )
    resp = call_groq(api_key, model, PARAPHRASE_MODES[mode],
                     f"TEXT TO PARAPHRASE:\n\"\"\"\n{text}\n\"\"\"{fmt_note}", stream=False)
    return resp.json()["choices"][0]["message"]["content"].strip()

def grammar_check(api_key, model, text):
    resp = call_groq(api_key, model, GRAMMAR_SYSTEM,
                     f"TEXT TO PROOFREAD:\n\"\"\"\n{text}\n\"\"\"", max_tokens=3000, stream=False)
    raw = resp.json()["choices"][0]["message"]["content"].strip()
    raw = re.sub(r"^```(?:json)?\s*","",raw); raw=re.sub(r"\s*```$","",raw)
    try: return json.loads(raw)
    except: return {"corrected":raw,"issues":[]}

# ══════════════════════════════════════════════════════════════════════════════
# RESEARCH TOOLS — Citation, Summariser, Literature Review
# ══════════════════════════════════════════════════════════════════════════════

CITATION_STYLES = {
    "APA 7th":     "apa7",
    "MLA 9th":     "mla9",
    "Chicago 17th":"chicago17",
    "Harvard":     "harvard",
    "Vancouver":   "vancouver",
    "IEEE":        "ieee",
}

# ── CrossRef DOI metadata fetch ────────────────────────────────────────────

def fetch_doi_metadata(doi_or_url: str) -> dict:
    """Fetch metadata from CrossRef for a DOI or DOI URL."""
    doi = doi_or_url.strip()
    # Extract raw DOI from URL forms
    for prefix in ["https://doi.org/", "http://doi.org/",
                   "https://dx.doi.org/", "http://dx.doi.org/"]:
        if doi.startswith(prefix):
            doi = doi[len(prefix):]
            break
    if doi.startswith("doi:"):
        doi = doi[4:]
    doi = doi.strip("/")
    try:
        r = requests.get(
            f"https://api.crossref.org/works/{doi}",
            headers={"User-Agent": "Zodha Research Writing Pro/1.0 (zodha@gmail.com)"},
            timeout=15
        )
        r.raise_for_status()
        msg = r.json().get("message", {})
        authors = msg.get("author", [])
        author_list = []
        for a in authors:
            given  = a.get("given", "")
            family = a.get("family", "")
            name   = f"{family}, {given}" if given else family
            if name.strip(): author_list.append(name)
        year = ""
        pub  = msg.get("published", msg.get("published-print", msg.get("published-online", {})))
        dp   = pub.get("date-parts", [[]])
        if dp and dp[0]: year = str(dp[0][0])
        journal = ""
        cp = msg.get("container-title", [])
        if cp: journal = cp[0]
        return {
            "doi":      doi,
            "title":    msg.get("title", [""])[0],
            "authors":  author_list,
            "year":     year,
            "journal":  journal,
            "volume":   msg.get("volume", ""),
            "issue":    msg.get("issue", ""),
            "pages":    msg.get("page", ""),
            "publisher":msg.get("publisher", ""),
            "url":      f"https://doi.org/{doi}",
            "type":     msg.get("type", "journal-article"),
        }
    except Exception as e:
        return {"error": str(e)}

def format_citation(meta: dict, style: str) -> str:
    """Format metadata into a citation string for the given style."""
    if "error" in meta:
        return f"Error fetching metadata: {meta['error']}"

    authors = meta.get("authors", [])
    title   = meta.get("title", "Untitled")
    year    = meta.get("year", "n.d.")
    journal = meta.get("journal", "")
    vol     = meta.get("volume", "")
    issue   = meta.get("issue", "")
    pages   = meta.get("pages", "")
    doi     = meta.get("doi", "")
    pub     = meta.get("publisher", "")
    url     = meta.get("url", "")

    def apa_authors(au_list):
        if not au_list: return "Unknown Author"
        if len(au_list) == 1: return au_list[0]
        if len(au_list) <= 20:
            return ", ".join(au_list[:-1]) + ", & " + au_list[-1]
        return ", ".join(au_list[:19]) + ", ... " + au_list[-1]

    def mla_author_format(au_list):
        if not au_list: return "Unknown Author"
        if len(au_list) == 1: return au_list[0]
        parts = au_list[0].split(", ")
        first_inv = f"{parts[0]}, {parts[1]}" if len(parts) == 2 else au_list[0]
        if len(au_list) == 2:
            p2 = au_list[1].split(", ")
            second = f"{p2[1]} {p2[0]}" if len(p2)==2 else au_list[1]
            return f"{first_inv}, and {second}"
        return f"{first_inv}, et al"

    vol_issue = f"{vol}({issue})" if vol and issue else vol or issue
    doi_str   = f"https://doi.org/{doi}" if doi else url

    if style == "apa7":
        au  = apa_authors(authors)
        ji  = f"*{journal}*" if journal else pub or "Unknown Source"
        vi  = f", *{vol_issue}*" if vol_issue else ""
        pg  = f", {pages}" if pages else ""
        doi_link = (f"\n  https://doi.org/{doi}") if doi else ""
        return f"{au} ({year}). {title}. {ji}{vi}{pg}.{doi_link}"

    elif style == "mla9":
        au  = mla_author_format(authors)
        jtl = f"*{journal}*" if journal else ""
        vi  = f"vol. {vol}" if vol else ""
        is_ = f"no. {issue}" if issue else ""
        vi_is = ", ".join(filter(None, [vi, is_]))
        pg  = f"pp. {pages}" if pages else ""
        do  = f"doi:{doi}" if doi else url
        parts_list = list(filter(None, [jtl, vi_is, year, pg, do]))
        return '%s "%s." %s.' % (au, title, ", ".join(parts_list))

    elif style == "chicago17":
        if not authors: auth = "Unknown Author"
        elif len(authors) == 1: auth = authors[0]
        elif len(authors) <= 3: auth = "; ".join(authors)
        else: auth = f"{authors[0]} et al."
        vi  = f"{vol}" + (f", no. {issue}" if issue else "")
        pg  = f": {pages}" if pages else ""
        doi_link = f" https://doi.org/{doi}." if doi else ""
        return f'{auth}. "{title}." *{journal or pub}* {vi} ({year}){pg}.{doi_link}'

    elif style == "harvard":
        if not authors: au = "Anon"
        else:
            parts = []
            for a in authors:
                sp = a.split(", ")
                initials = "".join(n[0].upper()+"." for n in sp[1].split()) if len(sp)>1 else ""
                parts.append(f"{sp[0]}, {initials}" if initials else sp[0])
            au = " and ".join(parts) if len(parts)<=2 else f"{parts[0]} et al."
        pg  = f", pp.{pages}" if pages else ""
        vi  = f", {vol}({issue})" if vol and issue else (f", {vol}" if vol else "")
        doi_link = f". Available at: https://doi.org/{doi}" if doi else ""
        return f"{au} ({year}) '{title}', *{journal or pub}*{vi}{pg}{doi_link}."

    elif style == "vancouver":
        van_au = []
        for a in authors[:6]:
            sp = a.split(", ")
            initials = "".join(n[0].upper() for n in sp[1].split()) if len(sp)>1 else ""
            van_au.append(f"{sp[0]} {initials}" if initials else sp[0])
        if len(authors) > 6: van_au.append("et al")
        au  = ", ".join(van_au) + "." if van_au else "Unknown."
        vi  = f";{vol}" + (f"({issue})" if issue else "")
        pg  = f":{pages}" if pages else ""
        doi_link = f" doi: {doi}" if doi else ""
        return f"{au} {title}. {journal or pub}. {year}{vi}{pg}.{doi_link}"

    elif style == "ieee":
        ieee_au = []
        for a in authors:
            sp = a.split(", ")
            initials = "".join(n[0].upper()+"." for n in sp[1].split()) if len(sp)>1 else ""
            ieee_au.append(f"{initials} {sp[0]}" if initials else sp[0])
        au  = ", ".join(ieee_au[:6])
        if len(authors) > 6: au += " et al."
        vi  = f", vol. {vol}" if vol else ""
        is_ = f", no. {issue}" if issue else ""
        pg  = f", pp. {pages}" if pages else ""
        doi_link = f", doi: {doi}" if doi else ""
        return f'{au}, "{title}," *{journal or pub}*{vi}{is_}{pg}, {year}{doi_link}.'

    return f"{', '.join(authors)} ({year}). {title}. {journal}."


# ── Article summariser (via OpenRouter) ──────────────────────────────────────────

SUMMARISE_SYSTEM = """You are a research analyst. Extract structured information from the academic article text provided.
Return a JSON object with exactly these keys:
- "title": full article title
- "authors": list of author names
- "year": publication year (string)
- "journal": journal or conference name
- "doi": DOI if present (or empty string)
- "main_objectives": 2-3 sentence summary of the research objectives
- "methodology": 2-3 sentence description of methods used
- "major_findings": 3-5 bullet points as a list of strings, each 1-2 sentences
- "keywords": list of up to 6 keywords
- "limitations": 1-2 sentence summary of limitations (or empty string)
- "conclusion": 1-2 sentence overall conclusion
Return ONLY valid JSON."""

def summarise_article(api_key: str, model: str, text: str) -> dict:
    """Summarise an academic article using OpenRouter (Claude Haiku 4.5)."""
    # Truncate to ~3000 words to stay within context
    words = text.split()
    if len(words) > 3000:
        text = " ".join(words[:3000]) + "\n[...truncated for processing...]"
    user_msg = f"ARTICLE TEXT:\n\"\"\"\n{text}\n\"\"\""
    resp = call_groq(api_key, model, SUMMARISE_SYSTEM, user_msg, max_tokens=1500, stream=False)
    raw  = resp.json()["choices"][0]["message"]["content"].strip()
    raw  = re.sub(r"^```(?:json)?\s*", "", raw)
    raw  = re.sub(r"\s*```$", "", raw)
    try:
        return json.loads(raw)
    except Exception:
        return {"error": "Could not parse article summary.", "raw": raw}


# ── PDF text extractor ─────────────────────────────────────────────────────

def extract_pdf_text(uploaded_file) -> str:
    """Extract text from an uploaded PDF file."""
    if not PDF_OK:
        return ""
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            pages = []
            for page in pdf.pages[:40]:  # max 40 pages
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)
    except Exception as e:
        return f"[PDF extraction error: {e}]"


# ── Literature Review Excel exporter ───────────────────────────────────────

def build_literature_excel(articles: list) -> bytes:
    """Build a formatted Excel workbook from a list of article summary dicts."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Literature Review"

    # Colour palette
    C_HEADER_BG  = "1A1A2E"   # dark navy
    C_HEADER_FG  = "C9A84C"   # gold
    C_ALT_BG     = "FAF7F2"   # cream
    C_WHITE      = "FFFFFF"
    C_BORDER     = "D4C9B5"
    C_ACCENT     = "4A7C59"   # sage green

    thin_border = Border(
        left=Side(style="thin", color=C_BORDER),
        right=Side(style="thin", color=C_BORDER),
        top=Side(style="thin", color=C_BORDER),
        bottom=Side(style="thin", color=C_BORDER),
    )
    thick_bottom = Border(bottom=Side(style="medium", color=C_HEADER_BG))

    # Title row
    ws.merge_cells("A1:J1")
    title_cell = ws["A1"]
    title_cell.value = "Literature Review Summary — Zodha Research Writing Pro Analytics"
    title_cell.font      = Font(name="Arial", size=14, bold=True, color=C_HEADER_FG)
    title_cell.fill      = PatternFill("solid", fgColor=C_HEADER_BG)
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    # Date row
    ws.merge_cells("A2:J2")
    date_cell = ws["A2"]
    date_cell.value = f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}"
    date_cell.font  = Font(name="Arial", size=9, color="888888")
    date_cell.alignment = Alignment(horizontal="right")
    ws.row_dimensions[2].height = 16

    # Column headers
    headers = ["#", "Title", "Authors", "Year", "Journal / Source",
               "Main Objectives", "Methodology", "Major Findings",
               "Keywords", "DOI / URL"]
    col_widths = [4, 35, 25, 6, 25, 40, 35, 50, 25, 30]

    for col_idx, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
        cell = ws.cell(row=3, column=col_idx)
        cell.value     = hdr
        cell.font      = Font(name="Arial", size=10, bold=True, color=C_WHITE)
        cell.fill      = PatternFill("solid", fgColor=C_ACCENT)
        cell.alignment = Alignment(horizontal="center", vertical="center",
                                   wrap_text=True)
        cell.border    = thin_border
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    ws.row_dimensions[3].height = 22

    # Data rows
    for row_num, art in enumerate(articles, start=1):
        row_idx   = row_num + 3
        bg_color  = C_WHITE if row_num % 2 == 0 else C_ALT_BG
        row_fill  = PatternFill("solid", fgColor=bg_color)

        findings = art.get("major_findings", [])
        if isinstance(findings, list):
            findings_str = "\n".join(f"• {f}" for f in findings)
        else:
            findings_str = str(findings)

        authors = art.get("authors", [])
        if isinstance(authors, list):
            authors_str = "; ".join(authors)
        else:
            authors_str = str(authors)

        keywords = art.get("keywords", [])
        kw_str = "; ".join(keywords) if isinstance(keywords, list) else str(keywords)

        doi = art.get("doi", "") or ""
        url_val = f"https://doi.org/{doi}" if doi else art.get("url","")

        values = [
            row_num,
            art.get("title",""),
            authors_str,
            art.get("year",""),
            art.get("journal",""),
            art.get("main_objectives",""),
            art.get("methodology",""),
            findings_str,
            kw_str,
            url_val,
        ]

        for col_idx, val in enumerate(values, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value     = val
            cell.fill      = row_fill
            cell.border    = thin_border
            cell.alignment = Alignment(vertical="top", wrap_text=True,
                                       horizontal="center" if col_idx in (1,4) else "left")
            cell.font      = Font(name="Arial", size=9)
            if col_idx == 1:
                cell.font = Font(name="Arial", size=9, bold=True, color=C_ACCENT)
            if col_idx == 2:
                cell.font = Font(name="Arial", size=9, bold=True, color=C_HEADER_BG)
            if col_idx == 4:
                cell.font = Font(name="Arial", size=9, bold=True)

        ws.row_dimensions[row_idx].height = max(60, 15 * len(findings_str.split("\n")))

    # Freeze header rows
    ws.freeze_panes = "A4"

    # Auto-filter
    ws.auto_filter.ref = f"A3:J{3 + len(articles)}"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()



def _extract_file_text(uploaded_file) -> str:
    """Extract text from PDF, DOCX, or TXT upload."""
    if not uploaded_file:
        return ""
    ext = uploaded_file.name.split(".")[-1].lower()
    try:
        if ext == "txt":
            return uploaded_file.read().decode("utf-8", errors="ignore")
        elif ext == "pdf" and PDF_OK:
            import io as _io2
            with pdfplumber.open(_io2.BytesIO(uploaded_file.read())) as _pdf2:
                return "\n".join(p.extract_text() or "" for p in _pdf2.pages[:25])
        elif ext in ("docx", "doc"):
            try:
                import docx as _docx2
                import io as _io2
                doc = _docx2.Document(_io2.BytesIO(uploaded_file.read()))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        return f"[File read error: {e}]"
    return ""

# ══════════════════════════════════════════════════════════════════════════════
# LAYER 3 — STATISTICAL & METHODOLOGY SUPPORT PROMPTS
# ══════════════════════════════════════════════════════════════════════════════

SPSS_INTERPRET_SYSTEM = """You are a senior research methodologist and statistician specialising
in social science and management research. A researcher has pasted raw SPSS/AMOS/SmartPLS output.
Your task: explain it in clear, publication-ready academic prose that can be inserted directly
into a thesis or journal paper Results section.

Return a JSON object with exactly these keys:
- "test_type": name of the statistical test detected (e.g. "Multiple Linear Regression", "CFA", "SEM", "t-test")
- "plain_english": 2-3 sentence plain explanation a non-statistician would understand
- "academic_prose": 3-5 sentences of formal academic writing ready to paste into a Results section.
  Include specific values (F, t, p, R², β, RMSEA, CFI etc.) exactly as given.
- "interpretation": what the results mean for the research hypothesis (supported / not supported / partial)
- "apa_table": if a table is present, reproduce key stats in APA 7th format as a string
- "assumptions_check": any statistical assumptions the researcher should verify for this test
- "limitations": 1-2 sentence note on what this result cannot conclude

Return ONLY valid JSON."""

METHODOLOGY_SYSTEM = """You are an expert research design consultant with deep expertise in
quantitative, qualitative, and mixed-method social science research.
A researcher has described their study. Recommend the most suitable research methodology.

Return a JSON object with exactly these keys:
- "recommended_approach": primary methodology (Quantitative / Qualitative / Mixed Methods)
- "design": specific design (e.g. "Cross-sectional survey", "Longitudinal panel", "Grounded theory", "SEM-based causal model")
- "rationale": 3-4 sentences explaining why this methodology fits the research objectives
- "data_collection": list of 3-5 recommended data collection methods
- "analysis_techniques": list of 4-6 statistical or analytical techniques suitable for this study
- "software": recommended software (SPSS / AMOS / SmartPLS / NVivo / R / Python)
- "sample_size": recommended sample size with justification (mention Hair et al. or Roscoe rule if relevant)
- "sampling_strategy": e.g. "Stratified random sampling", "Purposive sampling"
- "alternative_designs": list 2 alternative methodologies with 1-line pros/cons each
- "conceptual_framework_hint": 2-3 sentence suggestion for the theoretical framework

Return ONLY valid JSON."""

HYPOTHESIS_SYSTEM = """You are a senior academic researcher specialising in hypothesis development
for management, finance, social science, and business research.
Generate testable, well-grounded research hypotheses based on the topic and variables provided.

Return a JSON object with exactly these keys:
- "research_question": the overarching research question derived from the input
- "null_hypotheses": list of H0 statements (at least 3)
- "alternate_hypotheses": list of H1/H2/H3... corresponding alternate hypotheses
- "directional_rationale": for each hypothesis, 1-2 sentences of theoretical justification
  (cite plausible theories — TAM, Agency Theory, Stakeholder Theory, etc. where applicable)
- "variable_classification": object with "independent", "dependent", "moderating", "mediating"
  variable lists derived from the input
- "suggested_scale": for each variable, suggest a validated measurement scale if applicable
- "expected_relationships": a summary of expected directional relationships (+/-)

Return ONLY valid JSON."""

VARIABLE_OPERATIONALISE_SYSTEM = """You are a psychometrics and scale development expert.
Given a research variable, operationalise it with validated measurement items suitable
for a 5-point Likert scale survey instrument for Indian academic research.

Return a JSON object with exactly these keys:
- "variable": the variable name
- "construct_definition": formal academic definition (2-3 sentences)
- "dimensions": list of sub-dimensions/facets of this construct
- "measurement_items": list of 5-7 survey items (statements for Likert scale)
- "scale_source": name of the original validated scale this is adapted from (e.g. "Davis 1989 TAM", "Fornell & Larcker 1981")
- "response_format": the Likert scale description (e.g. "1=Strongly Disagree to 5=Strongly Agree")
- "reliability_note": how to check reliability (Cronbach alpha threshold, AVE, CR)
- "reverse_coded_items": list of item numbers that should be reverse coded (if any)

Return ONLY valid JSON."""


# ── Layer 3 functions ────────────────────────────────────────────────────────

def interpret_spss_output(api_key: str, model: str, spss_text: str) -> dict:
    user_msg = f"SPSS/AMOS/SmartPLS OUTPUT TO INTERPRET:\n\"\"\"\n{spss_text[:4000]}\n\"\"\""
    resp = call_groq(api_key, model, SPSS_INTERPRET_SYSTEM, user_msg, max_tokens=2000, stream=False)
    raw  = resp.json()["choices"][0]["message"]["content"].strip()
    raw  = re.sub(r"^```(?:json)?\s*","",raw); raw=re.sub(r"\s*```$","",raw)
    try: return json.loads(raw)
    except: return {"error": "Parse error", "raw": raw}

def recommend_methodology(api_key: str, model: str, description: str) -> dict:
    user_msg = f"RESEARCH DESCRIPTION:\n\"\"\"\n{description}\n\"\"\""
    resp = call_groq(api_key, model, METHODOLOGY_SYSTEM, user_msg, max_tokens=2000, stream=False)
    raw  = resp.json()["choices"][0]["message"]["content"].strip()
    raw  = re.sub(r"^```(?:json)?\s*","",raw); raw=re.sub(r"\s*```$","",raw)
    try: return json.loads(raw)
    except: return {"error": "Parse error", "raw": raw}

def generate_hypotheses(api_key: str, model: str, topic: str, variables: str) -> dict:
    user_msg = f"RESEARCH TOPIC:\n{topic}\n\nKEY VARIABLES MENTIONED:\n{variables}"
    resp = call_groq(api_key, model, HYPOTHESIS_SYSTEM, user_msg, max_tokens=2000, stream=False)
    raw  = resp.json()["choices"][0]["message"]["content"].strip()
    raw  = re.sub(r"^```(?:json)?\s*","",raw); raw=re.sub(r"\s*```$","",raw)
    try: return json.loads(raw)
    except: return {"error": "Parse error", "raw": raw}

def operationalise_variable(api_key: str, model: str, variable: str, context: str) -> dict:
    user_msg = f"VARIABLE TO OPERATIONALISE: {variable}\nRESEARCH CONTEXT: {context}"
    resp = call_groq(api_key, model, VARIABLE_OPERATIONALISE_SYSTEM, user_msg, max_tokens=1500, stream=False)
    raw  = resp.json()["choices"][0]["message"]["content"].strip()
    raw  = re.sub(r"^```(?:json)?\s*","",raw); raw=re.sub(r"\s*```$","",raw)
    try: return json.loads(raw)
    except: return {"error": "Parse error", "raw": raw}


# ══════════════════════════════════════════════════════════════════════════════
# LAYER 5 — JOURNAL & PUBLICATION SUPPORT PROMPTS
# Tools: Journal Matcher, Cover Letter Writer, Reviewer Response Drafter,
#         Research Contribution Statement
# ══════════════════════════════════════════════════════════════════════════════

JOURNAL_MATCH_SYSTEM = """You are a senior academic publishing consultant with expertise in
Scopus, ABDC, ABS, Web of Science, and Indian journal indices (UGC CARE).
Given an abstract and research details, recommend the most suitable journals.

Return a JSON object with exactly these keys:
- "research_domain": detected primary field and sub-field
- "recommendations": list of 6 journal objects, each with:
    "name": journal name
    "publisher": publisher name
    "index": indexing (Scopus / ABDC A / ABDC B / ABDC C / ABS / Web of Science / UGC CARE)
    "impact_factor": approximate IF or quartile (Q1/Q2/Q3) or "N/A"
    "scope_fit": 1-2 sentence explanation of why this journal fits
    "acceptance_rate": approximate % or "not published"
    "turnaround": typical review time in weeks
    "open_access": yes/no/hybrid
    "submission_url": provide the official journal page URL if known, else "search: [journal name] submissions"
- "avoid": list of 2-3 predatory or low-quality journals in this space to avoid
- "strategy_note": 2-3 sentences on the submission strategy (where to aim first, backup plan)

Return ONLY valid JSON."""

COVER_LETTER_SYSTEM = """You are an expert academic editor who has helped hundreds of authors
publish in top-tier journals. Write a compelling, professional cover letter for journal submission.

Output a single JSON object with:
- "cover_letter": the full cover letter text (400-600 words) formatted for direct use
- "key_contributions": bullet list of 3-4 key contributions of the paper
- "novelty_statement": 2-sentence statement of what is new/original about this work
- "ethical_statements": standard statements (conflicts of interest, funding, author contributions)

Return ONLY valid JSON."""

REVIEWER_RESPONSE_SYSTEM = """You are a senior academic who has reviewed for top journals and
helped authors navigate the revision process. Draft a point-by-point response to reviewer comments.

Return a JSON object with:
- "response_header": professional opening paragraph thanking reviewers (3-4 sentences)
- "responses": list of objects for each comment, each with:
    "reviewer": "Reviewer 1" or "Reviewer 2" etc.
    "comment_summary": brief summary of the reviewer's comment
    "response": full professional response (2-5 sentences)
    "manuscript_change": what was changed in the manuscript (or "No change required — see explanation")
- "closing_paragraph": professional closing (2-3 sentences)
- "revision_summary": 1-paragraph executive summary of all changes made

Return ONLY valid JSON."""

CONTRIBUTION_STATEMENT_SYSTEM = """You are an expert in academic writing and research positioning.
Write a Research Contribution Statement suitable for a journal submission or thesis.

Return a JSON object with:
- "theoretical_contribution": 2-3 sentences on contribution to theory
- "methodological_contribution": 2-3 sentences on methodological innovation (if any)
- "practical_contribution": 2-3 sentences on practical/policy implications
- "contextual_novelty": 1-2 sentences on geographic/sector novelty (e.g. Indian context)
- "gap_addressed": 2-3 sentences on the research gap this work fills
- "full_statement": a complete 150-200 word Research Contribution paragraph ready to use in a paper

Return ONLY valid JSON."""


# ── Layer 5 functions ────────────────────────────────────────────────────────

def match_journals(api_key: str, model: str, abstract: str, keywords: str, field: str) -> dict:
    user_msg = (f"ABSTRACT:\n{abstract[:2000]}\n\n"
                f"KEYWORDS: {keywords}\nFIELD: {field}")
    resp = call_groq(api_key, model, JOURNAL_MATCH_SYSTEM, user_msg, max_tokens=2500, stream=False)
    raw  = resp.json()["choices"][0]["message"]["content"].strip()
    raw  = re.sub(r"^```(?:json)?\s*","",raw); raw=re.sub(r"\s*```$","",raw)
    try: return json.loads(raw)
    except: return {"error": "Parse error", "raw": raw}

def write_cover_letter(api_key: str, model: str, title: str, abstract: str,
                        journal: str, authors: str) -> dict:
    user_msg = (f"PAPER TITLE: {title}\nJOURNAL TARGET: {journal}\n"
                f"AUTHORS: {authors}\nABSTRACT:\n{abstract[:1500]}")
    resp = call_groq(api_key, model, COVER_LETTER_SYSTEM, user_msg, max_tokens=1500, stream=False)
    raw  = resp.json()["choices"][0]["message"]["content"].strip()
    raw  = re.sub(r"^```(?:json)?\s*","",raw); raw=re.sub(r"\s*```$","",raw)
    try: return json.loads(raw)
    except: return {"error": "Parse error", "raw": raw}

def draft_reviewer_response(api_key: str, model: str, comments: str, abstract: str) -> dict:
    user_msg = (f"PAPER ABSTRACT:\n{abstract[:1000]}\n\n"
                f"REVIEWER COMMENTS:\n\"\"\"\n{comments[:3000]}\n\"\"\"")
    resp = call_groq(api_key, model, REVIEWER_RESPONSE_SYSTEM, user_msg, max_tokens=2500, stream=False)
    raw  = resp.json()["choices"][0]["message"]["content"].strip()
    raw  = re.sub(r"^```(?:json)?\s*","",raw); raw=re.sub(r"\s*```$","",raw)
    try: return json.loads(raw)
    except: return {"error": "Parse error", "raw": raw}

def write_contribution_statement(api_key: str, model: str, title: str,
                                   objectives: str, findings: str, context: str) -> dict:
    user_msg = (f"PAPER TITLE: {title}\nRESEARCH OBJECTIVES: {objectives}\n"
                f"KEY FINDINGS: {findings}\nCONTEXT/SETTING: {context}")
    resp = call_groq(api_key, model, CONTRIBUTION_STATEMENT_SYSTEM, user_msg, max_tokens=1200, stream=False)
    raw  = resp.json()["choices"][0]["message"]["content"].strip()
    raw  = re.sub(r"^```(?:json)?\s*","",raw); raw=re.sub(r"\s*```$","",raw)
    try: return json.loads(raw)
    except: return {"error": "Parse error", "raw": raw}


# ══════════════════════════════════════════════════════════════════════════════
# LITERATURE REVIEW CHAPTER WRITER
# ══════════════════════════════════════════════════════════════════════════════

LIT_REVIEW_WRITER_SYSTEM = """You are a senior academic writer with expertise in producing
comprehensive, publication-quality literature review chapters for PhD dissertations and
journal articles. You write in flowing scholarly prose — no bullet points, no headers
unless requested, no numbered lists in the body.

Given structured summaries of research articles, write a complete Literature Review chapter
following these conventions:
- Thematically organised paragraphs grouping related studies
- Critical synthesis — compare, contrast, and evaluate studies; do NOT just summarise each one
- Identify research gaps, contradictions, and agreements across studies
- Appropriate scholarly hedging and academic register
- Smooth transitions between paragraphs and sections
- Every claim attributed to a specific author using the citation style specified
- References section at the end in the exact citation style specified

CITATION STYLES:
- APA 7th:    (Author, Year) in-text;  Author, A. (Year). Title. Journal, Vol(Issue), pages. https://doi.org/xxx
- MLA 9th:    (Author page) in-text;   Author. "Title." Journal, vol. X, no. Y, Year, pp. Z-Z.
- Chicago 17th: (Author Year) in-text; Author. "Title." Journal X, no. Y (Year): Z-Z.
- Harvard:    (Author Year) in-text;   Author (Year) 'Title', Journal, Vol(Issue), pp. Z-Z.
- Vancouver:  [1] superscript in-text; 1. Author. Title. Journal. Year;Vol(Issue):pages.
- IEEE:       [1] superscript in-text; [1] Author, "Title," Journal, vol. X, no. Y, pp. Z-Z, Year.

Return ONLY the formatted chapter text — no preamble, no meta-commentary.
Start directly with the chapter content. Include a REFERENCES section at the end."""

def write_literature_review_chapter(api_key: str, model: str, articles: list,
                                     citation_style: str, research_topic: str,
                                     word_limit: int, extra_instructions: str) -> str:
    """Generate a full literature review chapter from article summaries."""

    # Build structured article list for the prompt
    articles_block = ""
    for i, art in enumerate(articles, 1):
        authors = art.get("authors", [])
        if isinstance(authors, list):
            authors_str = "; ".join(authors) if authors else "Unknown"
        else:
            authors_str = str(authors)
        findings = art.get("major_findings", [])
        if isinstance(findings, list):
            findings_str = " ".join(f"({j}) {f}" for j,f in enumerate(findings, 1))
        else:
            findings_str = str(findings)
        articles_block += f"""
ARTICLE {i}:
  Title:       {art.get("title", "Unknown")}
  Authors:     {authors_str}
  Year:        {art.get("year", "n.d.")}
  Journal:     {art.get("journal", "")}
  DOI:         {art.get("doi", "")}
  Objectives:  {art.get("main_objectives", "")}
  Methodology: {art.get("methodology", "")}
  Findings:    {findings_str}
  Keywords:    {", ".join(art.get("keywords", [])) if isinstance(art.get("keywords",[]), list) else ""}
  Limitations: {art.get("limitations", "")}
  Conclusion:  {art.get("conclusion", "")}
"""

    user_msg = f"""RESEARCH TOPIC / CHAPTER FOCUS:
{research_topic}

CITATION STYLE: {citation_style}
TARGET WORD COUNT: approximately {word_limit} words
EXTRA INSTRUCTIONS: {extra_instructions if extra_instructions.strip() else "None"}

ARTICLES TO SYNTHESISE:
{articles_block}

Write the complete literature review chapter now. Organise thematically, synthesise critically,
cite every claim using {citation_style} style in-text, and include a full References section at the end.
Do NOT use bullet points or numbered lists in the body text."""

    # Use long-context model for large reviews
    resp = call_groq(api_key, model, LIT_REVIEW_WRITER_SYSTEM, user_msg,
                     max_tokens=4096, stream=False)
    return resp.json()["choices"][0]["message"]["content"].strip()


def build_literature_review_docx(chapter_text: str, topic: str, citation_style: str,
                                  article_count: int) -> bytes:
    """Build a Word document from the literature review chapter text using python-docx."""
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io, re

        doc = DocxDoc()

        # ── Page margins ────────────────────────────────────────────
        section = doc.sections[0]
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

        # ── Styles ──────────────────────────────────────────────────
        normal_style   = doc.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)

        # ── Title page ──────────────────────────────────────────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("LITERATURE REVIEW")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1E, 0x5C, 0x22)  # Zodha green
        run.font.name = "Times New Roman"

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = subtitle.add_run(topic if topic.strip() else "Systematic Literature Review")
        sr.bold = True
        sr.font.size = Pt(13)
        sr.font.name = "Times New Roman"

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = meta.add_run(
            f"Citation Style: {citation_style}  |  Articles Reviewed: {article_count}  |  "
            f"Generated by Zodha Research Writing Pro"
        )
        mr.font.size = Pt(10)
        mr.font.color.rgb = RGBColor(0x3D, 0x5E, 0x40)
        mr.font.name = "Times New Roman"

        doc.add_paragraph()  # spacer

        # ── Parse and render chapter ─────────────────────────────────
        # Split into lines and handle formatting
        lines = chapter_text.split("\n")
        i = 0
        in_references = False

        while i < len(lines):
            line = lines[i].strip()
            i += 1

            if not line:
                continue

            # Detect section headings (ALL CAPS lines or lines starting with #)
            is_heading = (
                (line.isupper() and len(line) > 4 and len(line) < 80) or
                line.startswith("# ") or line.startswith("## ") or
                line.upper() == line and 5 < len(line) < 60
            )
            is_references = any(kw in line.upper() for kw in
                                 ["REFERENCES", "BIBLIOGRAPHY", "WORKS CITED"])

            if is_references:
                in_references = True
                ref_para = doc.add_paragraph()
                ref_run  = ref_para.add_run(line.lstrip("# "))
                ref_run.bold = True
                ref_run.font.size = Pt(13)
                ref_run.font.color.rgb = RGBColor(0x1E, 0x5C, 0x22)
                ref_run.font.name = "Times New Roman"
                continue

            if is_heading and not in_references:
                h_para = doc.add_paragraph()
                h_run  = h_para.add_run(line.lstrip("# "))
                h_run.bold = True
                h_run.font.size = Pt(13)
                h_run.font.color.rgb = RGBColor(0x1E, 0x5C, 0x22)
                h_run.font.name = "Times New Roman"
                continue

            # Normal paragraph
            p = doc.add_paragraph()
            if not in_references:
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                # Double spacing for body
                p.paragraph_format.line_spacing = Pt(24)
            else:
                # Hanging indent for references
                p.paragraph_format.left_indent     = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = Pt(24)

            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        # ── Footer ──────────────────────────────────────────────────
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Zodha Research Writing Pro · Generated Literature Review")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x9A, 0x8A, 0x7A)
        footer_run.font.name = "Times New Roman"

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # python-docx not available — return plain text as bytes
        return chapter_text.encode("utf-8")


def parse_excel_to_articles(excel_bytes, api_key: str = "", model: str = "") -> list:
    """
    Parse a Zodha Literature Review Excel file back into article dicts.
    Cleans up placeholder values (CrossRef context snippets, DOI-only findings).
    If api_key provided, re-enriches articles that have junk placeholder data.
    """
    if not XLSX_OK:
        return []
    try:
        from openpyxl import load_workbook
        import io as _io3
        wb = load_workbook(_io3.BytesIO(excel_bytes))
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))

        # Find header row (row index with "#" in col 0)
        header_row_idx = None
        for i, row in enumerate(rows):
            if row and str(row[0]).strip() == "#":
                header_row_idx = i
                break
        if header_row_idx is None:
            for i, row in enumerate(rows):
                if row and any(str(c).strip().lower() == "title" for c in row if c):
                    header_row_idx = i
                    break
        if header_row_idx is None:
            return []

        headers = [str(h).strip().lower() if h else "" for h in rows[header_row_idx]]

        def get_col(row, name_fragments):
            for frag in name_fragments:
                for j, h in enumerate(headers):
                    if frag in h:
                        val = row[j] if j < len(row) else None
                        if val and str(val).strip() not in ("None","NULL",""):
                            return str(val).strip()
            return ""

        def is_placeholder(text: str) -> bool:
            """Detect junk CrossRef placeholder values."""
            if not text:
                return True
            t = text.strip()
            if t.startswith("Retrieved via CrossRef"):
                return True
            if t.startswith("• See: https://doi.org/"):
                return True
            if t == "See: https://doi.org/" or "See: https://doi.org/" in t:
                return True
            if len(t) < 10:
                return True
            return False

        articles = []
        for row in rows[header_row_idx + 1:]:
            if not row or not any(row):
                continue
            title = get_col(row, ["title"])
            if not title or title == "#":
                continue

            authors_raw = get_col(row, ["author"])
            if ";" in authors_raw:
                authors = [a.strip() for a in authors_raw.split(";") if a.strip()]
            elif "," in authors_raw and len(authors_raw) > 40:
                authors = [authors_raw]  # single author with comma in name
            elif authors_raw:
                authors = [a.strip() for a in authors_raw.split(",") if a.strip()]
            else:
                authors = []

            findings_raw = get_col(row, ["finding", "major finding"])
            # Clean bullet points and split on newlines
            findings = [f.lstrip("•●-– ").strip()
                        for f in findings_raw.split("\n") if f.strip()]
            findings = [f for f in findings if not is_placeholder(f)]

            objectives_raw = get_col(row, ["objective"])
            objectives = "" if is_placeholder(objectives_raw) else objectives_raw

            methodology_raw = get_col(row, ["methodology", "method"])
            methodology = "" if is_placeholder(methodology_raw) else methodology_raw

            doi_raw = get_col(row, ["doi", "url"])
            doi = doi_raw.replace("https://doi.org/","").replace("http://doi.org/","").strip()
            # Remove anything after whitespace in DOI
            doi = doi.split()[0] if doi else ""

            articles.append({
                "title":           title,
                "authors":         authors,
                "year":            get_col(row, ["year"]),
                "journal":         get_col(row, ["journal", "source"]),
                "doi":             doi,
                "main_objectives": objectives,
                "methodology":     methodology,
                "major_findings":  findings,
                "keywords":        [k.strip() for k in get_col(row, ["keyword"]).split(";") if k.strip()],
                "limitations":     "",
                "conclusion":      "",
                "_needs_enrichment": not objectives or not findings,
            })
        return articles
    except Exception:
        return []


def extract_dois_from_text(text: str) -> list:
    """Extract all DOIs from plain text using regex patterns."""
    import re
    patterns = [
        r'10\.\d{4,9}/[-._;()/:A-Z0-9a-z]+',           # raw DOI
        r'https?://(?:dx\.)?doi\.org/(10\.\d{4,9}/[-._;()/:A-Z0-9a-z]+)',  # URL form
        r'DOI:\s*(10\.\d{4,9}/[-._;()/:A-Z0-9a-z]+)',  # "DOI: ..." form
    ]
    found = set()
    for pat in patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            doi = m.group(1) if m.lastindex else m.group(0)
            # Clean trailing punctuation
            doi = doi.rstrip('.,;)')
            if doi.startswith('10.'):
                found.add(doi)
    return list(found)


def parse_word_references(file_bytes, filename: str) -> tuple:
    """
    Extract DOIs and plain references from a Word (.docx) or text file.
    Returns (doi_list, plain_refs_list, raw_text).
    """
    raw_text = ""
    ext = filename.split(".")[-1].lower()
    try:
        if ext in ("docx", "doc"):
            try:
                import docx as _docx_mod
                import io as _io4
                doc = _docx_mod.Document(_io4.BytesIO(file_bytes))
                raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                raw_text = file_bytes.decode("utf-8", errors="ignore")
        else:
            raw_text = file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        raw_text = ""

    dois = extract_dois_from_text(raw_text)

    # Also parse plain reference lines (lines that look like citations)
    import re
    plain_refs = []
    for line in raw_text.split("\n"):
        line = line.strip()
        # Skip very short lines and headings
        if len(line) < 30:
            continue
        # Looks like a reference: starts with author pattern or number
        if re.match(r'^(\d+\.?\s|[A-Z][a-z]+,\s*[A-Z]\.?|\[\d+\])', line):
            plain_refs.append(line)

    return dois, plain_refs, raw_text


def _ai_summarise_from_meta(api_key: str, model: str, meta: dict) -> dict:
    """
    Ask the AI to generate objectives, methodology, findings, keywords
    for an article given only its CrossRef metadata (title, authors, journal, abstract if any).
    Returns a dict with the enriched fields.
    """
    title   = meta.get("title","")
    authors = "; ".join(meta.get("authors",[])) if meta.get("authors") else ""
    year    = meta.get("year","")
    journal = meta.get("journal","")
    doi     = meta.get("doi","")

    sys_prompt = (
        "You are an academic research analyst. Based on the bibliographic details of a research article, "
        "infer its likely content and return a JSON object with these keys:\n"
        "- \"main_objectives\": 2 sentences describing what the study likely investigated\n"
        "- \"methodology\": 1-2 sentences describing the likely research design and methods\n"
        "- \"major_findings\": list of 3-4 plausible key findings (strings)\n"
        "- \"keywords\": list of 5-6 relevant keywords\n"
        "- \"limitations\": 1 sentence on likely limitations\n"
        "- \"conclusion\": 1 sentence conclusion\n"
        "Base your response on the title, authors, journal, and year. "
        "Do NOT invent specific statistics. Write in academic prose. "
        "Return ONLY valid JSON."
    )
    user_msg = (
        f"Title: {title}\n"
        f"Authors: {authors}\n"
        f"Year: {year}\n"
        f"Journal: {journal}\n"
        f"DOI: {doi}\n\n"
        "Generate the structured summary for this article."
    )
    try:
        resp = call_groq(api_key, model, sys_prompt, user_msg, max_tokens=800, stream=False)
        raw  = resp.json()["choices"][0]["message"]["content"].strip()
        raw  = re.sub(r"^```(?:json)?\s*","",raw)
        raw  = re.sub(r"\s*```$","",raw)
        parsed = json.loads(raw)
        return {
            "main_objectives": parsed.get("main_objectives",""),
            "methodology":     parsed.get("methodology",""),
            "major_findings":  parsed.get("major_findings",[]) if isinstance(parsed.get("major_findings",[]),list) else [str(parsed.get("major_findings",""))],
            "keywords":        parsed.get("keywords",[]) if isinstance(parsed.get("keywords",[]),list) else [],
            "limitations":     parsed.get("limitations",""),
            "conclusion":      parsed.get("conclusion",""),
        }
    except Exception:
        return {
            "main_objectives": f"Study on {title[:80]}",
            "methodology":     "Refer to the original article for methodology details.",
            "major_findings":  [f"See full article at https://doi.org/{doi}" if doi else "See original article"],
            "keywords":        [],
            "limitations":     "",
            "conclusion":      "",
        }


def fetch_articles_from_dois(api_key: str, model: str, dois: list,
                               plain_refs: list, raw_text: str,
                               progress_cb=None) -> list:
    """
    For each DOI: fetch CrossRef metadata then AI-enrich with objectives,
    methodology, findings, keywords.
    For plain refs without DOIs: AI parses and enriches each one.
    Returns fully populated article dicts ready for LR chapter writing.
    """
    articles = []
    total_steps = len(dois) + (1 if plain_refs else 0)

    # ── Process DOIs: CrossRef metadata + AI enrichment ──────────────
    for i, doi in enumerate(dois):
        if progress_cb:
            progress_cb(i, total_steps,
                        f"📡 Fetching DOI {i+1}/{len(dois)}: {doi[:45]}…")
        meta = {}
        try:
            meta = fetch_doi_metadata(doi)
        except Exception:
            meta = {"error": "fetch_failed"}

        if "error" not in meta and meta.get("title"):
            # AI-enrich with objectives, methodology, findings, keywords
            if progress_cb:
                progress_cb(i, total_steps,
                            f"🤖 AI enriching {i+1}/{len(dois)}: {meta.get('title','')[:40]}…")
            enriched = _ai_summarise_from_meta(api_key, model, meta) if api_key else {
                "main_objectives": "", "methodology": "",
                "major_findings": [], "keywords": [], "limitations": "", "conclusion": ""
            }
            articles.append({
                "title":           meta.get("title",""),
                "authors":         meta.get("authors",[]),
                "year":            meta.get("year",""),
                "journal":         meta.get("journal",""),
                "doi":             doi,
                "main_objectives": enriched["main_objectives"],
                "methodology":     enriched["methodology"],
                "major_findings":  enriched["major_findings"],
                "keywords":        enriched["keywords"],
                "limitations":     enriched["limitations"],
                "conclusion":      enriched["conclusion"],
            })
        else:
            # CrossRef failed — still store with empty fields for manual editing
            articles.append({
                "title":           f"Article (DOI: {doi})",
                "authors":         [],
                "year":            "",
                "journal":         "",
                "doi":             doi,
                "main_objectives": "",
                "methodology":     "",
                "major_findings":  [],
                "keywords":        [],
                "limitations":     "",
                "conclusion":      "",
            })

    # ── Process plain references via AI ──────────────────────────────
    if plain_refs and api_key:
        if progress_cb:
            progress_cb(len(dois), total_steps,
                        f"🤖 AI parsing {len(plain_refs)} plain references…")

        # Batch parse all plain refs in one call for speed
        refs_block = "\n".join(f"{j+1}. {r}" for j,r in enumerate(plain_refs[:50]))
        sys_prompt = (
            "You are an academic reference parser AND research analyst. "
            "Extract structured data AND generate research summaries from the reference list.\n"
            "Return a JSON array where each element has:\n"
            "  \"title\", \"authors\" (list of strings), \"year\", \"journal\", "
            "\"doi\" (or empty string),\n"
            "  \"main_objectives\" (2 sentences inferring what the study investigated),\n"
            "  \"methodology\" (1-2 sentences on likely research design),\n"
            "  \"major_findings\" (list of 3 plausible key findings),\n"
            "  \"keywords\" (list of 5 keywords)\n"
            "Return ONLY valid JSON array. No markdown."
        )
        user_msg = f"Parse and enrich these references:\n\n{refs_block}"
        try:
            resp = call_groq(api_key, model, sys_prompt, user_msg,
                             max_tokens=4096, stream=False)
            raw  = resp.json()["choices"][0]["message"]["content"].strip()
            raw  = re.sub(r"^```(?:json)?\s*","",raw)
            raw  = re.sub(r"\s*```$","",raw)
            parsed_refs = json.loads(raw)
            existing_dois = {a.get("doi","") for a in articles}
            for ref in parsed_refs:
                if not ref.get("title"):
                    continue
                if ref.get("doi","") and ref["doi"] in existing_dois:
                    continue
                findings = ref.get("major_findings",[])
                if not isinstance(findings, list):
                    findings = [str(findings)] if findings else []
                authors = ref.get("authors",[])
                if not isinstance(authors, list):
                    authors = [str(authors)] if authors else []
                keywords = ref.get("keywords",[])
                if not isinstance(keywords, list):
                    keywords = [str(keywords)] if keywords else []
                articles.append({
                    "title":           ref.get("title",""),
                    "authors":         authors,
                    "year":            str(ref.get("year","")),
                    "journal":         ref.get("journal",""),
                    "doi":             ref.get("doi",""),
                    "main_objectives": ref.get("main_objectives",""),
                    "methodology":     ref.get("methodology",""),
                    "major_findings":  findings,
                    "keywords":        keywords,
                    "limitations":     "",
                    "conclusion":      "",
                })
        except Exception:
            pass

    return articles


def build_journal_excel(jm_result: dict) -> bytes:
    """Build a formatted Excel workbook from Journal Matcher results."""
    if not XLSX_OK:
        return b""
    import io as _io_jm
    from openpyxl import Workbook as JWB
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = JWB()
    ws = wb.active
    ws.title = "Journal Recommendations"

    # ── Colors ─────────────────────────────────────────────────────
    HDR_FILL  = PatternFill("solid", fgColor="0F3312")
    ALT_FILL  = PatternFill("solid", fgColor="F0F7F0")
    WHT_FILL  = PatternFill("solid", fgColor="FFFFFF")
    GRN_FILL  = PatternFill("solid", fgColor="1E5C22")
    THIN      = Side(style="thin", color="3A8C3F")
    BORDER    = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

    def hdr_font(sz=11): return Font(name="Calibri", bold=True, color="FFFFFF", size=sz)
    def body_font(bold=False): return Font(name="Calibri", bold=bold, color="0A1E0B", size=10)
    def green_font(): return Font(name="Calibri", bold=True, color="1E5C22", size=10)

    # ── Title row ───────────────────────────────────────────────────
    ws.merge_cells("A1:J1")
    ws["A1"] = f"Journal Recommendations — {jm_result.get('research_domain','')}"
    ws["A1"].font = Font(name="Calibri", bold=True, color="FFFFFF", size=14)
    ws["A1"].fill = PatternFill("solid", fgColor="0F3312")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 28

    # ── Header row ──────────────────────────────────────────────────
    headers = ["#", "Journal Name", "Publisher", "Index", "Impact Factor",
               "Open Access", "Acceptance Rate", "Turnaround (wks)", "Scope Fit", "Submission URL"]
    col_widths = [4, 38, 24, 14, 12, 12, 16, 16, 50, 40]

    for col, (hdr, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=2, column=col, value=hdr)
        cell.font    = hdr_font()
        cell.fill    = PatternFill("solid", fgColor="1E5C22")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border  = BORDER
        ws.column_dimensions[chr(64+col)].width = w
    ws.row_dimensions[2].height = 20

    # ── Data rows ───────────────────────────────────────────────────
    recs = jm_result.get("recommendations", [])
    for i, j in enumerate(recs, 1):
        row   = i + 2
        fill  = ALT_FILL if i % 2 == 0 else WHT_FILL
        data  = [
            i,
            j.get("name", ""),
            j.get("publisher", ""),
            j.get("index", ""),
            str(j.get("impact_factor", "")),
            j.get("open_access", ""),
            str(j.get("acceptance_rate", "")),
            str(j.get("turnaround", "")),
            j.get("scope_fit", ""),
            j.get("submission_url", ""),
        ]
        for col, val in enumerate(data, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.fill      = fill
            cell.border    = BORDER
            cell.alignment = Alignment(vertical="top", wrap_text=True,
                                        horizontal="center" if col in (1,4,5,6,7,8) else "left")
            if col == 1:
                cell.font = green_font()
            elif col == 2:
                cell.font = body_font(bold=True)
            elif col == 4:
                # Index badge — highlight by tier
                tier_colors = {"ABDC A":"1A4A1C","ABDC B":"2D5C1E","Scopus":"0C3A7C",
                               "ABDC C":"5C4A00","UGC CARE":"5C1A1A","ABS":"3A1A5C"}
                tier_fill   = tier_colors.get(str(val), "1E5C22")
                cell.fill   = PatternFill("solid", fgColor=tier_fill)
                cell.font   = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
            else:
                cell.font = body_font()
        ws.row_dimensions[row].height = max(20, min(60, len(str(data[8]))//3))

    # ── Strategy note ───────────────────────────────────────────────
    strategy = jm_result.get("strategy_note", "")
    if strategy:
        note_row = len(recs) + 4
        ws.merge_cells(f"A{note_row}:J{note_row}")
        ws[f"A{note_row}"] = f"📌 Submission Strategy: {strategy}"
        ws[f"A{note_row}"].font      = Font(name="Calibri", italic=True, color="1E5C22", size=10)
        ws[f"A{note_row}"].alignment = Alignment(wrap_text=True, vertical="top")
        ws[f"A{note_row}"].fill      = PatternFill("solid", fgColor="F0F7F0")
        ws.row_dimensions[note_row].height = max(20, len(strategy)//6)

    # ── Avoid list ─────────────────────────────────────────────────
    avoid = jm_result.get("avoid", [])
    if avoid:
        avoid_row = len(recs) + 6
        ws.merge_cells(f"A{avoid_row}:J{avoid_row}")
        ws[f"A{avoid_row}"] = "⚠️ Journals to avoid: " + " | ".join(str(a) for a in avoid)
        ws[f"A{avoid_row}"].font      = Font(name="Calibri", color="CC3333", size=10)
        ws[f"A{avoid_row}"].alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[avoid_row].height = 20

    # Freeze header rows
    ws.freeze_panes = "A3"

    buf = _io_jm.BytesIO()
    wb.save(buf)
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG & CSS
# ══════════════════════════════════════════════════════════════════════════════

st.set_page_config(page_title="Zodha Research Writing Pro", page_icon="🔬",
                   layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700;900&family=DM+Sans:wght@300;400;500;600&family=DM+Mono:wght@400;500&display=swap');
:root{--ink:#1a2e1b;--cream:#f0f7f0;--gold:#3a8c3f;--gold-lt:#a8d5aa;--rust:#1e5c22;--slate:#3d5e40;--border:#b0d4b2;--green-bright:#4caf50;--green-dark:#1e5c22;--white:#ffffff;}
html,body,[class*="css"]{font-family:'DM Sans',sans-serif;background-color:#ffffff!important;color:var(--ink);}
#MainMenu,footer,header{visibility:hidden;}
.block-container{padding-top:1.2rem!important;max-width:1400px!important;}
.hero-banner{background:linear-gradient(135deg,#1a2e1b 0%,#1e5c22 50%,#2d7a32 100%);border-radius:16px;padding:1.5rem 2.5rem 1.3rem;margin-bottom:1.5rem;position:relative;overflow:hidden;}
.hero-title{font-family:'Playfair Display',serif;font-size:2.2rem;font-weight:900;color:#ffffff;margin:0 0 0.2rem;position:relative;}
.hero-sub{font-size:0.9rem;color:rgba(255,255,255,0.6);margin:0;position:relative;}
.hero-badge{position:absolute;top:1.2rem;right:2rem;background:rgba(255,255,255,0.15);border:1px solid rgba(255,255,255,0.5);color:#ffffff;border-radius:20px;padding:0.3rem 0.9rem;font-size:0.75rem;font-family:'DM Mono',monospace;letter-spacing:1px;text-transform:uppercase;}
.card-title{font-family:'Playfair Display',serif;font-size:1.05rem;font-weight:700;color:#ffffff!important;margin-bottom:0.8rem;padding-bottom:0.5rem;border-bottom:2px solid #4caf50;}
.score-ring-wrap{display:flex;flex-direction:column;align-items:center;gap:0.3rem;}
.score-ring{width:80px;height:80px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-family:'Playfair Display',serif;font-size:1.5rem;font-weight:900;}
.score-label{font-size:0.68rem;font-weight:600;text-transform:uppercase;letter-spacing:1px;color:#3d5e40;text-align:center;}
.metric-row{display:flex;flex-wrap:wrap;gap:0.5rem;margin-top:0.4rem;}
.metric-chip{background:#f0f7f0;border:1px solid #b0d4b2;border-radius:7px;padding:0.35rem 0.7rem;font-size:0.78rem;color:#2d5c30;}
.metric-chip b{color:#1a2e1b;font-weight:600;}
.output-box{background:white;border:1.5px solid #3a8c3f;border-radius:10px;padding:1rem 1.2rem;height:380px;overflow-y:auto;overflow-x:hidden;font-family:'DM Sans',sans-serif;font-size:0.9rem;line-height:1.75;color:#1a2e1b;word-break:break-word;box-sizing:border-box;text-align:justify;}
.output-box-sm{background:white;border:1.5px solid #3a8c3f;border-radius:10px;padding:1rem 1.2rem;height:320px;overflow-y:auto;font-family:'DM Sans',sans-serif;font-size:0.9rem;line-height:1.75;color:#1a2e1b;word-break:break-word;text-align:justify;}
.wc-badge{display:inline-block;background:#f0f7f0;border:1px solid #b0d4b2;border-radius:6px;padding:0.2rem 0.6rem;font-family:'DM Mono',monospace;font-size:0.75rem;color:#2d5c30;margin-top:0.3rem;}
textarea{font-family:'DM Sans',sans-serif!important;font-size:0.9rem!important;line-height:1.65!important;border-radius:10px!important;border:1.5px solid #b0d4b2!important;background:white!important;color:#1a2e1b!important;}
div[data-testid="stButton"]>button[kind="primary"]{background:linear-gradient(135deg,#1e5c22,#3a8c3f)!important;color:#ffffff!important;border:1.5px solid #3a8c3f!important;border-radius:10px!important;font-family:'DM Sans',sans-serif!important;font-weight:600!important;transition:all 0.25s!important;}
div[data-testid="stButton"]>button[kind="primary"]:hover{transform:translateY(-2px)!important;box-shadow:0 6px 20px rgba(26,26,46,0.3)!important;}
div[data-testid="stRadio"] label{border:1px solid rgba(58,140,63,0.4)!important;border-radius:8px!important;padding:0.4rem 0.9rem!important;background:#f0f7f0!important;}
div[data-testid="stRadio"] label>div,div[data-testid="stRadio"] label span,div[data-testid="stRadio"] label p{color:#1a2e1b!important;font-weight:500!important;}
[data-testid="stSidebar"]{background:#1a2e1b!important;border-right:1px solid rgba(58,140,63,0.3)!important;}
[data-testid="stSidebar"] p,[data-testid="stSidebar"] span:not([data-testid]),[data-testid="stSidebar"] .stMarkdown p{color:rgba(255,255,255,0.85)!important;}
[data-testid="stSidebar"] h3{color:#4caf50!important;font-family:'Playfair Display',serif!important;}
[data-testid="stSidebar"] input{background:rgba(255,255,255,0.1)!important;color:white!important;border:1px solid rgba(76,175,80,0.5)!important;border-radius:8px!important;}
[data-testid="stSidebar"] input::placeholder{color:rgba(255,255,255,0.35)!important;}
[data-testid="column"]{overflow:hidden!important;min-width:0!important;}
section.main .stMarkdown p{color:#1a2e1b!important;}
section.main strong,section.main b{color:#1a2e1b!important;}
div[data-testid="stAlert"] p{color:#1a2e1b!important;}
.improvement-banner{background:linear-gradient(135deg,#1a2e1b,#1e5c22);border:1px solid rgba(74,124,89,0.5);border-radius:12px;padding:1rem 1.5rem;display:flex;align-items:center;gap:1rem;margin:1rem 0;}
.improvement-banner .score-delta{font-family:'Playfair Display',serif;font-size:2rem;font-weight:900;white-space:nowrap;}
.improvement-banner .score-desc{font-size:0.88rem;color:rgba(255,255,255,0.8);line-height:1.5;}
/* Admin panel */
.admin-hero{background:linear-gradient(135deg,#0a1e0b 0%,#1a2e1b 50%,#1e5c22 100%);border-radius:16px;padding:1.5rem 2.5rem;margin-bottom:1.5rem;position:relative;}
.admin-hero-title{font-family:'Playfair Display',serif;font-size:2rem;font-weight:900;color:#4caf50;}
.admin-badge{position:absolute;top:1.2rem;right:2rem;background:rgba(76,175,80,0.15);border:1px solid rgba(76,175,80,0.4);color:#4caf50;border-radius:20px;padding:0.3rem 0.9rem;font-size:0.75rem;font-family:'DM Mono',monospace;text-transform:uppercase;letter-spacing:1px;}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════════
for k,v in [("output_text",""),("paraphrase_out",""),
             ("grammar_corrected",""),("grammar_issues",[]),
             ("admin_provider", _DEFAULT_PROVIDER),("platform_or_key",""),
             ("clear_input",False),("clear_para",False),("clear_gram",False),
             ("clear_spss",False),("clear_meth",False),("clear_hyp",False),
             ("clear_jm",False),("clear_cl",False),("clear_rv",False),("clear_cs",False),
             ("show_admin",False),("admin_auth",False),
             ("spss_result",{}),("method_result",{}),("hyp_result",{}),("var_result",{}),
             ("lit_chapter_text",""),("lit_chapter_topic",""),("lit_chapter_style","APA 7th"),
             ("journal_result",{}),("coverletter_result",{}),
             ("reviewer_result",{}),("contrib_result",{})]:
    if k not in st.session_state: st.session_state[k]=v

if st.session_state.clear_input:
    st.session_state.input_text=""
    st.session_state.output_text=""
    st.session_state.clear_input=False

if st.session_state.clear_para:
    st.session_state.para_input=""
    st.session_state.paraphrase_out=""
    st.session_state.clear_para=False

if st.session_state.clear_gram:
    st.session_state.gram_input=""
    st.session_state.grammar_corrected=""
    st.session_state.grammar_issues=[]
    st.session_state.clear_gram=False

if st.session_state.clear_spss:
    st.session_state.spss_input_paste=""
    st.session_state.spss_result={}
    st.session_state.clear_spss=False

if st.session_state.clear_meth:
    st.session_state.meth_input_paste=""
    st.session_state.method_result={}
    st.session_state.clear_meth=False

if st.session_state.clear_hyp:
    st.session_state.hyp_topic=""
    st.session_state.hyp_vars=""
    st.session_state.hyp_result={}
    st.session_state.clear_hyp=False

if st.session_state.clear_jm:
    st.session_state.jm_abstract_paste=""
    st.session_state.journal_result={}
    st.session_state.clear_jm=False

if st.session_state.clear_cl:
    st.session_state.cl_abstract_paste=""
    st.session_state.cl_title=""
    st.session_state.cl_journal=""
    st.session_state.cl_authors=""
    st.session_state.coverletter_result={}
    st.session_state.clear_cl=False

if st.session_state.clear_rv:
    st.session_state.rv_comments_paste=""
    st.session_state.rv_abstract_paste=""
    st.session_state.reviewer_result={}
    st.session_state.clear_rv=False

if st.session_state.clear_cs:
    for k in ["cs_title","cs_obj","cs_findings","cs_context"]:
        st.session_state[k]=""
    st.session_state.contrib_result={}
    st.session_state.clear_cs=False

# ══════════════════════════════════════════════════════════════════════════════
# ADMIN PANEL (hidden, password-protected overlay)
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.show_admin:
    if not st.session_state.admin_auth:
        st.markdown("""<div class="admin-hero">
          <div class="admin-badge">Admin</div>
          <div class="admin-hero-title">🛡️ Admin Access</div>
          <div style="color:rgba(255,255,255,0.5);font-size:0.88rem;">Zodha Research Solutions</div>
        </div>""", unsafe_allow_html=True)
        _, cc, _ = st.columns([1,1.2,1])
        with cc:
            with st.container(border=True):
                pw = st.text_input("Admin Password", type="password",
                                   placeholder="Enter admin password", label_visibility="collapsed")
                if st.button("🔐 Enter", type="primary", use_container_width=True):
                    if pw == ADMIN_PASSWORD:
                        st.session_state.admin_auth = True
                        st.rerun()
                    else:
                        st.error("❌ Incorrect password.")
            if st.button("← Back to App", use_container_width=True):
                st.session_state.show_admin = False
                st.rerun()
        st.stop()

    # ── ADMIN DASHBOARD ──────────────────────────────────────────────────────
    st.markdown("""<div class="admin-hero">
      <div class="admin-badge">Admin · Zodha</div>
      <div class="admin-hero-title">🛡️ Admin Dashboard</div>
      <div style="color:rgba(255,255,255,0.5);font-size:0.88rem;">Zodha Research Writing Pro · Platform Settings</div>
    </div>""", unsafe_allow_html=True)

    with st.sidebar:
        st.markdown('<div style="color:#4caf50;font-family:Playfair Display,serif;font-weight:700;font-size:1rem;padding:0.5rem 0;">🛡️ Admin Mode</div>', unsafe_allow_html=True)
        if st.button("← Exit Admin", use_container_width=True):
            st.session_state.show_admin = False
            st.session_state.admin_auth = False
            st.rerun()

    adm_tab1, adm_tab2 = st.tabs(["⚙️ API Key Settings", "📊 App Info"])

    with adm_tab1:
        st.markdown('<div class="card-title">⚙️ API Key & Model Settings</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Select provider, enter API key(s) and choose model. Settings apply to all users.</p>', unsafe_allow_html=True)

        # ── Provider selector ─────────────────────────────────────
        st.markdown('<p style="color:#ffffff;font-weight:600;margin-top:0.5rem;">🔀 AI Provider</p>', unsafe_allow_html=True)
        current_provider = st.session_state.get("admin_provider", _DEFAULT_PROVIDER)
        selected_provider = st.radio(
            "Provider",
            ["openrouter", "groq"],
            index=0 if current_provider == "openrouter" else 1,
            format_func=lambda x: "🌐 OpenRouter (Claude Haiku / free models)" if x == "openrouter" else "⚡ Groq (Llama / Mixtral / Gemma)",
            label_visibility="collapsed",
            horizontal=True,
        )

        # ── API Keys ──────────────────────────────────────────────
        if selected_provider == "openrouter":
            st.markdown('<p style="color:#ffffff;font-weight:600;margin-top:1rem;">🔑 OpenRouter API Key</p>', unsafe_allow_html=True)
            st.caption("Get free key → openrouter.ai/keys · Free models available · Claude Haiku $1/$5 per M tokens")
            stored_or  = st.session_state.get("platform_or_key",  os.environ.get("OPENROUTER_API_KEY", ""))
            new_or_key = st.text_input("OR Key", value=stored_or, type="password",
                                        placeholder="sk-or-v1-...", label_visibility="collapsed")
            if stored_or:
                masked = stored_or[:8] + "••••••••" + stored_or[-4:] if len(stored_or) > 12 else "••••••••"
                st.markdown(f'<div style="font-size:0.78rem;color:#6fcf97;">Active key: {masked}</div>', unsafe_allow_html=True)
            new_key = new_or_key
        else:
            st.markdown('<p style="color:#ffffff;font-weight:600;margin-top:1rem;">🔑 Groq API Key</p>', unsafe_allow_html=True)
            st.caption("Get free key → console.groq.com · Always free · Rate limits apply")
            stored_gq  = st.session_state.get("platform_groq_key", os.environ.get("GROQ_API_KEY", ""))
            new_gq_key = st.text_input("Groq Key", value=stored_gq, type="password",
                                        placeholder="gsk_...", label_visibility="collapsed")
            if stored_gq:
                masked = stored_gq[:8] + "••••••••" + stored_gq[-4:] if len(stored_gq) > 12 else "••••••••"
                st.markdown(f'<div style="font-size:0.78rem;color:#6fcf97;">Active key: {masked}</div>', unsafe_allow_html=True)
            new_key = new_gq_key

        # ── Model selector — shows correct list per provider ──────
        st.markdown('<p style="color:#ffffff;font-weight:600;margin-top:1.2rem;">🤖 Default AI Model</p>', unsafe_allow_html=True)
        model_dict    = OPENROUTER_MODELS if selected_provider == "openrouter" else GROQ_MODELS
        current_model = st.session_state.get("admin_model_choice", list(model_dict.keys())[0])
        if current_model not in model_dict:
            current_model = list(model_dict.keys())[0]
        selected_model = st.selectbox(
            "Default Model",
            list(model_dict.keys()),
            index=list(model_dict.keys()).index(current_model),
            format_func=lambda x: model_dict[x],
            label_visibility="collapsed",
        )
        model_tips = {
            # OpenRouter
            "anthropic/claude-haiku-4-5":              "⭐ Best AI-detection bypass — Claude token patterns fool GPTZero & Originality.ai",
            "anthropic/claude-3-5-haiku":              "🏆 Strongest quality + undetection — ideal for long academic texts",
            "anthropic/claude-3-haiku":                "⚡ Ultra cheap Claude — still beats Llama for undetection",
            "google/gemini-flash-1.5":                 "💡 Google Gemini — good quality, free tier available on OpenRouter",
            "meta-llama/llama-3.3-70b-instruct:free":  "🆓 Completely free · GPT-4 level quality · 200 req/day limit",
            "mistralai/mistral-7b-instruct:free":      "🆓 Completely free · Use for grammar/paraphrase only",
            # Groq
            "llama-3.3-70b-versatile": "🏆 Best Groq model — fast, free, GPT-4 level quality",
            "llama-3.1-8b-instant":    "⚡ Ultra fast — good for grammar checking and paraphrasing",
            "mixtral-8x7b-32768":      "📄 Long context — best for large document processing",
            "gemma2-9b-it":            "⚖️ Balanced — solid all-rounder for mixed tasks",
        }
        st.caption(model_tips.get(selected_model, ""))

        if st.button("💾 Save Settings", type="primary", use_container_width=True):
            st.session_state.admin_provider     = selected_provider
            st.session_state.admin_model_choice = selected_model
            # Store key in the correct slot
            if selected_provider == "openrouter":
                st.session_state.platform_or_key   = new_key
                st.session_state.platform_groq_key = st.session_state.get("platform_groq_key", "")
            else:
                st.session_state.platform_groq_key = new_key
                st.session_state.platform_or_key   = st.session_state.get("platform_or_key", "")
            st.success(f"✅ Saved — {selected_provider.upper()} · {model_dict[selected_model]}")

        st.markdown("---")
        st.markdown("**To persist permanently (Streamlit Cloud):**")
        st.code('''# Set EITHER or BOTH in Streamlit Cloud → Settings → Secrets:
OPENROUTER_API_KEY = "sk-or-v1-your_key_here"   # openrouter.ai/keys
GROQ_API_KEY       = "gsk_your_key_here"          # console.groq.com
ADMIN_PASSWORD     = "your_admin_password"
# Switch provider anytime in Admin panel — no redeploy needed''', language="toml")

    with adm_tab2:
        st.markdown('<div class="card-title">📊 App Info</div>', unsafe_allow_html=True)
        info_items = [
            ("Version", "v6.2 · Zodha"),
            ("Mode", "Single-user · No login required"),
            ("Tools", "Humanizer · Paraphraser · Grammar Checker · Research Tools · Statistics Suite · Publication Suite"),
            ("Backend", f"{st.session_state.get('admin_provider', _DEFAULT_PROVIDER).upper()} API"),
            ("OpenRouter models", "Claude Haiku 4.5 · Claude 3.5 Haiku · Gemini Flash · Llama 3.3 70B (free)"),
            ("Groq models", "Llama 3.3 70B · Llama 3.1 8B · Mixtral 8x7B · Gemma 2 9B"),
            ("Max Words/Session", "Unlimited"),
        ]
        for label, val in info_items:
            st.markdown(f"""<div style="background:white;border:1px solid #d4c9b5;border-radius:8px;
                 padding:0.6rem 1rem;margin-bottom:0.4rem;display:flex;justify-content:space-between;">
              <span style="color:#5a6a7a;font-size:0.85rem;">{label}</span>
              <span style="font-weight:600;color:#1a2e1b;font-size:0.85rem;">{val}</span>
            </div>""", unsafe_allow_html=True)

    st.stop()

# ══════════════════════════════════════════════════════════════════════════════
# MAIN APP — No login, opens directly
# ══════════════════════════════════════════════════════════════════════════════

# ── SIDEBAR ────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown(f'''<div style="display:flex;align-items:center;gap:0.7rem;margin-bottom:1rem;padding-bottom:0.8rem;border-bottom:1px solid rgba(76,175,80,0.3);">
      <img src="data:image/png;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAUDBAQEAwUEBAQFBQUGBwwIBwcHBw8LCwkMEQ8SEhEPERETFhwXExQaFRERGCEYGh0dHx8fExciJCIeJBweHx7/2wBDAQUFBQcGBw4ICA4eFBEUHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh7/wAARCAGGBOoDASIAAhEBAxEB/8QAHQABAAICAwEBAAAAAAAAAAAAAAgJBgcBBAUDAv/EAGMQAAEDAwEDBQgIEQgHBwMFAQABAgMEBQYRBwghEhMxQVEiYXGBkbGy0RQyN3J0dZShFRYYIzQ2QlJTVVZic5KTs8EJFzM1VILS4SQnQ0ZkZYQlOGODoqPCJkXwRFeFlaTT/8QAGgEBAAMBAQEAAAAAAAAAAAAAAAQFBgMBAv/EADQRAQACAgEDAgMGBQMFAAAAAAABAgMEEQUhMRJBIlGxEzJhkaHRI0JxgeEUFcEGM1Lw8f/aAAwDAQACEQMRAD8Ai4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB6+NY3esiquYtVDJPp7eTTRjO+53Qh5MxEcyeHkHLUVy6NRVVepDc+M7GadvJlyG4ukdwVYKXgnT0K5U8yGwrNi+O2ZrUttmpYXN4pI5nLk/WdqpW5+r62H35/oi5NzFTtzyjdbsVyO4ta6jstdK1y6I5IlRPKpklv2SZjVKvPQUdEidc9QnHxN1UkMrnKmnKXQ4KzJ/wBQWn7lEW3UJ/lq0guxK+pHyvo7ZXLp7Vqzar5Y0Q+X8zF//Gls8r/8JvQEW3W9mfHH5OU72Vo9mxO+vZyvo7ZWL969ZtfmjVDzK/ZFmFMiLDHRViKv+wqE1T9ZEJBHPiOlOu54+9ES9r1C8eYhFq7Yfk9q5Xs6y1kbU6XJHym+VNUPDexzHK17Va5OlFTRSYKOcnQqoebdbDY7s1W3S0UdVqmnKdHo5PA5NFJuLr2Of+5WYSKdQpP3oRPBvDJNjdrqFfLYbhLRvXi2Co7tngR3Snj1NWZTiN/xqREutBJHE5dGTt7qN3gcnDxdJb4NvDnj+HblMx5qZPuy8IAEh0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAP0jHqmqMcqeA/JYju0WOy1GxHGpqi1UU0r6RFc+SBrlVfCqAV4c3J947yDm5PwbvIWr/S7YPxJbvkzPUPpdsH4kt3yZnqAqo5uT8G/yDmpPwb/IWrpj1hT/AOy275Mz1HP0v2L8S275Mz1AVT81J+Df5BzUn4N/kLWPpfsP4lt3yZnqH0v2H8S275Mz1AVT81J+Df5BzUn4N/kLWPpfsP4lt3yZnqH0v2H8S275Mz1AVT81J+Df5BzUn4N/kLWPpfsP4lt3yZnqH0v2H8S275Mz1AVT81J+Df5BzUn4N/kLWPpfsX4lt3yZnqH0vWH8S275Mz1AVT83J+Dd5Bzcn4N/kLV/pdsH4lt3yZnqH0u2D8SW75Mz1AVUc3J+Dd5DhWPRNVa5PEWsfS7YPxJbvkzPUYBvBY5YmbIMimitFDFKykc5r2QNaqKnfRAK5QAAAAAAAAAAAAAAADsW6hq7jVspKGnkqJ5F0axiaqp6uF4tc8puaUlDHyY26LNO5O4jb2r3+xCROI4vacZoG01uhRZF4yzuT65IvfXqTvJwIG7v49SvfvPycM+xXDHfywbB9klLStZWZO9tTP0pSRu7hvvnda95OHfNo0sEFJTMpqSGOngYmjY42o1qeJD9omgMntdQzbM/FPb5eynzbN8k957AAITgHJwAAGir0H65uTTXku08B7FZn2e8S/IGip0opyvA8mOPLzhwAAB+Zo4p6d9PUQxzwSJo+ORqOa5O+in6B7W01nms93sTMTzDWObbJKCva+sxyRlDVdPsaT+if3mr0tXyp4DTF4tdfaK59FcqWSmnZ0temnjTtQlqePlmOWzJrY6iuUSKun1qZqd3Evai/wAC/wBLrVqzFM/ePn+6wwb0x8OT80VgZFnGIXXFK5Iq1nOU0irzFSxO4kT+C94x009bRaOY8LSJiY5gAB69AAAAAAAAAAB+4IZZ5WxQRvlkdwRrG6qviO8lhvjuiz3BfBTP9RnO7LGyTbNZGyNRycpy6KmvUWOpS0ydFPCngYgFVf0vX/8AElx+TP8AUc/S7f8A8SXL5K/1FqiU8Cf7CP8AVQcxB+Bj/VQCqr6Xb/8AiS5fJn+ofS9f/wASXH5M/wBRarzEH4GP9VB7Hp/wEf6qAVVfS9fvxLcfkz/UPpfv34luPyZ/qLVPY1P+Ai/UQexqf8BF+ogFVf0vX78S3H5M/wBR+VsN8Tps9wT/AKZ/qLVvY1N/Z4v1EOFpKVemmhX+4gFU0tpusTVdLbaxjUTVVdC5ET5jpqx6dLXJ4i2Ga2W6ZisloKV7XJoqOiaqL8x5VdhOIV0XNVeNWqZmuvJdTN9QFWQLJ7zsQ2XXWNzZ8QoIlcuvKgasa+VprrKN0vB65ZJLLcLhbJHe1ZyucY3y8fnAg6CQmZ7qec2lj5rJU0t4ibxRrV5EnkU0lkuMZBjVY6kvlpq6CVq6Kksaoi+BehQPIAAAAAAAAO/DZbxMxHw2qukaqaorYHKi/MfXEKaOsy2z0kzUdHPXQRvRetHSNRfOWmUFtoaWjhggo4I2MYjURsaJoiIBVl9L1/8AxJcfkz/UPpdv/wCJLl8lf6i1VKenTogi/VQcxB+Bj/VQCqr6Xr/+JLj8mf6h9L1//Elx+TP9Rar7Hp/wEf6qD2NT/gIv1UAqqXH78nTZbj8mf6jpVdLU0knN1VPLA/TXkyMVq+RS2H2NT/gIv1EISb+VPDBtHtiwxsj5VDqvJTTXugI5gAAAAB9qSkqqyRY6SmmqHomvJjYrl08CHxJGbhkMU20e6c7G1/JoNU5Sa/dIBoVMfvq9FmuPyZ/qOfpev/4kuPyZ/qLVPYtN/Z4v1EOfY1P+Ai/VQCqr6Xb/APiS5fJn+ofS9fvxLcfkz/UWq+x4PwMf6qHHsan/AAEX6iAVV/S9fvxLcfkz/Ufl1hvjUVzrPcGonSq07/UWrexqb+zxfqIfKspKV1JK11PEqKxfuE7AKnnIrXK1yKipwVFOD1cxajctu7WoiIlbMiInv1PKAAAAAAAAA/UMck0jY4o3SPcuiNamqr4j0G2C+uTVtmuCp3qZ/qN77j2ExX3OKrIq6mbLS2yPSPlt1asrvUhN5tLStTRtPCid5iAVV/S9f/xJcfkz/UPpev34luPyZ/qLVfY9P+Aj/VQex4PwEf6qAVVfS7f/AMSXL5K/1HP0u3/8SXL5K/1FqvMQfgY/1UHMQfgY/wBVAKqVx6/p02S4/Jn+o+U9mu8EayT2utiY1NVc+ByInlQtaWngXphj/VQ8+/2K3Xez1duqaOB8dRC6N2sadaaAVTg9nOLFUYzl91sNSxzZKKpfEnKTRVai9yvjTRTxgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAFke7EmmwvFu/RNK3CyTdj47C8V+AtA2SAAAAAAAAAAAAAAAAAABgW8H7jeTfAnmemA7wq6bGsl+BOArQAAAAAAAAAAAAADKdnuGV2WXLksR0FBEutRUKnBE7E7XKdrZxgVblU3sqVzqa2Rv0fMqcXr1tb2r3+okLarfR2q3xW+3wNgpoU0axvnXtUq+odSrrR6a97fT+qLs7MYo4jy6+O2S3WC2st9rg5qFvFVXi569auXrU9EAx+TJfLabXnmZUt7zeeZADq3W40Npt8lwudVHS00fS969K9iJ1r3kPmmO2S0VpHMy8rWbTxDtdeh0b1eLVZYFnutfBSNTqe7ul8DU4qaizTa9WVPOUmNRrRwrq1ap6ayuT83qb5zXtBQ33JblzdLBWXKrkXVVTV6+FV6vGaHV6F29Waf7LLFoe+SW3sh2yWmmV0dkt01c9OHOzrzbPEiaqvzGG1+1zLKhzlgdR0rV6EjhRdPGup7Vg2K10jWSX66w0aL7aGnTnZE8ftfOZxZtmWFW1rVfa5LjK3ReXVzOVFX3rVRPEupLtk6dqduImfzdptrYuzSFZnGXVjORNfq3k666MfyfNoec+9X10nOOutxV/as79fOSmobbaKB/Lt9ktdG/TRHQ0jGuROzXTU7vOL97H+zb6jn/vmvXtWs/o+P8AX448QitS5blFH/Q324s17ZnL5z16HafmVLwW5pOn/jRNd/AkVWUtBWta2vtlvrEb0c/Sxv08GqGOXbZ7hNzR3OWBlHI7/aUUrolT+7qrfmPqvV9TJ2vX84fUbmG3mGCWDbQ/lpHfbS1zFTRZaR2jvG1eC+U2Lj+X45ftG226RulX/Yy9xJ5F6fEa8v2xR6IslgvLJeyKrbyF/WTVDWl/sF8xutSC60NRRypxY9U7l3fa5OC+JT7to6O3HOPz+H7PZ18GaPh/RKtdUXRU0U5I+4btSvlmWOmuTludEnDSVfrjE/Nd0+JTduM5Basjt3s21VKStRdJI3cHxr2OT+JR7nS8ut8XmvzQM+pfF38w9QAFYiupebZQXm3SW+5Uzaimk6Wr0ovai9Sp2ke9peC1mKVyzQ8uotUq/WZ9OLfzXdi+ckefKtpqatpJaOsgZPTzN5Mkb01Ryf8A51lr07qV9afTbvX6JettTinifCIYM82mbParGXrcKHl1Npe7RH9LoV6mv/gpgZsMeSuSsWrPMSua2i0cwAA+30AAAAAAAA2Zuxe7RZPfO9EshK3d2L3aLH793mLIgAAAAAAAAAAAAAAAAB5WRY7ZMhoX0V5ttNWwPTRWyxop6oAipth3VKKeCe6YDULBUIiu9gTLqx3ea7qInZHYrtjt2mtd6oJqKshXR8cjdF8XanfLXDX+2PZTjW0qyvprnTNhr2NVaasjTSSN3Vx607ygVpAy3algF/2eZJJZ75TqmuroJ2p3Eze1F/gYkAAAHt4B9veP/GdN+9aWop0IVW4IumcWFey5U/71pajH/Rt8CAfoAAAAAIQ7/C/6yLUn/Af/ACUm8Qe39l12mWxOygT0lAjmAAAAAEj9wb3Sbr8Xr6SEcCR24Ov+su6fF6+kgE3wAAAAA+VX9iy+8XzH1PjW8KOZf/Dd5gKrsz+268fDZvTU8k9XMeOWXdf+Nl9NTygAAAAAAERVXROKgyfZVj8mUbQrLZI2q5KiqYj9OpqLqvzATo3SsSdiux+3LUQ83V3DWql1TRdHe1RfFobePjRU7KWjhpok0ZFG1jU7yJofYAAAAAAAACEW/div0NzuhyWCFGw3KDkSuTrkb2+IjeWF73OGrlmyStlp41fWWxfZcKJpqvJ9snk1K9AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAFkW7CipsLxbX+xNK3Ta+I7wO0XF8dorDaq6lZRUcaRwtdToqo1O+BYsCACb0O1X8YUPyVpz9VFtV/t9D8laBP4EAfqo9qv9vofkrT8rvQ7Vl/+40Sf9K0CwAFfv1T+1f8AGdH8lac/VP7V/wAZ0fyVoFgIK/F3ntq/40pPkrTZe73tS2x7Ss2ho1r6ZtppFSSvm9itROR96i9qgS4BwcgAAAAAA0NvibRbbjWBVGMNVs1zu0asbEi+0Z1uU2btVzi1bP8AD6u/3SRv1tqpDFykR0r+pqdpW/tCy665tlVXkF3kV01Q7VrNdWxt6mp3kAx4AAAAAAAAAADYWzHZ3U3+SO53Vj4LUnFvU6dexOxvap6Oy3ZotxjhveQMcyjVeVDTLwdMidbuxvnN2xsZHG2ONjWMaiI1rU0REToREKbqXVI14+zx97fRC2duMfw18vnRU1PR0kVJSwshgibyY42JojUPqAZK1pvPqt5U8zNp5nyHJwp4GbZZbMTt3P1sjZKp7VWnpWu7uRe1exvfXxHTBgvnvFKRzL6x47ZJ9NX2y7JrXi1sWtuUmr3IvMU7fbyu7E7E7VI65flF3yu5JPXyK5qLpDTs9oxF6kTt75+Lzc71mORJNPzlVWVD0ZDDGnBqdTWp1Ibn2cbNqPHHx3K6LHWXVG6tboixwL3vvnd/q6jV48ev0zF6ref1n+i3rTHq05nyxPZ9somrEZcMn52lp9UWOkbwkk98v3KfP4DclqoaK1USUVspIqSnb9xEmmq9qr0qvhPvqqqqr0rxXUGf3OpZdmeOeK/JXZtq+WflAACuRgHIA4AADq0PxWQU9bSOpK6niqqd3TFK3lN+foP2cn3S9sc81niXtbTWeY7NT55sjgmjkr8UdzcqaufQyO4L7xy+ZfKausV2vGJX72RT8unqoXciWGRFRHJ1tcnYSoMT2i4PQZfS883kUt2jbpFUaaJInU1/anYvSngNHodX9f8ADz/n+6y192LfBkdvBsut2WW72TSpzFTGic/TOdqrF7U7U75kJFf/ALdw7JFRUmoLjSSaKnb/AAc1fIqG/dnmbW/LKPkIrYLnG3WanXhr+cztTvdRH6n0r7P+Lh8e8fJ8bWp6fjp4ZWcnAKGFe/M8UU8EkE8bJYZGq17Hpq1yL1KaL2pbN5bO6S8WNjpravdSxdL6f1t7/Ub2C6dCoiovBUVNUVCw0OoX1LfOvvCRr7NsM/gh8Dbe1TZotOk18xyBVgTV9RSM4rH2uYnW3vdRqRUVF0Xgps8OamekXpPMLul63r6qgAOr7AAAAAGy92L3aLH793mLIit7dh92iye+f6JZCAAAAAADjVO1A5dGqveK9ttO0rO6DapkdDQ5Rc6elhrnsiiZMqNY3sQCwjlN++TynOqdqFYn86W0PXX6b7t8oU+jNq+0Zi6tzC7ft1As4BXdad4zavQLGn0wJUMYiJyZoGO18K6amyMM3vLzDOyPKrHT1MP3UtL3Dk8SgTJBg+zTanh20ClR9hukbqhE1fSyLyZWf3V/gZwAAAAAAYJtq2c2raNiM9rrIY0rGNV1JUK3uon9XHsK4sux+54vkNXY7tA6CrpZFY5qp09ip3lLWCMO/Bs2bcrCzPLZT61dCiMrUYntotfbL4PMBDAAAexg/wButi+Maf8AeNLUov6JnvUKrcIXTNLGv/Maf940tSg4ws96nmA/YAAAAAQd39PdNt3wFPSUnEQc38/dOt3wFPSUCOoAAAAASM3CF/1m3P4vX0kI5ki9wr3Trl8AX0kAnEAAAAAHwuH2DP8Ao3eY+517h9gVH6N3mAquy77art8Ml9NTyz1Mt+2m6/DJfTU8sAAAAAAEodwnD0rMhumYVMesdExKanVU4K93Fyp4E08pF5EVVRE6VLJN23EEw3ZLabfJHyKqeP2TUcNF5b+Oi+BNEA2SAAAAAAAAAAOtc6SKvt1RRTNR0c8To3IvWipoVf7TsflxfPLxY5Wq32NUvazVOluuqfMWkkLN/PEUoMrtuV08ekdfGsMyp0ctvR5UX5gIygAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAActarnI1qKrlXRETrA9PFLDcsmv9JZLTA6erqpEYxqdXfXvIWQbFtnlr2cYbT2ihjRapzUfWTrxdLJpx49nYa03QdkSYhYG5Xe4E+jNxjRY43tTWnj6vGvSpIQAAAAAAHVutfSWu3T3CumZBTQMV8kjl0RqIdleCaqQ43ytsLrhVyYFj1UqUsS/9oTRu/pHfeap1J1gau3kdqlXtJzORYJHMsdC5Y6GLo5XbIvfXzaGqwAAAAAAAAfuGKSeVkMMbpJHryWtamqqvYiAflqK5yNaiqqroiJ1m4tl2zJ0b4b1kcTehHwUbk1X3z/Uejsq2cttCxXm/QNfcE7qGnforYexV/O83hNneFSg6l1aMfOLDPf3n5K/a3PT8FPJ2d5NAcAy8zM95VM/iAGEbSdoFHi0T6GiVtReHN4N01bBr0K7v97ykjV1cmzf00j/Driw2y24q720TNKLEbeqaMqbnKn1imVeCfnP06u91mg1+j2a5Kq/XK2vqXeJqeZGocWe23zM8j5iFZausqHq+WaRVVGJ1ucvUif5ISGwjErXidvWGiRZamVE5+pcndPVOpOxveNPM6/S8PHm0/nP+FpM49Sn4urs/wi3YlRtciMqbo9v16qVPa69LWdidWvSplXWAZbZ2cmxf13lVZctstubABynFTjEOflwde4VtHb4OfrquGlj++lejUNfbSNp0FlmfbLFzdTXNTSSZe6jhXsTqcvzGlbvdbneKx1TcayaqmevFXuVfEidRe6fRLZI9WaeI+XusMOjNo5v2b+ue1HEKFXNZWS1jkTogiVU18K6HkP2zY+j9G2y4ub29wnzamsse2f5ffWNloLJU8wq6c9MiRRp/edoZF/Mnmv8Ayr5fH6y0/wBr0qdrR+cpcauCvmGfW7ariFW5GyVFTSKv4aLh5U1MrtF3td3i522V1PVN01Xm36qnhTpQ0DetlmcWpj5JLLJVQsbynSUj2zN0/uqqmJ01RXWyrSSCWekqI16WqrHNU55Oi62SOcc8fq+LaOK0fD2S5ODUWzjams0sdryaREc7RkVZp0r0aP8AWbdRUVEVFRUVNUVF6TO7mjk1bcX8e0q3Ngvhnu5OFAIbi8HOMUtuXWz2LW/WqqNF9jVaJ3Ua9i9rV608hHm62++4VkaRSrJSVtO7lRSxrwenU5q9aKSlPIy3HbZlFpW33ONNUReYqETu4XdqL2dqF703qk4uMWX7vz+Sfq7fp+G/hj2zLPabJ6RtHWuigu8fBWa6JOn3zU7e1DOCLmV49ecNvqQVXKjex3LpqmJV5MidTmr/AA6jbmzHaRDfXRWm8K2G5cnSObXRs6p5nec7dR6VFo+2wfl+zps6nPx42xjkHBnFY5RVRdUXRTVW1TZo2u52947EjKlNXVFI1OEn5zO/2obVCKqKiouikzT3cmrf1V8e8O+DPbFPMIfPa5j1Y9qtc1dFRU0VFOCQG0/ZzBkTX3SzRxU91RNXx8GsqP4I7v8AWaEq6aekqpKWqifDNE5WvY9NFaqdSm01trHs09dJXeLLXLX1VfIAEh0AABs3dg92ey+F/mLICt/dg92ezeF/olkAAAAAABw72q+ArI268dsGUfGEnnLNn+0XwFY+3FddruUfGMnnAwwAAAAB6GPXq52C7QXW0VclLVwORzJGLp/+IWIbu202m2k4RFVyK1l0pUSKti1+6T7pO8vSVvm3d1HOJMO2rUEMsqtoLo9KWduvDV3Bi+Xh4wLEAcIqKiKnQpyAAAA6d6t9NdrRV2ysjbJT1ULopGqmqKipodwAVabSsbmxHOrvj0zVRaOpcxmvWzpavkVDHSSu/pirLfmtsyiBnJZcYOZmVOt7OhfIvzEagPWwv7cbL8YQfvGlqdP9jx+8TzFVmGfbhZfjCD940tTpvsaL3ieYD6AAAAABB3fz9023fAf/AJKTiIPb+qf6y7b8B/8AkoEcwAAAAAkZuFe6Zcl/4D/5IRzJGbhXumXL4B/8kAnCAAAAAHXuX9X1H6N3mOwda5/1dU/oneYCq7LPtouvwyX01PMPSyr7Z7p8Ll9NTzQAAAAADO9geILm21SzWV7eVSpMk9Vx/wBkxdVTx8E8ZZhExscbY2Jo1qIiJ2IRL3BcQVG3bM6iNUR3+iUyq3pROLlRfm8RLYAAAB06q50NNcKagnqY46mq5XMRuXRX8np0O4Q1247U5qPebs0tPVq2gscrYJNOKd2v1zh0dCgTKB8aOeOqpIamJyOjlYj2qnWipqh9gAAAGqN6rE1yvZBc44Y0fVUKeyoe51XuelE8Wptc+NZBHVUktNK1HRysVjkXoVFTQCphU0XRekGWbXsZfh+0i92ByKjKapdzSr1xu7pvzKhiYAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAkTuf7Hlyu9szHIKbWy0L9aaJ7eFTKnQvvW/OvgNabDNnNdtIzantMLHsoIlSStnRODI9ejXtXoQsdxqy2/HbFSWW1U7YKOkiSKJjepEQD0WNaxqNaiI1E0RE6jkAAAAABhu2DPbXs7wyqv1xeivRORTwp7aWRehEQDAN6ra7HgeNOs9nqY1v1cxWsRF1WBi9L1TzEB6meWpqJKieR0ksjlc97l1Vyr0qerm2SXLLcnrsguszpKqrlV66rwanU1O8icDxgAAAAAAAd2y2uuvNyit9ugdPUSro1qdXfXsTvgfCipaitq46SkhfNPK5GsYxNVcpv7Zls9psbjZcrijai7OTh1sp0Xqb2u7/kO9s5wajxOl5+RWVF0kbpJOicGIvS1ne7V6zL0Mz1Pq3POLDP9Z/ZV7W5/JQABnZVocpqq6IgRFVUREVVNVbWtoyUCzWCwStdUqisqatq8I+1jO/2r5CZp6WTavxXx7y7YMFs1uI8O7tR2jw2SF9qsU8c1zdwkmavKbTp3upX+Y1PieOXnM72+OBznKruXU1UqqrWIvWq9a97rO1s7wqvy64quroLfEqLUVLk6PzW9rlJD2Gz22xW1lutVM2Cnauq9bnu63OXrU0ebPh6Zi+zp3t/75Wd8lNWnpr5dfFMetuM2ltvtsenBOdmcndzO7V/gnUesAZXNmvmvN7zzMqi95yTzYOTg4e5jGOfI9rGNTVz3LojU7VU51ibTxHl5EczxDl3BFVVRERNVVehE7TTu1TaVqktjxyXueLaisavT+azvd/yHn7V9oyXdklksUj20HK0mn6Fn0XoTsb5zWBq+m9KjDxky/e+XyW+rqej4reXesVrrb5eKe2UMfOVNQ/ktRV8qqvYnSSGwnZ9YcYp2SPhjuVzTi+qmbq1i9jGr0eFdVXvEdLVX1druENfQzOhqIXcpj29RIXZrnVLlFHzFS+OC6Rp3cWuiSJ9831Ejqs7MYucPj3+bptzkinwM2klllXWSRzvCp+TgGNtaZnvKkmZny/Ub3xuRzHuaqdaLoeLl2K2HK6V0V1pWxVOi83WwsRsrHdq/fJ3lPYMX2gZnb8So0WREqK6ZPrNO1eOn3zuxCZoW2PtYjBPd315y+qIo0Fm+M1uJ319rrJYplRqPjliXuXtXoXtTwKZnsq2iy26WOz32dX0K9zDO7phXsVetvmNf5Bd66+XWa5XCVZJ5V17zU6kROpEOgbTJhrmx+jJHK7tSL19NkvoZY5oWTQyMkjenKa9i6oqdqKh+yPuzDaFPjr2Wy5cqe1udw63QKvWne7xv6kqIKumjqaaZs0MrUcx7V1RyGO3+nX1bc+a/P91Lsa1sU8x4fQAFcjPPyKyW3IbU+2XWDnYXcWPT28Tvvmr1KR1zzELnh90ayZVkpXu1pqpnBHonmcnYSbOrdbfQ3a3TW65UzKilmTRzHdS9SovUqdpcdO6pbXn0ZO9fom621OOeLeGsNlm0mKohisuRVHJqGqjaeqf0PT7169S982wRy2k4FXYpVeyYVdVWuVy81OicWfmv7F85lGyfaQsSwWHIZ05lERlPVvXizsa9ezv9RYdQ6bXYr9tg8/VJ2NWMsevG3McBFRURUVHIqaoqLqigzExMTxKpnmPIYZtKwOjyqmdV06Mp7uxvcS9DZU+9f/BTNDjThod9bavrX9dHTFltin1VRIutvrbVXy0Nwp309REujmPTRUOqSZ2g4XQZbRIj1bT3CJNIKlE6vvXdqeYjrkFnr7FdJrbcoHQzxLouvQ5OpUXrRTaae7Tap6q+feF5hz1zV5h0AATHZs3dg92ey+F/mLICt7dh92iye+d6JZCAAAAAAcO9qvgKxtuSabXsoT/mMnnLOXe1XwFZG3X3YMo+MZPOBhQAAAAAfWinkpayCqicrZIZGyNVOpUXVD5AC0/Z7dvo7hFmu2mi1VHHIvHXirU1PeNUbple6v2FWB73ve6Fj4VV35rlQ2uAAAAAAR936bO2u2TwXFGayUNY1yLp0I7gpBQsi3m7dHctimQxPRV5unWVNO1vErdA9bDOOYWX4wg/eNLU6b7Gi94nmKrcITXNLGn/ADGn/eNLUqf7Hj96nmA+gAAAAAQs357TdbhtHtslBbK2qY2h0V0MDnoi8pexCaZ83wxPfy3xMc7tVqKoFVyYtk69GOXdf+ik9Rz9KmUfk3ePkUnqLUUijTojZ5DnkM+8b5AKrfpUyn8m7x8ik9Q+lPKfybvHyGT1FqXIZ943yDkM+9b5AKrFxXJ06ccvHyKT1EhNxey3e37RrnNcLXXUka0GiOnp3MRV5ScEVUJnLHGvSxvkOWsY1dWsa1e8gH6AAAAADrXT+ran9E7zHZOtdP6tqf0TvMBVblX2zXT4XL6anmnpZV9s10+Fy+mp5oAAAD9wRvmmZDGmr3uRrU7VVdEPwbO3YsVdle1+00zo+XT0jvZU/Dho3o+fQCdGw3Fo8P2X2SytaiSMp2yTKidMju6d86mbnDWo1qNRNERNEOQAAA8bN71BjuJXO9VDmtZSUz5e6XRFVE4IVd5HdKi9X6tu1S5Vmq53Su1XXpXUmjv0Zelq2f0+MwPVKi6yor9F6I2rqvl6CDwFj27Blbcs2P2epc9XVNJH7FqNV48pnDXxpops8h3uBZOkN3vWKTzojZ2JVQMVelycHaeLQmIAAAAAAQ53+sSbTXm0ZhTxaNqmrS1LkThym8Wqvi18hFgsi3l8STMNkF5oI4kfVU8Xsqm4aqj2d1w8KIqeMrdcitVUVFRU4KigAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA7+PWiuv16pbRbYHz1dVIkcbGprqqnQJs7nWx+OwWeHOb/Sf9q1jNaOORvGCNfutF6HKnzAbT2E7N7ds4wqntsMbHXCVqSVs+ndPeqdHgQ2CAAAAAA4cqNRVVdETiqgdW8XGjtFsqLlcJ2QUtPGskj3LojWomqldW8PtQrdpeYvqGudHaaRVjoodeGn3699TZ++JtkS910mD43VqtBTu0r5mLwlen3CL2J1kYwAAAAAAAAPRx6y3K/XJlBbKZ00rl46e1Yna5epCRmA4hb8TtiRQtbLXSNT2RUqnFy9idjTzdjUNjhw+F9oVj55E1q3L/SLJ2L2InUhmyGX6v1G82nBTtHv+P8AhVbmzPM447AOTgz6uDlOOvFERE1VVXRETtU/E0sUET5p5WQxMTV73uRrWp2qq9BozadtLnuyz2exPdDbV7iSZOD6j1N73X1lhodPvtW+VY8yk6+tbNP4PS2q7SucSSyY1UuSPVUqati6K781i9navWY5sz2fVeUz+z7gstJaGL3U2ndzL96zX516EO1st2dVF+dFeLsxYbS1+rWrqj6jTqT83tXyG+oo2QwRU8EbIoYmoyONiaNY1OpC8293F0/H9jh8/T+qfmzV16+inl8LVb6K1UEVBbqdtPTQt0YxvnVete1TtHByZbJktktNrTzMqi1ptPMuAcnTu9yorTbpa+4Ttgp4k1c53X3kTrXvDHS2S0VrHeStZtPEPvVVEFJSy1VVKyGCJvKkkeuiNQ0NtO2jVN+fLa7U5ae1IujlTg+dU617E7x520XPa7KpfYsTXUtsjfqyFF4v7HP7V8xhhrundMrrR6797fRc62rGL4p8gALdMD60dTPR1MdTTSuimjdymPauiop8gBv/AGXbQosgYy2XRzIro1NGu6Gz+Dsd3jYXQRHtcFbVXGnp7dHLJVySIkLYvbK7XhoSjs1rv8WGNt9dc4UyJYHI2Z3dRsk+5a5U6V7XdGpm+p9NxxaL0mK8zx/lWbWrXmLRPHLwNo+eUWMU0lLTKye6ub3EfSkev3TvUR8utwrLpXS11fO+eeVdXPcp2Mnobvbr9V0l9jmjuDJF5/nV1VXduvWi9Op5pc6mnj1aemn5p2HDXFXioACW6hmezfOqzF61IKh0lRa5F0kh14s/Ob2L3uswwHxkx1yVmto5iXlqxaOJS3tFxortQR19vnbPTyJq1zfN3lO2ad3d7bkD6yeuWZafHmIvP84mqTSacGxp992r0InT1G4fB0GL6lp11cvFZ5if0Ue1hjFfiJAAVyM/FTBBVU0tJVwsnppm8mSJ6atcnfNE7VdnEthkku9jZJPaHcXs6X0y9i9rexfKb5PlVVVPSUss1W+JlOjF53ndOQrdOKLr0p3iz6bv5Ne0V8xPt+yVq7FsdvT5ho7ZZtHfZWss97V81vVUSKXXV0Hrb3uo3nBLFPBHNDI2SKRqOY9q6o5F6FQinlEltlyCtks8SxULpnLCxepuvV3jbe7pFktVBW88rm47TxuXnZ9dGy6cGR9qr1onDrLvqnTqZqzlr2tH6p21rRePXHaW1TgAyKnDwszxa2ZVa1pK5qRztRfY9S1urol/i3tQ90L0HXDnvgvF6TxL7x5LY7eqqKmVY7c8bubqG5wKx3TG9PayN7Wr1nkkmdp9PY6rEalL85jI4kV1PJ92yTq5Pbr2EZ3acpeSuqa8Db6O1/qsUZOOF7gy/a09XDZe7D7tFk9870SyEre3YvdosnvneiWQkx2AAAAAHDvar4Csjbt7sGU/GMnnLN3e1XwFZO3f3Ysp+MZAMJAAAAAAABP7cmqFn2GUbFbpzVXO1O/3Ru80VuPJ/qPg+HT+kb1AAAAAAMO2100tXsqyOCFvKkdQSclO3uSsNektK2mTxU2z++TTvRkbKKVXOVeCdypVs72y+ED2MF+3axfGVP8AvGlqEH9BH71PMVX4L9u9i+Mqf940tQg/oI/ep5gPoAAAAAAAAAAAAAAAAAAAAAHWuv8AVtT+id5jsnWuv9WVP6J3mAqtyr7Zrp8Ll9NTzT0sq+2a6fC5fTU80AAABNHcMxJlHilxyueHSetl5mFypx5tvTp4yGtupZa6vp6OBqulnkbGxE61VdCz/ZbjcGJYFabFTsRqU9O1H6db1TVV8oGTgAAcKqIiqvQhyYnteyaLEdnV5vsj0a6CmdzWvW9U0anlAg5va5g3LNrtcynl5dHbE9iRaLw5Tfbr5eHiNQn1rKiWrrJqqd6vlmkdI9yrxVyrqqnyAzPYjlE2H7UbHeo5ORG2qbFUa9CxPXku18S6+Is3hkbLEyVi6te1HIveUqWaqtcjkXRUXVCyPdryh2WbIbNcJXq6ohj9jzKq6qrmcNQNkgAAAAPxPG2aF8T0RWvarVTvKVm7dsafim1S+WlYljiSoWWFOpWO4oWakRt/nD2NdaczpokR7v8ARKlyJ0p0tVfnAiWAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAABmex3ALptEzKlslBG5sHKR9VPp3MUadK+HsA2Ruk7IHZtkTckvdOv0Bt8iOa1ycKmVOKN8CdZPBjWsY1jGo1rU0RE6EQ8jC8ctuJ41RWG0wNhpaWNGNRE6V61XvqeyAAAAAADQm9ltfjwrHX45Zptb5cI1bymO407F6XL3+w2Ltm2gW3Z1hVVe617XVHJVlJBrxlkVOCeDtK4MzyS6ZZklZfrxOs1XVPVzl6mp1NTvIB5M0kk0r5ZXue96q5znLqqqvWfkAAAAAAAAAD1cYv8Acsducddbpla5q93GvtJE60chIzCsrtuUW9J6N/Inaic7A5e6Yv8AFO+RePQx+8V9juUdwt8yxzM8jk7FTrQgb2hj269+1vaUfPr1zR38pY6nyramnoqOasq52QU8LeVJI9dEahheNbS8errGtbcqtlFPC369CqauVdPuU69TU+0PN7hltdzMXOwW1jvrFKi9K9TnadLvN1FDq9GyXyTGXtEfqr8OlebcX8Q7e07P6jJploKFHU9qjfq1q+2mX753e7EPd2V7M1rGxXzJIVbSKnKp6ReDpuxzuxve6z0Nl+zGOGKnvWSw8qZyo+ChcnBqdTpE7evk+XsNtLqvT0k7e6jTVp9hr+fp/lI2NmMUejH5fljWMY2ONjWMa1Gta1NEaidCIh+gcGXmZtPM+VVM8yA5PHyzIrdjNpfX3GTRVTSKFF7uV3YifxOmHDfNaK0jmZe0pa9vTV9sjvVvsFrkuNznSKFvBqdLnr1NanWpHbP8xuGV3JZJXOioo1/0enReDU7V7XL2nVzLKLplVy9k10n1tiqkEDfaxoq9CJ2982Lso2aIrYb/AJND3C6PpaJycX9jpE6k7E6zW6uph6fjm+Se/wA/+IXGLDTWp6rPM2X7M5bqsF6yKKSG18HxQa8l9T/Fre/19XaeptY2UxUtJLkeHxvfQsTlVdBqrn035zVXi5nzobgkcr3au8nUh9KSompZ2ywv5Lk4eFOtF7UKz/fb/berj4PkjRvz9pz7IcA37tX2WxXxk+RYjTxxVrUWSttzO5R/DVXxJ52+Q0G9jo3uY9qtc1dFRU0VFNLizUzUi9J5hZ0vW8c1cHdslqr71cobdbKaSpqZV0axia+PvIfTGrNWZBfKW0UDWrUVD+S1XLo1qdKqq9iIiqSTwbEbbhlvdT0L0qK6ZNKms00V/wCa3sb5zht7mPVp6refaHPNmrirzLqbO8Ht+HUbZO4qbxKzSep6Wx69LI+xO1evwGUKnE5Q4MXtbWTZvNrypMuW2W3qs8nMsXtWZWlaG5o2GrY3/RK5G93E7sd98xezq6iNeW47dMYvElsukCxyN4sentZG9Tmr1opKo6GTWG0ZTaUtl6iVzWa+x6lv9JTqvWnanahbdM6r9nxizT2+aZq7fHw38Img9zOMZrsTyGa0VzmSK3R8U0a6tljXocnqPDRFVdE4qaiJiY5hahs7ZJszkvvNX6/sfT2RjtY417l9YqdTfze13kPV2T7K2TQRZDl9O9lI7R1JQrwdUfnO60Z86m4ZZOXyEaxkUcbUZHGxNGxtToaidiFT1HqddaPRTvb6ImztRijivkV0bYo6enp4qamibyYYIm8lkbexE/j1n4ByZG97ZLTa095UtrTeeZcAKebkV7t9htj6+4zJHE1OCfdPXsROtRjx2yWitY5krWbT6Y8uxdrjR2qgkrq+dsFPGmrnu83hI/7S88qsmqXUdIroLXGvcs14yr9871Hn59mVfldaiy6w0cS/WYEXgnfXtUz3Y7srSqgiyjL6dY7ZpyqSjfwfVr1KqdKM7/Wa3S0MejT7TLPxfP5LnBr1wV9VvLy9kOy9+RNjv2QSPo7Ex/ctRPrlWqfcs73a43xNOxYIqSlhZS0UDeTT08SaNjb2J3+/0qpxUzrMkbGxxwwRMSOGGJOSyNqdDWp1IfEpeo9Ttsz6a9q/VB2dqcs+mvgABUoYvBDycov9ux62urbjMjG6dwxF7p69iIdXN8rt+L2x1RUvR87k0hgRe6evqI65XkNxyS6OrrhKrl6I40XuY07ELnpnS5z/AMTJ2r9U3V1JyfFfw7WcZZcMpuSz1LlZTMcvMQIvBifxXvmPAGtrWKR6axxC4iIiOIbL3Yvdosfv3eYsiK3d2P3aLH793mLIj6egAAAADh3tV8BWTt592PKfjGQs2d7VfAVlbevdkyn4xkAwgAAAAAACJqqInWBYTubwLBsHtOsfI5yWV/R06u6TchguwO1JZdkGN0KNcitomPcjunVyar5zOgAAAAADBtvkrYtj+SucuiewZE+YrLLB98i7LbNidwja9WPq5GQJoumuq8SvgD2MG+3axfGVP+8aWowf0EfvU8xVbhPDM7Iv/Maf940tSp/seP3ieYD6AAAAAABEze62nZphe0GkoMcvD6OnkpEe5iNRUVde+BLMFc6bwu1dP95Xfsmn7TeI2sJ/vJ/7DfUBYqCutN4rayn+8n/sN9QXeK2s/lJ/7DfUBYoCul28PtYd05K7xQt9RuXdB2mZnmefV1Dkd4krII6NZGNVqJo7lInUBLEAAAAAOrdv6rqv0TvMdo6l5/qmr/Qu8wFV2UfbLc/hcvpKecehk3HI7kv/ABUnpKeeAAAG5tz3EUyba9R1VREr6S1NWqk7OUntU8vHxFgxHfcZw91m2d1GSVUXIqLvLrHqnHmm8E8q6r5CRIAAACLe/wA5Z7Gx60YjA9vLrJVqKhOtGM00TxqvzKSkPFveK45e6ltTd7JQV0zW8lr54GvVE7OKAVWAtE/m7wb8lLP8kZ6j9Js/whOjFbP8kZ6gKuSW+4FlDlbe8Tnl7lvJq4GL3+Dv4EkFwDCfyVs/yNnqO5ZsUxuzVS1dqsdBRTq3krJBA1jtOzVEA9oAAAAAMA3gMXbluyq82zmkkmbAssPee3ihn5+ZGNkjcxyatcmioBUrIx0cjmPRUc1VRUXqU4Nj7yGJLh+1q70EcaspZ5VqKfVOHJfx4ePU1wAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB27Pbqy73OnttvgfUVVRIkcUbE1VyqWLbvOzGk2a4VDRvRsl1qUSWtm048pU9qneToNT7mmx5bZRx59kVMxauob/2fBIzuoW/hF1616u8SkAAAAAAB0Mgu1BYrPU3a51DKekpo1kke9dEREO69zWMV71RrUTVVXqIQb322CTJbxLhtjqNLTRyaVMkbuE8idWvYgGu94HafX7S8zlrFe+O00yrHQ0+vBG/fL+cvqNbgAAAAAAAAAAAAAAA2ZsDorHUXueor3MkuMKItJDJ0J2vTtVOrs6TWZ9qGqqKKrjqqWV0U0buUx7V0VFOeWk5KTWJ45fN6+qsxEpeaqqq5VVVXpUGvNl20KLII0tt1fHDc2+1VE0bMne7/AHjYSdBhNrVya2T03/8AqgzYrY7cWADEtoOd27FaV8CK2ouj2axU6dDexz16k73Sp86+tk2L+ikPMWK2S3FXbznMLbidEktV9eqpE1hpmro5/fXsTvkdsqyC45JdpLjcZeU93BjE9rG3qa1Ow+VwrLrkV5dUVDpq2uqX6IiJqqr1IiJ1d5Dalm2HVNfissj7lzORqznYaJze4cn4NXff/N1Gx1dXDpViJnvPv811hw0wRx7uhsCx6z1tVUXq4LHU1NK9G09M7RUYv4RydenV1am7ZHuker3uVzlXiqkVLbW3nEr+skaS0dbTvVksT0VNdF4tcnYb/wABzW25VS8mNUgro2ostO/p8Le1NSt63rZr/wASJ5rHt8kXexXt8UeGVg48YMyq+X7hlkhmbLE9WPauqKi9Bg21fZxSZZTy3vHaZtNfmJyqilZwZWJ1uanU/vdCmbnLHPY9HscrXNXVFTgqKTdLeyatuY8e8O+vnnDPbwiXQVdyx+8tqIecpK2leqKjm6OaqcFRUXyaG7sJ2oWq6xR013c2hrdNFc7hG9e1F6jJ89wWw53rU1UqWq98nRK1rNWTL1c6icdfzk4+E0PmWzvK8Wmd9ELZJJTIujKuBOcikTtRyfxNTMavUsfz+sLbjFs1SShmimjbLDIyRjk1RzV1RUP3qRPtt7vNrd/oNxqqbT7lkionkPabtEzBrUal5l4JpqrW6+YrL9Ann4L9vxhEt07v8NklVVETlKqIidamHZjtDsdhhfHHOyurNF5MMTtURfzlToNEXPJ8gubVZW3armavS1ZFRPIh9caxDJMkqWw2i0VVSrul/I0Y3vq5eCEjX6Hixz6sk8/R1x6FKzzaeXwye+XDJby+4Vyo6V/csYxODG9TUNwbJdmENsjgyPKo0fVKnOUltc3Xk9j5eztRvlPb2c7NbZhskV0ucsNzvrU1axreVBSO7UX7tydvQnVqZnLI+WR0kjlc9yqrlXpVTzqHVa4Y+ywefobG3GOPTTy5qJZJ5lllernL2n4BwZaZm08yqZmZnmQeA5MdzbLLdi9udNVSNfUub9Zp0Xunr/BO+dMOC+a8UpHMvaY7ZLemr75bkluxq2rWXCTiq6RxN9s9exCO+aZPccpuzqqqcrYkXkwQNXuWJ6++fDIb1dcpvHsiqc6aaR3JiiYnBuvQ1qG79lezGmxZkF/yVsNTeFRH0tAqcptL1o+TXgruxvV1mt1tXD07F67z395/4hc4sVNanMvN2SbK6Wip6fJs0pnue7SShtbuHLTTg+Xr5PQqN4a+A2lV1EtTMssrtV6ERE0RqdSInUhxUzzVU7p6iV8kjl1Vzl1VT5Gd3+oX27ceK/JW7GzbNP4AByV6MGKbQM0ocWo9HaTVsjV5qFF+dexDpbR89ocapn0tK9lRc3IqNjRdUj77vUR/u1wrLpXy11dM6aeVdXOcpoOmdJ9fGXNHb2hY6up6vivHZ9b/AHivvdxkrrhM6WV666dTU7ETqQ6ABp4iIjiFr4AAejZe7F7tFj9+7zFkRW7ux+7RY/fu8xZEAAAAAAcL0KVl7ffdlyr4xkLNHe1XwFZe3z3Zcq+MZAMHAAAAADJdl+OVOWZ/ZrDTRuf7JqmJLyU9rGior18mpjRL7cb2bT0jJ8+utOsbpWLDQtcnHk/dO8YEqaGnZSUcNNGmjImIxqd5E0PsAAAAAA/Mj2xxue9URrU1VV6kAih/KAZE1tHYcYjevKe91TKiL1JwTVPCpEM2RvJ5YmYbXbvcIZVkpYH+xaddeHJZw4ePU1uB6+FfblZPjCn/AHjS1Km+x4/eJ5iq7CE1zSxp/wAxp/3jS1Gn+x4/eJ5gPoAAAAAEGd/Bf9aVEn/Ap51JzEFt+1ddq1InZQt86gR7AAAAACRW4Wum064p20C+khHUkRuHLptSru/QL6SATlAAAAADp3r+qKz9C/zKdw6d6/qir/Qv8ygVW5J9sNx+FSekp0Dv5J9sNx+FSekp0AB6OMWmovuRW+zUrdZqyoZCxPfLoecb73IsTkvm1N17lhV1JaIFfyurnXcGp5OUBN7E7RT2HGrdZqVjWQ0dOyFqJ0cERD1AAAAAAAAAAAAAAAAAAAAAilv84m+W32nL6eLVIHLTVConQi8WqvjIflnm2jGWZdszvdjc3V81M50Xee3um/OiFY9TDJT1EkErVbJG5WuavSiouigfMAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAN3bqeyN2f5Ql4u8LvoDbno6RFTRKh/UzXs6FU15sqwi55/mVHj9ua5Ekcizy6apFH1uUskwHFLVheLUeP2eFI6amYjddOL3dbl7VVQPagijghZDCxrI2NRrWomiIiH0AAAAAAcLxRUAj1vf7XExKxLilkqUS818a865q8YIl4a95V6iDL3Oe5XOVXOVdVVelSwfKN3LA8lv9Xe7vLdJ62rkV8j1qPmTsRDz27rGy5E4w3Jf+pUCA4J8rus7LvwFx+Uqfh26tswXojuSf8AUAQJBPZN1bZgnTHcl/6g5+pX2X/grl8pAgQCeq7q2zFehlyT/qP8h9Srsx+8uXygCBQJ7fUrbMPwdy+UGObTN23Z5YMBvV5t8dw9lUdI+WLlz6pqiapwAhYAAAAAAAD9RSSRSNlie5j2rq1zV0VFN3bMtpUVxSG0Xx6R1mnJjqF4Nk7EXsU0eEVUXVOCkfZ1sexT0XhzyYq5K8WSH2mZ/SY3Suo7dLHUXaRODdeU2FPvnd/sQ0PGy6ZDekY1Jq2vq5NE63Pcp0kR0siJrq5y6aqvWpJjZziVrxazRPhjjqLlOxHT1at1VNU4tYvU3q7VIk/YdMwdo/eZcZ+z1cbr7M8GpMPpEqpnx1N6lb9clRO5p0+8Yvb2u8hmCOc1yPa5Uci6ouvFAvE4Mns7eTZyeu/9vwVGXNfLbmXlbRMMt20a2t5yWChyOBNIKxyaNqU+8kVOvsd5SNldS33DcmfS1Mc1vulFJoqLwVF/ii/OhKhOCop1svxux57am26+oymuETeTR3RrO7Zp0Mk++Z5i86b1aJiMWef7/un625z8GRiezfO6PJqVlNUvjgujU0fFrokn5zfV1GakY8xxbIsByJtJco3U9RG7nKeoidqyRuvB7HJ0obP2b7Tae58i231zKes0RGT9DZPD2KfPUek885cEf2/Z5s6fPx4/ybNOThFReKcU6gZ2YmFZMcOT70tbVUyK2GZzWu9s1dFavhReB1we0vak81niXtbTWeYl+K+24vdV1vOI2asdx7tsKwvTxsVE8qKeRLgWzKWoSoXEpWPTT63HXvSNdO9pr857QJ1OqbVY49bvG3mj3daksOFUD0ktuFWiGVOiSVHzKnicunzHpS1tS+FIUkSOFOiONqMYngRE0OsDlm3s+aOLXl8X2Ml/MmhycAiOXIAYBtL2hUuPwvoLa9lRcnIqcOLYe+vf7xJ1tXJs39FI/aHTFhtltxV6m0HNqDFaPkryaivkT61Ai/O7sQj/AFE16yzIWpyZ6+4VcnJjjYiuVVXoa1Ow4oqW9ZXf2U9OyouFxq36InFzlVfMhJLZ/hFr2fUfKjc2tyKaPk1FX0tptU4xxd/tcaulMHTMPM+f1lcVrj1ad3R2YbO6TBKdtxvEMFXksiasavdMoE06upZO/wBXhMqke+SRZJHK5zl1VV46nGqqqqvFVODMbm7k2r828e0KrPntltzPgACqicVIUR7OAa62obQ4bFHJa7U9stxc3Rz04pBr/HvHR2qbRo6KOSz2KVr6pdWzTt6I+8nfNJyyPlkdJI9z3uXVznLqqqabpnSeOMuaP6R+601dTj47v1UzzVNQ+oqJXyyyLynvcuqqp8wDRLIAAAAAbK3ZPdosXv3eYsjK3N2T3aLF+kd5iyMAAAAAA4XoUrL2++7LlXxjIWaLxQ0vk+7bs/yLIa6+XBbgtVWzLNLyJtE5S9iAV9gnl9Snsz7bn8o/yC7qWzPtunyj/ICBp3LTa7ldqlKa2UFTWzLp3EESvX5iftp3atlNC1iSWWWrc37qady6+HRTZGMYjjWM0qU1istFQRp1QxI1V8YET9hO7HdKyupb7nTEpaNipI2g+7k7z+xO8TFoKSmoaOKko4WQQRNRrI2JojUTqPuAAAAAAAam3o9oMGDbNKxsUyNulxYtPSNReKKqaK7wIhsrILvQWK0VN1uVQyClp41e97l6kQrk2+bR6raRnNRdF5TLfAqxUUSr7ViL0r316QNevc571e5Vc5y6qq9KqcAAezgnHOLCn/Mqf960tRh/oWe9TzFV+BfbzYPjOm/etLUIf6JnvUA/YAAAAAQV37U/1r0vwFvnUnUQX37vdVpPgLfOoEegAAAAAkPuH+6nW/AXekhHgkRuGp/rRr/gC+kgE5QAAAAA6t2TW2VSf+E7zHaOvcf6vqP0bvMBVXlCaZLc07KuX0lPOPTy37abr8Ml9NTzABP7czxRcd2SwV08SMqbq9ahy8nReT0N18RCDZ5j82VZrabBCmq1lSyNy9jde6XyaloNit0Fps9JbKViMhpoWxMaidCImgHdAAAA8TOr5BjmI3O9VD0Yylp3P1Xt04AYBmO8Fs9xbIaqxXKsqFq6V3IlSOJXI1ezU8b6qPZhr9k1v7BSCmR3SovV+rrtVPV81XO+VyquvSp0AJ+fVRbMP7VW/sFOU3odl/8AbK39gpAIAT+Teg2XL/8AraxP+nU/Sbz2y3+31af+QpX+ALH8C25YHmuQx2Ky1s76yRquY18StRUTpNnFWOzvIqnFM1tV+pnq11JUNc7Relmujk8mpaDYrhBdrNSXKmej4qmFsjXJ1oqagd0AAAABw5Ec1WqmqLwUri3ocWXFNsd3pWQ81TVbkq6fRNEVr+nT+8iljxGDf3xNK3GLVltPFrNQyrTzuT8G/o8ioBDEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAABLz+Tzpqd1Dl1W6FizslpWNkVO6RqtkVU18SEsSKn8nj/U2Y/CKT0ZSVYAAAAAAAAAAAAAAAAAAADDdtvuS5P8XS+iZkYdts9yXJ/i6X0QKwwAAAAAAAAAANk7L9os9olZa7zM6W3uXRkrl1dCv+E1sDlmw0zUml45iXzelbxxZL6CaOeGOeF7XxyNRzXNXVFRetD6Eetmm0Grx6aK33Bzp7UrtFTpdFr1t73eN/UNXTV1LHVUkzJoZE5THsXVFQxu/0++rbnzX5/updjWnFPMeH2OTgFciuLxbrNk9jdYMkgdLSrxgqGJ9dpX/fNXs7U6yOe0zZ3e8HrWvnb7Ktky601fCirHInZr1O7xI07VPVt9hy26upoq+3T8JqWdvKY7wa9C98vOndWnDxjy94+ifrbk0+G/hHrAtqFfZ0jobwjq2iRURHqv1yNP4obpx++Wu+0aVVsq452KnFqL3TfCnShhmc7Dqe5Ndctn1QiuXVZLTVSokrV/8ADd0OTvLoppqrpchxK7rFUw11pro10VHIsbv80LbY6fr7tfXSeJ+cf8peTWx549VfKVZwaCsW1zIKJGsr4oLhGnW7uH+VPUZ1Z9rmNVaI2tZVUEmmqq9nKZr2IrdV8qIUebo2zj+7HMfgg30stfHdsM5MfpMzxaqRqw3yj1VNUR0mi/OdpMjsKpr9GKL9s0hTp5480n8nD7DJ/wCMvWB4VVl2M039Ne6JvDo51F8xj122sYrRx/6NJUV0munJij0Tyu0PunT9nJPak/R9V1stp7VZ6p1Lrc6C1Urqq41cVNEn3UjtNfB2mlr3tivNTqy20UFEz75y847+CGC11fesir2+yJqqvqXroxiauXwIiFrr9CtM85p4/ol49CZ73lsfPNq7545KDHEWNjtWuqnJ3Sp+anV4TBcNxa+5rfUobXBJPK9eVPO/2kTetz3dSdJsLA9iFxq2xXPM6n6CW9e6Sm01qpk7Eb9z4VNw0EVvs1qSzY9QsttuRdXMYur5l++e7pcvzFll2dbp2P00jv8AL90m+XFrV4h5mEYxa8CtT6G0ztqrhP8AZdw5OiuT7xnY3zno66+EAyu1t5Nm/ruqMuW2W3Ng5OD8TyRwxOlle2ONqauc5dERCPWs2niPLnEczw/T3NYxz3qiNamqqq8ENO7UtpPK52zY/Nw4tnqW9fajV/ieftT2jSXR8tnskqsoU7mWdq8Zu8n5vnNYmq6b0qMXGTL976LfV1PR8V/Llyq5yucqqq9KqcAF6ngAAAAAAANl7snu0WP37vMWRFbm7J7tFj9+7zFkYAAAAAAAAAAAAAAAAAA4A5Ole7pQWa2T3K51UdNSwNV8kkjtEREMN2obW8NwChfJdbjHLV6LzdJCqOkcvg6vGQg21bZ8l2k10kc0j6G0Iv1qijfwVO1y9agZHvNbbqnP7i6yWWR8Ngp38NF0WocnWve7xo0AAAAPawL7ebB8Z0/71pahD/RM96hVfgn28WH4yp/3rS1CH+iZ71AP2AAAAAEF9+73VaP4C3zqToIL793uqUfwFPOoEegAAAAAkRuHL/rRr/gC+khHckPuH+6lW/AV9JAJzAAAAAB8Lh9gz/o3eY+58K/7Cn/Ru8wFVuXfbVdvhsvpqeWepl/213b4bN6anloiqqIiaqvQBJXcQxJtxzGvyeohV0Vvj5uFy9HOO6fmJsGp91bD0xHZHbmTMVtZXp7Kn1TRUV3QniTQ2wAAAAjrvz5b9CNntNj0Emk90m0ciLx5tvFf4Eiiv/fMyd1/2vT0LH8qntUaU7URdU5S8XfwA0kAAAAAAAAWD7neTLkWxuihlejqi2yOpX8eKInFvzKhXwSa3B8mSizC74zNKqR18CTwtV3Dls4LonaqKnkAmmAAAAAGLbWMcjyzZ7ebFIiKtTTORi6a6ORNUXymUnCoioqL0KBUzW00tHWT0k7VZLBI6N7V6nIuiofE23vY4k3Ftr9xWCNGUtwX2VEiJwRXe2TympAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAmF/J4/1NmPwik9GUlWV8buG2mm2T0V6p6izTXFblJC9FjkRvI5COTr98bcbvjWvrw+t/btAlWCK31Y1o/I+u+UMP0m+NZuvD6/9uwCU4IspvjWTrw+4fKGHP1Y1i/I+4/t2ASlBFv6saw/kfcv27B9WNYPyPuX7dgEpARc+rGsH5H3P9uwfVjY/wDkfc/27AJRgi59WNj/AOR9z/bsH1Y2P/khc/27AJRgi59WNj/5IXP9uwfVjY/+R9z/AG7AJRgi2u+NYerD7l+3YcfVjWL8j7j+3YBKUw7bXx2TZP8AF0voqaK+rGsf5HXH5Qw8bOt6yzZHh91sUWK18D66lfA2R0zFRquTTVQIogAAAAAAAAAAAABmGzvOq7FqlYX8qot8i93Cq+1/Ob2KYeD4vSuSs1tHMS8tWLRxKWdiu1BerdHX2+obNC9OlF4tXsXsU7xFzDcpueL3BKiify4nL9dgcvcvT+C98kHhuW2rJqBk1JK1lQjfrtO5e6YvX4U75k+odKtg+PH3r9FPs6k4/ir4ZABqCmQnKK5rkc1VRU6FQ71VW010ofofkVror3Raac3VR6uan5r07pPKdAHbDsZMM80nh90yWxzzWWJXvYxs9uz3SWe6XOwSuVV5qdEqIU16kXg5ETv6mGXTd7zJiOfZK20XtidCU9U1r9O+12mhuA51VF1RVRS3xddy17Xrz+ibTqFo+9HKPFz2P7S7fKkc2H3SXVNeVTxc83ys1Q8tdn2dJwXEb38ik9RKSC4V8DeTDWVEbdddGyKiHZ+j96/GlV+0UmR1/H70l1jqNfki7Q7K9otbLHHBhl6+uLo1z6VzG+NyoiJ4zJ7fu/57I7W6sttmi+/q6xia+BGqqm9JrpcpuUklfUvR3Siyrop1Fc5y6ucq+FTnfr8fyU/V5bqMe1WA2bYVi1E5JMgyqouCp0wW2HkIv99/qNhWSCw41TpT4pj9HatE0Wp5POVL++sjuKeLQ+RyV2fq+zl7RPEfgjZN3Jfx2fqeWWeVZZpHSPcuqucuqqfgHJWTM27ok8zPMuDkHl5HfrZYKJ1XcqlsTdO5b90/vIh94sV8torSOZfVaTaeIh3K+spqCkkq6uZsUEbVc57l4IhoLabn9TkNQ6htz3wW1i6cF0Wbvr3u8dDaBnFwyiqWNFdT0DV+twovT33dqmImu6d0yutHrv3t9FxrasYo5nyAAtkwAAAAAAAAAAGy92NNdtFj9+7zFkRVzsuyhMOzagyBaf2QlMqqrNdNdU0JHN3xV17rD18VQnqAlwCJjd8WD7rD5vFUIfv6sWk/I+f5QgEsAROXfGpPyPn+UIPqxqb8j5/lCeoCWIImLvjU/Vh83yhPUcLvjRdWHy/KE9QEtARFk3xZP9niH61R/kdZ++HcV15OJwp2a1C+oCYYIVSb4GVLryMZtzexVncv8Dw67es2jzq/mIrbTo5e50iVeT84E8F4dJ07jdbbbollrq+mpmJ0rJIjfOV2X3bvtPu6SNlySeBknS2BEZp4DBbvkF8u7+VdLvW1i/8AizOcgE+c33jNmuNseyK6/RWqbqnM0acvinUruhCOm0/ejzDIo5aHG4m2KjdqiyN7qZyeHoQj8APvX1tXX1T6qtqZaieRdXySOVzlXwqfAAAAAAAA9bC3cjMbK7suEC/+40tSpV1pol7WIvzFTlHUSUlZDVRLpJDI2Ri99F1Qk1aN7y801BBT1WM080kbEa57ZlTlKidOmgEzgRCj3xajTu8RavgqP8j6fViv/JBflCeoCXIIhv3xZ/ucRb46j/I+L98S4L7XEoU8M/8AkBMIgvv3e6nR/Ak86mRfVhXTrxSD9uvqNJ7a9o1TtLyeO91NAyidHCkSRtfyuvpAwQAAAAAJEbhya7Ua9eyhX0kI7me7EdpFTsyyea9U1AytWWBYljc/k9aLr8wFmAIet3xLh91iUPiqF9R9G74lVp3WIsXwVH+QEvgRD+rFm/JBvyj/ACOF3xajqxFvyj/ICXp8LgulDOv/AIbvMREfviV6+0xOJPDP/kdau3vbrPSSwsxeBjntVvK59eGviAjjl/HK7sqf22X01Pd2LY07Ldp1ksiIislqEdLqmqchvdO+ZDFblVOrbhUVj0Rrp5XSKidSqupm2xLaIuzTJZr7FaIrjO+BYo+cfyeb1VFVU4d4CyulhjpqaKniajWRsRjUTqRE0PqQ8j3xLin9JiMC+CoX1HYbvizfdYgniqP8gJdgiN9WK/T7UF+UJ6j4Tb4lYqfWsSjT30/+QEqM0vdPjmK3K91T+RFSU75VXwIVc5HdKi936uu9U9z5qud8rlXp4rrobi2ubxl/z3Fp8e+hkNvpqhU51zJFVXIi66GjgAAAAAAAABlGyfJZMQ2h2bIGOVG01S1ZNOti8HJ5FMXAFslrrILjbqeupno+GeNsjHIvBUVNTskBtke8lkmDWCKx1VBHdqODhCskitexvZr2Gw274rdO6w9+veqE9QEtgRKXfFZ1YfJ8oT1H4dvir9ziDvHUJ6gJcAiE/fFqfuMQYnhqP8j4P3w7kvtMUhTwz/5AZZv54olbhVFlcMacu3TNjmd18h6o1P8A1KhCo35ta3kK/aBs/uOJ1GPw0rK1YlWZsvKVvIlZJ0adfJ08ZoMAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB2LfW1dvqmVVFUSQTMXVr2LoqHXAG5ML2uRuaylySPkv6PZMbeH95PUbWoK2kr6VlVRVEdRC9NWvY7VFIinpWS+3ayzpLba6aDRdVa13cr4U6FKjb6Phz/ABU+Gf0Q82lTJ3jtKV4NOY5tke1rIb7buV99PTrov6q/wU2HZsyxq6tRaW7QcpU1VkjkY5PEpns/S9nD/LzH4K7JqZae3LIAcNc1yatcip3j9ECazHlG4n3cAA8AHJwAA1RDyrxkdjtLFdX3OnhVOlvLRXeROJ0x4cmSeKVmX1XHa3aIerqfieaKCJZZ5GxMamrnOXRETt4mrcl2w0ELXRWOjfUydCSzdyxO/p0r8xq/I8svt/ev0QrnujVdUiZ3LE8Rca3Q8t++WfTH6puLQvbvfs25mW1a2W5H01lRtfU8U5xF+ttXw9Zpa+3m43utdV3KqfPI5eGq8Gp2InUh0AaPW08WtXjHH9/dZYsNMUcVgABJdQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAACKqLqi6AAepbMhvltejqG7VkOi66JKvJXwp0KZJQbVMvpkVJKuCp1/Cwpw8mhg4Pi+Ol/vREvma1t5hsui2x36PleyqGjn16OTq3T51O6zbTXad3ZoFXvSKanBHnQ1p80j8nxODHP8raz9tFw+4s9OnhkU8uo2u5S9zuabRQoq8NItVTyqa9B7XS16+KR+RGHHHirILrmuUXNvIqrzU8jo5Ma82mn93Q8F73yOVz3uc5elVXVT8gkxWK9oh0iIjwAA9egAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAO9YbVXXy80lotsSTVlXIkULFcjeU5ehNV4Gzvqb9r/wCS6fK4v8RjGwr3YMW+MY/OWcAV2ru4bX0/3X//ANUX+I/Lt3Pa8iKv0quXwVMX+InnW5ridFVSUtXkFvhnjcrXsfO1FavYpzRZpidZO2ClyG2yyu9q1tQ3VfnArVy7BcvxKV0eRY9cLeiO5POSRKsar2I9NWr4lMcLYLtbLbebfJQ3KkgrKWZujo5WI5rkIA70+yyn2cZjHJaUf9BrkjpKdruPNOReLNezjwA06AANhYbsX2jZfYIL7YLD7Lt86uSOX2RG3XkqqLwVUXpRT2vqcNr/AOS6fK4v8RK/c39wGyfpJ/3rjY+VZZjmKxQy5Dd6W2sncrYlnejUeqdKIBApN2/bAv8Auu35ZF/iC7t22BE+1hq/9ZF/iJqptg2Zqv252j5Qh+v53tmn5aWf5QgEANoGyzOMDoIK/KLP7Bp55ObjdzzH6u0107lVMKJb76+b4lk+DWmlsF/oLjPFXct7IJUcrW8leJEgAfaipaquqo6Wjp5aioldyY4omK5zl7EROKnxJy7neyq1WPD6bM7jSNmvFxZy4nSoi8zF1clOpV6dQI+4zu07Vr3RtqltFLbWOajmpXVKMcqL3kRVTwLod+4bq+1akpXTMp7RVuToigrNXr+s1E+cm1muaYzhlE2ryO7U9DG9dGI93dO8CdKnj4vtb2e5JOtPa8moZJU+4e/kKvg1Arry3CsrxOpfBkVgr7erXK3lyxLzblT716dy7xKY+WoZtFjdTitbLk8NJPaGQrJPz6IrOSnHUrFzKez1WVXKosFM+ltb6hy0sTl1VrNeAHkn6iY6WVsbE1c9Uaidqqfk7Nq/rOl/TM9JANn0+7rtcngjmixjVkjUc1fZcXFFTVPuj6fU37YPyXT5XF/iLBLEutkoV/4eP0UPAvm0fBrHc5bbdsnttHWRac5DLMjXN16NUAg2m7dtfX/dhqf9ZF/iOV3bdr6Jr9LLF/6yL/ETW/ne2aflpZ/lCD+d7Zprp9Odn+UIBXNm2KXzDL9JY8io/YlfGxr3Rc41+iOTVOLVVDxDbm9terTkG2WtuVlr4K6jfSwtbNC9HNVUbxTU1GAAAHoY1Z67Ib/QWO2Rc5WV07IIW6LpynLpquiLoidKr1Iim6W7qG1RU11sSeGtd/gMk3DsKWuyiuzSqicsNAxaemVUTkrI5O6XwonnJoAVT5dYLji+R1tguzGMraKVY5UY7VuveXrQ8olDv4YQtHf6HNaSJ6xVjUgqnJ0NentV8aEXgAAA/cET5pmQxN5T5HI1qdqrwQ2rDu67XZWNe3Fl5Lk1RVqov8RrXHk1v9vT/io/SQtapPsSH3jfMBXn9Tftf/JdPlcX+I/K7uW15P8AdZfFUxf4ietbmOLUVVJS1d+t8M8buS9j52orV7FTU/EOb4jNIkceRWxz3cERKhvH5wK2sv2fZriUjm5DjdfRNbprKsfLj4/nt1b85jBbFXUdBdaF9NVwQ1VNM3kua9qOa5FIL72+yKkwG9U99sMbmWe4vcixdUEvTyU7yp0AaFAAA79js13vlalFZrZV3CpXoipoXSO8OiIZTsX2c3TaVmENmouVDStVH1dTydUiZ616ELCtnGz/ABjArNHbbBbooVRE5ydURZJV7XO6QIX43ut7ULqxJK2moLSxeqpqEV/kbr5zu3XdO2lUrXPo57PWo1nK0SoVjlXsRFT+JMrKs5xLF+F+v1DQv6mSSojl8XSdbFtpOD5NK2Gy5HQVMzuiNJUR/kUCuTN8Cy/C6pafJbDV0HZI5vKjd4Hpq1fBqY0Ws5PbLLdrJU0l+paapoHMVZWztRW6Jx14laW12ixa37QLnS4dVyVVoZIvNOcnBq9bWr1tRehQMTO3ZrdV3e7UtroIudq6uVsMLOUicp7l0RNV4JxOoZXsd91bFvjWn/eIBlq7um19E1+lR3yqL/Ea1v8Aaa+xXmrs90g5itpJFimj5SLyXJ0pqnAteKzt4JNNtWWIn4ykAwQzLZ9swzXPaWpqsWtHs6KlekcrueYzkuVNdO6VOow0mV/J8fapkvw5noIBHPNNj20PDrHJeshsDqOgic1r5eeY5EVy6JwRVXpVDASwbfRbrsBvC9k9N++aV8gAAAAAAA3xue7MaHOMtnvF7p+ftdqVruad7WWVeLUXtROnQDB8C2N7RM1hSpsuPTpSKqJ7JqVSKNdetOVxcnfRFM4+pR2qf8j+WL/gJy11XasftDqiqlp6Cgp2cXLoxjEQwOh277LqyvZRxZRStke7kor9Wt18K8AIMZtsi2h4g5VvGM1nMaKvP07eej0TrVW68nx6GCqioqoqaKhbFR1dvutHz1LPT1lO9PbMcj2qhAne+kwdm0ZaDErdDTVNMipcJINEjfIvHTROtOsDSYAAHfsVku99rW0dmtlXcKh3RHTwuevzIZfsP2Z3PabljbVSOWnooUR9ZUqmqRs16E/OXqLBNnOz/F8Bs7bdj1ujgTT65M7upJF7VcvECFmP7ru1O5xNlqqOgtbVXi2pqUV6J26N1T5zsXzdW2nUELpaRltuKNTXkRVGj18CKmnzkz8tz/EMV4X2/UdG/TXkPkTleQ6mKbUMEyiVsNmyOiqJndEfOI1y+JQK3ssxTI8Ur3UOQ2ert8zV0TnY1Rrveu6F8SnilqmVYzYsptcluvlup62nkTRUkYiqnfReogjvK7F6jZtdkuNs5yewVT9Inu4rC77xy+ZQNMgAD0sasd0yO9U9ns1KtVXVC6RRI5EVy+FeBsOTd52uRxrI/FJEaiar/pEX+I/O6m3lbcrAn/iOX/0qWK1n2JN7xfMBU3VQS01TLTzN5MsT1Y9uvQ5F0VD5np5WmmU3ZOytm9NTzAAAAG06Hd92sVtFFWU2MOfDMxHsX2THxRU1T7o1YnShaphH2n2j4HF6KAVi5ni19w69us2RUDqGubG2RYnOR3crrouqKqdSnim9N+H3cH/F0PneaLA9HGrJc8jvdNZbPTLVV1U7kQxIqJyl6eleBsd+7rtdZGsjsXVGomq/6TH/AIjq7rKIu3XGteqdfRUsZrfsOb3i+YCp6vpZ6Gtmo6pnNzwPWORuuujkXRUPie3nv27Xr4dL6aniAZlgOzDNs6o5qvGLK+uggfzcj0lY1EdprpxVDK27t+15U+1lqeGri/xG49x/K8csWBXamvF4o6GZ1fy2tmlRqqnJTjxJCN2iYQ5dEye1r/1DQIMLu3bXk/3ab8rj9Z+Hbue15v8Auuq+Cpj9ZYdS1ENVTx1FPK2WGRqOY9q6o5O1Dyr/AJXjthmZDeLxR0Mj05TWzSI1VTxgQCdu77XG9OKP+UR/4jWt8tddZbtU2q5QLBWUsixzRqqLyXJ1aoWWVW1bZ7TsV0mV2zgnVMild21i50152kX+6UciSU9TWyPjenQ5uvBQMYPcwnE79md7SzY5QrW1yxul5pHtb3Kaarqqp2oeGb33Gfdu8Fsn9JgGM1u7/tYo6OWrqMWkZDExXvd7Ij4InSvtjVxaxmCa4pdU/wCDl9FSqdelQAAAGZYLsuzvNXa4/jtXPCnTPInNxJ/edoi+LU31uqbA4bjT0+a5nSJJTPRH0NFInB3Y9ydnYhLdraG10WjWwUlNE3oREY1qIBB2l3TNpEkCvnqrNC/qYk6u18ehimYbvm1DGqZ9VPY0r6diK5z6KRJOSng4L5EUmrdNs+zS3VraOpyug51X8hUa/lI1e+qdBleO5HYcko/ZNludLXwqnFYpEd5QKq5opIZXRSxuje1dHNcmiovfQ/JYbtx2FYzn9rnqaKmitl8aiviqomInOO09q9OtFICZNZLljl8qrLdqd1PWUsiskYvb2p3gPNAAGc4HslzzOLQ+641ZFraNkqxLJzzG90mmqaKuvWh+M62U57hNvZcMjsE1JSudyOdR7XtRe/yVXQlpuHJpseqV7bnL6LTdWaY3bMsxussN2hSWlqo1Y7tavUqd9AKqwZrtk2e3XZzmFRZa9rn06uV1JUcnRs0fUqd9OhTCgNi4rsS2l5PZYLzZscfPQ1DeVFKszG8pO3RVMdz3B8nwa4Q0GT211DPPGskbVka7lNRdOlqqWG7vcbY9jWMNamiewWKRo/lAW6Zxj7u2hf6SARmAAH0poJqmdkFPE+aV66MYxquc5exEQ2viW7rtTyOjbVx2SO3QubymLXy80rv7uiqnjRDfW5bssttBicWdXejjnuNfq6j5xuvMxIuiOTvu6dew3xmmY43h1C2syG6QUMTuDOWvFy9iJ1gQjrN1narT075W01qqFamqRxVernd5NWonzmsMxwbLcQqXQZFYa2gVqonOPj1jXVNeD01aviUsKxjbLs5yKr9iW7JaRZ+pkruQq+DUynKILBXY/VOvsNJU2xIldNzyI5nJROK8QKqgZBtGqMfqs1uk2LUi0lnWoclLGrlXudenj0IvUhj4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAZpsK92DFvjGLzlnClZGwr3YcW+MY/OWcAVjbc3L/PDlnFf60m9JTDY5JI5GyRyOY9q6tc1dFRe1CZOfbqVRk2ZXbIWZnHTJcKp9RzK0Cu5HKXXTXlpqdKy7nFHFXsfeM0lqqRPbR09GkT1X3yudw8QG090i+XS/7E7XWXepfU1EcksCSPXVysY5UbqvXwMA/lBI2/SJj8nJTlpc1br16c28kJiOPWrFMdpLFZqdKeipGciNuvzqvWpELfqzujvOQ0GH2+RkzbY5Zqp7XaokqpojfEirr4QIzAACwzc49wCx/pKj96415/KBQTz2DGEghkl0qpdUY1V07hOw2Huc+4BYv0lR+9cbcnp4J0RJ4Y5dOjltRdAKoPofcP7DVfsneofQ+v/sNT+yd6i1r6H0H9ip/2aD6H0H9ip/2aAVPSRvierJGOY5Olrk0VD8m0t6tjI9ud/ZHG2NqPZo1qaJ7VDVoAtH2WVFPVbOrBNTSMkidQxaOYuqL3KFXBKHdX2+27GrXHh2ZTLBQxrpRVmiqkaL9y/vdigZZvlbJsvzC80GTY3E+4x09LzEtG12jmaKq8pqLwXXXj4EIh3Wy3qzVKxXK2VtDMzRVSWJzFTv8AEtLsl8s97oo6y03KkrqeRNWPhlR6KniP1dbNabrA+C422lq43po5ssSORfKBWxXbVc3rsGkw6tvU9TbJHNVUkXV3JToZr97w6DCCde1rdhxHIqSorcVZ9A7tyVVjGL/o8juxzerwoQkyOzXDH75V2a607qespJFjljd1KgHnnZtP9aUn6dnpIdY7Vo/rWk/Ts9JALVrFwslCn/Dx+ihX9vZUdZLt1vz4qSd7VWPRzY1VF7hCwKx/1LQ/B4/RQ+slHSSvV8lLC9y9KuYiqBVF9Drh/Yar9k71BbfXomq0VSn/AJTvUWtfQ+h/sdP+zQ6OQUFC2xV7ko6fVKd/+zT71QKqgfWt+zJ/0jvOfIAfSlglqamKngYr5ZXoxjU6VVV0RD5m5t0LCXZbtVpqyoiV1BaU9kyqqcFd9wnl4+ICZuwnDYMG2ZWmyMYjZ+ZSWpd1ulcmrl8qniR7W6OTb87Zy1YVgbRqqyovH2R08jxN+dTP8vvVLjeLXK91bkbBQ0z5nd/kpwTx9BWpBll5i2lJm3ImStWv9lqjtV6Xa8nXhrw4AWFbcMOgznZtdbHIxrpliWWmcv3MrU1av8PGVnVlPLSVc1LOxWSwvVj2qnFFRdFQtSxC9U2R4vbr3SrrDWU7ZUTs1TinlII74eFOxTatPXQQq2gvDfZMLtEROX0PangXRfGBpYAAd/HPtgt3wqP0kLWaT7Eh943zFU2OfbDbvhUXpIWs0v2LF7xPMBWlvAuVdtOWcV/rKQwVHKi6oq+UmftD3VZssza75GmaR0iXCpdPzK29X8jXq15aa+Q8u1bm9OytY66ZvJPSp7ZlPRJG9fA5XOT5gM23JMjvV92YzwXZ887aGpWKnnlVVV7NNdNV6dOg+++5S08+xSeeViOkgq4nRqv3Kqui/MbWwTE7LhWN01hsVMkFHAnWurnL1uVetVIxb9e0K31UdHglum52eGVKitcx3BnDuWL2r194CJoATpAn7uZ4fDjuyemuzmN9l3j/AEh7k6eR9ynkPd3ldpv82uDLV0iNfdK1ywUbV6EdpxcvgQzPZxDFBgViigjbHG2gh0a1NETuEIvfyhrn/RbEGaryFgqVVOrXlRgRhv14ud9us90u1ZNV1c71fJJI7VVU6tNUT0s7J6aaSGVi6tfG5WuaveVD5gDaFy265/cNnrsOq7o6SF68l9Uq/Xnx/eKvZ3zV4AAyzY57q2LfGtP+8QxMyzY57q2LfGtP+8QC0ErV3gKGtftoyt7KOoc1bjIqKkaqi/MWUp0HwfRUb3q99LA5y8VVY0VVAqj+h9f/AGGp/ZO9RMf+T/p54MTyRJ4ZIlWuZoj2qmvcJ2klPYFD/Y6f9mh9YYIYEVIYo40Xp5LUTUDT2+d/3fr1+npv3zCvcsI3zv8Au/Xv9NTfvmFe4AAAAAAJm/yfssS4jf4ke3nEq2qrdeOnJIZGyt33alVbMMv9nuhdU2yqRI6yFvtlb1Ob30Am1vHYRdM+2ZVdjs1QkVakjZY2Odo2Xk/cqpAHK8AzLFqh8N8x6vpeQ7k84sSujVerRycFLF8E2l4XmlI2exXylmeqIr4XPRsjF7FavEyqenpqqPkzwxTMXqc1HIBWRgW0nMsGdLHYbvUU8MjVa+nc5Vj1Xr5K9CmKVtTPWVc1XUyOlnmer5HuXVXOVdVUsY2gbCdnWYwSLUWWOgrH6q2qo0SN6KvXw4L4yEm3DZXedl+RpQ1r0qqCo1dSVbU0R7dehexydYGvQD9Q8ZWe+QCwrdKwunxTZPQ1SwIyuuiJU1D1Tul19qniQ/G9TtSk2dYYyG2ORLzclWOmX8G1PbP8XV3zZWDRxxYdaI4kRGNo4tET3qEPd/uWodtItEL+VzDLfrH2aq7j/ACPV4ulwvFwluF0rJquqmdynyyuVyqp16aeamnZPTyvhlYurXscqOavaiofMATY3ONr9yyxlTiOSVXsivpIklpZ3r3czNdFRe+nDym6drmLUuY7PrtY6pjXc7TuWJVT2r0TVqp40IJ7p8s0W3Ow8y9zVc97XadaK1dULFZURYnoqaorV1Aqbr6aSjrZ6SVNJIZHMcnfRdD4mRbTWtZtDyBrURGpcJkRE9+pjoG1d0/3c7F75/oqWJ1X2NL7xfMV2bp/u5WL3z/RUsUqfseX3i+YCqrLvtru/wAOm9NTyz1Mv+2y7/Dpv3inlgAAATpLVcJ+1C0fA4vRQqqTpLVcJ+0+0fA4vRQCD+/H7t7vi2D0nmije+/Kmm29fiyD0nmiANn7rPu643+nd6KljFb9hze8XzFdG6smu3bHP0zvRUsXrfsOb9G7zAVZ599u16+HS+mp4h7WeLrmt6X/AI6X01PFABFXtAQC0HY9p/NfjmnR9D4vRQijv9e6DZ/gS+kSt2N+5ZjfxfF6KEUt/v3QLP8AAl9ICNgAAG99xr3bv/4yf0mGiDfG4wmu25fiuf0mATiyv7WLp8El9FSqVektayr7Wbp8El9FSqVelQBnOwfFo8x2qWWxzacxJNzkqL1sYnKVPmMGJBbh1PTzbYauSaNr3w2qV8SqntXctiap39FVPGBOWkggo6SKmgY2KGFiMY1qaI1qIQY3sNr92yPLazFrRWyU9moJFiekL9OfenSqqnV3ic1w19gVGnTzTvMpVLfHPfeq171VXLUPVVX3ygdMybZ1m9/wW/w3ax10sKtciyxI7uJW9aOTrMZAFo+zHLKTNsItuR0eiMqokV7dfaPTg5PEpGrf4w2nidac0pYuTJI5aSqVreC8NWuVfKnjMw3Camqm2W3GGaV7oYLk5sLVXg1Fa1VRPGqqevvusY7YrOrkTVtXEqa9uoEBgABOvcP9x2p+NJfRYb9lngikZHLNGx8i6Ma5yIrvAaB3DvcdqfjSX0WnT3275cMbo8TvNrnfDVU1c6RiovBdETgveUDZO33ZnQ7SsKmtzkZFcoEWWin04tenUveXoK58hs9wsF5qrRdad1PWUsixysXqVPOhZHsT2i2zaThlPeKJebqmIkdZTr0xSInFPB2Kau3wtkMOS2J+Y2KkRLxQt1qGxt41ESd5OtOkDaOwD3HMY+AM8xGr+UET/wCs8dX/AIKT0mkl9gfKTY7jCORUclAxFRU6OBGn+UE+3HHPgUvpNAjCAALLN3OWOXYjibo3IqJbo0XTqVE4mrt8nZblubvtl5xuJKyOgheyakR+j3KqovKai8FXga03VtvNDhVB9KOWc421LIr6Wrbq7mFXpa5PvdePAmLjuT4/kVEysst3o66F6ao6KVHfMBV7eLBfbHPzd1tNdQSJx+vQuZ4+J7385uauwqfEJb5Uy2qZU5THu1cife8rp0XsLLbrZrTdqd0Fyt1LVxPTRzZYkcip4zSO1bdjwvI6GWpxmL6BXREVzFi4wvXsc3q8KAQPB6uW4/c8XyGssV4gWCspJFZI3qXvp2op5QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAZpsK92DFvjGLzlnJWPsJTXbDiyf8xj85ZuoEfMw3psZxrKblYKnHbrLNQVL6d72OZyXK1dNU1U3Ds6zC0ZzitJkFmkV0FQ3VWO9tG7ra7voV0bc112w5Z8aT+kpsLdG2qJhGW/QK7TOSzXR6M1VeEMqrojvAvQoEjN767ZzZdnC12I1HselRysuMkSLzzGLwRWr1J2r0kAZpZJ5nzTSPklkcrnveurnKvFVVetS1+vpKK7WyWjq4o6mkqolY9jk1a9rkK6t4vZlNs0zmSihR77TV6y0EjvvdeLFXrVANZAACwzc59wCxfpKj964xbfbzDJsRsePS43eaq2SVFTI2Z0DkRXojUVEUync49wCxfpKj964zLaVs5xfaFTUlPk1I+ojpHq+JGvVuiqmi9AEAU207VE6M3u37RPUP56tqn5b3X9dvqJkfU0bKPxLN+3ccpu1bKE/+ySr/wCe4CAuQXm6X+6zXW81stbWzaLJNJ7Z2nadAlTvX7IsGwTZ/T3THba+mq31bY1csqu7nRe0isAMoxzZ/l2Q43XZBZ7NPV0FCqJM9icf7qdLtOvQ8OxxUM94pIbnUPpqJ8zWzysbylYxV4qieAsw2Qw4fDgdBSYZUU1Tao4ka10aoquXrV3fXvgVq2a9XzHq7n7Tcq221DHoqrDK5i8pF4aonT4FNtYJvLbRbBVxJdK9L3Ro7WSOoROWqd5yEsc92EbOcwnkq62ytpayTVXT0jubcqrx1XTgphWM7p+B2y5Mq7hW3C6MYqK2GVyNaqovXp0gb0xm6NvWPW+7sjWJtZTsnRi9LeUmuhCXfvt9PR7XqWphbyX1ltjkl77kc5vmRCcMbKK02xsbebpaOli0TVdGsY1PMiFde85m9Pne1auuFA7l0FKiUtM7XVHtbrq5O8q6gawOzaf61pP07PSQ6x2bT/WlJ+nZ6SAWr2P+paH4PH6KEKN5bajn+P7YrzarLlVwoqGFWc3DE5Ea3VqKvUTYsf8AUtD8Hj9FDXuZ7DNn2W5FUX6822WWtqNOcc2VURdE0TgBCH+evar+XF1/Xb6j8y7Z9qMsT4pM1uj2PRWuarm6Ki+ImT9TTsp/E037dxym7Vsp/Esv7dwFe7nK5yucuqquqqcG597fB8cwPOaC2Y3SLS00tEkr2q9Xau5SprxNMAE4roWB7n2EpimyyCtqIObr7s72RKqpx5P3KeTzkMth+GzZ1tKtViYjuYdKktS5E15MTeLvV4yy6CKnt1uZExGxU9NEjUToRrWoB95GMkYrJGtc1elFTVD4ewaL+yU/7NCAm1TbvnlbtAvEthyWsoLYypdHTQw6I1GNXkovHrXTVfCYx/PZtT/LS5/rN9QFlLGMY1GMajWp0IiaIhp3e6wf6cNlVTU00HOXG0qtVBondK1E7tvk8xGfY9t6zei2iWl2S5HV3C1TTJDURzKnJRHcOVwTqXQnsqQ1dIqLyZIZmeFHNVAKmF4LooM/3gMLkwXahdbO1itpJJVqKRdOCxPXVE8XR4jAAO/jf2w274VH6SFrNL9ixe8TzFU2OfbBbvhUfpIWs0v2LF7xPMBoPNt6LGcWyy5Y7VY/dJprfUOgfJGrOS5U601U2rsrzu0bQ8SgyGzq5kb3KySF6or4nJ0tdp1lfG8D7tOWfGUnnMk3YNqc2zrNGQVszvoFcHJHVsVeEbuhJETtTr7wEtd6Stzu2bOJ7lhFSkLoF1rFY3WRItOKt76FeVbVVNbVSVVZPJPPI5XSSSOVznL2qqlrzm0tyt6tcjKilqY+heLXtcnqK+t6PZfNs+ziSpo4HJZLi5ZKV6JwY7rYvg6gNQAACxbdUyeDJNjdoRsz5Kmhj9i1HLfyncpvDVfCh4u+Fs7q82wGOvtNOs9ztLlljY32z419s1O/w18RF7do2sy7Ncr5uuV0ljrlRtWxOKxr1PRO91lgVgvFrv8AaILpaayGso6hvKZJG5HNVAKpJY5IpXRSscx7V0c1yaKi9in5LH892FbOcyrVrrlZvY9U5dXy0j+ac/w6HmYxu37LbFcW10dpnrXsXVrauZZGIvgAg/S7OczqcLmzCGx1LrPC7kum5PFe+idKp3zEy16rgtVFZZYKmKmgt0cSo9jmokaM04pp0aFau2uTEJdotzfhKSJaVkXk8r2vL+65H5uvQBhZlmxz3V8W+Naf94hiZlmxv3V8W+Naf94gFoBAXbZta2j2faxklstmX3Klo6avfHDCxycljU6ETgT7Kzt4T3a8t+MpAPt/PbtW/Le6frN9RKfcozDJcvxq/VGS3ipuc0FYxkTplRVa1WIuiaJ2kFiZf8nx9qOSr/xzPQQDPd87/u/Xv9PTfvmFe5YTvm/9329/pqb98wr2AAAAAAB7GH4xe8tvcVmsFC+rrJdeSxvBEROtV6EQ8cmvuQUuBU2LPnt1winyeo1WtjlVGyRoi8GtT73vgQ8vFrvmLXp9HcKeqtlwp3Kmi6sc1e1FTzoZrie3LaZjj4EpsmqqqCFqMbBVLzjdPHx8epPnOtnmH5rAseRWWmq36aNl5OkjfA5OJp6s3SMEluiVEFzukFLy0VafloqadacpU14gZ/u57SqjadhMl3rKFtJVU0608yMdq1yoiLqnlMN366Snl2PRVb42rNDcIkjcqcU111Nt7PMKsGB4+2y49S+x6ZHctyq7Vz3ffKvWpG/fuzyilo6LBaKVks6SJU1atdryNPatXv8AWBEYNVUVFTpTiABZJu2ZTBlWyOz1TJGOnp4Up52ovtXN4cTCd8nZhWZpi1Pf7JTunutpRdYm+2lhXi5E76dJHPdh2tu2bZO+mubnvsVeqNqGpx5p3VIn8Sflju9rv1riuNprYK2kmbqySJ6OavkAqjmjkhldFKxzJGKqOa5NFRexT8ljmd7Btm+YXB1wuFndT1T11kkpJFiV69/Q8/GN2/ZbYriyujtM9ZIxeU1tXOsjUXwAai3INmNcy7TZ7eqJ8MDIubtySJor3L7Z+nYicE8Kko85vdLjuI3O81kqRxUtM96qq6cdOCeU9GaWgtNvdJK+CjpIGaqqqjGManzIQq3tdtkGXTriWL1LnWmB/wDpU7V7mocnUna1PnAj7e62S5XisuEqqr6iZ8qqvfXU6YAG1t073crF75/oqWJ1H2PJ7xfMV2bpvu6WH3z/AEVLFnIjmq1ehU0UCqfL/tsvHw6b01PLLEKrdz2WVVZNVT2WR8s0jpHqs7uLnLqp813bdk/4hf8At3esCvME4drO7/szsezm/Xm32qaGro6KSaF3PuVEciapwIPAE6S0jZdcqW77PLDcaNyugmoYnMVU0XTkp1FW5Mrcj2n0NRYUwC61TYq2lVz6FZHac5H0q1NetOzsAxPf3xWqhy215ZFG99NU0yU0rkbwY5iqrdV7/KXyEYS1nJrBaMltMtqvdDDW0kqaOjkTVPF2GoX7rWyx1Us/sS4Iiu5XNpVLyU72nYBHncpxWtvO1uG9JC/2DaY3SSS9CctyaNb4enyE5sorG2/HLjWvdyWw0z3qvZo1TqYXiWP4daG2vHrbDRUycVRicXL2qvSqmk98fanRWDEp8OtlU193uLeTKjHa8zF169ir1AQnvdU+uvNbWSO5T5p3vVdOnVVU6YAAAAWgbHPctxv4vi9FCKW/37oNn+BL6RK3Y77luN/F8XooRS3+/dCs/wABX0gI2AAAb53Fvdud8Vz+lGaGN87i3u3O+Kp/SjAnDlf2sXT4JL6KlUq9KlrWV/axdPgkvoqVSr0gDZu7DkcOMbZrLW1Mjo6eZ600io7RO7TRNe9roayP1DI+GVksTlY9io5rk6UVALaVRskapwVrk8qKV37zuza44PtAraplK5bRcJHT00zUVWt1Xi1V6lRSRm63tyo8rtNNi+S1kcN8p2pHFJI5GpVInRp+d3jemQ2K0ZDbn2+82+nrqV/THKxHIBVId/H7Pcb9d6e1WulkqaqoejGMY3VdVJ33Ddd2V1dZLUpQ18CSLrzcVUrWN8CGc7PdlmE4Lq/H7NFDUKmjp393IvjUBsRwxuBbN7Xjqqjp4o1fO5OuRy8p3zqaR3+Mrp4ccteIwytdU1U/smZmmvJYzo49XFUN57VNouO7O8elul7q2pJoqQUzV1kmd1IifxK59pWYXLOswrciua6SVD+4jReEbE9q1PAgGNgACdW4d7jtT8aS+i08L+UCbriWPu7Kt6f+k97cP9xyp+NJfRYeL/KAp/8ARdhX/jHeiBHDYRtMuWzTMoLjC98lsncjK+mReEjO1O+nShY1j92t2QWSmuttnZU0dVEj2ORdUVFToUqjJEbom2OTFLyzEL9Up9Bax/1iSR32PIv/AMVAnDTwQ00LYaeJkUbehjG6IniIafygn24458Cl9JCZrHNexr2ORzXJqip1oQy/lBU/+scbX/gpfSaBGEA+1AynkrYY6uZ0NO56JJI1vKVrdeK6dYGQYlgWWZXa7hc7DZ56ylt7OVO9ifMnavXoh41DXXWyVzn0VXWW6qYujlikdG9FRehdNCyXYdR4TQ7PaClwealntyRor3xuRznvVE5Sv6+V26nQ2g7ENnuayPqLlZmU9Y7VVqKVebeq9/TpAiHgW8jtGxypYlfcfo1Sa93FVJq7TvOQnVgeRU+WYhbcipI3xw10DZWsemit16UNI2TdLwaiurKqsuVxr4GLr7He5GtXj1qiam/6OnobPa46anZHS0dLGjWNTg1jUQCEu/lS08G1GgnihYySahRZHNTRXqjlRNe0juba3qs6o852oz1FseySgoWexYZW/wC00VdXeDXoNSgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB7OEX1+MZbbMgjp0qH0FQ2dIldojtOrUk/JvlsWDRmAuSXT2zrnqmvg5siOAPXzW+PyXLbpkEkCU77hUvqFiR2qMVy66ankIqoqKiqip0KgAEhcH3qMqxzF6Ky1FnpLm6kjSNtRLI5HuROjXTvaHibYdvcu0zFvoNeMToopY3pJT1Uc7uVC7tRNOPDVDSoAAADe2yjeQvGz/BaPFqPHqOrjpVerZpJXIq8pyu6ETvmSP3wcrX2uMWpP/MeRlAEmU3wcs68ZtS/33hd8HLPyZtP67yMwA3Htn29XfabjENkuFlpKJsUyS85DI5VVU6tFNOAADJcEzzLMIrFqcavNRRcpUWSNq6xyaffNXgpjQAkvje97lVJzDL5jtvuLGJpK6GR0L5OHTroqIuveMgq98prqaRtLgTo5lTuHSXNHNRe+iRpr5SJAA2ptR2857ntNJQVday322Re6paTVqOTsc7pVO8arAAH7ppVhqI5kRFWN6ORF69F1PwAJL0u93ktNRw00eMW1UijaxFWV3HRND9Lvg5X1Yza/2jyMwAk43fCyj7rFrWv/AJrz9N3w8k17rFbaqd6ZxGEAZ9tv2l1e1DJKe9VlthoHwQcykcb1ciprrrqpgIAGzdgu1WPZZcq+4sx+O6VNVGkbXun5vm29OicF6TP803rMhyDG7hZoMepKH2ZE6LnmTuVzEXgunBOJHMAFVVVVXpUAAEVUVFRdFTihJjBd7O6WHGaK0XPFornJSxNiSobVrGr0RNEVU5K8SM4A2lt+2tR7Va6grHY5Fap6Nrmcts6yOe1epeCGrQAPtQVC0ldBVNajlhkbIiL16LqSlpd8WtipGRPwmF8jW6cpK5UTycgioAPazrIJMqzC6ZFLTtp5LhUOndE12qMVerU8UADf+y/eeyPDcUpMeqrPT3eKkbyIZZJnNejOpq8F106DjatvGR7Q8QqLBdcJpG8vuoZ/ZSq6F/U5O5NAgAAABmuzXajmez6Zy47dnx07+L6WXu4XL28lehe+hhQAlbY98ethpOTecKiqqjX29NXLE3TwKx3nPtct8mR9KrbdgqQT6po+a4c43TwJGnnImADZG1LbXnW0JjqW63BKW3K7X2HS6sjX33W7xmtwAB6mI3h+PZRbL5HC2d9BVR1DY3Lojla5F018R5YAlqzfKckCI7A0WXT2yXHhr4ObIy55kEmV5jdcjlp20z7hUOnWJrtUZr1a9Z4gAG5N3rbi7ZPb7jQux5LrDWzJKqpU805qo3TT2qmmwBIPbbvJLtGwOqxWHFPoZHUvje+Z1bzqpyHo5NE5CdhHwAAAAAAAHbtFzuFnuEVwtdbPR1US6slherXN8aHUAG+sO3p9odkpEpbkyivTW8EkqGq2TTwt6fIbAi3y4EjakmASOfonKVt0REVe8nNERQBIHPt6rOL7DLS2OkprFTyapymLzkyNXq5Soia99ENCV1XVV1XJV1lRLUVEruVJJI5XOcvaqqfEAAAAM12a7Ucz2fTOXHbq+KneuslNInLid/dXo8RhQAlTYN8a5wwObfMNp6yThyXUtYsKJ26orXan2vW+PVS0SstGEx0tTrwkqK7nW6e9RjfORQAGx9pu2nPM+jfS3a5+x7e52vsOmTkR+PrXxmuAAAAAyTZnltTg+ZUWS0lNHUzUiqrYnqqNdqmnUbyn3wMvcv1rHLSxO+56/wASNIAkl9V9mv5P2f8A9frC73ubacLBZ0/X9ZG0Ab4zPeey/KMWuGP1lntcMFdC6GR8fK5SIqdWqmhwAB9KaealqI6inlfFNG5HMexdFaqdaKfMAb62f70meY9BHSXiOnv1MzRNZ15EvJTq5SfxQ2F9WZS//t9N/wD2qf8A/IiGAJBZzvV5zemS09ipaWxQPVURzV52VGr+cqImvf0NC3OurLnXzV9wqZamqner5ZZHauc5etVOuAAAAAACSeJb110x/GLfZI8TpZ0ooGwpI6pcnKRqaa6aGrNt+06t2oZBTXastsNAtPDzTY43q5FTXXXVTX4AAAAZ5sK2h/zY5wmS/Qv6JotM+ndBz3Nro5WrrytF+97DAwBKzIN8D6IWeqoaXBlgfPE6PlyXHlo3VNNdEjQimvSAAAAH7p5pqedk9PK+KWNyOY9jtHNVOhUVOg3bgO83tDxqmbSV8sF9p2po32Wi843+8nT4zR4Al3HvlwJG1JMAlV+ndKl0REVf2RjWY73WUXCJ0WOWGks6OTRJJpVne1e9wanzEaQB6+WZNfsqur7nf7nUV9S9V7qV+qNTsanQid5DyAAAAA3bsO3gK3ZhicuPw49BcWPqXT846dWKiuRE000XsPP267cLjtTtlHb6qy09vipZVlascqvVdU00XVDUQAHLXK1yOaqoqLqip1HAAkJgu9RlmOYzSWastNLdn0zObbUzSuR7mp0a6dJr/bhtUuO1O60FfcLdT0LqKJ0bGxOVyKjlRevwGuwAAAGQYXmeT4bXezMbvNVb5F9skb+4f4W9Cm9MU3ustoWxR3+xUN1Y3274nrC93zKieQjWAJb1e+Ux1NI2lwJ8c6tXm3vuaOa1epVRI01TxmoNp+37Pc6o326oqo7bb3po+no9W8tOxztdVTvGpwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAf//Z" style="height:38px;width:auto;object-fit:contain;border-radius:4px;" alt="Zodha"/>
      <div>
        <div style="font-family:Playfair Display,serif;font-size:0.85rem;color:#4caf50;font-weight:700;line-height:1.2;">Research Writing Pro</div>
        <div style="font-size:0.62rem;color:rgba(255,255,255,0.4);letter-spacing:0.8px;text-transform:uppercase;">Zodha Research Solutions</div>
      </div>
    </div>''', unsafe_allow_html=True)

    st.markdown("### ⚙️ Style")
    style = st.selectbox("Style", list(STYLE_PROMPTS.keys()), label_visibility="collapsed")
    st.markdown("### 🔧 Intensity")
    intensity = st.radio("Intensity", ["Light","Moderate","Deep"], index=1, label_visibility="collapsed")

    # ── Resolve provider + model + key (admin-controlled) ────────────────────
    _active_provider = st.session_state.get("admin_provider", _DEFAULT_PROVIDER)
    _active_model_dict = OPENROUTER_MODELS if _active_provider == "openrouter" else GROQ_MODELS
    model_choice = st.session_state.get("admin_model_choice", list(_active_model_dict.keys())[0])
    if model_choice not in _active_model_dict:
        model_choice = list(_active_model_dict.keys())[0]

    # Resolve correct API key for the active provider
    if _active_provider == "openrouter":
        env_key  = os.environ.get("OPENROUTER_API_KEY", "")
        groq_key = st.session_state.get("platform_or_key", env_key)
        _provider_label = f"🌐 OpenRouter · {_active_model_dict.get(model_choice, model_choice).split(" · ")[0]}"
    else:
        env_key  = os.environ.get("GROQ_API_KEY", "")
        groq_key = st.session_state.get("platform_groq_key", env_key)
        _provider_label = f"⚡ Groq · {_active_model_dict.get(model_choice, model_choice).split(" · ")[0]}"

    if groq_key:
        st.markdown(f'<div style="font-size:0.72rem;color:#6fcf97;margin-top:0.3rem;">✅ {_provider_label} ready</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div style="font-size:0.72rem;color:#e8c97a;margin-top:0.3rem;">⚠️ {_active_provider.title()} key not configured</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("""<div style="font-size:0.72rem;color:rgba(255,255,255,0.35);line-height:1.7;">
    • Humanizer · Paraphraser · Grammar<br>
    • Research Tools · Statistics Suite<br>
    • Publication Suite · Literature Review<br>
    • 8-dimension humanness scoring
    </div>""", unsafe_allow_html=True)

    st.markdown("---")
    if st.button("🛡️ Admin", use_container_width=True, key="open_admin"):
        st.session_state.show_admin = True
        st.rerun()

# ── HERO ───────────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="hero-banner">
  <div class="hero-badge">v6.2 · Zodha</div>
  <div style="display:flex;align-items:center;gap:1.5rem;position:relative;">
    <img src="data:image/png;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAUDBAQEAwUEBAQFBQUGBwwIBwcHBw8LCwkMEQ8SEhEPERETFhwXExQaFRERGCEYGh0dHx8fExciJCIeJBweHx7/2wBDAQUFBQcGBw4ICA4eFBEUHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh7/wAARCAGGBOoDASIAAhEBAxEB/8QAHQABAAICAwEBAAAAAAAAAAAAAAgJBgcBBAUDAv/EAGMQAAEDAwEDBQgIEQgHBwMFAQABAgMEBQYRBwghEhMxQVEiYXGBkbGy0RQyN3J0dZShFRYYIzQ2QlJTVVZic5KTs8EJFzM1VILS4SQnQ0ZkZYQlOGODoqPCJkXwRFeFlaTT/8QAGgEBAAMBAQEAAAAAAAAAAAAAAAQFBgMBAv/EADQRAQACAgEDAgMGBQMFAAAAAAABAgMEEQUhMRJBIlGxEzJhkaHRI0JxgeEUFcEGM1Lw8f/aAAwDAQACEQMRAD8Ai4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB6+NY3esiquYtVDJPp7eTTRjO+53Qh5MxEcyeHkHLUVy6NRVVepDc+M7GadvJlyG4ukdwVYKXgnT0K5U8yGwrNi+O2ZrUttmpYXN4pI5nLk/WdqpW5+r62H35/oi5NzFTtzyjdbsVyO4ta6jstdK1y6I5IlRPKpklv2SZjVKvPQUdEidc9QnHxN1UkMrnKmnKXQ4KzJ/wBQWn7lEW3UJ/lq0guxK+pHyvo7ZXLp7Vqzar5Y0Q+X8zF//Gls8r/8JvQEW3W9mfHH5OU72Vo9mxO+vZyvo7ZWL969ZtfmjVDzK/ZFmFMiLDHRViKv+wqE1T9ZEJBHPiOlOu54+9ES9r1C8eYhFq7Yfk9q5Xs6y1kbU6XJHym+VNUPDexzHK17Va5OlFTRSYKOcnQqoebdbDY7s1W3S0UdVqmnKdHo5PA5NFJuLr2Of+5WYSKdQpP3oRPBvDJNjdrqFfLYbhLRvXi2Co7tngR3Snj1NWZTiN/xqREutBJHE5dGTt7qN3gcnDxdJb4NvDnj+HblMx5qZPuy8IAEh0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAP0jHqmqMcqeA/JYju0WOy1GxHGpqi1UU0r6RFc+SBrlVfCqAV4c3J947yDm5PwbvIWr/S7YPxJbvkzPUPpdsH4kt3yZnqAqo5uT8G/yDmpPwb/IWrpj1hT/AOy275Mz1HP0v2L8S275Mz1AVT81J+Df5BzUn4N/kLWPpfsP4lt3yZnqH0v2H8S275Mz1AVT81J+Df5BzUn4N/kLWPpfsP4lt3yZnqH0v2H8S275Mz1AVT81J+Df5BzUn4N/kLWPpfsP4lt3yZnqH0v2H8S275Mz1AVT81J+Df5BzUn4N/kLWPpfsX4lt3yZnqH0vWH8S275Mz1AVT83J+Dd5Bzcn4N/kLV/pdsH4lt3yZnqH0u2D8SW75Mz1AVUc3J+Dd5DhWPRNVa5PEWsfS7YPxJbvkzPUYBvBY5YmbIMimitFDFKykc5r2QNaqKnfRAK5QAAAAAAAAAAAAAAADsW6hq7jVspKGnkqJ5F0axiaqp6uF4tc8puaUlDHyY26LNO5O4jb2r3+xCROI4vacZoG01uhRZF4yzuT65IvfXqTvJwIG7v49SvfvPycM+xXDHfywbB9klLStZWZO9tTP0pSRu7hvvnda95OHfNo0sEFJTMpqSGOngYmjY42o1qeJD9omgMntdQzbM/FPb5eynzbN8k957AAITgHJwAAGir0H65uTTXku08B7FZn2e8S/IGip0opyvA8mOPLzhwAAB+Zo4p6d9PUQxzwSJo+ORqOa5O+in6B7W01nms93sTMTzDWObbJKCva+sxyRlDVdPsaT+if3mr0tXyp4DTF4tdfaK59FcqWSmnZ0temnjTtQlqePlmOWzJrY6iuUSKun1qZqd3Evai/wAC/wBLrVqzFM/ePn+6wwb0x8OT80VgZFnGIXXFK5Iq1nOU0irzFSxO4kT+C94x009bRaOY8LSJiY5gAB69AAAAAAAAAAB+4IZZ5WxQRvlkdwRrG6qviO8lhvjuiz3BfBTP9RnO7LGyTbNZGyNRycpy6KmvUWOpS0ydFPCngYgFVf0vX/8AElx+TP8AUc/S7f8A8SXL5K/1FqiU8Cf7CP8AVQcxB+Bj/VQCqr6Xb/8AiS5fJn+ofS9f/wASXH5M/wBRarzEH4GP9VB7Hp/wEf6qAVVfS9fvxLcfkz/UPpfv34luPyZ/qLVPY1P+Ai/UQexqf8BF+ogFVf0vX78S3H5M/wBR+VsN8Tps9wT/AKZ/qLVvY1N/Z4v1EOFpKVemmhX+4gFU0tpusTVdLbaxjUTVVdC5ET5jpqx6dLXJ4i2Ga2W6ZisloKV7XJoqOiaqL8x5VdhOIV0XNVeNWqZmuvJdTN9QFWQLJ7zsQ2XXWNzZ8QoIlcuvKgasa+VprrKN0vB65ZJLLcLhbJHe1ZyucY3y8fnAg6CQmZ7qec2lj5rJU0t4ibxRrV5EnkU0lkuMZBjVY6kvlpq6CVq6Kksaoi+BehQPIAAAAAAAAO/DZbxMxHw2qukaqaorYHKi/MfXEKaOsy2z0kzUdHPXQRvRetHSNRfOWmUFtoaWjhggo4I2MYjURsaJoiIBVl9L1/8AxJcfkz/UPpdv/wCJLl8lf6i1VKenTogi/VQcxB+Bj/VQCqr6Xr/+JLj8mf6h9L1//Elx+TP9Rar7Hp/wEf6qD2NT/gIv1UAqqXH78nTZbj8mf6jpVdLU0knN1VPLA/TXkyMVq+RS2H2NT/gIv1EISb+VPDBtHtiwxsj5VDqvJTTXugI5gAAAAB9qSkqqyRY6SmmqHomvJjYrl08CHxJGbhkMU20e6c7G1/JoNU5Sa/dIBoVMfvq9FmuPyZ/qOfpev/4kuPyZ/qLVPYtN/Z4v1EOfY1P+Ai/VQCqr6Xb/APiS5fJn+ofS9fvxLcfkz/UWq+x4PwMf6qHHsan/AAEX6iAVV/S9fvxLcfkz/Ufl1hvjUVzrPcGonSq07/UWrexqb+zxfqIfKspKV1JK11PEqKxfuE7AKnnIrXK1yKipwVFOD1cxajctu7WoiIlbMiInv1PKAAAAAAAAA/UMck0jY4o3SPcuiNamqr4j0G2C+uTVtmuCp3qZ/qN77j2ExX3OKrIq6mbLS2yPSPlt1asrvUhN5tLStTRtPCid5iAVV/S9f/xJcfkz/UPpev34luPyZ/qLVfY9P+Aj/VQex4PwEf6qAVVfS7f/AMSXL5K/1HP0u3/8SXL5K/1FqvMQfgY/1UHMQfgY/wBVAKqVx6/p02S4/Jn+o+U9mu8EayT2utiY1NVc+ByInlQtaWngXphj/VQ8+/2K3Xez1duqaOB8dRC6N2sadaaAVTg9nOLFUYzl91sNSxzZKKpfEnKTRVai9yvjTRTxgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAFke7EmmwvFu/RNK3CyTdj47C8V+AtA2SAAAAAAAAAAAAAAAAAABgW8H7jeTfAnmemA7wq6bGsl+BOArQAAAAAAAAAAAAADKdnuGV2WXLksR0FBEutRUKnBE7E7XKdrZxgVblU3sqVzqa2Rv0fMqcXr1tb2r3+okLarfR2q3xW+3wNgpoU0axvnXtUq+odSrrR6a97fT+qLs7MYo4jy6+O2S3WC2st9rg5qFvFVXi569auXrU9EAx+TJfLabXnmZUt7zeeZADq3W40Npt8lwudVHS00fS969K9iJ1r3kPmmO2S0VpHMy8rWbTxDtdeh0b1eLVZYFnutfBSNTqe7ul8DU4qaizTa9WVPOUmNRrRwrq1ap6ayuT83qb5zXtBQ33JblzdLBWXKrkXVVTV6+FV6vGaHV6F29Waf7LLFoe+SW3sh2yWmmV0dkt01c9OHOzrzbPEiaqvzGG1+1zLKhzlgdR0rV6EjhRdPGup7Vg2K10jWSX66w0aL7aGnTnZE8ftfOZxZtmWFW1rVfa5LjK3ReXVzOVFX3rVRPEupLtk6dqduImfzdptrYuzSFZnGXVjORNfq3k666MfyfNoec+9X10nOOutxV/as79fOSmobbaKB/Lt9ktdG/TRHQ0jGuROzXTU7vOL97H+zb6jn/vmvXtWs/o+P8AX448QitS5blFH/Q324s17ZnL5z16HafmVLwW5pOn/jRNd/AkVWUtBWta2vtlvrEb0c/Sxv08GqGOXbZ7hNzR3OWBlHI7/aUUrolT+7qrfmPqvV9TJ2vX84fUbmG3mGCWDbQ/lpHfbS1zFTRZaR2jvG1eC+U2Lj+X45ftG226RulX/Yy9xJ5F6fEa8v2xR6IslgvLJeyKrbyF/WTVDWl/sF8xutSC60NRRypxY9U7l3fa5OC+JT7to6O3HOPz+H7PZ18GaPh/RKtdUXRU0U5I+4btSvlmWOmuTludEnDSVfrjE/Nd0+JTduM5Basjt3s21VKStRdJI3cHxr2OT+JR7nS8ut8XmvzQM+pfF38w9QAFYiupebZQXm3SW+5Uzaimk6Wr0ovai9Sp2ke9peC1mKVyzQ8uotUq/WZ9OLfzXdi+ckefKtpqatpJaOsgZPTzN5Mkb01Ryf8A51lr07qV9afTbvX6JettTinifCIYM82mbParGXrcKHl1Npe7RH9LoV6mv/gpgZsMeSuSsWrPMSua2i0cwAA+30AAAAAAAA2Zuxe7RZPfO9EshK3d2L3aLH793mLIgAAAAAAAAAAAAAAAAB5WRY7ZMhoX0V5ttNWwPTRWyxop6oAipth3VKKeCe6YDULBUIiu9gTLqx3ea7qInZHYrtjt2mtd6oJqKshXR8cjdF8XanfLXDX+2PZTjW0qyvprnTNhr2NVaasjTSSN3Vx607ygVpAy3algF/2eZJJZ75TqmuroJ2p3Eze1F/gYkAAAHt4B9veP/GdN+9aWop0IVW4IumcWFey5U/71pajH/Rt8CAfoAAAAAIQ7/C/6yLUn/Af/ACUm8Qe39l12mWxOygT0lAjmAAAAAEj9wb3Sbr8Xr6SEcCR24Ov+su6fF6+kgE3wAAAAA+VX9iy+8XzH1PjW8KOZf/Dd5gKrsz+268fDZvTU8k9XMeOWXdf+Nl9NTygAAAAAAERVXROKgyfZVj8mUbQrLZI2q5KiqYj9OpqLqvzATo3SsSdiux+3LUQ83V3DWql1TRdHe1RfFobePjRU7KWjhpok0ZFG1jU7yJofYAAAAAAAACEW/div0NzuhyWCFGw3KDkSuTrkb2+IjeWF73OGrlmyStlp41fWWxfZcKJpqvJ9snk1K9AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAFkW7CipsLxbX+xNK3Ta+I7wO0XF8dorDaq6lZRUcaRwtdToqo1O+BYsCACb0O1X8YUPyVpz9VFtV/t9D8laBP4EAfqo9qv9vofkrT8rvQ7Vl/+40Sf9K0CwAFfv1T+1f8AGdH8lac/VP7V/wAZ0fyVoFgIK/F3ntq/40pPkrTZe73tS2x7Ss2ho1r6ZtppFSSvm9itROR96i9qgS4BwcgAAAAAA0NvibRbbjWBVGMNVs1zu0asbEi+0Z1uU2btVzi1bP8AD6u/3SRv1tqpDFykR0r+pqdpW/tCy665tlVXkF3kV01Q7VrNdWxt6mp3kAx4AAAAAAAAAADYWzHZ3U3+SO53Vj4LUnFvU6dexOxvap6Oy3ZotxjhveQMcyjVeVDTLwdMidbuxvnN2xsZHG2ONjWMaiI1rU0REToREKbqXVI14+zx97fRC2duMfw18vnRU1PR0kVJSwshgibyY42JojUPqAZK1pvPqt5U8zNp5nyHJwp4GbZZbMTt3P1sjZKp7VWnpWu7uRe1exvfXxHTBgvnvFKRzL6x47ZJ9NX2y7JrXi1sWtuUmr3IvMU7fbyu7E7E7VI65flF3yu5JPXyK5qLpDTs9oxF6kTt75+Lzc71mORJNPzlVWVD0ZDDGnBqdTWp1Ibn2cbNqPHHx3K6LHWXVG6tboixwL3vvnd/q6jV48ev0zF6ref1n+i3rTHq05nyxPZ9somrEZcMn52lp9UWOkbwkk98v3KfP4DclqoaK1USUVspIqSnb9xEmmq9qr0qvhPvqqqqr0rxXUGf3OpZdmeOeK/JXZtq+WflAACuRgHIA4AADq0PxWQU9bSOpK6niqqd3TFK3lN+foP2cn3S9sc81niXtbTWeY7NT55sjgmjkr8UdzcqaufQyO4L7xy+ZfKausV2vGJX72RT8unqoXciWGRFRHJ1tcnYSoMT2i4PQZfS883kUt2jbpFUaaJInU1/anYvSngNHodX9f8ADz/n+6y192LfBkdvBsut2WW72TSpzFTGic/TOdqrF7U7U75kJFf/ALdw7JFRUmoLjSSaKnb/AAc1fIqG/dnmbW/LKPkIrYLnG3WanXhr+cztTvdRH6n0r7P+Lh8e8fJ8bWp6fjp4ZWcnAKGFe/M8UU8EkE8bJYZGq17Hpq1yL1KaL2pbN5bO6S8WNjpravdSxdL6f1t7/Ub2C6dCoiovBUVNUVCw0OoX1LfOvvCRr7NsM/gh8Dbe1TZotOk18xyBVgTV9RSM4rH2uYnW3vdRqRUVF0Xgps8OamekXpPMLul63r6qgAOr7AAAAAGy92L3aLH793mLIit7dh92iye+f6JZCAAAAAADjVO1A5dGqveK9ttO0rO6DapkdDQ5Rc6elhrnsiiZMqNY3sQCwjlN++TynOqdqFYn86W0PXX6b7t8oU+jNq+0Zi6tzC7ft1As4BXdad4zavQLGn0wJUMYiJyZoGO18K6amyMM3vLzDOyPKrHT1MP3UtL3Dk8SgTJBg+zTanh20ClR9hukbqhE1fSyLyZWf3V/gZwAAAAAAYJtq2c2raNiM9rrIY0rGNV1JUK3uon9XHsK4sux+54vkNXY7tA6CrpZFY5qp09ip3lLWCMO/Bs2bcrCzPLZT61dCiMrUYntotfbL4PMBDAAAexg/wButi+Maf8AeNLUov6JnvUKrcIXTNLGv/Maf940tSg4ws96nmA/YAAAAAQd39PdNt3wFPSUnEQc38/dOt3wFPSUCOoAAAAASM3CF/1m3P4vX0kI5ki9wr3Trl8AX0kAnEAAAAAHwuH2DP8Ao3eY+517h9gVH6N3mAquy77art8Ml9NTyz1Mt+2m6/DJfTU8sAAAAAAEodwnD0rMhumYVMesdExKanVU4K93Fyp4E08pF5EVVRE6VLJN23EEw3ZLabfJHyKqeP2TUcNF5b+Oi+BNEA2SAAAAAAAAAAOtc6SKvt1RRTNR0c8To3IvWipoVf7TsflxfPLxY5Wq32NUvazVOluuqfMWkkLN/PEUoMrtuV08ekdfGsMyp0ctvR5UX5gIygAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAActarnI1qKrlXRETrA9PFLDcsmv9JZLTA6erqpEYxqdXfXvIWQbFtnlr2cYbT2ihjRapzUfWTrxdLJpx49nYa03QdkSYhYG5Xe4E+jNxjRY43tTWnj6vGvSpIQAAAAAAHVutfSWu3T3CumZBTQMV8kjl0RqIdleCaqQ43ytsLrhVyYFj1UqUsS/9oTRu/pHfeap1J1gau3kdqlXtJzORYJHMsdC5Y6GLo5XbIvfXzaGqwAAAAAAAAfuGKSeVkMMbpJHryWtamqqvYiAflqK5yNaiqqroiJ1m4tl2zJ0b4b1kcTehHwUbk1X3z/Uejsq2cttCxXm/QNfcE7qGnforYexV/O83hNneFSg6l1aMfOLDPf3n5K/a3PT8FPJ2d5NAcAy8zM95VM/iAGEbSdoFHi0T6GiVtReHN4N01bBr0K7v97ykjV1cmzf00j/Driw2y24q720TNKLEbeqaMqbnKn1imVeCfnP06u91mg1+j2a5Kq/XK2vqXeJqeZGocWe23zM8j5iFZausqHq+WaRVVGJ1ucvUif5ISGwjErXidvWGiRZamVE5+pcndPVOpOxveNPM6/S8PHm0/nP+FpM49Sn4urs/wi3YlRtciMqbo9v16qVPa69LWdidWvSplXWAZbZ2cmxf13lVZctstubABynFTjEOflwde4VtHb4OfrquGlj++lejUNfbSNp0FlmfbLFzdTXNTSSZe6jhXsTqcvzGlbvdbneKx1TcayaqmevFXuVfEidRe6fRLZI9WaeI+XusMOjNo5v2b+ue1HEKFXNZWS1jkTogiVU18K6HkP2zY+j9G2y4ub29wnzamsse2f5ffWNloLJU8wq6c9MiRRp/edoZF/Mnmv8Ayr5fH6y0/wBr0qdrR+cpcauCvmGfW7ariFW5GyVFTSKv4aLh5U1MrtF3td3i522V1PVN01Xm36qnhTpQ0DetlmcWpj5JLLJVQsbynSUj2zN0/uqqmJ01RXWyrSSCWekqI16WqrHNU55Oi62SOcc8fq+LaOK0fD2S5ODUWzjams0sdryaREc7RkVZp0r0aP8AWbdRUVEVFRUVNUVF6TO7mjk1bcX8e0q3Ngvhnu5OFAIbi8HOMUtuXWz2LW/WqqNF9jVaJ3Ua9i9rV608hHm62++4VkaRSrJSVtO7lRSxrwenU5q9aKSlPIy3HbZlFpW33ONNUReYqETu4XdqL2dqF703qk4uMWX7vz+Sfq7fp+G/hj2zLPabJ6RtHWuigu8fBWa6JOn3zU7e1DOCLmV49ecNvqQVXKjex3LpqmJV5MidTmr/AA6jbmzHaRDfXRWm8K2G5cnSObXRs6p5nec7dR6VFo+2wfl+zps6nPx42xjkHBnFY5RVRdUXRTVW1TZo2u52947EjKlNXVFI1OEn5zO/2obVCKqKiouikzT3cmrf1V8e8O+DPbFPMIfPa5j1Y9qtc1dFRU0VFOCQG0/ZzBkTX3SzRxU91RNXx8GsqP4I7v8AWaEq6aekqpKWqifDNE5WvY9NFaqdSm01trHs09dJXeLLXLX1VfIAEh0AABs3dg92ey+F/mLICt/dg92ezeF/olkAAAAAABw72q+ArI268dsGUfGEnnLNn+0XwFY+3FddruUfGMnnAwwAAAAB6GPXq52C7QXW0VclLVwORzJGLp/+IWIbu202m2k4RFVyK1l0pUSKti1+6T7pO8vSVvm3d1HOJMO2rUEMsqtoLo9KWduvDV3Bi+Xh4wLEAcIqKiKnQpyAAAA6d6t9NdrRV2ysjbJT1ULopGqmqKipodwAVabSsbmxHOrvj0zVRaOpcxmvWzpavkVDHSSu/pirLfmtsyiBnJZcYOZmVOt7OhfIvzEagPWwv7cbL8YQfvGlqdP9jx+8TzFVmGfbhZfjCD940tTpvsaL3ieYD6AAAAABB3fz9023fAf/AJKTiIPb+qf6y7b8B/8AkoEcwAAAAAkZuFe6Zcl/4D/5IRzJGbhXumXL4B/8kAnCAAAAAHXuX9X1H6N3mOwda5/1dU/oneYCq7LPtouvwyX01PMPSyr7Z7p8Ll9NTzQAAAAADO9geILm21SzWV7eVSpMk9Vx/wBkxdVTx8E8ZZhExscbY2Jo1qIiJ2IRL3BcQVG3bM6iNUR3+iUyq3pROLlRfm8RLYAAAB06q50NNcKagnqY46mq5XMRuXRX8np0O4Q1247U5qPebs0tPVq2gscrYJNOKd2v1zh0dCgTKB8aOeOqpIamJyOjlYj2qnWipqh9gAAAGqN6rE1yvZBc44Y0fVUKeyoe51XuelE8Wptc+NZBHVUktNK1HRysVjkXoVFTQCphU0XRekGWbXsZfh+0i92ByKjKapdzSr1xu7pvzKhiYAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAkTuf7Hlyu9szHIKbWy0L9aaJ7eFTKnQvvW/OvgNabDNnNdtIzantMLHsoIlSStnRODI9ejXtXoQsdxqy2/HbFSWW1U7YKOkiSKJjepEQD0WNaxqNaiI1E0RE6jkAAAAABhu2DPbXs7wyqv1xeivRORTwp7aWRehEQDAN6ra7HgeNOs9nqY1v1cxWsRF1WBi9L1TzEB6meWpqJKieR0ksjlc97l1Vyr0qerm2SXLLcnrsguszpKqrlV66rwanU1O8icDxgAAAAAAAd2y2uuvNyit9ugdPUSro1qdXfXsTvgfCipaitq46SkhfNPK5GsYxNVcpv7Zls9psbjZcrijai7OTh1sp0Xqb2u7/kO9s5wajxOl5+RWVF0kbpJOicGIvS1ne7V6zL0Mz1Pq3POLDP9Z/ZV7W5/JQABnZVocpqq6IgRFVUREVVNVbWtoyUCzWCwStdUqisqatq8I+1jO/2r5CZp6WTavxXx7y7YMFs1uI8O7tR2jw2SF9qsU8c1zdwkmavKbTp3upX+Y1PieOXnM72+OBznKruXU1UqqrWIvWq9a97rO1s7wqvy64quroLfEqLUVLk6PzW9rlJD2Gz22xW1lutVM2Cnauq9bnu63OXrU0ebPh6Zi+zp3t/75Wd8lNWnpr5dfFMetuM2ltvtsenBOdmcndzO7V/gnUesAZXNmvmvN7zzMqi95yTzYOTg4e5jGOfI9rGNTVz3LojU7VU51ibTxHl5EczxDl3BFVVRERNVVehE7TTu1TaVqktjxyXueLaisavT+azvd/yHn7V9oyXdklksUj20HK0mn6Fn0XoTsb5zWBq+m9KjDxky/e+XyW+rqej4reXesVrrb5eKe2UMfOVNQ/ktRV8qqvYnSSGwnZ9YcYp2SPhjuVzTi+qmbq1i9jGr0eFdVXvEdLVX1druENfQzOhqIXcpj29RIXZrnVLlFHzFS+OC6Rp3cWuiSJ9831Ejqs7MYucPj3+bptzkinwM2klllXWSRzvCp+TgGNtaZnvKkmZny/Ub3xuRzHuaqdaLoeLl2K2HK6V0V1pWxVOi83WwsRsrHdq/fJ3lPYMX2gZnb8So0WREqK6ZPrNO1eOn3zuxCZoW2PtYjBPd315y+qIo0Fm+M1uJ319rrJYplRqPjliXuXtXoXtTwKZnsq2iy26WOz32dX0K9zDO7phXsVetvmNf5Bd66+XWa5XCVZJ5V17zU6kROpEOgbTJhrmx+jJHK7tSL19NkvoZY5oWTQyMkjenKa9i6oqdqKh+yPuzDaFPjr2Wy5cqe1udw63QKvWne7xv6kqIKumjqaaZs0MrUcx7V1RyGO3+nX1bc+a/P91Lsa1sU8x4fQAFcjPPyKyW3IbU+2XWDnYXcWPT28Tvvmr1KR1zzELnh90ayZVkpXu1pqpnBHonmcnYSbOrdbfQ3a3TW65UzKilmTRzHdS9SovUqdpcdO6pbXn0ZO9fom621OOeLeGsNlm0mKohisuRVHJqGqjaeqf0PT7169S982wRy2k4FXYpVeyYVdVWuVy81OicWfmv7F85lGyfaQsSwWHIZ05lERlPVvXizsa9ezv9RYdQ6bXYr9tg8/VJ2NWMsevG3McBFRURUVHIqaoqLqigzExMTxKpnmPIYZtKwOjyqmdV06Mp7uxvcS9DZU+9f/BTNDjThod9bavrX9dHTFltin1VRIutvrbVXy0Nwp309REujmPTRUOqSZ2g4XQZbRIj1bT3CJNIKlE6vvXdqeYjrkFnr7FdJrbcoHQzxLouvQ5OpUXrRTaae7Tap6q+feF5hz1zV5h0AATHZs3dg92ey+F/mLICt7dh92iye+d6JZCAAAAAAcO9qvgKxtuSabXsoT/mMnnLOXe1XwFZG3X3YMo+MZPOBhQAAAAAfWinkpayCqicrZIZGyNVOpUXVD5AC0/Z7dvo7hFmu2mi1VHHIvHXirU1PeNUbple6v2FWB73ve6Fj4VV35rlQ2uAAAAAAR936bO2u2TwXFGayUNY1yLp0I7gpBQsi3m7dHctimQxPRV5unWVNO1vErdA9bDOOYWX4wg/eNLU6b7Gi94nmKrcITXNLGn/ADGn/eNLUqf7Hj96nmA+gAAAAAQs357TdbhtHtslBbK2qY2h0V0MDnoi8pexCaZ83wxPfy3xMc7tVqKoFVyYtk69GOXdf+ik9Rz9KmUfk3ePkUnqLUUijTojZ5DnkM+8b5AKrfpUyn8m7x8ik9Q+lPKfybvHyGT1FqXIZ943yDkM+9b5AKrFxXJ06ccvHyKT1EhNxey3e37RrnNcLXXUka0GiOnp3MRV5ScEVUJnLHGvSxvkOWsY1dWsa1e8gH6AAAAADrXT+ran9E7zHZOtdP6tqf0TvMBVblX2zXT4XL6anmnpZV9s10+Fy+mp5oAAAD9wRvmmZDGmr3uRrU7VVdEPwbO3YsVdle1+00zo+XT0jvZU/Dho3o+fQCdGw3Fo8P2X2SytaiSMp2yTKidMju6d86mbnDWo1qNRNERNEOQAAA8bN71BjuJXO9VDmtZSUz5e6XRFVE4IVd5HdKi9X6tu1S5Vmq53Su1XXpXUmjv0Zelq2f0+MwPVKi6yor9F6I2rqvl6CDwFj27Blbcs2P2epc9XVNJH7FqNV48pnDXxpops8h3uBZOkN3vWKTzojZ2JVQMVelycHaeLQmIAAAAAAQ53+sSbTXm0ZhTxaNqmrS1LkThym8Wqvi18hFgsi3l8STMNkF5oI4kfVU8Xsqm4aqj2d1w8KIqeMrdcitVUVFRU4KigAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA7+PWiuv16pbRbYHz1dVIkcbGprqqnQJs7nWx+OwWeHOb/Sf9q1jNaOORvGCNfutF6HKnzAbT2E7N7ds4wqntsMbHXCVqSVs+ndPeqdHgQ2CAAAAAA4cqNRVVdETiqgdW8XGjtFsqLlcJ2QUtPGskj3LojWomqldW8PtQrdpeYvqGudHaaRVjoodeGn3699TZ++JtkS910mD43VqtBTu0r5mLwlen3CL2J1kYwAAAAAAAAPRx6y3K/XJlBbKZ00rl46e1Yna5epCRmA4hb8TtiRQtbLXSNT2RUqnFy9idjTzdjUNjhw+F9oVj55E1q3L/SLJ2L2InUhmyGX6v1G82nBTtHv+P8AhVbmzPM447AOTgz6uDlOOvFERE1VVXRETtU/E0sUET5p5WQxMTV73uRrWp2qq9BozadtLnuyz2exPdDbV7iSZOD6j1N73X1lhodPvtW+VY8yk6+tbNP4PS2q7SucSSyY1UuSPVUqati6K781i9navWY5sz2fVeUz+z7gstJaGL3U2ndzL96zX516EO1st2dVF+dFeLsxYbS1+rWrqj6jTqT83tXyG+oo2QwRU8EbIoYmoyONiaNY1OpC8293F0/H9jh8/T+qfmzV16+inl8LVb6K1UEVBbqdtPTQt0YxvnVete1TtHByZbJktktNrTzMqi1ptPMuAcnTu9yorTbpa+4Ttgp4k1c53X3kTrXvDHS2S0VrHeStZtPEPvVVEFJSy1VVKyGCJvKkkeuiNQ0NtO2jVN+fLa7U5ae1IujlTg+dU617E7x520XPa7KpfYsTXUtsjfqyFF4v7HP7V8xhhrundMrrR6797fRc62rGL4p8gALdMD60dTPR1MdTTSuimjdymPauiop8gBv/AGXbQosgYy2XRzIro1NGu6Gz+Dsd3jYXQRHtcFbVXGnp7dHLJVySIkLYvbK7XhoSjs1rv8WGNt9dc4UyJYHI2Z3dRsk+5a5U6V7XdGpm+p9NxxaL0mK8zx/lWbWrXmLRPHLwNo+eUWMU0lLTKye6ub3EfSkev3TvUR8utwrLpXS11fO+eeVdXPcp2Mnobvbr9V0l9jmjuDJF5/nV1VXduvWi9Op5pc6mnj1aemn5p2HDXFXioACW6hmezfOqzF61IKh0lRa5F0kh14s/Ob2L3uswwHxkx1yVmto5iXlqxaOJS3tFxortQR19vnbPTyJq1zfN3lO2ad3d7bkD6yeuWZafHmIvP84mqTSacGxp992r0InT1G4fB0GL6lp11cvFZ5if0Ue1hjFfiJAAVyM/FTBBVU0tJVwsnppm8mSJ6atcnfNE7VdnEthkku9jZJPaHcXs6X0y9i9rexfKb5PlVVVPSUss1W+JlOjF53ndOQrdOKLr0p3iz6bv5Ne0V8xPt+yVq7FsdvT5ho7ZZtHfZWss97V81vVUSKXXV0Hrb3uo3nBLFPBHNDI2SKRqOY9q6o5F6FQinlEltlyCtks8SxULpnLCxepuvV3jbe7pFktVBW88rm47TxuXnZ9dGy6cGR9qr1onDrLvqnTqZqzlr2tH6p21rRePXHaW1TgAyKnDwszxa2ZVa1pK5qRztRfY9S1urol/i3tQ90L0HXDnvgvF6TxL7x5LY7eqqKmVY7c8bubqG5wKx3TG9PayN7Wr1nkkmdp9PY6rEalL85jI4kV1PJ92yTq5Pbr2EZ3acpeSuqa8Db6O1/qsUZOOF7gy/a09XDZe7D7tFk9870SyEre3YvdosnvneiWQkx2AAAAAHDvar4Csjbt7sGU/GMnnLN3e1XwFZO3f3Ysp+MZAMJAAAAAAABP7cmqFn2GUbFbpzVXO1O/3Ru80VuPJ/qPg+HT+kb1AAAAAAMO2100tXsqyOCFvKkdQSclO3uSsNektK2mTxU2z++TTvRkbKKVXOVeCdypVs72y+ED2MF+3axfGVP8AvGlqEH9BH71PMVX4L9u9i+Mqf940tQg/oI/ep5gPoAAAAAAAAAAAAAAAAAAAAAHWuv8AVtT+id5jsnWuv9WVP6J3mAqtyr7Zrp8Ll9NTzT0sq+2a6fC5fTU80AAABNHcMxJlHilxyueHSetl5mFypx5tvTp4yGtupZa6vp6OBqulnkbGxE61VdCz/ZbjcGJYFabFTsRqU9O1H6db1TVV8oGTgAAcKqIiqvQhyYnteyaLEdnV5vsj0a6CmdzWvW9U0anlAg5va5g3LNrtcynl5dHbE9iRaLw5Tfbr5eHiNQn1rKiWrrJqqd6vlmkdI9yrxVyrqqnyAzPYjlE2H7UbHeo5ORG2qbFUa9CxPXku18S6+Is3hkbLEyVi6te1HIveUqWaqtcjkXRUXVCyPdryh2WbIbNcJXq6ohj9jzKq6qrmcNQNkgAAAAPxPG2aF8T0RWvarVTvKVm7dsafim1S+WlYljiSoWWFOpWO4oWakRt/nD2NdaczpokR7v8ARKlyJ0p0tVfnAiWAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAABmex3ALptEzKlslBG5sHKR9VPp3MUadK+HsA2Ruk7IHZtkTckvdOv0Bt8iOa1ycKmVOKN8CdZPBjWsY1jGo1rU0RE6EQ8jC8ctuJ41RWG0wNhpaWNGNRE6V61XvqeyAAAAAADQm9ltfjwrHX45Zptb5cI1bymO407F6XL3+w2Ltm2gW3Z1hVVe617XVHJVlJBrxlkVOCeDtK4MzyS6ZZklZfrxOs1XVPVzl6mp1NTvIB5M0kk0r5ZXue96q5znLqqqvWfkAAAAAAAAAD1cYv8Acsducddbpla5q93GvtJE60chIzCsrtuUW9J6N/Inaic7A5e6Yv8AFO+RePQx+8V9juUdwt8yxzM8jk7FTrQgb2hj269+1vaUfPr1zR38pY6nyramnoqOasq52QU8LeVJI9dEahheNbS8errGtbcqtlFPC369CqauVdPuU69TU+0PN7hltdzMXOwW1jvrFKi9K9TnadLvN1FDq9GyXyTGXtEfqr8OlebcX8Q7e07P6jJploKFHU9qjfq1q+2mX753e7EPd2V7M1rGxXzJIVbSKnKp6ReDpuxzuxve6z0Nl+zGOGKnvWSw8qZyo+ChcnBqdTpE7evk+XsNtLqvT0k7e6jTVp9hr+fp/lI2NmMUejH5fljWMY2ONjWMa1Gta1NEaidCIh+gcGXmZtPM+VVM8yA5PHyzIrdjNpfX3GTRVTSKFF7uV3YifxOmHDfNaK0jmZe0pa9vTV9sjvVvsFrkuNznSKFvBqdLnr1NanWpHbP8xuGV3JZJXOioo1/0enReDU7V7XL2nVzLKLplVy9k10n1tiqkEDfaxoq9CJ2982Lso2aIrYb/AJND3C6PpaJycX9jpE6k7E6zW6uph6fjm+Se/wA/+IXGLDTWp6rPM2X7M5bqsF6yKKSG18HxQa8l9T/Fre/19XaeptY2UxUtJLkeHxvfQsTlVdBqrn035zVXi5nzobgkcr3au8nUh9KSompZ2ywv5Lk4eFOtF7UKz/fb/berj4PkjRvz9pz7IcA37tX2WxXxk+RYjTxxVrUWSttzO5R/DVXxJ52+Q0G9jo3uY9qtc1dFRU0VFNLizUzUi9J5hZ0vW8c1cHdslqr71cobdbKaSpqZV0axia+PvIfTGrNWZBfKW0UDWrUVD+S1XLo1qdKqq9iIiqSTwbEbbhlvdT0L0qK6ZNKms00V/wCa3sb5zht7mPVp6refaHPNmrirzLqbO8Ht+HUbZO4qbxKzSep6Wx69LI+xO1evwGUKnE5Q4MXtbWTZvNrypMuW2W3qs8nMsXtWZWlaG5o2GrY3/RK5G93E7sd98xezq6iNeW47dMYvElsukCxyN4sentZG9Tmr1opKo6GTWG0ZTaUtl6iVzWa+x6lv9JTqvWnanahbdM6r9nxizT2+aZq7fHw38Img9zOMZrsTyGa0VzmSK3R8U0a6tljXocnqPDRFVdE4qaiJiY5hahs7ZJszkvvNX6/sfT2RjtY417l9YqdTfze13kPV2T7K2TQRZDl9O9lI7R1JQrwdUfnO60Z86m4ZZOXyEaxkUcbUZHGxNGxtToaidiFT1HqddaPRTvb6ImztRijivkV0bYo6enp4qamibyYYIm8lkbexE/j1n4ByZG97ZLTa095UtrTeeZcAKebkV7t9htj6+4zJHE1OCfdPXsROtRjx2yWitY5krWbT6Y8uxdrjR2qgkrq+dsFPGmrnu83hI/7S88qsmqXUdIroLXGvcs14yr9871Hn59mVfldaiy6w0cS/WYEXgnfXtUz3Y7srSqgiyjL6dY7ZpyqSjfwfVr1KqdKM7/Wa3S0MejT7TLPxfP5LnBr1wV9VvLy9kOy9+RNjv2QSPo7Ex/ctRPrlWqfcs73a43xNOxYIqSlhZS0UDeTT08SaNjb2J3+/0qpxUzrMkbGxxwwRMSOGGJOSyNqdDWp1IfEpeo9Ttsz6a9q/VB2dqcs+mvgABUoYvBDycov9ux62urbjMjG6dwxF7p69iIdXN8rt+L2x1RUvR87k0hgRe6evqI65XkNxyS6OrrhKrl6I40XuY07ELnpnS5z/AMTJ2r9U3V1JyfFfw7WcZZcMpuSz1LlZTMcvMQIvBifxXvmPAGtrWKR6axxC4iIiOIbL3Yvdosfv3eYsiK3d2P3aLH793mLIj6egAAAADh3tV8BWTt592PKfjGQs2d7VfAVlbevdkyn4xkAwgAAAAAACJqqInWBYTubwLBsHtOsfI5yWV/R06u6TchguwO1JZdkGN0KNcitomPcjunVyar5zOgAAAAADBtvkrYtj+SucuiewZE+YrLLB98i7LbNidwja9WPq5GQJoumuq8SvgD2MG+3axfGVP+8aWowf0EfvU8xVbhPDM7Iv/Maf940tSp/seP3ieYD6AAAAAABEze62nZphe0GkoMcvD6OnkpEe5iNRUVde+BLMFc6bwu1dP95Xfsmn7TeI2sJ/vJ/7DfUBYqCutN4rayn+8n/sN9QXeK2s/lJ/7DfUBYoCul28PtYd05K7xQt9RuXdB2mZnmefV1Dkd4krII6NZGNVqJo7lInUBLEAAAAAOrdv6rqv0TvMdo6l5/qmr/Qu8wFV2UfbLc/hcvpKecehk3HI7kv/ABUnpKeeAAAG5tz3EUyba9R1VREr6S1NWqk7OUntU8vHxFgxHfcZw91m2d1GSVUXIqLvLrHqnHmm8E8q6r5CRIAAACLe/wA5Z7Gx60YjA9vLrJVqKhOtGM00TxqvzKSkPFveK45e6ltTd7JQV0zW8lr54GvVE7OKAVWAtE/m7wb8lLP8kZ6j9Js/whOjFbP8kZ6gKuSW+4FlDlbe8Tnl7lvJq4GL3+Dv4EkFwDCfyVs/yNnqO5ZsUxuzVS1dqsdBRTq3krJBA1jtOzVEA9oAAAAAMA3gMXbluyq82zmkkmbAssPee3ihn5+ZGNkjcxyatcmioBUrIx0cjmPRUc1VRUXqU4Nj7yGJLh+1q70EcaspZ5VqKfVOHJfx4ePU1wAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB27Pbqy73OnttvgfUVVRIkcUbE1VyqWLbvOzGk2a4VDRvRsl1qUSWtm048pU9qneToNT7mmx5bZRx59kVMxauob/2fBIzuoW/hF1616u8SkAAAAAAB0Mgu1BYrPU3a51DKekpo1kke9dEREO69zWMV71RrUTVVXqIQb322CTJbxLhtjqNLTRyaVMkbuE8idWvYgGu94HafX7S8zlrFe+O00yrHQ0+vBG/fL+cvqNbgAAAAAAAAAAAAAAA2ZsDorHUXueor3MkuMKItJDJ0J2vTtVOrs6TWZ9qGqqKKrjqqWV0U0buUx7V0VFOeWk5KTWJ45fN6+qsxEpeaqqq5VVVXpUGvNl20KLII0tt1fHDc2+1VE0bMne7/AHjYSdBhNrVya2T03/8AqgzYrY7cWADEtoOd27FaV8CK2ouj2axU6dDexz16k73Sp86+tk2L+ikPMWK2S3FXbznMLbidEktV9eqpE1hpmro5/fXsTvkdsqyC45JdpLjcZeU93BjE9rG3qa1Ow+VwrLrkV5dUVDpq2uqX6IiJqqr1IiJ1d5Dalm2HVNfissj7lzORqznYaJze4cn4NXff/N1Gx1dXDpViJnvPv811hw0wRx7uhsCx6z1tVUXq4LHU1NK9G09M7RUYv4RydenV1am7ZHuker3uVzlXiqkVLbW3nEr+skaS0dbTvVksT0VNdF4tcnYb/wABzW25VS8mNUgro2ostO/p8Le1NSt63rZr/wASJ5rHt8kXexXt8UeGVg48YMyq+X7hlkhmbLE9WPauqKi9Bg21fZxSZZTy3vHaZtNfmJyqilZwZWJ1uanU/vdCmbnLHPY9HscrXNXVFTgqKTdLeyatuY8e8O+vnnDPbwiXQVdyx+8tqIecpK2leqKjm6OaqcFRUXyaG7sJ2oWq6xR013c2hrdNFc7hG9e1F6jJ89wWw53rU1UqWq98nRK1rNWTL1c6icdfzk4+E0PmWzvK8Wmd9ELZJJTIujKuBOcikTtRyfxNTMavUsfz+sLbjFs1SShmimjbLDIyRjk1RzV1RUP3qRPtt7vNrd/oNxqqbT7lkionkPabtEzBrUal5l4JpqrW6+YrL9Ann4L9vxhEt07v8NklVVETlKqIidamHZjtDsdhhfHHOyurNF5MMTtURfzlToNEXPJ8gubVZW3armavS1ZFRPIh9caxDJMkqWw2i0VVSrul/I0Y3vq5eCEjX6Hixz6sk8/R1x6FKzzaeXwye+XDJby+4Vyo6V/csYxODG9TUNwbJdmENsjgyPKo0fVKnOUltc3Xk9j5eztRvlPb2c7NbZhskV0ucsNzvrU1axreVBSO7UX7tydvQnVqZnLI+WR0kjlc9yqrlXpVTzqHVa4Y+ywefobG3GOPTTy5qJZJ5lllernL2n4BwZaZm08yqZmZnmQeA5MdzbLLdi9udNVSNfUub9Zp0Xunr/BO+dMOC+a8UpHMvaY7ZLemr75bkluxq2rWXCTiq6RxN9s9exCO+aZPccpuzqqqcrYkXkwQNXuWJ6++fDIb1dcpvHsiqc6aaR3JiiYnBuvQ1qG79lezGmxZkF/yVsNTeFRH0tAqcptL1o+TXgruxvV1mt1tXD07F67z395/4hc4sVNanMvN2SbK6Wip6fJs0pnue7SShtbuHLTTg+Xr5PQqN4a+A2lV1EtTMssrtV6ERE0RqdSInUhxUzzVU7p6iV8kjl1Vzl1VT5Gd3+oX27ceK/JW7GzbNP4AByV6MGKbQM0ocWo9HaTVsjV5qFF+dexDpbR89ocapn0tK9lRc3IqNjRdUj77vUR/u1wrLpXy11dM6aeVdXOcpoOmdJ9fGXNHb2hY6up6vivHZ9b/AHivvdxkrrhM6WV666dTU7ETqQ6ABp4iIjiFr4AAejZe7F7tFj9+7zFkRW7ux+7RY/fu8xZEAAAAAAcL0KVl7ffdlyr4xkLNHe1XwFZe3z3Zcq+MZAMHAAAAADJdl+OVOWZ/ZrDTRuf7JqmJLyU9rGior18mpjRL7cb2bT0jJ8+utOsbpWLDQtcnHk/dO8YEqaGnZSUcNNGmjImIxqd5E0PsAAAAAA/Mj2xxue9URrU1VV6kAih/KAZE1tHYcYjevKe91TKiL1JwTVPCpEM2RvJ5YmYbXbvcIZVkpYH+xaddeHJZw4ePU1uB6+FfblZPjCn/AHjS1Km+x4/eJ5iq7CE1zSxp/wAxp/3jS1Gn+x4/eJ5gPoAAAAAEGd/Bf9aVEn/Ap51JzEFt+1ddq1InZQt86gR7AAAAACRW4Wum064p20C+khHUkRuHLptSru/QL6SATlAAAAADp3r+qKz9C/zKdw6d6/qir/Qv8ygVW5J9sNx+FSekp0Dv5J9sNx+FSekp0AB6OMWmovuRW+zUrdZqyoZCxPfLoecb73IsTkvm1N17lhV1JaIFfyurnXcGp5OUBN7E7RT2HGrdZqVjWQ0dOyFqJ0cERD1AAAAAAAAAAAAAAAAAAAAAilv84m+W32nL6eLVIHLTVConQi8WqvjIflnm2jGWZdszvdjc3V81M50Xee3um/OiFY9TDJT1EkErVbJG5WuavSiouigfMAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAN3bqeyN2f5Ql4u8LvoDbno6RFTRKh/UzXs6FU15sqwi55/mVHj9ua5Ekcizy6apFH1uUskwHFLVheLUeP2eFI6amYjddOL3dbl7VVQPagijghZDCxrI2NRrWomiIiH0AAAAAAcLxRUAj1vf7XExKxLilkqUS818a865q8YIl4a95V6iDL3Oe5XOVXOVdVVelSwfKN3LA8lv9Xe7vLdJ62rkV8j1qPmTsRDz27rGy5E4w3Jf+pUCA4J8rus7LvwFx+Uqfh26tswXojuSf8AUAQJBPZN1bZgnTHcl/6g5+pX2X/grl8pAgQCeq7q2zFehlyT/qP8h9Srsx+8uXygCBQJ7fUrbMPwdy+UGObTN23Z5YMBvV5t8dw9lUdI+WLlz6pqiapwAhYAAAAAAAD9RSSRSNlie5j2rq1zV0VFN3bMtpUVxSG0Xx6R1mnJjqF4Nk7EXsU0eEVUXVOCkfZ1sexT0XhzyYq5K8WSH2mZ/SY3Suo7dLHUXaRODdeU2FPvnd/sQ0PGy6ZDekY1Jq2vq5NE63Pcp0kR0siJrq5y6aqvWpJjZziVrxazRPhjjqLlOxHT1at1VNU4tYvU3q7VIk/YdMwdo/eZcZ+z1cbr7M8GpMPpEqpnx1N6lb9clRO5p0+8Yvb2u8hmCOc1yPa5Uci6ouvFAvE4Mns7eTZyeu/9vwVGXNfLbmXlbRMMt20a2t5yWChyOBNIKxyaNqU+8kVOvsd5SNldS33DcmfS1Mc1vulFJoqLwVF/ii/OhKhOCop1svxux57am26+oymuETeTR3RrO7Zp0Mk++Z5i86b1aJiMWef7/un625z8GRiezfO6PJqVlNUvjgujU0fFrokn5zfV1GakY8xxbIsByJtJco3U9RG7nKeoidqyRuvB7HJ0obP2b7Tae58i231zKes0RGT9DZPD2KfPUek885cEf2/Z5s6fPx4/ybNOThFReKcU6gZ2YmFZMcOT70tbVUyK2GZzWu9s1dFavhReB1we0vak81niXtbTWeYl+K+24vdV1vOI2asdx7tsKwvTxsVE8qKeRLgWzKWoSoXEpWPTT63HXvSNdO9pr857QJ1OqbVY49bvG3mj3daksOFUD0ktuFWiGVOiSVHzKnicunzHpS1tS+FIUkSOFOiONqMYngRE0OsDlm3s+aOLXl8X2Ml/MmhycAiOXIAYBtL2hUuPwvoLa9lRcnIqcOLYe+vf7xJ1tXJs39FI/aHTFhtltxV6m0HNqDFaPkryaivkT61Ai/O7sQj/AFE16yzIWpyZ6+4VcnJjjYiuVVXoa1Ow4oqW9ZXf2U9OyouFxq36InFzlVfMhJLZ/hFr2fUfKjc2tyKaPk1FX0tptU4xxd/tcaulMHTMPM+f1lcVrj1ad3R2YbO6TBKdtxvEMFXksiasavdMoE06upZO/wBXhMqke+SRZJHK5zl1VV46nGqqqqvFVODMbm7k2r828e0KrPntltzPgACqicVIUR7OAa62obQ4bFHJa7U9stxc3Rz04pBr/HvHR2qbRo6KOSz2KVr6pdWzTt6I+8nfNJyyPlkdJI9z3uXVznLqqqabpnSeOMuaP6R+601dTj47v1UzzVNQ+oqJXyyyLynvcuqqp8wDRLIAAAAAbK3ZPdosXv3eYsjK3N2T3aLF+kd5iyMAAAAAA4XoUrL2++7LlXxjIWaLxQ0vk+7bs/yLIa6+XBbgtVWzLNLyJtE5S9iAV9gnl9Snsz7bn8o/yC7qWzPtunyj/ICBp3LTa7ldqlKa2UFTWzLp3EESvX5iftp3atlNC1iSWWWrc37qady6+HRTZGMYjjWM0qU1istFQRp1QxI1V8YET9hO7HdKyupb7nTEpaNipI2g+7k7z+xO8TFoKSmoaOKko4WQQRNRrI2JojUTqPuAAAAAAAam3o9oMGDbNKxsUyNulxYtPSNReKKqaK7wIhsrILvQWK0VN1uVQyClp41e97l6kQrk2+bR6raRnNRdF5TLfAqxUUSr7ViL0r316QNevc571e5Vc5y6qq9KqcAAezgnHOLCn/Mqf960tRh/oWe9TzFV+BfbzYPjOm/etLUIf6JnvUA/YAAAAAQV37U/1r0vwFvnUnUQX37vdVpPgLfOoEegAAAAAkPuH+6nW/AXekhHgkRuGp/rRr/gC+kgE5QAAAAA6t2TW2VSf+E7zHaOvcf6vqP0bvMBVXlCaZLc07KuX0lPOPTy37abr8Ml9NTzABP7czxRcd2SwV08SMqbq9ahy8nReT0N18RCDZ5j82VZrabBCmq1lSyNy9jde6XyaloNit0Fps9JbKViMhpoWxMaidCImgHdAAAA8TOr5BjmI3O9VD0Yylp3P1Xt04AYBmO8Fs9xbIaqxXKsqFq6V3IlSOJXI1ezU8b6qPZhr9k1v7BSCmR3SovV+rrtVPV81XO+VyquvSp0AJ+fVRbMP7VW/sFOU3odl/8AbK39gpAIAT+Teg2XL/8AraxP+nU/Sbz2y3+31af+QpX+ALH8C25YHmuQx2Ky1s76yRquY18StRUTpNnFWOzvIqnFM1tV+pnq11JUNc7Relmujk8mpaDYrhBdrNSXKmej4qmFsjXJ1oqagd0AAAABw5Ec1WqmqLwUri3ocWXFNsd3pWQ81TVbkq6fRNEVr+nT+8iljxGDf3xNK3GLVltPFrNQyrTzuT8G/o8ioBDEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAABLz+Tzpqd1Dl1W6FizslpWNkVO6RqtkVU18SEsSKn8nj/U2Y/CKT0ZSVYAAAAAAAAAAAAAAAAAAADDdtvuS5P8XS+iZkYdts9yXJ/i6X0QKwwAAAAAAAAAANk7L9os9olZa7zM6W3uXRkrl1dCv+E1sDlmw0zUml45iXzelbxxZL6CaOeGOeF7XxyNRzXNXVFRetD6Eetmm0Grx6aK33Bzp7UrtFTpdFr1t73eN/UNXTV1LHVUkzJoZE5THsXVFQxu/0++rbnzX5/updjWnFPMeH2OTgFciuLxbrNk9jdYMkgdLSrxgqGJ9dpX/fNXs7U6yOe0zZ3e8HrWvnb7Ktky601fCirHInZr1O7xI07VPVt9hy26upoq+3T8JqWdvKY7wa9C98vOndWnDxjy94+ifrbk0+G/hHrAtqFfZ0jobwjq2iRURHqv1yNP4obpx++Wu+0aVVsq452KnFqL3TfCnShhmc7Dqe5Ndctn1QiuXVZLTVSokrV/8ADd0OTvLoppqrpchxK7rFUw11pro10VHIsbv80LbY6fr7tfXSeJ+cf8peTWx549VfKVZwaCsW1zIKJGsr4oLhGnW7uH+VPUZ1Z9rmNVaI2tZVUEmmqq9nKZr2IrdV8qIUebo2zj+7HMfgg30stfHdsM5MfpMzxaqRqw3yj1VNUR0mi/OdpMjsKpr9GKL9s0hTp5480n8nD7DJ/wCMvWB4VVl2M039Ne6JvDo51F8xj122sYrRx/6NJUV0munJij0Tyu0PunT9nJPak/R9V1stp7VZ6p1Lrc6C1Urqq41cVNEn3UjtNfB2mlr3tivNTqy20UFEz75y847+CGC11fesir2+yJqqvqXroxiauXwIiFrr9CtM85p4/ol49CZ73lsfPNq7545KDHEWNjtWuqnJ3Sp+anV4TBcNxa+5rfUobXBJPK9eVPO/2kTetz3dSdJsLA9iFxq2xXPM6n6CW9e6Sm01qpk7Eb9z4VNw0EVvs1qSzY9QsttuRdXMYur5l++e7pcvzFll2dbp2P00jv8AL90m+XFrV4h5mEYxa8CtT6G0ztqrhP8AZdw5OiuT7xnY3zno66+EAyu1t5Nm/ruqMuW2W3Ng5OD8TyRwxOlle2ONqauc5dERCPWs2niPLnEczw/T3NYxz3qiNamqqq8ENO7UtpPK52zY/Nw4tnqW9fajV/ieftT2jSXR8tnskqsoU7mWdq8Zu8n5vnNYmq6b0qMXGTL976LfV1PR8V/Llyq5yucqqq9KqcAF6ngAAAAAAANl7snu0WP37vMWRFbm7J7tFj9+7zFkYAAAAAAAAAAAAAAAAAA4A5Ole7pQWa2T3K51UdNSwNV8kkjtEREMN2obW8NwChfJdbjHLV6LzdJCqOkcvg6vGQg21bZ8l2k10kc0j6G0Iv1qijfwVO1y9agZHvNbbqnP7i6yWWR8Ngp38NF0WocnWve7xo0AAAAPawL7ebB8Z0/71pahD/RM96hVfgn28WH4yp/3rS1CH+iZ71AP2AAAAAEF9+73VaP4C3zqToIL793uqUfwFPOoEegAAAAAkRuHL/rRr/gC+khHckPuH+6lW/AV9JAJzAAAAAB8Lh9gz/o3eY+58K/7Cn/Ru8wFVuXfbVdvhsvpqeWepl/213b4bN6anloiqqIiaqvQBJXcQxJtxzGvyeohV0Vvj5uFy9HOO6fmJsGp91bD0xHZHbmTMVtZXp7Kn1TRUV3QniTQ2wAAAAjrvz5b9CNntNj0Emk90m0ciLx5tvFf4Eiiv/fMyd1/2vT0LH8qntUaU7URdU5S8XfwA0kAAAAAAAAWD7neTLkWxuihlejqi2yOpX8eKInFvzKhXwSa3B8mSizC74zNKqR18CTwtV3Dls4LonaqKnkAmmAAAAAGLbWMcjyzZ7ebFIiKtTTORi6a6ORNUXymUnCoioqL0KBUzW00tHWT0k7VZLBI6N7V6nIuiofE23vY4k3Ftr9xWCNGUtwX2VEiJwRXe2TympAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAmF/J4/1NmPwik9GUlWV8buG2mm2T0V6p6izTXFblJC9FjkRvI5COTr98bcbvjWvrw+t/btAlWCK31Y1o/I+u+UMP0m+NZuvD6/9uwCU4IspvjWTrw+4fKGHP1Y1i/I+4/t2ASlBFv6saw/kfcv27B9WNYPyPuX7dgEpARc+rGsH5H3P9uwfVjY/wDkfc/27AJRgi59WNj/AOR9z/bsH1Y2P/khc/27AJRgi59WNj/5IXP9uwfVjY/+R9z/AG7AJRgi2u+NYerD7l+3YcfVjWL8j7j+3YBKUw7bXx2TZP8AF0voqaK+rGsf5HXH5Qw8bOt6yzZHh91sUWK18D66lfA2R0zFRquTTVQIogAAAAAAAAAAAABmGzvOq7FqlYX8qot8i93Cq+1/Ob2KYeD4vSuSs1tHMS8tWLRxKWdiu1BerdHX2+obNC9OlF4tXsXsU7xFzDcpueL3BKiify4nL9dgcvcvT+C98kHhuW2rJqBk1JK1lQjfrtO5e6YvX4U75k+odKtg+PH3r9FPs6k4/ir4ZABqCmQnKK5rkc1VRU6FQ71VW010ofofkVror3Raac3VR6uan5r07pPKdAHbDsZMM80nh90yWxzzWWJXvYxs9uz3SWe6XOwSuVV5qdEqIU16kXg5ETv6mGXTd7zJiOfZK20XtidCU9U1r9O+12mhuA51VF1RVRS3xddy17Xrz+ibTqFo+9HKPFz2P7S7fKkc2H3SXVNeVTxc83ys1Q8tdn2dJwXEb38ik9RKSC4V8DeTDWVEbdddGyKiHZ+j96/GlV+0UmR1/H70l1jqNfki7Q7K9otbLHHBhl6+uLo1z6VzG+NyoiJ4zJ7fu/57I7W6sttmi+/q6xia+BGqqm9JrpcpuUklfUvR3Siyrop1Fc5y6ucq+FTnfr8fyU/V5bqMe1WA2bYVi1E5JMgyqouCp0wW2HkIv99/qNhWSCw41TpT4pj9HatE0Wp5POVL++sjuKeLQ+RyV2fq+zl7RPEfgjZN3Jfx2fqeWWeVZZpHSPcuqucuqqfgHJWTM27ok8zPMuDkHl5HfrZYKJ1XcqlsTdO5b90/vIh94sV8torSOZfVaTaeIh3K+spqCkkq6uZsUEbVc57l4IhoLabn9TkNQ6htz3wW1i6cF0Wbvr3u8dDaBnFwyiqWNFdT0DV+twovT33dqmImu6d0yutHrv3t9FxrasYo5nyAAtkwAAAAAAAAAAGy92NNdtFj9+7zFkRVzsuyhMOzagyBaf2QlMqqrNdNdU0JHN3xV17rD18VQnqAlwCJjd8WD7rD5vFUIfv6sWk/I+f5QgEsAROXfGpPyPn+UIPqxqb8j5/lCeoCWIImLvjU/Vh83yhPUcLvjRdWHy/KE9QEtARFk3xZP9niH61R/kdZ++HcV15OJwp2a1C+oCYYIVSb4GVLryMZtzexVncv8Dw67es2jzq/mIrbTo5e50iVeT84E8F4dJ07jdbbbollrq+mpmJ0rJIjfOV2X3bvtPu6SNlySeBknS2BEZp4DBbvkF8u7+VdLvW1i/8AizOcgE+c33jNmuNseyK6/RWqbqnM0acvinUruhCOm0/ejzDIo5aHG4m2KjdqiyN7qZyeHoQj8APvX1tXX1T6qtqZaieRdXySOVzlXwqfAAAAAAAA9bC3cjMbK7suEC/+40tSpV1pol7WIvzFTlHUSUlZDVRLpJDI2Ri99F1Qk1aN7y801BBT1WM080kbEa57ZlTlKidOmgEzgRCj3xajTu8RavgqP8j6fViv/JBflCeoCXIIhv3xZ/ucRb46j/I+L98S4L7XEoU8M/8AkBMIgvv3e6nR/Ak86mRfVhXTrxSD9uvqNJ7a9o1TtLyeO91NAyidHCkSRtfyuvpAwQAAAAAJEbhya7Ua9eyhX0kI7me7EdpFTsyyea9U1AytWWBYljc/k9aLr8wFmAIet3xLh91iUPiqF9R9G74lVp3WIsXwVH+QEvgRD+rFm/JBvyj/ACOF3xajqxFvyj/ICXp8LgulDOv/AIbvMREfviV6+0xOJPDP/kdau3vbrPSSwsxeBjntVvK59eGviAjjl/HK7sqf22X01Pd2LY07Ldp1ksiIislqEdLqmqchvdO+ZDFblVOrbhUVj0Rrp5XSKidSqupm2xLaIuzTJZr7FaIrjO+BYo+cfyeb1VFVU4d4CyulhjpqaKniajWRsRjUTqRE0PqQ8j3xLin9JiMC+CoX1HYbvizfdYgniqP8gJdgiN9WK/T7UF+UJ6j4Tb4lYqfWsSjT30/+QEqM0vdPjmK3K91T+RFSU75VXwIVc5HdKi936uu9U9z5qud8rlXp4rrobi2ubxl/z3Fp8e+hkNvpqhU51zJFVXIi66GjgAAAAAAAABlGyfJZMQ2h2bIGOVG01S1ZNOti8HJ5FMXAFslrrILjbqeupno+GeNsjHIvBUVNTskBtke8lkmDWCKx1VBHdqODhCskitexvZr2Gw274rdO6w9+veqE9QEtgRKXfFZ1YfJ8oT1H4dvir9ziDvHUJ6gJcAiE/fFqfuMQYnhqP8j4P3w7kvtMUhTwz/5AZZv54olbhVFlcMacu3TNjmd18h6o1P8A1KhCo35ta3kK/aBs/uOJ1GPw0rK1YlWZsvKVvIlZJ0adfJ08ZoMAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB2LfW1dvqmVVFUSQTMXVr2LoqHXAG5ML2uRuaylySPkv6PZMbeH95PUbWoK2kr6VlVRVEdRC9NWvY7VFIinpWS+3ayzpLba6aDRdVa13cr4U6FKjb6Phz/ABU+Gf0Q82lTJ3jtKV4NOY5tke1rIb7buV99PTrov6q/wU2HZsyxq6tRaW7QcpU1VkjkY5PEpns/S9nD/LzH4K7JqZae3LIAcNc1yatcip3j9ECazHlG4n3cAA8AHJwAA1RDyrxkdjtLFdX3OnhVOlvLRXeROJ0x4cmSeKVmX1XHa3aIerqfieaKCJZZ5GxMamrnOXRETt4mrcl2w0ELXRWOjfUydCSzdyxO/p0r8xq/I8svt/ev0QrnujVdUiZ3LE8Rca3Q8t++WfTH6puLQvbvfs25mW1a2W5H01lRtfU8U5xF+ttXw9Zpa+3m43utdV3KqfPI5eGq8Gp2InUh0AaPW08WtXjHH9/dZYsNMUcVgABJdQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAACKqLqi6AAepbMhvltejqG7VkOi66JKvJXwp0KZJQbVMvpkVJKuCp1/Cwpw8mhg4Pi+Ol/vREvma1t5hsui2x36PleyqGjn16OTq3T51O6zbTXad3ZoFXvSKanBHnQ1p80j8nxODHP8raz9tFw+4s9OnhkU8uo2u5S9zuabRQoq8NItVTyqa9B7XS16+KR+RGHHHirILrmuUXNvIqrzU8jo5Ma82mn93Q8F73yOVz3uc5elVXVT8gkxWK9oh0iIjwAA9egAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAO9YbVXXy80lotsSTVlXIkULFcjeU5ehNV4Gzvqb9r/wCS6fK4v8RjGwr3YMW+MY/OWcAV2ru4bX0/3X//ANUX+I/Lt3Pa8iKv0quXwVMX+InnW5ridFVSUtXkFvhnjcrXsfO1FavYpzRZpidZO2ClyG2yyu9q1tQ3VfnArVy7BcvxKV0eRY9cLeiO5POSRKsar2I9NWr4lMcLYLtbLbebfJQ3KkgrKWZujo5WI5rkIA70+yyn2cZjHJaUf9BrkjpKdruPNOReLNezjwA06AANhYbsX2jZfYIL7YLD7Lt86uSOX2RG3XkqqLwVUXpRT2vqcNr/AOS6fK4v8RK/c39wGyfpJ/3rjY+VZZjmKxQy5Dd6W2sncrYlnejUeqdKIBApN2/bAv8Auu35ZF/iC7t22BE+1hq/9ZF/iJqptg2Zqv252j5Qh+v53tmn5aWf5QgEANoGyzOMDoIK/KLP7Bp55ObjdzzH6u0107lVMKJb76+b4lk+DWmlsF/oLjPFXct7IJUcrW8leJEgAfaipaquqo6Wjp5aioldyY4omK5zl7EROKnxJy7neyq1WPD6bM7jSNmvFxZy4nSoi8zF1clOpV6dQI+4zu07Vr3RtqltFLbWOajmpXVKMcqL3kRVTwLod+4bq+1akpXTMp7RVuToigrNXr+s1E+cm1muaYzhlE2ryO7U9DG9dGI93dO8CdKnj4vtb2e5JOtPa8moZJU+4e/kKvg1Arry3CsrxOpfBkVgr7erXK3lyxLzblT716dy7xKY+WoZtFjdTitbLk8NJPaGQrJPz6IrOSnHUrFzKez1WVXKosFM+ltb6hy0sTl1VrNeAHkn6iY6WVsbE1c9Uaidqqfk7Nq/rOl/TM9JANn0+7rtcngjmixjVkjUc1fZcXFFTVPuj6fU37YPyXT5XF/iLBLEutkoV/4eP0UPAvm0fBrHc5bbdsnttHWRac5DLMjXN16NUAg2m7dtfX/dhqf9ZF/iOV3bdr6Jr9LLF/6yL/ETW/ne2aflpZ/lCD+d7Zprp9Odn+UIBXNm2KXzDL9JY8io/YlfGxr3Rc41+iOTVOLVVDxDbm9terTkG2WtuVlr4K6jfSwtbNC9HNVUbxTU1GAAAHoY1Z67Ib/QWO2Rc5WV07IIW6LpynLpquiLoidKr1Iim6W7qG1RU11sSeGtd/gMk3DsKWuyiuzSqicsNAxaemVUTkrI5O6XwonnJoAVT5dYLji+R1tguzGMraKVY5UY7VuveXrQ8olDv4YQtHf6HNaSJ6xVjUgqnJ0NentV8aEXgAAA/cET5pmQxN5T5HI1qdqrwQ2rDu67XZWNe3Fl5Lk1RVqov8RrXHk1v9vT/io/SQtapPsSH3jfMBXn9Tftf/JdPlcX+I/K7uW15P8AdZfFUxf4ietbmOLUVVJS1d+t8M8buS9j52orV7FTU/EOb4jNIkceRWxz3cERKhvH5wK2sv2fZriUjm5DjdfRNbprKsfLj4/nt1b85jBbFXUdBdaF9NVwQ1VNM3kua9qOa5FIL72+yKkwG9U99sMbmWe4vcixdUEvTyU7yp0AaFAAA79js13vlalFZrZV3CpXoipoXSO8OiIZTsX2c3TaVmENmouVDStVH1dTydUiZ616ELCtnGz/ABjArNHbbBbooVRE5ydURZJV7XO6QIX43ut7ULqxJK2moLSxeqpqEV/kbr5zu3XdO2lUrXPo57PWo1nK0SoVjlXsRFT+JMrKs5xLF+F+v1DQv6mSSojl8XSdbFtpOD5NK2Gy5HQVMzuiNJUR/kUCuTN8Cy/C6pafJbDV0HZI5vKjd4Hpq1fBqY0Ws5PbLLdrJU0l+paapoHMVZWztRW6Jx14laW12ixa37QLnS4dVyVVoZIvNOcnBq9bWr1tRehQMTO3ZrdV3e7UtroIudq6uVsMLOUicp7l0RNV4JxOoZXsd91bFvjWn/eIBlq7um19E1+lR3yqL/Ea1v8Aaa+xXmrs90g5itpJFimj5SLyXJ0pqnAteKzt4JNNtWWIn4ykAwQzLZ9swzXPaWpqsWtHs6KlekcrueYzkuVNdO6VOow0mV/J8fapkvw5noIBHPNNj20PDrHJeshsDqOgic1r5eeY5EVy6JwRVXpVDASwbfRbrsBvC9k9N++aV8gAAAAAAA3xue7MaHOMtnvF7p+ftdqVruad7WWVeLUXtROnQDB8C2N7RM1hSpsuPTpSKqJ7JqVSKNdetOVxcnfRFM4+pR2qf8j+WL/gJy11XasftDqiqlp6Cgp2cXLoxjEQwOh277LqyvZRxZRStke7kor9Wt18K8AIMZtsi2h4g5VvGM1nMaKvP07eej0TrVW68nx6GCqioqoqaKhbFR1dvutHz1LPT1lO9PbMcj2qhAne+kwdm0ZaDErdDTVNMipcJINEjfIvHTROtOsDSYAAHfsVku99rW0dmtlXcKh3RHTwuevzIZfsP2Z3PabljbVSOWnooUR9ZUqmqRs16E/OXqLBNnOz/F8Bs7bdj1ujgTT65M7upJF7VcvECFmP7ru1O5xNlqqOgtbVXi2pqUV6J26N1T5zsXzdW2nUELpaRltuKNTXkRVGj18CKmnzkz8tz/EMV4X2/UdG/TXkPkTleQ6mKbUMEyiVsNmyOiqJndEfOI1y+JQK3ssxTI8Ur3UOQ2ert8zV0TnY1Rrveu6F8SnilqmVYzYsptcluvlup62nkTRUkYiqnfReogjvK7F6jZtdkuNs5yewVT9Inu4rC77xy+ZQNMgAD0sasd0yO9U9ns1KtVXVC6RRI5EVy+FeBsOTd52uRxrI/FJEaiar/pEX+I/O6m3lbcrAn/iOX/0qWK1n2JN7xfMBU3VQS01TLTzN5MsT1Y9uvQ5F0VD5np5WmmU3ZOytm9NTzAAAAG06Hd92sVtFFWU2MOfDMxHsX2THxRU1T7o1YnShaphH2n2j4HF6KAVi5ni19w69us2RUDqGubG2RYnOR3crrouqKqdSnim9N+H3cH/F0PneaLA9HGrJc8jvdNZbPTLVV1U7kQxIqJyl6eleBsd+7rtdZGsjsXVGomq/6TH/AIjq7rKIu3XGteqdfRUsZrfsOb3i+YCp6vpZ6Gtmo6pnNzwPWORuuujkXRUPie3nv27Xr4dL6aniAZlgOzDNs6o5qvGLK+uggfzcj0lY1EdprpxVDK27t+15U+1lqeGri/xG49x/K8csWBXamvF4o6GZ1fy2tmlRqqnJTjxJCN2iYQ5dEye1r/1DQIMLu3bXk/3ab8rj9Z+Hbue15v8Auuq+Cpj9ZYdS1ENVTx1FPK2WGRqOY9q6o5O1Dyr/AJXjthmZDeLxR0Mj05TWzSI1VTxgQCdu77XG9OKP+UR/4jWt8tddZbtU2q5QLBWUsixzRqqLyXJ1aoWWVW1bZ7TsV0mV2zgnVMild21i50152kX+6UciSU9TWyPjenQ5uvBQMYPcwnE79md7SzY5QrW1yxul5pHtb3Kaarqqp2oeGb33Gfdu8Fsn9JgGM1u7/tYo6OWrqMWkZDExXvd7Ij4InSvtjVxaxmCa4pdU/wCDl9FSqdelQAAAGZYLsuzvNXa4/jtXPCnTPInNxJ/edoi+LU31uqbA4bjT0+a5nSJJTPRH0NFInB3Y9ydnYhLdraG10WjWwUlNE3oREY1qIBB2l3TNpEkCvnqrNC/qYk6u18ehimYbvm1DGqZ9VPY0r6diK5z6KRJOSng4L5EUmrdNs+zS3VraOpyug51X8hUa/lI1e+qdBleO5HYcko/ZNludLXwqnFYpEd5QKq5opIZXRSxuje1dHNcmiovfQ/JYbtx2FYzn9rnqaKmitl8aiviqomInOO09q9OtFICZNZLljl8qrLdqd1PWUsiskYvb2p3gPNAAGc4HslzzOLQ+641ZFraNkqxLJzzG90mmqaKuvWh+M62U57hNvZcMjsE1JSudyOdR7XtRe/yVXQlpuHJpseqV7bnL6LTdWaY3bMsxussN2hSWlqo1Y7tavUqd9AKqwZrtk2e3XZzmFRZa9rn06uV1JUcnRs0fUqd9OhTCgNi4rsS2l5PZYLzZscfPQ1DeVFKszG8pO3RVMdz3B8nwa4Q0GT211DPPGskbVka7lNRdOlqqWG7vcbY9jWMNamiewWKRo/lAW6Zxj7u2hf6SARmAAH0poJqmdkFPE+aV66MYxquc5exEQ2viW7rtTyOjbVx2SO3QubymLXy80rv7uiqnjRDfW5bssttBicWdXejjnuNfq6j5xuvMxIuiOTvu6dew3xmmY43h1C2syG6QUMTuDOWvFy9iJ1gQjrN1narT075W01qqFamqRxVernd5NWonzmsMxwbLcQqXQZFYa2gVqonOPj1jXVNeD01aviUsKxjbLs5yKr9iW7JaRZ+pkruQq+DUynKILBXY/VOvsNJU2xIldNzyI5nJROK8QKqgZBtGqMfqs1uk2LUi0lnWoclLGrlXudenj0IvUhj4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAZpsK92DFvjGLzlnClZGwr3YcW+MY/OWcAVjbc3L/PDlnFf60m9JTDY5JI5GyRyOY9q6tc1dFRe1CZOfbqVRk2ZXbIWZnHTJcKp9RzK0Cu5HKXXTXlpqdKy7nFHFXsfeM0lqqRPbR09GkT1X3yudw8QG090i+XS/7E7XWXepfU1EcksCSPXVysY5UbqvXwMA/lBI2/SJj8nJTlpc1br16c28kJiOPWrFMdpLFZqdKeipGciNuvzqvWpELfqzujvOQ0GH2+RkzbY5Zqp7XaokqpojfEirr4QIzAACwzc49wCx/pKj96415/KBQTz2DGEghkl0qpdUY1V07hOw2Huc+4BYv0lR+9cbcnp4J0RJ4Y5dOjltRdAKoPofcP7DVfsneofQ+v/sNT+yd6i1r6H0H9ip/2aD6H0H9ip/2aAVPSRvierJGOY5Olrk0VD8m0t6tjI9ud/ZHG2NqPZo1qaJ7VDVoAtH2WVFPVbOrBNTSMkidQxaOYuqL3KFXBKHdX2+27GrXHh2ZTLBQxrpRVmiqkaL9y/vdigZZvlbJsvzC80GTY3E+4x09LzEtG12jmaKq8pqLwXXXj4EIh3Wy3qzVKxXK2VtDMzRVSWJzFTv8AEtLsl8s97oo6y03KkrqeRNWPhlR6KniP1dbNabrA+C422lq43po5ssSORfKBWxXbVc3rsGkw6tvU9TbJHNVUkXV3JToZr97w6DCCde1rdhxHIqSorcVZ9A7tyVVjGL/o8juxzerwoQkyOzXDH75V2a607qespJFjljd1KgHnnZtP9aUn6dnpIdY7Vo/rWk/Ts9JALVrFwslCn/Dx+ihX9vZUdZLt1vz4qSd7VWPRzY1VF7hCwKx/1LQ/B4/RQ+slHSSvV8lLC9y9KuYiqBVF9Drh/Yar9k71BbfXomq0VSn/AJTvUWtfQ+h/sdP+zQ6OQUFC2xV7ko6fVKd/+zT71QKqgfWt+zJ/0jvOfIAfSlglqamKngYr5ZXoxjU6VVV0RD5m5t0LCXZbtVpqyoiV1BaU9kyqqcFd9wnl4+ICZuwnDYMG2ZWmyMYjZ+ZSWpd1ulcmrl8qniR7W6OTb87Zy1YVgbRqqyovH2R08jxN+dTP8vvVLjeLXK91bkbBQ0z5nd/kpwTx9BWpBll5i2lJm3ImStWv9lqjtV6Xa8nXhrw4AWFbcMOgznZtdbHIxrpliWWmcv3MrU1av8PGVnVlPLSVc1LOxWSwvVj2qnFFRdFQtSxC9U2R4vbr3SrrDWU7ZUTs1TinlII74eFOxTatPXQQq2gvDfZMLtEROX0PangXRfGBpYAAd/HPtgt3wqP0kLWaT7Eh943zFU2OfbDbvhUXpIWs0v2LF7xPMBWlvAuVdtOWcV/rKQwVHKi6oq+UmftD3VZssza75GmaR0iXCpdPzK29X8jXq15aa+Q8u1bm9OytY66ZvJPSp7ZlPRJG9fA5XOT5gM23JMjvV92YzwXZ887aGpWKnnlVVV7NNdNV6dOg+++5S08+xSeeViOkgq4nRqv3Kqui/MbWwTE7LhWN01hsVMkFHAnWurnL1uVetVIxb9e0K31UdHglum52eGVKitcx3BnDuWL2r194CJoATpAn7uZ4fDjuyemuzmN9l3j/AEh7k6eR9ynkPd3ldpv82uDLV0iNfdK1ywUbV6EdpxcvgQzPZxDFBgViigjbHG2gh0a1NETuEIvfyhrn/RbEGaryFgqVVOrXlRgRhv14ud9us90u1ZNV1c71fJJI7VVU6tNUT0s7J6aaSGVi6tfG5WuaveVD5gDaFy265/cNnrsOq7o6SF68l9Uq/Xnx/eKvZ3zV4AAyzY57q2LfGtP+8QxMyzY57q2LfGtP+8QC0ErV3gKGtftoyt7KOoc1bjIqKkaqi/MWUp0HwfRUb3q99LA5y8VVY0VVAqj+h9f/AGGp/ZO9RMf+T/p54MTyRJ4ZIlWuZoj2qmvcJ2klPYFD/Y6f9mh9YYIYEVIYo40Xp5LUTUDT2+d/3fr1+npv3zCvcsI3zv8Au/Xv9NTfvmFe4AAAAAAJm/yfssS4jf4ke3nEq2qrdeOnJIZGyt33alVbMMv9nuhdU2yqRI6yFvtlb1Ob30Am1vHYRdM+2ZVdjs1QkVakjZY2Odo2Xk/cqpAHK8AzLFqh8N8x6vpeQ7k84sSujVerRycFLF8E2l4XmlI2exXylmeqIr4XPRsjF7FavEyqenpqqPkzwxTMXqc1HIBWRgW0nMsGdLHYbvUU8MjVa+nc5Vj1Xr5K9CmKVtTPWVc1XUyOlnmer5HuXVXOVdVUsY2gbCdnWYwSLUWWOgrH6q2qo0SN6KvXw4L4yEm3DZXedl+RpQ1r0qqCo1dSVbU0R7dehexydYGvQD9Q8ZWe+QCwrdKwunxTZPQ1SwIyuuiJU1D1Tul19qniQ/G9TtSk2dYYyG2ORLzclWOmX8G1PbP8XV3zZWDRxxYdaI4kRGNo4tET3qEPd/uWodtItEL+VzDLfrH2aq7j/ACPV4ulwvFwluF0rJquqmdynyyuVyqp16aeamnZPTyvhlYurXscqOavaiofMATY3ONr9yyxlTiOSVXsivpIklpZ3r3czNdFRe+nDym6drmLUuY7PrtY6pjXc7TuWJVT2r0TVqp40IJ7p8s0W3Ow8y9zVc97XadaK1dULFZURYnoqaorV1Aqbr6aSjrZ6SVNJIZHMcnfRdD4mRbTWtZtDyBrURGpcJkRE9+pjoG1d0/3c7F75/oqWJ1X2NL7xfMV2bp/u5WL3z/RUsUqfseX3i+YCqrLvtru/wAOm9NTyz1Mv+2y7/Dpv3inlgAAATpLVcJ+1C0fA4vRQqqTpLVcJ+0+0fA4vRQCD+/H7t7vi2D0nmije+/Kmm29fiyD0nmiANn7rPu643+nd6KljFb9hze8XzFdG6smu3bHP0zvRUsXrfsOb9G7zAVZ599u16+HS+mp4h7WeLrmt6X/AI6X01PFABFXtAQC0HY9p/NfjmnR9D4vRQijv9e6DZ/gS+kSt2N+5ZjfxfF6KEUt/v3QLP8AAl9ICNgAAG99xr3bv/4yf0mGiDfG4wmu25fiuf0mATiyv7WLp8El9FSqVektayr7Wbp8El9FSqVelQBnOwfFo8x2qWWxzacxJNzkqL1sYnKVPmMGJBbh1PTzbYauSaNr3w2qV8SqntXctiap39FVPGBOWkggo6SKmgY2KGFiMY1qaI1qIQY3sNr92yPLazFrRWyU9moJFiekL9OfenSqqnV3ic1w19gVGnTzTvMpVLfHPfeq171VXLUPVVX3ygdMybZ1m9/wW/w3ax10sKtciyxI7uJW9aOTrMZAFo+zHLKTNsItuR0eiMqokV7dfaPTg5PEpGrf4w2nidac0pYuTJI5aSqVreC8NWuVfKnjMw3Camqm2W3GGaV7oYLk5sLVXg1Fa1VRPGqqevvusY7YrOrkTVtXEqa9uoEBgABOvcP9x2p+NJfRYb9lngikZHLNGx8i6Ma5yIrvAaB3DvcdqfjSX0WnT3275cMbo8TvNrnfDVU1c6RiovBdETgveUDZO33ZnQ7SsKmtzkZFcoEWWin04tenUveXoK58hs9wsF5qrRdad1PWUsixysXqVPOhZHsT2i2zaThlPeKJebqmIkdZTr0xSInFPB2Kau3wtkMOS2J+Y2KkRLxQt1qGxt41ESd5OtOkDaOwD3HMY+AM8xGr+UET/wCs8dX/AIKT0mkl9gfKTY7jCORUclAxFRU6OBGn+UE+3HHPgUvpNAjCAALLN3OWOXYjibo3IqJbo0XTqVE4mrt8nZblubvtl5xuJKyOgheyakR+j3KqovKai8FXga03VtvNDhVB9KOWc421LIr6Wrbq7mFXpa5PvdePAmLjuT4/kVEysst3o66F6ao6KVHfMBV7eLBfbHPzd1tNdQSJx+vQuZ4+J7385uauwqfEJb5Uy2qZU5THu1cife8rp0XsLLbrZrTdqd0Fyt1LVxPTRzZYkcip4zSO1bdjwvI6GWpxmL6BXREVzFi4wvXsc3q8KAQPB6uW4/c8XyGssV4gWCspJFZI3qXvp2op5QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAZpsK92DFvjGLzlnJWPsJTXbDiyf8xj85ZuoEfMw3psZxrKblYKnHbrLNQVL6d72OZyXK1dNU1U3Ds6zC0ZzitJkFmkV0FQ3VWO9tG7ra7voV0bc112w5Z8aT+kpsLdG2qJhGW/QK7TOSzXR6M1VeEMqrojvAvQoEjN767ZzZdnC12I1HselRysuMkSLzzGLwRWr1J2r0kAZpZJ5nzTSPklkcrnveurnKvFVVetS1+vpKK7WyWjq4o6mkqolY9jk1a9rkK6t4vZlNs0zmSihR77TV6y0EjvvdeLFXrVANZAACwzc59wCxfpKj964xbfbzDJsRsePS43eaq2SVFTI2Z0DkRXojUVEUync49wCxfpKj964zLaVs5xfaFTUlPk1I+ojpHq+JGvVuiqmi9AEAU207VE6M3u37RPUP56tqn5b3X9dvqJkfU0bKPxLN+3ccpu1bKE/+ySr/wCe4CAuQXm6X+6zXW81stbWzaLJNJ7Z2nadAlTvX7IsGwTZ/T3THba+mq31bY1csqu7nRe0isAMoxzZ/l2Q43XZBZ7NPV0FCqJM9icf7qdLtOvQ8OxxUM94pIbnUPpqJ8zWzysbylYxV4qieAsw2Qw4fDgdBSYZUU1Tao4ka10aoquXrV3fXvgVq2a9XzHq7n7Tcq221DHoqrDK5i8pF4aonT4FNtYJvLbRbBVxJdK9L3Ro7WSOoROWqd5yEsc92EbOcwnkq62ytpayTVXT0jubcqrx1XTgphWM7p+B2y5Mq7hW3C6MYqK2GVyNaqovXp0gb0xm6NvWPW+7sjWJtZTsnRi9LeUmuhCXfvt9PR7XqWphbyX1ltjkl77kc5vmRCcMbKK02xsbebpaOli0TVdGsY1PMiFde85m9Pne1auuFA7l0FKiUtM7XVHtbrq5O8q6gawOzaf61pP07PSQ6x2bT/WlJ+nZ6SAWr2P+paH4PH6KEKN5bajn+P7YrzarLlVwoqGFWc3DE5Ea3VqKvUTYsf8AUtD8Hj9FDXuZ7DNn2W5FUX6822WWtqNOcc2VURdE0TgBCH+evar+XF1/Xb6j8y7Z9qMsT4pM1uj2PRWuarm6Ki+ImT9TTsp/E037dxym7Vsp/Esv7dwFe7nK5yucuqquqqcG597fB8cwPOaC2Y3SLS00tEkr2q9Xau5SprxNMAE4roWB7n2EpimyyCtqIObr7s72RKqpx5P3KeTzkMth+GzZ1tKtViYjuYdKktS5E15MTeLvV4yy6CKnt1uZExGxU9NEjUToRrWoB95GMkYrJGtc1elFTVD4ewaL+yU/7NCAm1TbvnlbtAvEthyWsoLYypdHTQw6I1GNXkovHrXTVfCYx/PZtT/LS5/rN9QFlLGMY1GMajWp0IiaIhp3e6wf6cNlVTU00HOXG0qtVBondK1E7tvk8xGfY9t6zei2iWl2S5HV3C1TTJDURzKnJRHcOVwTqXQnsqQ1dIqLyZIZmeFHNVAKmF4LooM/3gMLkwXahdbO1itpJJVqKRdOCxPXVE8XR4jAAO/jf2w274VH6SFrNL9ixe8TzFU2OfbBbvhUfpIWs0v2LF7xPMBoPNt6LGcWyy5Y7VY/dJprfUOgfJGrOS5U601U2rsrzu0bQ8SgyGzq5kb3KySF6or4nJ0tdp1lfG8D7tOWfGUnnMk3YNqc2zrNGQVszvoFcHJHVsVeEbuhJETtTr7wEtd6Stzu2bOJ7lhFSkLoF1rFY3WRItOKt76FeVbVVNbVSVVZPJPPI5XSSSOVznL2qqlrzm0tyt6tcjKilqY+heLXtcnqK+t6PZfNs+ziSpo4HJZLi5ZKV6JwY7rYvg6gNQAACxbdUyeDJNjdoRsz5Kmhj9i1HLfyncpvDVfCh4u+Fs7q82wGOvtNOs9ztLlljY32z419s1O/w18RF7do2sy7Ncr5uuV0ljrlRtWxOKxr1PRO91lgVgvFrv8AaILpaayGso6hvKZJG5HNVAKpJY5IpXRSscx7V0c1yaKi9in5LH892FbOcyrVrrlZvY9U5dXy0j+ac/w6HmYxu37LbFcW10dpnrXsXVrauZZGIvgAg/S7OczqcLmzCGx1LrPC7kum5PFe+idKp3zEy16rgtVFZZYKmKmgt0cSo9jmokaM04pp0aFau2uTEJdotzfhKSJaVkXk8r2vL+65H5uvQBhZlmxz3V8W+Naf94hiZlmxv3V8W+Naf94gFoBAXbZta2j2faxklstmX3Klo6avfHDCxycljU6ETgT7Kzt4T3a8t+MpAPt/PbtW/Le6frN9RKfcozDJcvxq/VGS3ipuc0FYxkTplRVa1WIuiaJ2kFiZf8nx9qOSr/xzPQQDPd87/u/Xv9PTfvmFe5YTvm/9329/pqb98wr2AAAAAAB7GH4xe8tvcVmsFC+rrJdeSxvBEROtV6EQ8cmvuQUuBU2LPnt1winyeo1WtjlVGyRoi8GtT73vgQ8vFrvmLXp9HcKeqtlwp3Kmi6sc1e1FTzoZrie3LaZjj4EpsmqqqCFqMbBVLzjdPHx8epPnOtnmH5rAseRWWmq36aNl5OkjfA5OJp6s3SMEluiVEFzukFLy0VafloqadacpU14gZ/u57SqjadhMl3rKFtJVU0608yMdq1yoiLqnlMN366Snl2PRVb42rNDcIkjcqcU111Nt7PMKsGB4+2y49S+x6ZHctyq7Vz3ffKvWpG/fuzyilo6LBaKVks6SJU1atdryNPatXv8AWBEYNVUVFTpTiABZJu2ZTBlWyOz1TJGOnp4Up52ovtXN4cTCd8nZhWZpi1Pf7JTunutpRdYm+2lhXi5E76dJHPdh2tu2bZO+mubnvsVeqNqGpx5p3VIn8Sflju9rv1riuNprYK2kmbqySJ6OavkAqjmjkhldFKxzJGKqOa5NFRexT8ljmd7Btm+YXB1wuFndT1T11kkpJFiV69/Q8/GN2/ZbYriyujtM9ZIxeU1tXOsjUXwAai3INmNcy7TZ7eqJ8MDIubtySJor3L7Z+nYicE8Kko85vdLjuI3O81kqRxUtM96qq6cdOCeU9GaWgtNvdJK+CjpIGaqqqjGManzIQq3tdtkGXTriWL1LnWmB/wDpU7V7mocnUna1PnAj7e62S5XisuEqqr6iZ8qqvfXU6YAG1t073crF75/oqWJ1H2PJ7xfMV2bpvu6WH3z/AEVLFnIjmq1ehU0UCqfL/tsvHw6b01PLLEKrdz2WVVZNVT2WR8s0jpHqs7uLnLqp813bdk/4hf8At3esCvME4drO7/szsezm/Xm32qaGro6KSaF3PuVEciapwIPAE6S0jZdcqW77PLDcaNyugmoYnMVU0XTkp1FW5Mrcj2n0NRYUwC61TYq2lVz6FZHac5H0q1NetOzsAxPf3xWqhy215ZFG99NU0yU0rkbwY5iqrdV7/KXyEYS1nJrBaMltMtqvdDDW0kqaOjkTVPF2GoX7rWyx1Us/sS4Iiu5XNpVLyU72nYBHncpxWtvO1uG9JC/2DaY3SSS9CctyaNb4enyE5sorG2/HLjWvdyWw0z3qvZo1TqYXiWP4daG2vHrbDRUycVRicXL2qvSqmk98fanRWDEp8OtlU193uLeTKjHa8zF169ir1AQnvdU+uvNbWSO5T5p3vVdOnVVU6YAAAAWgbHPctxv4vi9FCKW/37oNn+BL6RK3Y77luN/F8XooRS3+/dCs/wABX0gI2AAAb53Fvdud8Vz+lGaGN87i3u3O+Kp/SjAnDlf2sXT4JL6KlUq9KlrWV/axdPgkvoqVSr0gDZu7DkcOMbZrLW1Mjo6eZ600io7RO7TRNe9roayP1DI+GVksTlY9io5rk6UVALaVRskapwVrk8qKV37zuza44PtAraplK5bRcJHT00zUVWt1Xi1V6lRSRm63tyo8rtNNi+S1kcN8p2pHFJI5GpVInRp+d3jemQ2K0ZDbn2+82+nrqV/THKxHIBVId/H7Pcb9d6e1WulkqaqoejGMY3VdVJ33Ddd2V1dZLUpQ18CSLrzcVUrWN8CGc7PdlmE4Lq/H7NFDUKmjp393IvjUBsRwxuBbN7Xjqqjp4o1fO5OuRy8p3zqaR3+Mrp4ccteIwytdU1U/smZmmvJYzo49XFUN57VNouO7O8elul7q2pJoqQUzV1kmd1IifxK59pWYXLOswrciua6SVD+4jReEbE9q1PAgGNgACdW4d7jtT8aS+i08L+UCbriWPu7Kt6f+k97cP9xyp+NJfRYeL/KAp/8ARdhX/jHeiBHDYRtMuWzTMoLjC98lsncjK+mReEjO1O+nShY1j92t2QWSmuttnZU0dVEj2ORdUVFToUqjJEbom2OTFLyzEL9Up9Bax/1iSR32PIv/AMVAnDTwQ00LYaeJkUbehjG6IniIafygn24458Cl9JCZrHNexr2ORzXJqip1oQy/lBU/+scbX/gpfSaBGEA+1AynkrYY6uZ0NO56JJI1vKVrdeK6dYGQYlgWWZXa7hc7DZ56ylt7OVO9ifMnavXoh41DXXWyVzn0VXWW6qYujlikdG9FRehdNCyXYdR4TQ7PaClwealntyRor3xuRznvVE5Sv6+V26nQ2g7ENnuayPqLlZmU9Y7VVqKVebeq9/TpAiHgW8jtGxypYlfcfo1Sa93FVJq7TvOQnVgeRU+WYhbcipI3xw10DZWsemit16UNI2TdLwaiurKqsuVxr4GLr7He5GtXj1qiam/6OnobPa46anZHS0dLGjWNTg1jUQCEu/lS08G1GgnihYySahRZHNTRXqjlRNe0juba3qs6o852oz1FseySgoWexYZW/wC00VdXeDXoNSgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB7OEX1+MZbbMgjp0qH0FQ2dIldojtOrUk/JvlsWDRmAuSXT2zrnqmvg5siOAPXzW+PyXLbpkEkCU77hUvqFiR2qMVy66ankIqoqKiqip0KgAEhcH3qMqxzF6Ky1FnpLm6kjSNtRLI5HuROjXTvaHibYdvcu0zFvoNeMToopY3pJT1Uc7uVC7tRNOPDVDSoAAADe2yjeQvGz/BaPFqPHqOrjpVerZpJXIq8pyu6ETvmSP3wcrX2uMWpP/MeRlAEmU3wcs68ZtS/33hd8HLPyZtP67yMwA3Htn29XfabjENkuFlpKJsUyS85DI5VVU6tFNOAADJcEzzLMIrFqcavNRRcpUWSNq6xyaffNXgpjQAkvje97lVJzDL5jtvuLGJpK6GR0L5OHTroqIuveMgq98prqaRtLgTo5lTuHSXNHNRe+iRpr5SJAA2ptR2857ntNJQVday322Re6paTVqOTsc7pVO8arAAH7ppVhqI5kRFWN6ORF69F1PwAJL0u93ktNRw00eMW1UijaxFWV3HRND9Lvg5X1Yza/2jyMwAk43fCyj7rFrWv/AJrz9N3w8k17rFbaqd6ZxGEAZ9tv2l1e1DJKe9VlthoHwQcykcb1ciprrrqpgIAGzdgu1WPZZcq+4sx+O6VNVGkbXun5vm29OicF6TP803rMhyDG7hZoMepKH2ZE6LnmTuVzEXgunBOJHMAFVVVVXpUAAEVUVFRdFTihJjBd7O6WHGaK0XPFornJSxNiSobVrGr0RNEVU5K8SM4A2lt+2tR7Va6grHY5Fap6Nrmcts6yOe1epeCGrQAPtQVC0ldBVNajlhkbIiL16LqSlpd8WtipGRPwmF8jW6cpK5UTycgioAPazrIJMqzC6ZFLTtp5LhUOndE12qMVerU8UADf+y/eeyPDcUpMeqrPT3eKkbyIZZJnNejOpq8F106DjatvGR7Q8QqLBdcJpG8vuoZ/ZSq6F/U5O5NAgAAABmuzXajmez6Zy47dnx07+L6WXu4XL28lehe+hhQAlbY98ethpOTecKiqqjX29NXLE3TwKx3nPtct8mR9KrbdgqQT6po+a4c43TwJGnnImADZG1LbXnW0JjqW63BKW3K7X2HS6sjX33W7xmtwAB6mI3h+PZRbL5HC2d9BVR1DY3Lojla5F018R5YAlqzfKckCI7A0WXT2yXHhr4ObIy55kEmV5jdcjlp20z7hUOnWJrtUZr1a9Z4gAG5N3rbi7ZPb7jQux5LrDWzJKqpU805qo3TT2qmmwBIPbbvJLtGwOqxWHFPoZHUvje+Z1bzqpyHo5NE5CdhHwAAAAAAAHbtFzuFnuEVwtdbPR1US6slherXN8aHUAG+sO3p9odkpEpbkyivTW8EkqGq2TTwt6fIbAi3y4EjakmASOfonKVt0REVe8nNERQBIHPt6rOL7DLS2OkprFTyapymLzkyNXq5Soia99ENCV1XVV1XJV1lRLUVEruVJJI5XOcvaqqfEAAAAM12a7Ucz2fTOXHbq+KneuslNInLid/dXo8RhQAlTYN8a5wwObfMNp6yThyXUtYsKJ26orXan2vW+PVS0SstGEx0tTrwkqK7nW6e9RjfORQAGx9pu2nPM+jfS3a5+x7e52vsOmTkR+PrXxmuAAAAAyTZnltTg+ZUWS0lNHUzUiqrYnqqNdqmnUbyn3wMvcv1rHLSxO+56/wASNIAkl9V9mv5P2f8A9frC73ubacLBZ0/X9ZG0Ab4zPeey/KMWuGP1lntcMFdC6GR8fK5SIqdWqmhwAB9KaealqI6inlfFNG5HMexdFaqdaKfMAb62f70meY9BHSXiOnv1MzRNZ15EvJTq5SfxQ2F9WZS//t9N/wD2qf8A/IiGAJBZzvV5zemS09ipaWxQPVURzV52VGr+cqImvf0NC3OurLnXzV9wqZamqner5ZZHauc5etVOuAAAAAACSeJb110x/GLfZI8TpZ0ooGwpI6pcnKRqaa6aGrNt+06t2oZBTXastsNAtPDzTY43q5FTXXXVTX4AAAAZ5sK2h/zY5wmS/Qv6JotM+ndBz3Nro5WrrytF+97DAwBKzIN8D6IWeqoaXBlgfPE6PlyXHlo3VNNdEjQimvSAAAAH7p5pqedk9PK+KWNyOY9jtHNVOhUVOg3bgO83tDxqmbSV8sF9p2po32Wi843+8nT4zR4Al3HvlwJG1JMAlV+ndKl0REVf2RjWY73WUXCJ0WOWGks6OTRJJpVne1e9wanzEaQB6+WZNfsqur7nf7nUV9S9V7qV+qNTsanQid5DyAAAAA3bsO3gK3ZhicuPw49BcWPqXT846dWKiuRE000XsPP267cLjtTtlHb6qy09vipZVlascqvVdU00XVDUQAHLXK1yOaqoqLqip1HAAkJgu9RlmOYzSWastNLdn0zObbUzSuR7mp0a6dJr/bhtUuO1O60FfcLdT0LqKJ0bGxOVyKjlRevwGuwAAAGQYXmeT4bXezMbvNVb5F9skb+4f4W9Cm9MU3ustoWxR3+xUN1Y3274nrC93zKieQjWAJb1e+Ux1NI2lwJ8c6tXm3vuaOa1epVRI01TxmoNp+37Pc6o326oqo7bb3po+no9W8tOxztdVTvGpwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAf//Z"
         style="height:64px;width:auto;object-fit:contain;border-radius:6px;flex-shrink:0;"
         alt="Zodha Research Solutions" />
    <div>
      <div class="hero-title" style="font-size:1.9rem;">Research Writing Pro</div>
      <div class="hero-sub">Writing · Statistics · Publication · Citation · Literature Review · 8-dimension Scoring</div>
    </div>
  </div>
</div>""", unsafe_allow_html=True)

# ── TABS ───────────────────────────────────────────────────────────────────
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["✍️  Humanizer", "🔄  Paraphraser", "✅  Grammar Checker", "🔬  Research Tools", "📊  Statistics Suite", "📰  Publication Suite"])
scores_in = {}
scores_out = {}

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — HUMANIZER
# ══════════════════════════════════════════════════════════════════════════════
with tab1:
    col_in, col_out = st.columns([1,1], gap="large")

    with col_in:
        h1,h2 = st.columns([3,1])
        with h1: st.markdown('<div class="card-title">📄 Input Text</div>', unsafe_allow_html=True)
        with h2:
            components.html("""<!DOCTYPE html><html><body style="margin:0;padding:4px 0 0;background:transparent;">
<button onclick="if(navigator.clipboard&&navigator.clipboard.readText){navigator.clipboard.readText().then(function(t){
  var f=window.parent.document.querySelectorAll('textarea');
  for(var i=0;i<f.length;i++){var n=Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype,'value').set;
    n.call(f[i],t);f[i].dispatchEvent(new Event('input',{bubbles:true}));break;}
}).catch(function(){alert('Use Ctrl+V / Cmd+V in the text box');});}
else{alert('Use Ctrl+V / Cmd+V in the text box');}"
style="background:#1a2e1b;color:#6fcf97;border:1px solid #4a7c59;border-radius:7px;
       padding:0.3rem 0.8rem;cursor:pointer;font-size:0.76rem;font-family:'DM Sans',sans-serif;
       width:100%;white-space:nowrap;">📋 Paste</button>
</body></html>""", height=40, scrolling=False)

        input_text = st.text_area("Input", height=350,
            placeholder="Paste AI-generated text here (5000+ words supported)…",
            label_visibility="collapsed", key="input_text")

        wc_in = len(input_text.split()) if input_text.strip() else 0
        bc,cc = st.columns([3,1])
        with bc: st.markdown(f'<span class="wc-badge">📝 {wc_in:,} words</span>', unsafe_allow_html=True)
        with cc:
            if st.button("🗑️ Clear", key="clear_btn", use_container_width=True, disabled=(not input_text.strip())):
                st.session_state.clear_input=True; st.rerun()

        if input_text.strip():
            scores_in = compute_scores(input_text, style)
            st.markdown('<p style="color:#ffffff;font-weight:700;margin-top:0.6rem;">Before — Humanness</p>', unsafe_allow_html=True)
            r1,r2,r3,r4=st.columns(4)
            with r1: st.markdown(render_score_ring(scores_in["humanness"],"Humanness"),unsafe_allow_html=True)
            with r2: st.markdown(render_score_ring(scores_in["flesch"],"Readability"),unsafe_allow_html=True)
            with r3: st.markdown(render_score_ring(scores_in["ttr"],"Lexical Div"),unsafe_allow_html=True)
            with r4: st.markdown(render_score_ring(scores_in["sl_variation"],"Rhythm"),unsafe_allow_html=True)
            st.markdown(render_metrics(scores_in), unsafe_allow_html=True)

    with col_out:
        st.markdown('<div class="card-title">✨ Humanized Output</div>', unsafe_allow_html=True)
        output_text = st.session_state.output_text
        stream_placeholder = st.empty()

        if output_text.strip():
            import html as _html
            stream_placeholder.markdown(
                f'<div class="output-box">{_html.escape(output_text)}</div>',
                unsafe_allow_html=True)
            make_copy_btn("humanizer-out", output_text, "📋 Copy Text")
            scores_out = compute_scores(output_text, style)
            wc_out = scores_out.get("word_count",0)
            st.markdown(f'<span class="wc-badge">📝 {wc_out:,} words</span>', unsafe_allow_html=True)
            st.markdown('<p style="color:#ffffff;font-weight:700;margin-top:0.6rem;">After — Humanness</p>', unsafe_allow_html=True)
            r1,r2,r3,r4=st.columns(4)
            with r1: st.markdown(render_score_ring(scores_out["humanness"],"Humanness"),unsafe_allow_html=True)
            with r2: st.markdown(render_score_ring(scores_out["flesch"],"Readability"),unsafe_allow_html=True)
            with r3: st.markdown(render_score_ring(scores_out["ttr"],"Lexical Div"),unsafe_allow_html=True)
            with r4: st.markdown(render_score_ring(scores_out["sl_variation"],"Rhythm"),unsafe_allow_html=True)
            st.markdown(render_metrics(scores_out), unsafe_allow_html=True)
        else:
            stream_placeholder.markdown(
                '''<div style="background:white;border:1.5px dashed #d4c9b5;border-radius:10px;
                   min-height:380px;display:flex;align-items:center;justify-content:center;
                   flex-direction:column;gap:0.6rem;color:#9a8a7a;">
                 <div style="font-size:2rem;">✨</div>
                 <div style="font-size:0.9rem;">Humanized text streams here in real-time</div>
               </div>''', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    btn_c, info_c = st.columns([2,3])
    with btn_c:
        run_btn = st.button("⚡ Humanize (Streaming)", type="primary",
                            use_container_width=True, disabled=(not input_text.strip()))
    with info_c:
        if not groq_key:
            st.warning("⚙️ OpenRouter key not configured — set OPENROUTER_API_KEY in Streamlit secrets")
        elif not input_text.strip():
            st.info("📄 Paste text above to begin.")
        else:
            chunks = chunk_text(input_text)
            _model_label = _active_model_dict.get(model_choice, model_choice).split(" · ")[0]
            st.markdown(
                f'<div class="wc-badge">⚡ {len(chunks)} chunk(s) · {_model_label} · {intensity}</div>',
                unsafe_allow_html=True)

    if run_btn:
        if not groq_key:
            st.error("⚙️ OpenRouter key not configured. Set OPENROUTER_API_KEY in Streamlit secrets.")
        elif not input_text.strip():
            st.error("Please enter text.")
        else:
            try:
                chunks  = chunk_text(input_text)
                n       = len(chunks)
                results = []
                progress_bar = st.progress(0, text="Starting stream…")
                for i, chunk in enumerate(chunks):
                    progress_bar.progress(i/n, text=f"⚡ Streaming chunk {i+1}/{n}…")
                    if n > 1:
                        temp_ph = st.empty()
                        text_out = humanize_streaming(groq_key, model_choice, chunk, style, intensity, temp_ph)
                        temp_ph.empty()
                    else:
                        text_out = humanize_streaming(groq_key, model_choice, chunk, style, intensity, stream_placeholder)
                    results.append(text_out)
                progress_bar.progress(1.0, text="✅ Complete!")
                st.session_state.output_text = "\n\n".join(results)
                time.sleep(0.3)
                st.rerun()
            except requests.exceptions.HTTPError as e:
                if e.response.status_code == 429:
                    st.error("⚠️ Rate limit hit. Groq: wait 1 min. OpenRouter: wait or add $5 credits at openrouter.ai")
                elif e.response.status_code == 401:
                    st.error("❌ Invalid API key. Check your Groq key (console.groq.com) or OpenRouter key (openrouter.ai/keys) in the Admin panel.")
                else:
                    st.error(f"❌ HTTP Error: {e}")
            except Exception as e:
                st.error(f"❌ {str(e)}")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — PARAPHRASER
# ══════════════════════════════════════════════════════════════════════════════
with tab2:
    st.markdown('<div class="card-title">🔄 Paraphraser</div>', unsafe_allow_html=True)
    p1,p2 = st.columns([1,1], gap="large")
    with p1:
        st.markdown('<p style="color:#3a8c3f;font-weight:700;">Input</p>', unsafe_allow_html=True)
        para_input = st.text_area("Para",height=300,placeholder="Paste text to paraphrase…",
                                  label_visibility="collapsed",key="para_input")
        para_wc_col, para_clear_col = st.columns([3,1])
        with para_wc_col:
            if para_input.strip():
                st.markdown(f'<span class="wc-badge">📝 {len(para_input.split()):,} words</span>',unsafe_allow_html=True)
        with para_clear_col:
            if st.button("🗑️ Clear", key="clear_para_btn", use_container_width=True,
                         disabled=(not para_input.strip())):
                st.session_state.clear_para = True
                st.rerun()
        pm,pb = st.columns([2,1])
        with pm: para_mode=st.selectbox("Mode",list(PARAPHRASE_MODES.keys()),label_visibility="collapsed",key="para_mode")
        with pb: para_btn=st.button("🔄 Paraphrase",type="primary",use_container_width=True,disabled=(not para_input.strip()))
    with p2:
        st.markdown('<p style="color:#3a8c3f;font-weight:700;">Output</p>', unsafe_allow_html=True)
        para_out = st.session_state.paraphrase_out
        if para_out.strip():
            import html as _html
            _paras_para = [p.strip() for p in para_out.split("\n\n") if p.strip()]
            if not _paras_para: _paras_para=[para_out]
            _html_para  = "".join(f'<p style="margin:0 0 0.8em 0;">{_html.escape(p)}</p>' for p in _paras_para)
            st.markdown(f'<div class="output-box-sm">{_html_para}</div>',unsafe_allow_html=True)
            make_copy_btn("para-out",para_out,"📋 Copy")
            st.markdown(f'<span class="wc-badge">📝 {len(para_out.split()):,} words</span>',unsafe_allow_html=True)
        else:
            st.markdown('''<div style="background:white;border:1.5px dashed #d4c9b5;border-radius:10px;
                min-height:300px;display:flex;align-items:center;justify-content:center;
                flex-direction:column;gap:0.5rem;color:#9a8a7a;">
              <div style="font-size:2rem;">🔄</div><div style="font-size:0.9rem;">Paraphrased text here</div>
            </div>''', unsafe_allow_html=True)
    if para_btn:
        if not groq_key: st.warning("⚙️ Service key not configured. Contact admin or set GROQ_API_KEY in Streamlit secrets.")
        else:
            with st.spinner(f"Paraphrasing ({para_mode})…"):
                try:
                    result = paraphrase_text(groq_key,model_choice,para_input,para_mode)
                    st.session_state.paraphrase_out = result
                    st.rerun()
                except Exception as e: st.error(f"❌ {e}")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — GRAMMAR CHECKER
# ══════════════════════════════════════════════════════════════════════════════
with tab3:
    st.markdown('<div class="card-title">✅ Grammar & Style Checker</div>', unsafe_allow_html=True)
    g1,g2 = st.columns([1,1], gap="large")
    with g1:
        st.markdown('<p style="color:#3a8c3f;font-weight:700;">Input</p>', unsafe_allow_html=True)
        gram_input = st.text_area("Grammar",height=300,
            placeholder="Paste text to proofread…",label_visibility="collapsed",key="gram_input")
        gram_wc_col, gram_clear_col = st.columns([3,1])
        with gram_wc_col:
            if gram_input.strip():
                st.markdown(f'<span class="wc-badge">📝 {len(gram_input.split()):,} words</span>',unsafe_allow_html=True)
        with gram_clear_col:
            if st.button("🗑️ Clear", key="clear_gram_btn", use_container_width=True,
                         disabled=(not gram_input.strip())):
                st.session_state.clear_gram = True
                st.rerun()
        gram_btn = st.button("✅ Check Grammar",type="primary",use_container_width=True,disabled=(not gram_input.strip()))
    with g2:
        st.markdown('<p style="color:#3a8c3f;font-weight:700;">Corrected Text</p>', unsafe_allow_html=True)
        gram_corrected = st.session_state.grammar_corrected
        if gram_corrected.strip():
            import html as _html
            st.markdown(f'<div style="background:white;border:1.5px solid #4a7c59;border-radius:10px;padding:1rem 1.2rem;height:220px;overflow-y:auto;font-family:DM Sans,sans-serif;font-size:0.9rem;line-height:1.7;color:#1a2e1b;white-space:pre-wrap;word-break:break-word;">{_html.escape(gram_corrected)}</div>',unsafe_allow_html=True)
            make_copy_btn("gram-out",gram_corrected,"📋 Copy Corrected","#6fcf97","#1a3a2e","#4a7c59")
            issues = st.session_state.grammar_issues
            if issues:
                st.markdown(f'<p style="color:#3a8c3f;font-weight:700;margin-top:0.8rem;">⚠️ {len(issues)} Issue(s)</p>',unsafe_allow_html=True)
                type_colors={"grammar":"#e87a7a","spelling":"#e8a87a","punctuation":"#e8d47a",
                             "style":"#a8d47a","wordiness":"#7ab8e8","clarity":"#b87ae8"}
                for iss in issues:
                    t=iss.get("type","other"); tc=type_colors.get(t,"#aaa")
                    orig=_html.escape(iss.get("original","")); corr=_html.escape(iss.get("corrected",""))
                    expl=_html.escape(iss.get("explanation",""))
                    st.markdown(f'''<div style="background:white;border-left:4px solid {tc};border-radius:0 8px 8px 0;
                         padding:0.6rem 0.9rem;margin-bottom:0.4rem;">
                      <span style="background:{tc}22;color:{tc};border-radius:4px;padding:0.1rem 0.4rem;
                             font-size:0.68rem;font-weight:700;text-transform:uppercase;">{t}</span>
                      <div style="font-size:0.82rem;color:#5a6a7a;margin-top:0.3rem;">
                        <span style="color:#e87a7a;text-decoration:line-through;">{orig}</span>
                        <span style="margin:0 0.3rem;">→</span>
                        <span style="color:#4a7c59;font-weight:600;">{corr}</span>
                      </div>
                      <div style="font-size:0.75rem;color:#7a8a9a;margin-top:0.2rem;">{expl}</div>
                    </div>''', unsafe_allow_html=True)
            else:
                st.markdown('<p style="color:#6fcf97;font-weight:600;">✅ No issues — text looks great!</p>',unsafe_allow_html=True)
        else:
            st.markdown('''<div style="background:white;border:1.5px dashed #d4c9b5;border-radius:10px;
                min-height:300px;display:flex;align-items:center;justify-content:center;
                flex-direction:column;gap:0.5rem;color:#9a8a7a;">
              <div style="font-size:2rem;">✅</div><div style="font-size:0.9rem;">Grammar report here</div>
            </div>''', unsafe_allow_html=True)
    if gram_btn:
        if not groq_key: st.warning("⚙️ Service key not configured. Contact admin or set GROQ_API_KEY in Streamlit secrets.")
        else:
            with st.spinner("Checking grammar…"):
                try:
                    result=grammar_check(groq_key,model_choice,gram_input)
                    st.session_state.grammar_corrected=result.get("corrected","")
                    st.session_state.grammar_issues=result.get("issues",[])
                    st.rerun()
                except Exception as e: st.error(f"❌ {e}")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 4 — RESEARCH TOOLS
# ══════════════════════════════════════════════════════════════════════════════
with tab4:
    r_tab1, r_tab2, r_tab3, r_tab4 = st.tabs([
        "📖 Citation Manager",
        "📄 Article Summariser",
        "📚 Literature Review",
        "✍️ LR Chapter Writer"
    ])

    # ════════════════════════════════════════════════════════════════════════
    # R-TAB 1 — CITATION MANAGER
    # ════════════════════════════════════════════════════════════════════════
    with r_tab1:
        st.markdown('<div class="card-title">📖 Citation Manager</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Enter a DOI or DOI URL to fetch metadata and generate a formatted citation.</p>', unsafe_allow_html=True)

        ci_col1, ci_col2 = st.columns([3, 1])
        with ci_col1:
            doi_input = st.text_input(
                "DOI or URL",
                placeholder="e.g.  10.1016/j.jfineco.2021.01.004  or  https://doi.org/...",
                label_visibility="collapsed", key="doi_input"
            )
        with ci_col2:
            cite_style = st.selectbox("Style", list(CITATION_STYLES.keys()),
                                       label_visibility="collapsed", key="cite_style")

        fetch_btn = st.button("🔍 Fetch & Generate Citation", type="primary",
                              disabled=(not doi_input.strip()), key="fetch_cite")

        if "citation_meta"  not in st.session_state: st.session_state.citation_meta  = {}
        if "citation_text"  not in st.session_state: st.session_state.citation_text  = ""
        if "all_citations"  not in st.session_state: st.session_state.all_citations  = []

        if fetch_btn and doi_input.strip():
            with st.spinner("Fetching metadata from CrossRef…"):
                meta = fetch_doi_metadata(doi_input.strip())
                if "error" in meta:
                    st.error(f"❌ {meta['error']} — Check your DOI or internet connection.")
                else:
                    st.session_state.citation_meta = meta
                    st.session_state.citation_text = format_citation(meta, CITATION_STYLES[cite_style])
                    st.rerun()

        # Re-format if style changes without re-fetching
        if st.session_state.citation_meta and not fetch_btn:
            st.session_state.citation_text = format_citation(
                st.session_state.citation_meta, CITATION_STYLES[cite_style]
            )

        meta = st.session_state.citation_meta
        cit  = st.session_state.citation_text

        if meta and cit:
            # Metadata card
            import html as _html
            st.markdown('<p style="color:#3a8c3f;font-weight:700;margin-top:1rem;">📋 Article Metadata</p>', unsafe_allow_html=True)
            cols_meta = st.columns(3)
            meta_items = [
                ("Title",   meta.get("title","—")[:80]+"…" if len(meta.get("title",""))>80 else meta.get("title","—")),
                ("Authors", "; ".join(meta.get("authors",[]))[:70] or "—"),
                ("Year",    meta.get("year","—")),
                ("Journal", meta.get("journal","—")[:60] or "—"),
                ("Volume",  meta.get("volume","—") or "—"),
                ("Pages",   meta.get("pages","—") or "—"),
            ]
            for i, (label, val) in enumerate(meta_items):
                with cols_meta[i % 3]:
                    st.markdown(f"""<div style="background:#f5f0e8;border:1px solid #d4c9b5;
                         border-radius:8px;padding:0.6rem 0.8rem;margin-bottom:0.5rem;">
                      <div style="font-size:0.68rem;color:#9a8a7a;text-transform:uppercase;letter-spacing:0.8px;">{label}</div>
                      <div style="font-size:0.85rem;font-weight:600;color:#1a2e1b;margin-top:0.2rem;">{_html.escape(str(val))}</div>
                    </div>""", unsafe_allow_html=True)

            # Citation output
            st.markdown(f'<p style="color:#3a8c3f;font-weight:700;margin-top:0.5rem;">{cite_style} Citation</p>', unsafe_allow_html=True)
            # Render citation in a readable box
            st.markdown(f"""<div style="background:white;border:1.5px solid #3a8c3f;border-radius:10px;
                 padding:1rem 1.3rem;font-family:'DM Sans',sans-serif;font-size:0.92rem;
                 line-height:1.8;color:#1a2e1b;white-space:pre-wrap;">{_html.escape(cit)}</div>""",
                unsafe_allow_html=True)

            c_copy, c_add = st.columns([1,1])
            with c_copy:
                make_copy_btn("citation-out", cit, "📋 Copy Citation")
            with c_add:
                if st.button("➕ Add to List", use_container_width=True, key="add_to_list"):
                    entry = dict(meta)
                    entry["citation"] = cit
                    entry["style"]    = cite_style
                    st.session_state.all_citations.append(entry)
                    st.success(f"✅ Added! ({len(st.session_state.all_citations)} in list)")

        # Citation list
        if st.session_state.all_citations:
            st.markdown("---")
            st.markdown(f'<p style="color:#3a8c3f;font-weight:700;">📋 Citation List ({len(st.session_state.all_citations)} entries)</p>', unsafe_allow_html=True)

            # Style selector for re-formatting all
            new_style_all = st.selectbox("Re-format all as:", list(CITATION_STYLES.keys()),
                                          key="bulk_style")
            all_text = ""
            for i, entry in enumerate(st.session_state.all_citations, 1):
                formatted = format_citation(entry, CITATION_STYLES[new_style_all])
                all_text += f"{i}. {formatted}\n\n"
                st.markdown(f"""<div style="background:white;border:1px solid #d4c9b5;border-radius:8px;
                     padding:0.7rem 1rem;margin-bottom:0.4rem;font-size:0.85rem;color:#1a2e1b;
                     line-height:1.7;">
                  <span style="color:#3a8c3f;font-weight:700;">{i}.</span> {_html.escape(formatted)}
                </div>""", unsafe_allow_html=True)

            bc1, bc2 = st.columns([1,1])
            with bc1:
                make_copy_btn("all-citations", all_text, "📋 Copy All Citations")
            with bc2:
                if st.button("🗑️ Clear List", use_container_width=True, key="clear_citations"):
                    st.session_state.all_citations = []
                    st.rerun()

            st.download_button(
                "⬇️ Download Citations (.txt)",
                data=all_text, file_name="citations.txt",
                mime="text/plain", use_container_width=True
            )

    # ════════════════════════════════════════════════════════════════════════
    # R-TAB 2 — ARTICLE SUMMARISER
    # ════════════════════════════════════════════════════════════════════════
    with r_tab2:
        st.markdown('<div class="card-title">📄 Research Article Summariser</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Paste article text or upload a PDF to extract: Title · Authors · Year · Objectives · Methodology · Major Findings</p>', unsafe_allow_html=True)

        if "article_summary" not in st.session_state: st.session_state.article_summary = {}

        src_tab_a, src_tab_b = st.tabs(["✏️ Paste Text", "📎 Upload PDF"])

        with src_tab_a:
            article_text = st.text_area("Article text", height=250,
                placeholder="Paste the full abstract or body of the research article…",
                label_visibility="collapsed", key="article_paste")
            summ_btn_a = st.button("🔬 Summarise Article", type="primary",
                                    disabled=(not article_text.strip() or not groq_key),
                                    key="summ_paste")
            if not groq_key:
                st.caption("🔑 Add OpenRouter API key in Admin panel to enable.")

        with src_tab_b:
            if PDF_OK:
                uploaded_pdf = st.file_uploader("Upload PDF", type=["pdf"],
                                                 label_visibility="collapsed", key="article_pdf")
                summ_btn_b = st.button("🔬 Summarise PDF", type="primary",
                                        disabled=(uploaded_pdf is None or not groq_key),
                                        key="summ_pdf")
                if not groq_key:
                    st.caption("🔑 Add OpenRouter API key in Admin panel to enable.")
            else:
                st.warning("pdfplumber not installed. Add to requirements.txt.")
                uploaded_pdf = None
                summ_btn_b  = False

        # Process
        if (summ_btn_a and article_text.strip()) or (summ_btn_b and uploaded_pdf):
            text_to_summarise = ""
            if summ_btn_b and uploaded_pdf:
                with st.spinner("Extracting PDF text…"):
                    text_to_summarise = extract_pdf_text(uploaded_pdf)
                if text_to_summarise.startswith("[PDF extraction error"):
                    st.error(text_to_summarise)
                    text_to_summarise = ""
            else:
                text_to_summarise = article_text

            if text_to_summarise and groq_key:
                with st.spinner("Analysing article with AI…"):
                    try:
                        summary = summarise_article(groq_key, model_choice, text_to_summarise)
                        st.session_state.article_summary = summary
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ {e}")

        # Display summary
        s = st.session_state.article_summary
        if s and "error" not in s:
            import html as _html
            st.markdown("<br>", unsafe_allow_html=True)

            # Header info
            hc1, hc2 = st.columns([3, 1])
            with hc1:
                st.markdown(f"""<div style="background:white;border:1.5px solid #3a8c3f;
                     border-radius:12px;padding:1.2rem 1.5rem;">
                  <div style="font-family:'Playfair Display',serif;font-size:1.1rem;
                       font-weight:700;color:#1a2e1b;margin-bottom:0.5rem;">
                    {_html.escape(s.get("title","Unknown Title"))}</div>
                  <div style="font-size:0.85rem;color:#5a6a7a;">
                    <b>Authors:</b> {_html.escape(", ".join(s.get("authors",[])) or "—")}</div>
                  <div style="font-size:0.85rem;color:#5a6a7a;margin-top:0.2rem;">
                    <b>Journal:</b> {_html.escape(s.get("journal","—"))} &nbsp;|&nbsp;
                    <b>Year:</b> {_html.escape(s.get("year","—"))}
                    {(" &nbsp;|&nbsp; <b>DOI:</b> " + _html.escape(s.get("doi",""))) if s.get("doi") else ""}
                  </div>
                </div>""", unsafe_allow_html=True)
            with hc2:
                kw = s.get("keywords",[])
                if kw:
                    kw_html = " ".join(f'<span style="background:#f5f0e8;border:1px solid #d4c9b5;border-radius:12px;padding:0.2rem 0.6rem;font-size:0.72rem;color:#5a6a7a;">{_html.escape(k)}</span>' for k in kw)
                    st.markdown(f'<div style="line-height:2.2;">{kw_html}</div>', unsafe_allow_html=True)

            # Three info boxes
            bx1, bx2, bx3 = st.columns(3)
            box_items = [
                ("🎯 Main Objectives", "main_objectives", "#1a3a2e", "#6fcf97"),
                ("⚗️ Methodology",     "methodology",     "#1a0a2e", "#b87ae8"),
                ("🔍 Limitations",     "limitations",     "#2d1a0a", "#e8a87a"),
            ]
            for col, (title_b, key_b, bg_b, ac_b) in zip([bx1,bx2,bx3], box_items):
                val = s.get(key_b,"") or "Not specified."
                col.markdown(f"""<div style="background:{bg_b};border:1px solid {ac_b}44;
                     border-radius:10px;padding:1rem;height:160px;overflow-y:auto;">
                  <div style="font-size:0.78rem;font-weight:700;color:{ac_b};
                       text-transform:uppercase;letter-spacing:0.8px;margin-bottom:0.5rem;">{title_b}</div>
                  <div style="font-size:0.85rem;color:rgba(255,255,255,0.85);line-height:1.6;">
                    {_html.escape(str(val))}</div>
                </div>""", unsafe_allow_html=True)

            # Major findings
            st.markdown('<p style="color:#3a8c3f;font-weight:700;margin-top:1rem;">📊 Major Findings</p>', unsafe_allow_html=True)
            findings = s.get("major_findings",[])
            if isinstance(findings, list):
                for i, f in enumerate(findings, 1):
                    st.markdown(f"""<div style="background:white;border-left:4px solid #3a8c3f;
                         border-radius:0 8px 8px 0;padding:0.6rem 1rem;margin-bottom:0.4rem;
                         font-size:0.88rem;color:#1a2e1b;line-height:1.6;">
                      <b style="color:#3a8c3f;">{i}.</b> {_html.escape(str(f))}
                    </div>""", unsafe_allow_html=True)
            else:
                st.markdown(f'<div style="font-size:0.88rem;color:#1a2e1b;">{_html.escape(str(findings))}</div>', unsafe_allow_html=True)

            # Conclusion
            if s.get("conclusion"):
                st.markdown(f"""<div style="background:linear-gradient(135deg,#1a3a2e,#2d4a1e);
                     border-radius:10px;padding:1rem 1.3rem;margin-top:0.5rem;">
                  <div style="font-size:0.75rem;font-weight:700;color:#6fcf97;
                       text-transform:uppercase;letter-spacing:0.8px;margin-bottom:0.4rem;">
                    ✅ Conclusion</div>
                  <div style="font-size:0.88rem;color:rgba(255,255,255,0.85);line-height:1.6;">
                    {_html.escape(s.get("conclusion",""))}</div>
                </div>""", unsafe_allow_html=True)

            # Actions
            ac1, ac2 = st.columns([1,1])
            summary_text = f"""ARTICLE SUMMARY
{'='*50}
Title:   {s.get("title","")}
Authors: {", ".join(s.get("authors",[]))}
Year:    {s.get("year","")}
Journal: {s.get("journal","")}
DOI:     {s.get("doi","")}

MAIN OBJECTIVES:
{s.get("main_objectives","")}

METHODOLOGY:
{s.get("methodology","")}

MAJOR FINDINGS:
""" + "\n".join(f"• {f}" for f in (s.get("major_findings",[]) if isinstance(s.get("major_findings",[]),list) else [s.get("major_findings","")])) + f"""

LIMITATIONS:
{s.get("limitations","")}

CONCLUSION:
{s.get("conclusion","")}

KEYWORDS: {", ".join(s.get("keywords",[]))}
"""
            with ac1:
                make_copy_btn("article-summary", summary_text, "📋 Copy Summary")
            with ac2:
                # Add to literature review
                if st.button("➕ Add to Literature Review", use_container_width=True, key="add_to_lit"):
                    if "lit_review_articles" not in st.session_state:
                        st.session_state.lit_review_articles = []
                    st.session_state.lit_review_articles.append(dict(s))
                    st.success(f"✅ Added to Literature Review ({len(st.session_state.lit_review_articles)} articles)")

        elif s and "error" in s:
            st.error(f"❌ {s['error']}")
            if s.get("raw"):
                with st.expander("Raw response"):
                    st.code(s["raw"])

    # ════════════════════════════════════════════════════════════════════════
    # R-TAB 3 — LITERATURE REVIEW
    # ════════════════════════════════════════════════════════════════════════
    with r_tab3:
        st.markdown('<div class="card-title">📚 Literature Review Manager</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Upload multiple PDFs or add articles from the Summariser tab. Download a comprehensive literature review table as Excel.</p>', unsafe_allow_html=True)

        if "lit_review_articles" not in st.session_state:
            st.session_state.lit_review_articles = []

        # Bulk PDF upload
        st.markdown('<p style="color:#3a8c3f;font-weight:700;">📎 Bulk Upload PDFs</p>', unsafe_allow_html=True)
        if PDF_OK:
            uploaded_pdfs = st.file_uploader("Upload PDFs", type=["pdf"],
                                              accept_multiple_files=True,
                                              label_visibility="collapsed", key="lit_pdfs")
            if uploaded_pdfs and groq_key:
                if st.button(f"🔬 Process {len(uploaded_pdfs)} PDF(s)", type="primary", key="process_pdfs"):
                    progress = st.progress(0)
                    for i, pdf_file in enumerate(uploaded_pdfs):
                        st.markdown(f'<span class="wc-badge">Processing: {pdf_file.name}…</span>', unsafe_allow_html=True)
                        text = extract_pdf_text(pdf_file)
                        if not text.startswith("[PDF"):
                            try:
                                summary = summarise_article(groq_key, model_choice, text)
                                if "error" not in summary:
                                    # Tag with filename if no title
                                    if not summary.get("title"):
                                        summary["title"] = pdf_file.name.replace(".pdf","")
                                    st.session_state.lit_review_articles.append(summary)
                            except Exception as e:
                                st.warning(f"⚠️ Could not process {pdf_file.name}: {e}")
                        progress.progress((i+1)/len(uploaded_pdfs))
                    st.success(f"✅ Processed {len(uploaded_pdfs)} article(s). Total in review: {len(st.session_state.lit_review_articles)}")
                    st.rerun()
            elif uploaded_pdfs and not groq_key:
                st.warning("🔑 Add OpenRouter API key in Admin panel to process PDFs.")
        else:
            st.warning("Install pdfplumber: add `pdfplumber` to requirements.txt")

        st.markdown("---")

        # Current articles list
        arts = st.session_state.lit_review_articles
        if arts:
            st.markdown(f'<p style="color:#3a8c3f;font-weight:700;">📋 Articles in Review ({len(arts)})</p>', unsafe_allow_html=True)

            for i, art in enumerate(arts):
                import html as _html
                with st.expander(f"{i+1}. {art.get('title','Untitled')[:70]}… ({art.get('year','?')})"):
                    lc1, lc2, lc3 = st.columns([3,2,1])
                    with lc1:
                        st.markdown(f"""
                        **Authors:** {_html.escape(", ".join(art.get("authors",[]))[:80])}
                        **Journal:** {_html.escape(art.get("journal","—"))}
                        **DOI:** {_html.escape(art.get("doi","—") or "—")}
                        """)
                    with lc2:
                        st.markdown(f"**Objectives:** {art.get('main_objectives','')[:120]}…")
                    with lc3:
                        if st.button("🗑️ Remove", key=f"rm_art_{i}"):
                            st.session_state.lit_review_articles.pop(i)
                            st.rerun()

            st.markdown("<br>", unsafe_allow_html=True)

            # Download Excel
            if XLSX_OK:
                if st.button("📊 Generate Literature Review Excel", type="primary",
                             use_container_width=True, key="gen_excel"):
                    with st.spinner("Building Excel workbook…"):
                        excel_bytes = build_literature_excel(arts)
                    fname = f"literature_review_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
                    st.download_button(
                        label="⬇️  Download Literature Review (.xlsx)",
                        data=excel_bytes,
                        file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        key="dl_excel"
                    )
            else:
                st.warning("Install openpyxl: add `openpyxl` to requirements.txt")

            if st.button("🗑️ Clear All Articles", use_container_width=True, key="clear_lit"):
                st.session_state.lit_review_articles = []
                st.rerun()

        else:
            st.markdown("""<div style="background:white;border:1.5px dashed #d4c9b5;border-radius:10px;
                 min-height:200px;display:flex;align-items:center;justify-content:center;
                 flex-direction:column;gap:0.5rem;color:#9a8a7a;">
              <div style="font-size:2rem;">📚</div>
              <div style="font-size:0.9rem;">No articles yet — upload PDFs or add from the Summariser tab</div>
            </div>""", unsafe_allow_html=True)


    # ════════════════════════════════════════════════════════════════════════
    # R-TAB 4 — LITERATURE REVIEW CHAPTER WRITER
    # ════════════════════════════════════════════════════════════════════════
    with r_tab4:
        import html as _html
        st.markdown('''<div style="background:linear-gradient(135deg,#1a2e1b,#1e5c22);
            border-radius:12px;padding:1rem 1.5rem;margin-bottom:1.2rem;">
          <div style="font-family:'Playfair Display',serif;font-size:1.2rem;font-weight:700;color:#ffffff;">
            ✍️ Literature Review Chapter Writer</div>
          <div style="font-size:0.85rem;color:rgba(255,255,255,0.65);margin-top:0.3rem;">
            Synthesise all articles in your review into a complete, citation-rich scholarly chapter
            — download as Word document (.docx)
          </div>
        </div>''', unsafe_allow_html=True)

        arts_for_writer = st.session_state.get("lit_review_articles", [])

        # ── IMPORT PANEL (3 methods) ────────────────────────────────────────
        with st.expander("📥 Import Articles", expanded=(not arts_for_writer)):
            imp_tab_a, imp_tab_b, imp_tab_c = st.tabs([
                "📊 From Excel",
                "📝 From Word / DOI file",
                "🔗 Paste DOIs or References"
            ])

            # ── TAB A: Excel ─────────────────────────────────────────────────
            with imp_tab_a:
                st.markdown('<p style="color:rgba(255,255,255,0.75);font-size:0.85rem;">'
                            'Upload the <b>.xlsx</b> generated by the Literature Review Manager. '
                            'All columns (title, authors, objectives, methodology, findings) '
                            'are parsed automatically.</p>', unsafe_allow_html=True)
                excel_upload = st.file_uploader(
                    "Upload Excel", type=["xlsx"], key="lr_excel_import",
                    label_visibility="collapsed"
                )
                enrich_toggle = st.checkbox(
                    "🤖 AI-enrich missing fields (objectives, methodology, findings) using Groq",
                    value=True, key="excel_enrich_toggle",
                    help="For articles imported from DOI-only sources, AI fills in objectives, "
                         "methodology and findings. Takes ~2s per article."
                )
                if excel_upload:
                    if st.button("📥 Import from Excel", type="primary", key="import_excel_btn",
                                 use_container_width=True):
                        with st.spinner("Parsing Excel…"):
                            imported = parse_excel_to_articles(excel_upload.read())
                        if imported:
                            # Count how many need enrichment
                            needs_enrichment = [a for a in imported if a.get("_needs_enrichment")]
                            existing = st.session_state.get("lit_review_articles", [])
                            existing_titles = {a.get("title","").lower() for a in existing}
                            new_arts = [a for a in imported
                                        if a.get("title","").lower() not in existing_titles]

                            if enrich_toggle and needs_enrichment and groq_key:
                                prog = st.progress(0, f"AI enriching {len(needs_enrichment)} articles…")
                                enriched_count = 0
                                for idx_e, art in enumerate(new_arts):
                                    if art.get("_needs_enrichment") and art.get("title"):
                                        prog.progress(
                                            (idx_e+1)/len(new_arts),
                                            f"🤖 Enriching {idx_e+1}/{len(new_arts)}: {art['title'][:40]}…"
                                        )
                                        enriched = _ai_summarise_from_meta(
                                            groq_key, model_choice, {
                                                "title":   art.get("title",""),
                                                "authors": art.get("authors",[]),
                                                "year":    art.get("year",""),
                                                "journal": art.get("journal",""),
                                                "doi":     art.get("doi",""),
                                            }
                                        )
                                        if not art.get("main_objectives"):
                                            art["main_objectives"] = enriched["main_objectives"]
                                        if not art.get("methodology"):
                                            art["methodology"] = enriched["methodology"]
                                        if not art.get("major_findings"):
                                            art["major_findings"] = enriched["major_findings"]
                                        if not art.get("keywords"):
                                            art["keywords"] = enriched["keywords"]
                                        if not art.get("limitations"):
                                            art["limitations"] = enriched["limitations"]
                                        art.pop("_needs_enrichment", None)
                                        enriched_count += 1
                                prog.progress(1.0, f"✅ Enriched {enriched_count} articles")
                            else:
                                for art in new_arts:
                                    art.pop("_needs_enrichment", None)

                            st.session_state.lit_review_articles = existing + new_arts
                            st.success(f"✅ Imported {len(new_arts)} article(s) from Excel."
                                       + (f" AI-enriched {len(needs_enrichment)}." if enrich_toggle and needs_enrichment else ""))
                            st.rerun()
                        else:
                            st.error("❌ Could not parse. Make sure it's a Zodha Literature Review export.")

            # ── TAB B: Word / DOI file ───────────────────────────────────────
            with imp_tab_b:
                st.markdown('<p style="color:rgba(255,255,255,0.75);font-size:0.85rem;">'
                            'Upload a <b>Word document (.docx)</b> or <b>text file (.txt)</b> '
                            'containing references, a reference list, or DOIs. '
                            'DOIs are fetched from CrossRef; plain references are parsed by AI.</p>',
                            unsafe_allow_html=True)

                word_file = st.file_uploader(
                    "Upload Word / text file",
                    type=["docx","doc","txt"], key="lr_word_import",
                    label_visibility="collapsed",
                    help="Supports: reference list, bibliography, DOI list, any document with DOIs"
                )
                if word_file:
                    file_bytes_w = word_file.read()
                    dois_found, plain_refs_found, raw_txt = parse_word_references(
                        file_bytes_w, word_file.name
                    )
                    # Show preview
                    col_prev1, col_prev2 = st.columns(2)
                    with col_prev1:
                        st.markdown(f"""<div style="background:#f0f7f0;border-radius:8px;
                             padding:0.7rem;text-align:center;">
                          <div style="font-size:0.7rem;color:#3d5e40;text-transform:uppercase;">DOIs found</div>
                          <div style="font-size:1.6rem;font-weight:700;color:#1e5c22;">{len(dois_found)}</div>
                        </div>""", unsafe_allow_html=True)
                    with col_prev2:
                        st.markdown(f"""<div style="background:#f0f7f0;border-radius:8px;
                             padding:0.7rem;text-align:center;">
                          <div style="font-size:0.7rem;color:#3d5e40;text-transform:uppercase;">Plain refs found</div>
                          <div style="font-size:1.6rem;font-weight:700;color:#1e5c22;">{len(plain_refs_found)}</div>
                        </div>""", unsafe_allow_html=True)

                    if dois_found:
                        with st.expander(f"Preview {len(dois_found)} DOI(s)"):
                            for d in dois_found[:20]:
                                st.code(d, language=None)
                    if plain_refs_found:
                        with st.expander(f"Preview {len(plain_refs_found)} plain reference(s)"):
                            for r in plain_refs_found[:10]:
                                st.markdown(f'<div style="font-size:0.78rem;color:#1a2e1b;margin-bottom:0.3rem;">{r[:120]}</div>',
                                            unsafe_allow_html=True)

                    if dois_found or plain_refs_found:
                        if st.button("🔍 Fetch Metadata & Import",
                                     type="primary", use_container_width=True,
                                     key="import_word_btn",
                                     disabled=not groq_key and not dois_found):
                            prog_bar = st.progress(0, "Starting…")
                            def _prog(i, total, msg):
                                prog_bar.progress(min(1.0, (i+1)/max(total,1)), text=msg)

                            with st.spinner("Fetching metadata and parsing references…"):
                                fetched = fetch_articles_from_dois(
                                    groq_key, model_choice,
                                    dois_found, plain_refs_found, raw_txt,
                                    progress_cb=_prog
                                )
                            prog_bar.progress(1.0, "✅ Done")
                            if fetched:
                                existing = st.session_state.get("lit_review_articles", [])
                                existing_titles = {a.get("title","").lower() for a in existing}
                                new_arts = [a for a in fetched
                                            if a.get("title","").lower() not in existing_titles]
                                st.session_state.lit_review_articles = existing + new_arts
                                st.success(f"✅ Added {len(new_arts)} article(s). "
                                           f"Total: {len(st.session_state.lit_review_articles)}")
                                st.rerun()
                            else:
                                st.warning("⚠️ No articles could be extracted.")
                    else:
                        st.info("No DOIs or reference patterns detected in this file. "
                                "Try a file with a reference list or bibliography section.")

            # ── TAB C: Paste DOIs / References ───────────────────────────────
            with imp_tab_c:
                st.markdown('<p style="color:rgba(255,255,255,0.75);font-size:0.85rem;">'
                            'Paste a list of <b>DOIs</b> (one per line) or a <b>reference list</b> '
                            '(APA, MLA, Harvard, Vancouver etc.). '
                            'DOIs are fetched from CrossRef; plain references are parsed by AI.</p>',
                            unsafe_allow_html=True)

                pasted_refs = st.text_area(
                    "Paste DOIs or references",
                    height=200,
                    label_visibility="collapsed",
                    key="lr_paste_refs",
                    placeholder=(
                        "Paste DOIs (one per line):\n"
                        "10.1016/j.jfineco.2021.01.004\n"
                        "10.1086/261009\n\n"
                        "Or paste a full reference list:\n"
                        "Davis, F.D. (1989). Perceived usefulness, perceived ease of use... MIS Quarterly, 13(3), 319.\n"
                        "Fornell, C., & Larcker, D.F. (1981). Evaluating structural equation models...\n"
                    )
                )

                if pasted_refs.strip():
                    dois_pasted   = extract_dois_from_text(pasted_refs)
                    # Also treat each non-DOI line as a plain ref
                    import re as _re3
                    plain_pasted  = [
                        line.strip() for line in pasted_refs.split("\n")
                        if line.strip() and len(line.strip()) > 30
                        and not _re3.match(r'^10\.\d{4,9}/', line.strip())
                        and "doi.org" not in line.lower()
                    ]
                    st.markdown(f'<span class="wc-badge">DOIs: {len(dois_pasted)} · Plain refs: {len(plain_pasted)}</span>',
                                unsafe_allow_html=True)

                    if st.button("🔍 Fetch & Import", type="primary",
                                 use_container_width=True, key="import_paste_btn",
                                 disabled=(not dois_pasted and not plain_pasted)):
                        prog2 = st.progress(0, "Starting…")
                        def _prog2(i, total, msg):
                            prog2.progress(min(1.0, (i+1)/max(total,1)), text=msg)

                        with st.spinner("Processing references…"):
                            fetched2 = fetch_articles_from_dois(
                                groq_key, model_choice,
                                dois_pasted, plain_pasted, pasted_refs,
                                progress_cb=_prog2
                            )
                        prog2.progress(1.0, "✅ Done")
                        if fetched2:
                            existing = st.session_state.get("lit_review_articles", [])
                            existing_titles = {a.get("title","").lower() for a in existing}
                            new_arts = [a for a in fetched2
                                        if a.get("title","").lower() not in existing_titles]
                            st.session_state.lit_review_articles = existing + new_arts
                            st.success(f"✅ Added {len(new_arts)} article(s). "
                                       f"Total: {len(st.session_state.lit_review_articles)}")
                            st.rerun()
                        else:
                            st.warning("⚠️ Could not extract any articles. Check your input.")

        if not arts_for_writer:
            st.markdown('''<div style="background:white;border:1.5px dashed #b0d4b2;border-radius:10px;
                 min-height:160px;display:flex;align-items:center;justify-content:center;
                 flex-direction:column;gap:0.5rem;color:#3d5e40;padding:2rem;">
              <div style="font-size:2rem;">📚</div>
              <div style="font-size:0.95rem;font-weight:600;">No articles yet</div>
              <div style="font-size:0.85rem;color:#9a8a7a;text-align:center;">
                Import from Excel above, add via <b>Article Summariser</b>, or upload PDFs in
                the <b>Literature Review</b> tab.
              </div>
            </div>''', unsafe_allow_html=True)
        else:
            # ── Configuration row ───────────────────────────────────────────
            st.markdown(f'<p style="color:#ffffff;font-weight:700;">'
                        f'📋 {len(arts_for_writer)} article(s) ready for synthesis</p>',
                        unsafe_allow_html=True)

            # Article list preview
            with st.expander(f"View articles in review ({len(arts_for_writer)})"):
                for idx_a, a in enumerate(arts_for_writer, 1):
                    authors_preview = ", ".join(a.get("authors",[])[:2]) if isinstance(a.get("authors",[]),list) else ""
                    st.markdown(f'<div style="font-size:0.82rem;color:#1a2e1b;margin-bottom:0.2rem;">'
                                f'<b>{idx_a}.</b> {_html.escape(a.get("title","Untitled")[:70])} '
                                f'— <span style="color:#3a8c3f;">{_html.escape(authors_preview)}</span> '
                                f'({_html.escape(str(a.get("year","")))})</div>', unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)

            # ── Settings ────────────────────────────────────────────────────
            lw_c1, lw_c2 = st.columns([2, 1])
            with lw_c1:
                lw_topic = st.text_area(
                    "Chapter focus / research topic",
                    height=90,
                    placeholder="e.g. The role of financial literacy in women's investment behaviour "
                                "in emerging economies — focusing on India, Kerala, SHG context",
                    key="lw_topic", label_visibility="visible"
                )
                lw_extra = st.text_area(
                    "Additional instructions (optional)",
                    height=70,
                    placeholder="e.g. Organise thematically into: (1) Financial Literacy & Investment, "
                                "(2) Gender & Finance, (3) Methodology gaps. Emphasise Indian context. "
                                "Include a paragraph on research gaps.",
                    key="lw_extra", label_visibility="visible"
                )
            with lw_c2:
                lw_citation = st.selectbox(
                    "Citation style",
                    ["APA 7th", "MLA 9th", "Chicago 17th", "Harvard", "Vancouver", "IEEE"],
                    key="lw_citation"
                )
                lw_words = st.select_slider(
                    "Target word count",
                    options=[800, 1000, 1200, 1500, 2000, 2500, 3000],
                    value=1500, key="lw_words"
                )
                st.markdown(f'<div style="font-size:0.78rem;color:rgba(255,255,255,0.6);">'
                            f'~{lw_words} words · {lw_citation} · {len(arts_for_writer)} sources</div>',
                            unsafe_allow_html=True)

            lw_btn = st.button(
                f"✍️ Write Literature Review Chapter ({lw_words} words · {lw_citation})",
                type="primary", use_container_width=True, key="lw_btn",
                disabled=(not lw_topic.strip() or not groq_key)
            )
            if not groq_key:
                st.caption("⚙️ Service key not configured.")

            # ── Generate ────────────────────────────────────────────────────
            if lw_btn and lw_topic.strip() and groq_key:
                with st.spinner(f"Writing literature review chapter — synthesising {len(arts_for_writer)} articles…"):
                    try:
                        chapter = write_literature_review_chapter(
                            groq_key, model_choice, arts_for_writer,
                            lw_citation, lw_topic, lw_words, lw_extra
                        )
                        st.session_state.lit_chapter_text  = chapter
                        st.session_state.lit_chapter_topic = lw_topic
                        st.session_state.lit_chapter_style = lw_citation
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ {e}")

            # ── Display chapter ─────────────────────────────────────────────
            chapter_text = st.session_state.get("lit_chapter_text", "")
            if chapter_text:
                st.markdown("---")

                # Stats
                words_in_chapter = len(chapter_text.split())
                # Split body vs references
                ref_split = chapter_text.upper().find("REFERENCES")
                body_text = chapter_text[:ref_split] if ref_split > 0 else chapter_text
                ref_text  = chapter_text[ref_split:] if ref_split > 0 else ""
                ref_count = ref_text.count("\n") - 2 if ref_text else 0

                sc1, sc2, sc3 = st.columns(3)
                for col_s, label_s, val_s in [
                    (sc1, "Total words",     f"{words_in_chapter:,}"),
                    (sc2, "References",      f"~{max(0,ref_count)} entries"),
                    (sc3, "Citation style",  st.session_state.get("lit_chapter_style","APA 7th")),
                ]:
                    col_s.markdown(f'''<div style="background:#f0f7f0;border:1px solid #b0d4b2;
                         border-radius:8px;padding:0.7rem;text-align:center;">
                      <div style="font-size:0.7rem;color:#3d5e40;text-transform:uppercase;letter-spacing:0.8px;">{label_s}</div>
                      <div style="font-family:'Playfair Display',serif;font-size:1.3rem;font-weight:700;color:#1a2e1b;">{val_s}</div>
                    </div>''', unsafe_allow_html=True)

                st.markdown("<br>", unsafe_allow_html=True)

                # Chapter preview
                st.markdown('<p style="color:#ffffff;font-weight:700;">📄 Chapter Preview</p>', unsafe_allow_html=True)
                preview_lines = chapter_text.split("\n")[:60]
                preview_text  = "\n".join(preview_lines)
                if len(chapter_text.split("\n")) > 60:
                    preview_text += "\n\n[… scroll down in the full document …]"
                preview_style = (
                    'background:white;border:1.5px solid #3a8c3f;border-radius:10px;'
                    'padding:1.5rem 1.8rem;font-family:"Times New Roman",serif;font-size:0.92rem;'
                    'line-height:1.85;color:#1a2e1b;max-height:500px;overflow-y:auto;white-space:pre-wrap;'
                )
                st.markdown(
                    f'<div style="{preview_style}">' + _html.escape(preview_text) + '</div>',
                    unsafe_allow_html=True
                )

                # ── Download buttons ──────────────────────────────────────
                st.markdown("<br>", unsafe_allow_html=True)
                dl1, dl2, dl3 = st.columns([1, 1, 1])

                # TXT download
                with dl1:
                    fname_txt = f"literature_review_{datetime.now().strftime('%Y%m%d_%H%M')}.txt"
                    st.download_button(
                        "⬇️ Download as .txt",
                        data=chapter_text.encode("utf-8"),
                        file_name=fname_txt,
                        mime="text/plain",
                        use_container_width=True, key="dl_lr_txt"
                    )

                # Word download
                with dl2:
                    with st.spinner("Building Word document…"):
                        docx_bytes = build_literature_review_docx(
                            chapter_text,
                            st.session_state.get("lit_chapter_topic","Literature Review"),
                            st.session_state.get("lit_chapter_style","APA 7th"),
                            len(arts_for_writer)
                        )
                    fname_docx = f"literature_review_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                    if isinstance(docx_bytes, bytes) and docx_bytes[:4] == b'PK\x03\x04':
                        st.download_button(
                            "⬇️ Download as .docx",
                            data=docx_bytes,
                            file_name=fname_docx,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True, key="dl_lr_docx"
                        )
                    else:
                        # Fallback: txt
                        st.download_button(
                            "⬇️ Download as .txt (install python-docx for .docx)",
                            data=docx_bytes,
                            file_name=fname_txt,
                            mime="text/plain",
                            use_container_width=True, key="dl_lr_txt_fb"
                        )

                # Copy
                with dl3:
                    make_copy_btn("lr-chapter", chapter_text, "📋 Copy Text")

                # Regenerate
                if st.button("🔄 Regenerate Chapter", use_container_width=True, key="lw_regen"):
                    st.session_state.lit_chapter_text = ""
                    st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# COMPARISON PANEL
# ══════════════════════════════════════════════════════════════════════════════
if scores_in and scores_out:
    st.markdown("---")
    st.markdown('<div class="card-title" style="font-size:1.1rem;">📊 Before vs After</div>', unsafe_allow_html=True)
    delta=scores_out["humanness"]-scores_in["humanness"]  # style-aware scores
    delta_sign="+" if delta>=0 else ""
    delta_color="#6fcf97" if delta>=0 else "#e87a7a"
    st.markdown(f"""<div class="improvement-banner">
      <div class="score-delta" style="color:{delta_color};">{delta_sign}{delta:.1f}</div>
      <div class="score-desc"><b>Humanness Score Change</b><br>
        Before: <b>{scores_in['humanness']}</b> → After: <b>{scores_out['humanness']}</b><br>
        {style} · {intensity} · {OPENROUTER_MODELS.get(model_choice, model_choice.split('/')[-1])}</div>
    </div>""", unsafe_allow_html=True)
    metrics=[("Humanness","humanness","Higher = more natural"),
             ("Flesch","flesch","Higher = easier"),
             ("Lexical Div","ttr","Higher = richer vocab"),
             ("Rhythm Var","sl_variation","Higher = more varied"),
             ("Burstiness","burstiness","Higher = human rhythm"),
             ("Passive Voice","passive_score","Higher = active voice"),
             ("Transitions","transition_score","Higher = better flow"),
             ("Grade Level","grade_level","Lower = accessible")]
    c1,c2,c3,c4=st.columns(4)
    cols=[c1,c2,c3,c4]
    for idx,(label,key,hint) in enumerate(metrics):
        vb=scores_in.get(key,0); va=scores_out.get(key,0); d=va-vb
        sign="+" if d>=0 else ""
        cd="#6fcf97" if d>=0 else "#e87a7a"
        if key=="grade_level": cd="#6fcf97" if d<=0 else "#e87a7a"
        with cols[idx%4]:
            st.markdown(f"""<div style="background:white;border:1px solid #d4c9b5;border-radius:10px;
                  padding:0.8rem;margin-bottom:0.7rem;text-align:center;">
              <div style="font-size:0.68rem;color:#5a6a7a;text-transform:uppercase;letter-spacing:0.8px;margin-bottom:0.4rem;">{label}</div>
              <div style="display:flex;justify-content:space-around;align-items:center;">
                <div><div style="font-size:1.2rem;font-weight:700;color:#1a2e1b;">{vb}</div>
                     <div style="font-size:0.65rem;color:#9a8a7a;">Before</div></div>
                <div style="color:#3a8c3f;">→</div>
                <div><div style="font-size:1.2rem;font-weight:700;color:#1a2e1b;">{va}</div>
                     <div style="font-size:0.65rem;color:#9a8a7a;">After</div></div>
              </div>
              <div style="font-size:0.8rem;font-weight:700;color:{cd};margin-top:0.3rem;">{sign}{d:.1f}</div>
              <div style="font-size:0.63rem;color:#9a8a7a;margin-top:0.1rem;">{hint}</div>
            </div>""", unsafe_allow_html=True)
    dc,_,_ = st.columns([1,1,1])
    with dc:
        st.download_button("⬇️ Download Output", data=st.session_state.output_text,
                           file_name="humanized_output.txt", mime="text/plain",
                           use_container_width=True)


# ══════════════════════════════════════════════════════════════════════════════
# TAB 5 — STATISTICS SUITE (Layer 3)
# ══════════════════════════════════════════════════════════════════════════════
with tab5:
    import html as _html

    st.markdown("""<div style="background:linear-gradient(135deg,#1a2e1b,#1e5c22);border-radius:12px;
        padding:1rem 1.5rem;margin-bottom:1.2rem;">
      <div style="font-family:'Playfair Display',serif;font-size:1.3rem;font-weight:700;color:#fff;">
        📊 Statistics & Methodology Suite</div>
      <div style="font-size:0.85rem;color:rgba(255,255,255,0.65);margin-top:0.3rem;">
        SPSS/AMOS output interpretation · Methodology recommender · Hypothesis generator · Variable operationaliser
      </div>
    </div>""", unsafe_allow_html=True)

    s1, s2, s3, s4 = st.tabs([
        "📉 SPSS Interpreter",
        "🧪 Methodology Recommender",
        "🔬 Hypothesis Generator",
        "📏 Variable Operationaliser"
    ])

    # ── S1: SPSS OUTPUT INTERPRETER ─────────────────────────────────────────
    with s1:
        st.markdown('<div class="card-title">📉 SPSS / AMOS / SmartPLS Output Interpreter</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Paste raw statistical output — regression tables, SEM fit indices, factor loadings, ANOVA tables — and get publication-ready Results section prose.</p>', unsafe_allow_html=True)

        spss_src_a, spss_src_b = st.tabs(["✏️ Paste Output", "📎 Upload File"])
        spss_input = ""
        with spss_src_a:
            spss_input = st.text_area("SPSS output", height=200,
                placeholder="Paste your SPSS/AMOS/SmartPLS output here...\n\nExample:\nModel Summary\nR = .742  R² = .551  Adjusted R² = .538  F(3,146) = 59.87  p < .001\n\nCoefficients:\nBeta  t    Sig.\nFinancial Literacy  .412  6.23  .000\nRisk Perception    .287  4.11  .001",
                label_visibility="collapsed", key="spss_input_paste")
        with spss_src_b:
            spss_file = st.file_uploader(
                "Upload SPSS output file",
                type=["txt","csv","pdf","docx","doc"],
                label_visibility="collapsed", key="spss_upload",
                help="Supported: .txt  .csv  .pdf  .docx"
            )
            if spss_file:
                file_ext = spss_file.name.split(".")[-1].lower()
                if file_ext in ("txt","csv"):
                    spss_input = spss_file.read().decode("utf-8", errors="ignore")
                elif file_ext == "pdf" and PDF_OK:
                    import io as _io
                    with pdfplumber.open(_io.BytesIO(spss_file.read())) as _pdf:
                        spss_input = "\n".join(p.extract_text() or "" for p in _pdf.pages[:20])
                elif file_ext in ("docx","doc"):
                    try:
                        import docx as _docx
                        import io as _io
                        doc = _docx.Document(_io.BytesIO(spss_file.read()))
                        spss_input = "\n".join(p.text for p in doc.paragraphs)
                    except ImportError:
                        # fallback: read raw text bytes
                        spss_input = spss_file.read().decode("utf-8", errors="ignore")
                if spss_input.strip():
                    st.success(f"✅ Loaded: {spss_file.name} ({len(spss_input.split()):,} words)")
                    with st.expander("Preview extracted text"):
                        st.text(spss_input[:600] + ("…" if len(spss_input) > 600 else ""))

        # Use paste if filled, else use file
        spss_wc_col, spss_cl_col = st.columns([3,1])
        with spss_wc_col:
            spss_paste_val = st.session_state.get("spss_input_paste","")
            if spss_paste_val.strip():
                st.markdown(f'<span class="wc-badge">📝 {len(spss_paste_val.split()):,} words</span>', unsafe_allow_html=True)
        with spss_cl_col:
            if st.button("🗑️ Clear", key="clear_spss_btn", use_container_width=True,
                         disabled=(not st.session_state.get("spss_input_paste","").strip())):
                st.session_state.clear_spss = True
                st.rerun()

        spss_final = (st.session_state.get("spss_input_paste","") or spss_input).strip()

        spss_btn = st.button("🔍 Interpret Output", type="primary",
                             disabled=(not spss_final or not groq_key), key="spss_btn")
        if not groq_key: st.caption("⚙️ Service key not configured.")

        spss_res = st.session_state.spss_result
        if spss_btn and spss_final and groq_key:
            with st.spinner("Interpreting statistical output…"):
                try:
                    spss_res = interpret_spss_output(groq_key, model_choice, spss_final)
                    st.session_state.spss_result = spss_res
                    st.rerun()
                except Exception as e: st.error(f"❌ {e}")

        if spss_res and "error" not in spss_res:
            st.markdown(f'<p style="color:#ffffff;font-weight:700;margin-top:0.8rem;">Test detected: {_html.escape(spss_res.get("test_type",""))}</p>', unsafe_allow_html=True)

            col_a, col_b = st.columns(2)
            with col_a:
                st.markdown("""<div style="font-size:0.72rem;font-weight:700;color:#3a8c3f;
                    text-transform:uppercase;letter-spacing:0.8px;margin-bottom:0.4rem;">Plain English</div>""",
                    unsafe_allow_html=True)
                st.markdown(f"""<div style="background:#f0f7f0;border-left:4px solid #3a8c3f;
                    border-radius:0 8px 8px 0;padding:0.8rem 1rem;font-size:0.88rem;
                    color:#1a2e1b;line-height:1.6;">{_html.escape(spss_res.get("plain_english",""))}</div>""",
                    unsafe_allow_html=True)
                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown("""<div style="font-size:0.72rem;font-weight:700;color:#3a8c3f;
                    text-transform:uppercase;letter-spacing:0.8px;margin-bottom:0.4rem;">Interpretation</div>""",
                    unsafe_allow_html=True)
                interp = spss_res.get("interpretation","")
                i_color = "#1a4a1c" if "supported" in interp.lower() else "#4a1a1a"
                i_text  = "#6fcf97" if "supported" in interp.lower() else "#e87a7a"
                st.markdown(f"""<div style="background:{i_color};border-radius:8px;
                    padding:0.7rem 1rem;font-size:0.88rem;color:{i_text};font-weight:600;">
                    {_html.escape(interp)}</div>""", unsafe_allow_html=True)

            with col_b:
                st.markdown("""<div style="font-size:0.72rem;font-weight:700;color:#3a8c3f;
                    text-transform:uppercase;letter-spacing:0.8px;margin-bottom:0.4rem;">Academic prose (paste into Results section)</div>""",
                    unsafe_allow_html=True)
                prose = spss_res.get("academic_prose","")
                st.markdown(f"""<div style="background:white;border:1.5px solid #3a8c3f;
                    border-radius:10px;padding:1rem 1.2rem;font-size:0.88rem;
                    color:#1a2e1b;line-height:1.7;min-height:120px;">{_html.escape(prose)}</div>""",
                    unsafe_allow_html=True)
                make_copy_btn("spss-prose", prose, "📋 Copy Prose")

            if spss_res.get("apa_table"):
                st.markdown('<p style="color:#1e5c22;font-weight:700;margin-top:0.8rem;">APA-format table</p>', unsafe_allow_html=True)
                st.code(spss_res.get("apa_table",""), language=None)

            # Assumptions & limitations
            with st.expander("⚠️ Assumptions to check & limitations"):
                st.markdown(f"**Assumptions:** {_html.escape(spss_res.get('assumptions_check',''))}")
                st.markdown(f"**Limitations:** {_html.escape(spss_res.get('limitations',''))}")

            # Full copy
            full_spss = f"""TEST: {spss_res.get('test_type','')}

PLAIN ENGLISH:
{spss_res.get('plain_english','')}

ACADEMIC PROSE (for Results section):
{spss_res.get('academic_prose','')}

INTERPRETATION:
{spss_res.get('interpretation','')}

ASSUMPTIONS TO CHECK:
{spss_res.get('assumptions_check','')}

LIMITATIONS:
{spss_res.get('limitations','')}"""
            make_copy_btn("spss-full", full_spss, "📋 Copy Full Report")

    # ── S2: METHODOLOGY RECOMMENDER ─────────────────────────────────────────
    with s2:
        st.markdown('<div class="card-title">🧪 Research Methodology Recommender</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Describe your research topic, objectives, and context — get a detailed methodology recommendation with sample size, sampling strategy, and analysis techniques.</p>', unsafe_allow_html=True)

        meth_src_a, meth_src_b = st.tabs(["✏️ Paste Description", "📎 Upload File"])
        meth_input = ""
        with meth_src_a:
            meth_input = st.text_area("Research description", height=170,
                placeholder="Describe your research. Example:\n\nI am studying the impact of financial literacy and digital banking adoption on the investment behaviour of women micro-entrepreneurs in Kerala. I want to understand causal relationships between constructs and test a structural model. My target population is women SHG members across Thrissur district.",
                label_visibility="collapsed", key="meth_input_paste")
        with meth_src_b:
            meth_file = st.file_uploader(
                "Upload research proposal / synopsis",
                type=["txt","pdf","docx","doc"],
                label_visibility="collapsed", key="meth_upload",
                help="Upload a research proposal, synopsis, or concept note (PDF, Word, TXT)"
            )
            if meth_file:
                file_ext_m = meth_file.name.split(".")[-1].lower()
                if file_ext_m == "txt":
                    meth_input = meth_file.read().decode("utf-8", errors="ignore")
                elif file_ext_m == "pdf" and PDF_OK:
                    import io as _io
                    with pdfplumber.open(_io.BytesIO(meth_file.read())) as _pdf:
                        meth_input = "\n".join(p.extract_text() or "" for p in _pdf.pages[:15])
                elif file_ext_m in ("docx","doc"):
                    try:
                        import docx as _docx
                        import io as _io
                        doc = _docx.Document(_io.BytesIO(meth_file.read()))
                        meth_input = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    except ImportError:
                        meth_input = meth_file.read().decode("utf-8", errors="ignore")
                if meth_input.strip():
                    st.success(f"✅ Loaded: {meth_file.name} ({len(meth_input.split()):,} words)")
                    with st.expander("Preview"):
                        st.text(meth_input[:500] + ("…" if len(meth_input) > 500 else ""))

        meth_wc_col, meth_cl_col = st.columns([3,1])
        with meth_wc_col:
            meth_paste_val = st.session_state.get("meth_input_paste","")
            if meth_paste_val.strip():
                st.markdown(f'<span class="wc-badge">📝 {len(meth_paste_val.split()):,} words</span>', unsafe_allow_html=True)
        with meth_cl_col:
            if st.button("🗑️ Clear", key="clear_meth_btn", use_container_width=True,
                         disabled=(not st.session_state.get("meth_input_paste","").strip())):
                st.session_state.clear_meth = True
                st.rerun()

        meth_final = (st.session_state.get("meth_input_paste","") or meth_input).strip()

        meth_btn = st.button("🧪 Recommend Methodology", type="primary",
                             disabled=(not meth_final or not groq_key), key="meth_btn")

        meth_res = st.session_state.method_result
        if meth_btn and meth_final and groq_key:
            with st.spinner("Analysing research design…"):
                try:
                    meth_res = recommend_methodology(groq_key, model_choice, meth_final)
                    st.session_state.method_result = meth_res
                    st.rerun()
                except Exception as e: st.error(f"❌ {e}")

        if meth_res and "error" not in meth_res:
            # Primary recommendation banner
            st.markdown(f"""<div style="background:linear-gradient(135deg,#1a2e1b,#1e5c22);
                border-radius:12px;padding:1rem 1.5rem;margin:0.8rem 0;">
              <div style="font-size:0.75rem;color:rgba(255,255,255,0.5);text-transform:uppercase;letter-spacing:0.8px;">Recommended approach</div>
              <div style="font-family:'Playfair Display',serif;font-size:1.3rem;font-weight:700;color:#4caf50;margin-top:0.2rem;">
                {_html.escape(meth_res.get("recommended_approach",""))} — {_html.escape(meth_res.get("design",""))}</div>
              <div style="font-size:0.88rem;color:rgba(255,255,255,0.75);margin-top:0.5rem;line-height:1.6;">
                {_html.escape(meth_res.get("rationale",""))}</div>
            </div>""", unsafe_allow_html=True)

            mc1, mc2, mc3 = st.columns(3)
            with mc1:
                st.markdown('<p style="color:#ffffff;font-weight:700;font-size:0.82rem;">📊 Analysis techniques</p>', unsafe_allow_html=True)
                for tech in meth_res.get("analysis_techniques",[]):
                    st.markdown(f'<div style="background:#f0f7f0;border-radius:6px;padding:0.35rem 0.7rem;margin-bottom:0.3rem;font-size:0.82rem;color:#1a2e1b;">✓ {_html.escape(tech)}</div>', unsafe_allow_html=True)
            with mc2:
                st.markdown('<p style="color:#ffffff;font-weight:700;font-size:0.82rem;">🔢 Sample & sampling</p>', unsafe_allow_html=True)
                st.markdown(f'<div style="background:#f0f7f0;border-radius:8px;padding:0.7rem;font-size:0.82rem;color:#1a2e1b;line-height:1.6;"><b>Size:</b> {_html.escape(str(meth_res.get("sample_size","")))}<br><b>Strategy:</b> {_html.escape(meth_res.get("sampling_strategy",""))}<br><b>Software:</b> {_html.escape(meth_res.get("software",""))}</div>', unsafe_allow_html=True)
            with mc3:
                st.markdown('<p style="color:#ffffff;font-weight:700;font-size:0.82rem;">📥 Data collection</p>', unsafe_allow_html=True)
                for dc in meth_res.get("data_collection",[]):
                    st.markdown(f'<div style="background:#f0f7f0;border-radius:6px;padding:0.35rem 0.7rem;margin-bottom:0.3rem;font-size:0.82rem;color:#1a2e1b;">• {_html.escape(dc)}</div>', unsafe_allow_html=True)

            if meth_res.get("conceptual_framework_hint"):
                st.markdown(f"""<div style="background:#f0f7f0;border-left:4px solid #3a8c3f;border-radius:0 8px 8px 0;
                    padding:0.8rem 1rem;margin-top:0.8rem;font-size:0.85rem;color:#1a2e1b;line-height:1.6;">
                  <b style="color:#1e5c22;">💡 Conceptual framework hint:</b><br>
                  {_html.escape(meth_res.get("conceptual_framework_hint",""))}</div>""", unsafe_allow_html=True)

            with st.expander("🔄 Alternative methodologies"):
                for alt in meth_res.get("alternative_designs",[]):
                    if isinstance(alt, dict):
                        st.markdown(f"**{_html.escape(alt.get('name',''))}** — {_html.escape(alt.get('pros_cons',''))}")
                    else:
                        st.markdown(f"• {_html.escape(str(alt))}")

            meth_text = f"""RECOMMENDED: {meth_res.get('recommended_approach','')} — {meth_res.get('design','')}

RATIONALE:
{meth_res.get('rationale','')}

ANALYSIS TECHNIQUES:
{chr(10).join('• '+t for t in meth_res.get('analysis_techniques',[]))}

SAMPLE SIZE: {meth_res.get('sample_size','')}
SAMPLING: {meth_res.get('sampling_strategy','')}
SOFTWARE: {meth_res.get('software','')}

DATA COLLECTION:
{chr(10).join('• '+d for d in meth_res.get('data_collection',[]))}

CONCEPTUAL FRAMEWORK HINT:
{meth_res.get('conceptual_framework_hint','')}"""
            make_copy_btn("meth-full", meth_text, "📋 Copy Methodology Report")

    # ── S3: HYPOTHESIS GENERATOR ─────────────────────────────────────────────
    with s3:
        st.markdown('<div class="card-title">🔬 Hypothesis Generator</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Enter your research topic and key variables — get testable H0/H1 hypotheses with theoretical justification and variable classification.</p>', unsafe_allow_html=True)

        hc1, hc2 = st.columns([3, 2])
        with hc1:
            hyp_topic = st.text_area("Research topic", height=100,
                placeholder="e.g. Impact of ESG disclosure quality on firm valuation among BSE 500 companies during BRSR era",
                label_visibility="visible", key="hyp_topic")
        with hc2:
            hyp_vars = st.text_area("Key variables (one per line)", height=100,
                placeholder="ESG disclosure score\nFirm valuation (Tobin's Q)\nFirm size\nLeverage\nProfitability (ROA)",
                label_visibility="visible", key="hyp_vars")

        hyp_wc_col, hyp_cl_col = st.columns([3,1])
        with hyp_wc_col:
            if hyp_topic.strip():
                st.markdown(f'<span class="wc-badge">📝 {len(hyp_topic.split()):,} words</span>', unsafe_allow_html=True)
        with hyp_cl_col:
            if st.button("🗑️ Clear", key="clear_hyp_btn", use_container_width=True,
                         disabled=(not hyp_topic.strip())):
                st.session_state.clear_hyp = True
                st.rerun()
        hyp_btn = st.button("🔬 Generate Hypotheses", type="primary",
                            disabled=(not hyp_topic.strip() or not groq_key), key="hyp_btn")

        hyp_res = st.session_state.hyp_result
        if hyp_btn and hyp_topic.strip() and groq_key:
            with st.spinner("Developing research hypotheses…"):
                try:
                    hyp_res = generate_hypotheses(groq_key, model_choice, hyp_topic, hyp_vars)
                    st.session_state.hyp_result = hyp_res
                    st.rerun()
                except Exception as e: st.error(f"❌ {e}")

        if hyp_res and "error" not in hyp_res:
            st.markdown(f'<p style="color:#ffffff;font-weight:700;margin-top:0.8rem;">Research question: {_html.escape(hyp_res.get("research_question",""))}</p>', unsafe_allow_html=True)

            # Variable classification
            vc = hyp_res.get("variable_classification", {})
            if vc:
                vcols = st.columns(4)
                for col, (vtype, label) in zip(vcols, [("independent","Independent"),("dependent","Dependent"),("moderating","Moderating"),("mediating","Mediating")]):
                    vars_list = vc.get(vtype, [])
                    with col:
                        items_html = "".join(f'<div style="font-size:0.78rem;color:#1a2e1b;margin-bottom:0.2rem;">• {_html.escape(str(v))}</div>' for v in vars_list)
                        st.markdown(f'<div style="background:#f0f7f0;border-radius:8px;padding:0.6rem;"><div style="font-size:0.7rem;font-weight:700;color:#3a8c3f;text-transform:uppercase;margin-bottom:0.3rem;">{label}</div>{items_html}</div>', unsafe_allow_html=True)

            # Hypotheses
            h0s = hyp_res.get("null_hypotheses", [])
            h1s = hyp_res.get("alternate_hypotheses", [])
            rationale = hyp_res.get("directional_rationale", [])
            st.markdown('<p style="color:#ffffff;font-weight:700;margin-top:1rem;">Hypotheses</p>', unsafe_allow_html=True)

            hyp_text_parts = [f"RESEARCH QUESTION:\n{hyp_res.get('research_question','')}\n"]
            for i, (h0, h1) in enumerate(zip(h0s, h1s), 1):
                rat = rationale[i-1] if i-1 < len(rationale) else ""
                if isinstance(rat, dict): rat = rat.get("justification", str(rat))
                st.markdown(f"""<div style="background:white;border:1px solid #b0d4b2;
                    border-radius:10px;padding:0.8rem 1.1rem;margin-bottom:0.6rem;">
                  <div style="display:flex;gap:0.5rem;align-items:flex-start;">
                    <span style="background:#f0f7f0;color:#1e5c22;border-radius:20px;
                           padding:0.1rem 0.6rem;font-size:0.72rem;font-weight:700;
                           white-space:nowrap;margin-top:0.1rem;">H{i}</span>
                    <div style="flex:1;">
                      <div style="font-size:0.82rem;color:#9a8a7a;margin-bottom:0.2rem;">
                        <b>H0:</b> {_html.escape(str(h0))}</div>
                      <div style="font-size:0.85rem;color:#1a2e1b;font-weight:600;margin-bottom:0.3rem;">
                        <b>H{i}:</b> {_html.escape(str(h1))}</div>
                      <div style="font-size:0.78rem;color:#3d5e40;line-height:1.5;font-style:italic;">
                        {_html.escape(str(rat))}</div>
                    </div>
                  </div>
                </div>""", unsafe_allow_html=True)
                hyp_text_parts.append(f"H{i}:\nH0: {h0}\nH1: {h1}\nJustification: {rat}\n")

            make_copy_btn("hyp-full", "\n".join(hyp_text_parts), "📋 Copy All Hypotheses")

    # ── S4: VARIABLE OPERATIONALISER ─────────────────────────────────────────
    with s4:
        st.markdown('<div class="card-title">📏 Variable Operationaliser & Scale Developer</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Enter a research variable and context — get a construct definition, validated scale items, reliability guidance, and measurement source.</p>', unsafe_allow_html=True)

        vc1, vc2 = st.columns([1, 2])
        with vc1:
            var_name  = st.text_input("Variable name", placeholder="e.g. Financial Literacy", key="var_name")
            var_ctx   = st.text_area("Research context", height=120,
                placeholder="e.g. Measuring financial literacy of women micro-entrepreneurs in Kerala for a study on investment behaviour",
                label_visibility="visible", key="var_ctx")
            var_btn = st.button("📏 Operationalise", type="primary",
                                disabled=(not var_name.strip() or not groq_key), key="var_btn")

        var_res = st.session_state.var_result
        if var_btn and var_name.strip() and groq_key:
            with st.spinner(f"Operationalising '{var_name}'…"):
                try:
                    var_res = operationalise_variable(groq_key, model_choice, var_name, var_ctx)
                    st.session_state.var_result = var_res
                    st.rerun()
                except Exception as e: st.error(f"❌ {e}")

        with vc2:
            if var_res and "error" not in var_res:
                st.markdown(f"""<div style="background:#f0f7f0;border-radius:10px;padding:0.9rem 1.1rem;margin-bottom:0.8rem;">
                  <div style="font-size:0.72rem;font-weight:700;color:#3a8c3f;text-transform:uppercase;letter-spacing:0.8px;">Construct definition</div>
                  <div style="font-size:0.88rem;color:#1a2e1b;line-height:1.6;margin-top:0.3rem;">{_html.escape(var_res.get("construct_definition",""))}</div>
                  <div style="font-size:0.78rem;color:#3d5e40;margin-top:0.5rem;"><b>Source:</b> {_html.escape(var_res.get("scale_source",""))} &nbsp;|&nbsp; <b>Response:</b> {_html.escape(var_res.get("response_format",""))}</div>
                </div>""", unsafe_allow_html=True)

                st.markdown('<p style="color:#ffffff;font-weight:700;font-size:0.85rem;">Survey items (Likert scale)</p>', unsafe_allow_html=True)
                items     = var_res.get("measurement_items", [])
                rev_coded = var_res.get("reverse_coded_items", [])
                var_text_parts = [f"VARIABLE: {var_res.get('variable','')}\n",
                                   f"DEFINITION:\n{var_res.get('construct_definition','')}\n",
                                   f"SOURCE: {var_res.get('scale_source','')}\n",
                                   f"RESPONSE: {var_res.get('response_format','')}\n\nSURVEY ITEMS:"]
                for i, item in enumerate(items, 1):
                    is_rev = i in rev_coded or str(i) in rev_coded
                    badge  = ' <span style="background:#e87a7a22;color:#e87a7a;border-radius:4px;padding:0.05rem 0.4rem;font-size:0.65rem;">R</span>' if is_rev else ''
                    st.markdown(f"""<div style="background:white;border:1px solid #b0d4b2;border-radius:8px;
                        padding:0.5rem 0.9rem;margin-bottom:0.3rem;font-size:0.85rem;color:#1a2e1b;
                        display:flex;gap:0.6rem;align-items:center;">
                      <span style="color:#3a8c3f;font-weight:700;min-width:20px;">{i}.</span>
                      {_html.escape(str(item))}{badge}
                    </div>""", unsafe_allow_html=True)
                    var_text_parts.append(f"{i}. {item}{' [REVERSE CODED]' if is_rev else ''}")

                st.markdown(f'<div style="font-size:0.78rem;color:#3d5e40;margin-top:0.5rem;">📋 Reliability: {_html.escape(var_res.get("reliability_note",""))}</div>', unsafe_allow_html=True)
                make_copy_btn("var-full", "\n".join(var_text_parts), "📋 Copy Scale Items")


# ══════════════════════════════════════════════════════════════════════════════
# TAB 6 — PUBLICATION SUITE (Layer 5)
# ══════════════════════════════════════════════════════════════════════════════
with tab6:
    st.markdown("""<div style="background:linear-gradient(135deg,#1a2e1b,#1e5c22);border-radius:12px;
        padding:1rem 1.5rem;margin-bottom:1.2rem;">
      <div style="font-family:'Playfair Display',serif;font-size:1.3rem;font-weight:700;color:#fff;">
        📰 Publication Suite</div>
      <div style="font-size:0.85rem;color:rgba(255,255,255,0.65);margin-top:0.3rem;">
        Journal matcher · Cover letter writer · Reviewer response drafter · Research contribution statement
      </div>
    </div>""", unsafe_allow_html=True)

    import html as _html

    p1, p2, p3, p4 = st.tabs([
        "🎯 Journal Matcher",
        "✉️ Cover Letter Writer",
        "🔁 Reviewer Response",
        "💡 Contribution Statement"
    ])

    # ── P1: JOURNAL MATCHER ──────────────────────────────────────────────────
    with p1:
        st.markdown('<div class="card-title">🎯 Journal Matcher — Scopus / ABDC / UGC CARE</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Paste your abstract and keywords to get ranked journal recommendations with indexing, impact factors, acceptance rates, and submission strategy.</p>', unsafe_allow_html=True)

        jc1, jc2 = st.columns([2, 1])
        with jc1:
            jm_src_a, jm_src_b = st.tabs(["✏️ Paste Abstract", "📎 Upload Paper"])
            jm_abstract = ""
            with jm_src_a:
                jm_abstract = st.text_area("Abstract", height=160,
                    placeholder="Paste your paper abstract here…",
                    label_visibility="collapsed", key="jm_abstract_paste")
            with jm_src_b:
                jm_file = st.file_uploader("Upload paper", type=["txt","pdf","docx"],
                    label_visibility="collapsed", key="jm_file",
                    help="Upload PDF, Word, or TXT — abstract extracted automatically")
                if jm_file:
                    jm_text_extracted = _extract_file_text(jm_file)
                    if jm_text_extracted and not jm_text_extracted.startswith("["):
                        st.success(f"✅ Loaded: {jm_file.name}")
                        with st.expander("Preview"): st.text(jm_text_extracted[:400]+"…")
            jm_wc_col, jm_cl_col = st.columns([3,1])
            with jm_wc_col:
                jm_pv = st.session_state.get("jm_abstract_paste","")
                if jm_pv.strip():
                    st.markdown(f'<span class="wc-badge">📝 {len(jm_pv.split()):,} words</span>', unsafe_allow_html=True)
            with jm_cl_col:
                if st.button("🗑️ Clear", key="clear_jm_btn", use_container_width=True,
                             disabled=(not st.session_state.get("jm_abstract_paste","").strip())):
                    st.session_state.clear_jm = True
                    st.rerun()
            jm_abstract_final = (st.session_state.get("jm_abstract_paste","") or jm_abstract or
                                  (jm_text_extracted if "jm_text_extracted" in dir() and jm_text_extracted else ""))
        with jc2:
            jm_keywords = st.text_input("Keywords (comma-separated)", key="jm_keywords",
                placeholder="financial literacy, women, Kerala, investment")
            jm_field    = st.text_input("Research field", key="jm_field",
                placeholder="e.g. Finance / Management / Social Science")
            jm_btn = st.button("🎯 Find Matching Journals", type="primary",
                               disabled=(not jm_abstract_final.strip() or not groq_key), key="jm_btn")

        jm_res = st.session_state.journal_result
        if jm_btn and jm_abstract.strip() and groq_key:
            with st.spinner("Searching journal database…"):
                try:
                    jm_res = match_journals(groq_key, model_choice, jm_abstract_final, jm_keywords, jm_field)
                    st.session_state.journal_result = jm_res
                    st.rerun()
                except Exception as e: st.error(f"❌ {e}")

        if jm_res and "error" not in jm_res:
            st.markdown(f'<p style="color:#ffffff;font-weight:700;margin-top:0.8rem;">Domain detected: {_html.escape(jm_res.get("research_domain",""))}</p>', unsafe_allow_html=True)

            INDEX_COLORS = {
                "Scopus":   ("#e6f1fb","#0c447c"),
                "ABDC A":   ("#1a2e1b","#6fcf97"),
                "ABDC B":   ("#2d3a1e","#b5d97a"),
                "ABDC C":   ("#3a3a1e","#e8d47a"),
                "ABS":      ("#2d1a3a","#b87ae8"),
                "UGC CARE": ("#3a1a1a","#e87a7a"),
                "Web of Science": ("#1a1a3a","#7ab8e8"),
            }
            jm_text = [f"JOURNAL RECOMMENDATIONS — {jm_res.get('research_domain','')}\n"]
            for j in jm_res.get("recommendations", []):
                idx       = j.get("index","")
                bg, txt   = INDEX_COLORS.get(idx, ("#f0f7f0","#1e5c22"))
                st.markdown(f"""<div style="background:white;border:1px solid #b0d4b2;border-radius:10px;
                    padding:0.9rem 1.1rem;margin-bottom:0.6rem;">
                  <div style="display:flex;justify-content:space-between;align-items:flex-start;gap:1rem;">
                    <div style="flex:1;">
                      <div style="font-weight:700;color:#1a2e1b;font-size:0.92rem;">{_html.escape(j.get('name',''))}</div>
                      <div style="font-size:0.78rem;color:#3d5e40;margin-top:0.1rem;">{_html.escape(j.get('publisher',''))}</div>
                      <div style="font-size:0.82rem;color:#3d5e40;margin-top:0.4rem;line-height:1.5;">{_html.escape(j.get('scope_fit',''))}</div>
                    </div>
                    <div style="text-align:right;min-width:130px;">
                      <span style="background:{bg};color:{txt};border-radius:12px;padding:0.2rem 0.7rem;
                             font-size:0.7rem;font-weight:700;">{_html.escape(idx)}</span>
                      <div style="font-size:0.75rem;color:#5a6a7a;margin-top:0.4rem;">IF: {_html.escape(str(j.get('impact_factor','')))} &nbsp;|&nbsp; {_html.escape(j.get('open_access',''))}</div>
                      <div style="font-size:0.72rem;color:#9a8a7a;">Accept: {_html.escape(str(j.get('acceptance_rate','')))} &nbsp;|&nbsp; ~{_html.escape(str(j.get('turnaround','')))} wks</div>
                    </div>
                  </div>
                </div>""", unsafe_allow_html=True)
                jm_text.append(f"{j.get('name','')} [{idx}] IF:{j.get('impact_factor','')} | {j.get('scope_fit','')}")

            if jm_res.get("strategy_note"):
                st.markdown(f"""<div style="background:#f0f7f0;border-left:4px solid #3a8c3f;border-radius:0 8px 8px 0;
                    padding:0.8rem 1rem;margin-top:0.5rem;font-size:0.85rem;color:#1a2e1b;line-height:1.6;">
                  <b style="color:#1e5c22;">📌 Submission strategy:</b><br>{_html.escape(jm_res.get("strategy_note",""))}
                </div>""", unsafe_allow_html=True)

            if jm_res.get("avoid"):
                with st.expander("⚠️ Journals to avoid"):
                    for j in jm_res.get("avoid",[]):
                        st.markdown(f"• {_html.escape(str(j))}")
            make_copy_btn("jm-full", "\n".join(jm_text), "📋 Copy Journal List")
            if XLSX_OK and jm_res.get("recommendations"):
                from datetime import datetime as _dt
                jm_xl = build_journal_excel(jm_res)
                st.download_button(
                    label="📊 Download as Excel (.xlsx)",
                    data=jm_xl,
                    file_name=f"journal_recommendations_{_dt.now().strftime('%Y%m%d_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="dl_jm_excel"
                )

    # ── P2: COVER LETTER WRITER ──────────────────────────────────────────────
    with p2:
        st.markdown('<div class="card-title">✉️ Journal Cover Letter Writer</div>', unsafe_allow_html=True)
        cl1, cl2 = st.columns([1, 1])
        with cl1:
            cl_title   = st.text_input("Paper title", key="cl_title")
            cl_journal = st.text_input("Target journal", key="cl_journal", placeholder="e.g. Journal of Finance Research")
            cl_authors = st.text_input("Author(s)", key="cl_authors", placeholder="Dr. Niyas N, Dr. …")
            cl_abs_a, cl_abs_b = st.tabs(["✏️ Paste", "📎 Upload"])
            cl_abstract = ""
            with cl_abs_a:
                cl_abstract = st.text_area("Abstract", height=120, key="cl_abstract_paste",
                    placeholder="Paste abstract here…", label_visibility="collapsed")
            with cl_abs_b:
                cl_file = st.file_uploader("Upload paper/abstract", type=["txt","pdf","docx"],
                    label_visibility="collapsed", key="cl_file")
                if cl_file:
                    cl_extracted = _extract_file_text(cl_file)
                    if cl_extracted and not cl_extracted.startswith("["):
                        st.success(f"✅ {cl_file.name}")
            cl_wc_col, cl_cl_col = st.columns([3,1])
            with cl_wc_col:
                cl_pv = st.session_state.get("cl_abstract_paste","")
                if cl_pv.strip():
                    st.markdown(f'<span class="wc-badge">📝 {len(cl_pv.split()):,} words</span>', unsafe_allow_html=True)
            with cl_cl_col:
                if st.button("🗑️ Clear", key="clear_cl_btn", use_container_width=True,
                             disabled=(not st.session_state.get("cl_abstract_paste","").strip())):
                    st.session_state.clear_cl = True
                    st.rerun()
            cl_abstract_final = (st.session_state.get("cl_abstract_paste","") or cl_abstract or
                                  (cl_extracted if "cl_extracted" in dir() and cl_extracted else ""))
            cl_btn = st.button("✉️ Write Cover Letter", type="primary",
                               disabled=(not cl_title.strip() or not cl_abstract_final.strip() or not groq_key), key="cl_btn")
        with cl2:
            cl_res = st.session_state.coverletter_result
            if cl_btn and cl_title.strip() and groq_key:
                with st.spinner("Drafting cover letter…"):
                    try:
                        cl_res = write_cover_letter(groq_key, model_choice, cl_title, cl_abstract_final, cl_journal, cl_authors)
                        st.session_state.coverletter_result = cl_res
                        st.rerun()
                    except Exception as e: st.error(f"❌ {e}")
            if cl_res and "error" not in cl_res:
                letter = cl_res.get("cover_letter","")
                st.markdown(f"""<div style="background:white;border:1.5px solid #3a8c3f;border-radius:10px;
                    padding:1.2rem 1.4rem;font-size:0.88rem;color:#1a2e1b;line-height:1.75;
                    max-height:420px;overflow-y:auto;white-space:pre-wrap;">{_html.escape(letter)}</div>""",
                    unsafe_allow_html=True)
                make_copy_btn("cl-letter", letter, "📋 Copy Cover Letter")
                if cl_res.get("key_contributions"):
                    with st.expander("Key contributions summary"):
                        for kc in cl_res.get("key_contributions",[]):
                            st.markdown(f"• {_html.escape(str(kc))}")
            else:
                st.markdown("""<div style="background:white;border:1.5px dashed #b0d4b2;border-radius:10px;
                    min-height:300px;display:flex;align-items:center;justify-content:center;
                    flex-direction:column;gap:0.5rem;color:#9a8a7a;">
                  <div style="font-size:2rem;">✉️</div>
                  <div style="font-size:0.9rem;">Cover letter will appear here</div>
                </div>""", unsafe_allow_html=True)

    # ── P3: REVIEWER RESPONSE DRAFTER ────────────────────────────────────────
    with p3:
        st.markdown('<div class="card-title">🔁 Reviewer Response Letter Drafter</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Paste the reviewer comments you received. Get a structured, professional point-by-point response letter ready to submit.</p>', unsafe_allow_html=True)

        rc1, rc2 = st.columns([1, 1])
        with rc1:
            rv_abs_a, rv_abs_b = st.tabs(["✏️ Paste Abstract", "📎 Upload Paper"])
            rv_abstract = ""
            with rv_abs_a:
                rv_abstract = st.text_area("Original abstract (for context)", height=100, key="rv_abstract_paste",
                    placeholder="Paste your paper abstract…", label_visibility="collapsed")
            with rv_abs_b:
                rv_paper_file = st.file_uploader("Upload paper", type=["txt","pdf","docx"],
                    label_visibility="collapsed", key="rv_paper_file",
                    help="Upload your paper — abstract extracted automatically")
                if rv_paper_file:
                    rv_paper_text = _extract_file_text(rv_paper_file)
                    if rv_paper_text and not rv_paper_text.startswith("["):
                        st.success(f"✅ {rv_paper_file.name}")
            rv_abstract_final = (st.session_state.get("rv_abstract_paste","") or rv_abstract or
                                  (rv_paper_text if "rv_paper_text" in dir() and rv_paper_text else ""))

            rv_com_a, rv_com_b = st.tabs(["✏️ Paste Comments", "📎 Upload Comments File"])
            rv_comments = ""
            with rv_com_a:
                rv_comments = st.text_area("Reviewer comments", height=180, key="rv_comments_paste",
                    placeholder="Paste all reviewer comments here. Include both Reviewer 1 and Reviewer 2 comments.\n\nExample:\nREVIEWER 1:\n1. The sample size seems insufficient. Please justify.\n2. The literature review lacks recent papers (2020-2024).\n\nREVIEWER 2:\n1. The methodology section needs more detail on the SEM specification.",
                    label_visibility="collapsed")
            with rv_com_b:
                rv_comments_file = st.file_uploader("Upload reviewer comments", type=["txt","pdf","docx"],
                    label_visibility="collapsed", key="rv_comments_file",
                    help="Upload the reviewer decision letter (PDF, Word, TXT)")
                if rv_comments_file:
                    rv_comments_text = _extract_file_text(rv_comments_file)
                    if rv_comments_text and not rv_comments_text.startswith("["):
                        st.success(f"✅ {rv_comments_file.name}")
                        with st.expander("Preview"): st.text(rv_comments_text[:400]+"…")
            rv_wc_col, rv_cl_col = st.columns([3,1])
            with rv_wc_col:
                rv_pv = st.session_state.get("rv_comments_paste","")
                if rv_pv.strip():
                    st.markdown(f'<span class="wc-badge">📝 {len(rv_pv.split()):,} words</span>', unsafe_allow_html=True)
            with rv_cl_col:
                if st.button("🗑️ Clear", key="clear_rv_btn", use_container_width=True,
                             disabled=(not st.session_state.get("rv_comments_paste","").strip())):
                    st.session_state.clear_rv = True
                    st.rerun()
            rv_comments_final = (st.session_state.get("rv_comments_paste","") or rv_comments or
                                  (rv_comments_text if "rv_comments_text" in dir() and rv_comments_text else ""))

            rv_btn = st.button("🔁 Draft Response Letter", type="primary",
                               disabled=(not rv_comments_final.strip() or not groq_key), key="rv_btn")
        with rc2:
            rv_res = st.session_state.reviewer_result
            if rv_btn and rv_comments.strip() and groq_key:
                with st.spinner("Drafting reviewer response…"):
                    try:
                        rv_res = draft_reviewer_response(groq_key, model_choice, rv_comments_final, rv_abstract_final)
                        st.session_state.reviewer_result = rv_res
                        st.rerun()
                    except Exception as e: st.error(f"❌ {e}")

            if rv_res and "error" not in rv_res:
                st.markdown(f"""<div style="background:#f0f7f0;border-radius:8px;padding:0.8rem 1rem;
                    margin-bottom:0.8rem;font-size:0.85rem;color:#1a2e1b;line-height:1.6;">
                  {_html.escape(rv_res.get("response_header",""))}</div>""", unsafe_allow_html=True)

                rv_full = [rv_res.get("response_header",""), ""]
                type_colors_rv = {"grammar":"#e87a7a","spelling":"#e8a87a","comment":"#b0d4b2","point":"#b0d4b2"}
                for resp in rv_res.get("responses", []):
                    r_label   = resp.get("reviewer","")
                    r_comment = resp.get("comment_summary","")
                    r_response= resp.get("response","")
                    r_change  = resp.get("manuscript_change","")
                    st.markdown(f"""<div style="background:white;border:1px solid #b0d4b2;border-radius:10px;
                        padding:0.8rem 1rem;margin-bottom:0.5rem;">
                      <div style="display:flex;gap:0.5rem;align-items:center;margin-bottom:0.4rem;">
                        <span style="background:#f0f7f0;color:#1e5c22;border-radius:4px;
                               padding:0.1rem 0.6rem;font-size:0.7rem;font-weight:700;">{_html.escape(r_label)}</span>
                        <span style="font-size:0.78rem;color:#3d5e40;font-style:italic;">{_html.escape(r_comment)}</span>
                      </div>
                      <div style="font-size:0.85rem;color:#1a2e1b;line-height:1.6;margin-bottom:0.4rem;">{_html.escape(r_response)}</div>
                      <div style="font-size:0.78rem;color:#3a8c3f;"><b>Manuscript change:</b> {_html.escape(r_change)}</div>
                    </div>""", unsafe_allow_html=True)
                    rv_full.extend([f"[{r_label}] {r_comment}", f"Response: {r_response}", f"Change: {r_change}", ""])

                st.markdown(f"""<div style="background:#f0f7f0;border-radius:8px;padding:0.8rem 1rem;
                    font-size:0.85rem;color:#1a2e1b;line-height:1.6;">
                  {_html.escape(rv_res.get("closing_paragraph",""))}</div>""", unsafe_allow_html=True)
                make_copy_btn("rv-full", "\n".join(rv_full), "📋 Copy Response Letter")
            else:
                st.markdown("""<div style="background:white;border:1.5px dashed #b0d4b2;border-radius:10px;
                    min-height:300px;display:flex;align-items:center;justify-content:center;
                    flex-direction:column;gap:0.5rem;color:#9a8a7a;">
                  <div style="font-size:2rem;">🔁</div>
                  <div style="font-size:0.9rem;">Response letter will appear here</div>
                </div>""", unsafe_allow_html=True)

    # ── P4: CONTRIBUTION STATEMENT ───────────────────────────────────────────
    with p4:
        st.markdown('<div class="card-title">💡 Research Contribution Statement Generator</div>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(255,255,255,0.7);font-size:0.88rem;">Generate a publication-ready Research Contribution paragraph for thesis chapters, journal submissions, and conference papers.</p>', unsafe_allow_html=True)

        csc1, csc2 = st.columns([1, 1])
        with csc1:
            cs_title   = st.text_input("Paper / Chapter title", key="cs_title")
            cs_obj     = st.text_area("Research objectives (brief)", height=80, key="cs_obj",
                placeholder="e.g. To examine the effect of ESG disclosure on firm valuation…", label_visibility="visible")
            cs_findings= st.text_area("Key findings (brief)", height=80, key="cs_findings",
                placeholder="e.g. ESG disclosure positively moderates the relationship…", label_visibility="visible")
            cs_context = st.text_input("Study context / setting", key="cs_context",
                placeholder="e.g. BSE 500 firms, India, 2018-2023, BRSR era")
            cs_cl_col1, cs_cl_col2 = st.columns([3,1])
            with cs_cl_col2:
                if st.button("🗑️ Clear All", key="clear_cs_btn", use_container_width=True,
                             disabled=(not cs_title.strip())):
                    st.session_state.clear_cs = True
                    st.rerun()
            cs_btn = st.button("💡 Generate Contribution Statement", type="primary",
                               disabled=(not cs_title.strip() or not groq_key), key="cs_btn")
        with csc2:
            cs_res = st.session_state.contrib_result
            if cs_btn and cs_title.strip() and groq_key:
                with st.spinner("Writing contribution statement…"):
                    try:
                        cs_res = write_contribution_statement(groq_key, model_choice, cs_title, cs_obj, cs_findings, cs_context)
                        st.session_state.contrib_result = cs_res
                        st.rerun()
                    except Exception as e: st.error(f"❌ {e}")

            if cs_res and "error" not in cs_res:
                st.markdown(f"""<div style="background:white;border:1.5px solid #3a8c3f;border-radius:10px;
                    padding:1.1rem 1.3rem;font-size:0.9rem;color:#1a2e1b;line-height:1.75;
                    margin-bottom:0.8rem;">{_html.escape(cs_res.get("full_statement",""))}</div>""",
                    unsafe_allow_html=True)
                make_copy_btn("cs-statement", cs_res.get("full_statement",""), "📋 Copy Statement")

                cs_tabs = st.tabs(["Theoretical","Methodological","Practical","Gap addressed"])
                for tab_cs, key_cs in zip(cs_tabs, ["theoretical_contribution","methodological_contribution","practical_contribution","gap_addressed"]):
                    with tab_cs:
                        st.markdown(f'<div style="font-size:0.85rem;color:#1a2e1b;line-height:1.65;padding:0.5rem 0;">{_html.escape(cs_res.get(key_cs,""))}</div>', unsafe_allow_html=True)
            else:
                st.markdown("""<div style="background:white;border:1.5px dashed #b0d4b2;border-radius:10px;
                    min-height:300px;display:flex;align-items:center;justify-content:center;
                    flex-direction:column;gap:0.5rem;color:#9a8a7a;">
                  <div style="font-size:2rem;">💡</div>
                  <div style="font-size:0.9rem;">Contribution statement will appear here</div>
                </div>""", unsafe_allow_html=True)
